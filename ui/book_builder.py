
"""
Builder principale per la creazione di libri.
"""
import os
import sys
import re
import json
import time
import datetime
import shutil
import logging
import traceback
from datetime import datetime
from pathlib import Path

import gradio as gr
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from analysis.analyzers import analyze_market

from ai_interfaces.browser_manager import (
    setup_browser,
    check_login,
    create_fresh_chat,
    set_connection_status
)
from ai_interfaces.interaction_utils import get_input_box, clear_chat

from framework.crisp_framework    import CRISPFramework
from framework.crisp_utils        import replace_variables
from framework.analysis.market_analysis import filter_legacy_prompt_sections
from framework.formatters         import format_analysis_results_html, save_analysis_to_html

from .cooldown_manager    import CooldownManager
from .chat_manager        import ChatManager
from .database_manager    import DatabaseManager

# Import dei generatori
from generators.crisp_generator    import _generate_book_crisp   as generate_book_crisp
from generators.legacy_generator   import _generate_book_legacy  as generate_book_legacy
from generators.common_generator   import generate_book

# Import dei metodi di analisi
from analysis.analyzers import (
    _analyze_market_crisp,
    _analyze_market_legacy,
    resume_analysis,
    continue_analysis,
    _continue_analysis_crisp,
    complete_analysis,
    _complete_analysis_crisp,
    _complete_analysis_legacy
)

# Directory costante
ANALYSIS_DIR = Path("context_files")

class AIBookBuilder:
    def __init__(self):
        self._analysis_session_id = "AUTO"
        self.cooldown_manager = CooldownManager()
        self.chat_manager     = ChatManager(parent=self)
        set_connection_status(False)
        self.driver           = None
        self.log_history      = []
        self.current_analysis = None
        self.context_dir      = ANALYSIS_DIR
        ANALYSIS_DIR.mkdir(exist_ok=True)
        Path("output").mkdir(exist_ok=True)
        Path("debug").mkdir(exist_ok=True)

        self.db_manager = DatabaseManager(
            project_db_path="crisp_projects.db",
            log_callback=self.add_log
        )

        self.crisp    = CRISPFramework(
            prompt_dir="prompt_crisp",
            project_db_path="crisp_projects.db",
            driver=None
        )
        self.use_crisp = True

        self.book_types = [
            "Manuale (Non-Fiction)",
            "Ricettario",
            "Craft & Hobby",
            "Survival & Outdoor",
            "Test Study"
        ]
        self.markets = {
            "USA": "Amazon.com",
            "Italia": "Amazon.it",
            # …
        }

        # Prompt di analisi default
              
        self.default_analysis_prompt = """1) Analizza la concorrenza su {amazon_url} per la keyword {keyword} nel mercato {market}: elenca i primi 5 risultati Amazon con titolo, sottotitolo, BSR, prezzo, recensioni, formato, keyword usate nei titoli, pattern visivi delle copertine (colori, stile, elementi ricorrenti), call to action e benefit promessi nei titoli o sottotitoli; aggiungi dati da Google Trends, query emergenti e insight dai social (es. video virali, reel, post rilevanti); includi anche eventuali “claim” ricorrenti o promesse implicite nei testi di vendita; concludi con una tabella di sintesi e un commento operativo su cosa domina, cosa manca, e quali pattern emergono; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
2) Valuta la profittabilità e competitività della keyword {keyword} su {amazon_url} nel mercato {market}: considera vendite mensili stimate per range di BSR, prezzo medio per ebook e cartaceo, royalty KDP stimate per formato, numero e qualità dei competitor (copertine, recensioni, struttura, USP), livello di saturazione e difficoltà stimata nel posizionarsi (forza dei top 5 titoli), segnala se sono self-published o con editore; includi 3 tabelle: “BSR vs Vendite e Margini”, “Top 5 Competitor”, “Analisi Competitività” con punteggi; concludi con 3 bullet: "Alta opportunità se…", "Moderata se…", "Bassa se…"; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
3) Analizza i 3 migliori concorrenti per la keyword {keyword} su {amazon_url} nel mercato {market}: mostra per ciascuno titolo, sottotitolo, BSR, recensioni, prezzo, formato, numero di pagine (se disponibile), struttura interna (indice o sezioni visibili), copertina (stile, colori, elementi distintivi), USP percepita, e bonus offerti (esercizi, checklist, link, QR code); includi una tabella comparativa con righe = libri e colonne: BSR, prezzo, recensioni, pagine, bonus, punto di forza percepito; concludi con insight su ciò che li rende forti e ripetuti pattern utili da superare; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
4) Definisci una buyer persona sintetica per {keyword} nel mercato {market}: includi età, professione, livello di istruzione (se rilevante), obiettivi specifici, problemi percepiti, livello di consapevolezza del problema e grado di urgenza nella ricerca di una soluzione; indica dove cerca soluzioni (es. YouTube, libri, social), completa con 3 bullet: "Cosa cerca", "Cosa teme", "Cosa sogna"; assegna un nome simbolico e aggiungi una frase tipo "Vorrei un libro che…", più una possibile frase Google-style che digiterebbe per cercare il libro; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
5) Identifica i principali gap nei libri esistenti su {amazon_url} per {keyword}: analizza recensioni negative (1★–2★) per evidenziare frustrazioni comuni, bisogni insoddisfatti, contenuti confusi o trattati superficialmente; indica almeno 3 aree tematiche mancanti o trascurate e il valore che avrebbero se inserite in un nuovo libro; specifica se si tratta di carenze pratiche (esempi, esercizi), strutturali (ordine, chiarezza), o valoriali (tono, empatia); concludi con una tabella “Gap vs Opportunità” con colonne: Problema segnalato, Frequenza, Opportunità editoriale; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
6)  Genera 3 idee editoriali differenzianti per un nuovo libro su {keyword} nel mercato {market}: per ciascuna definisci titolo provvisorio, angolo editoriale, approccio (pratico, teorico, visuale…), target specifico e una USP sintetica (max 2 righe) che risponda ai bisogni emersi; includi una tabella comparativa con righe = idee e colonne: originalità, potenziale commerciale, compatibilità con la buyer persona, copertura dei gap, ognuna con voto da 1 a 5 stelle e punteggio totale finale; scegli la migliore idea e motiva la scelta sulla base del potenziale commerciale, dell'originalità, della compatibilità con la buyer persona e dei gap riscontrati; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
7) Definisci il tono di voce, lo stile narrativo e la struttura ideale del testo per il libro basato sull’idea editoriale selezionata: specifica persona narrativa, registro, lessico, ritmo e livello di approfondimento, indicando le regole di coerenza stilistica (es. uso della terza persona, assenza di bullet non richiesti, paragrafi sviluppati in blocchi coesi); descrivi come devono essere strutturati apertura, sviluppo e conclusione di ogni capitolo, inclusa l’integrazione fluida di riferimenti storici, pratici e scientifici senza frammentazioni; includi 2–3 esempi di paragrafi scritti nello stile indicato da usare come modello per la stesura del libro; scrivi in {lingua}; concludi con la parola FINE.
8) In base all’idea selezionata, proponi 3 titoli con relativo sottotitolo (titolo + sottotitolo per ciascuna variante), valuta i titoli con punteggio da 1 a 5 stelle in base a chiarezza, potere evocativo, potenziale di vendita, pertinenza tematica e compatibilità con la buyer persona e i gap individuati; scegli il migliore e motiva la scelta, poi crea un bonus testuale in PDF coerente con il titolo scelto e utile alla buyer persona. Genera 3 idee di indice coerenti con il titolo selezionato, scegli e motiva la migliore, poi sviluppa l’indice completo del libro con tutti i capitoli necessari, basandoti anche sulla lunghezza media dei libri concorrenti. L’indice deve essere scritto su una riga per gruppo di sottocapitoli: se un capitolo ha 2–3 sottocapitoli, vanno tutti su una riga; se ha 4–6 sottocapitoli, vanno distribuiti in 2 righe; se ha 7 o più, vanno divisi in 3 righe o più. Non spezzare mai i paragrafi dal proprio sottocapitolo. Ogni riga dell’indice sarà un’unità editoriale da mandare in produzione. Infine, scrivi l’indice completo del bonus scelto; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
 """
     

    # ———————————————————————————————————————————————
    # DELEGA DEI METODI DI ANALISI
    # ———————————————————————————————————————————————

    def analyze_market(self, book_type, keyword, language, market,
                       analysis_prompt=None, use_crisp=None):
        # se l’utente ha selezionato Legacy, forziamo use_crisp=False
        if self.selected_analysis_type == "Legacy":
            use_crisp = False

        return analyze_market(
            self,
            book_type, keyword, language, market,
            analysis_prompt=analysis_prompt,
            use_crisp=use_crisp
        )

    def _analyze_market_crisp(self, book_type, keyword, language, market, selected_phases=None):
        """Delega l'analisi CRISP al modulo esterno."""
        return _analyze_market_crisp(self, book_type, keyword, language, market, selected_phases)

    def _analyze_market_legacy(self, book_type, keyword, language, market, analysis_prompt):
        """Delega l'analisi legacy al modulo esterno."""
        return _analyze_market_legacy(self, book_type, keyword, language, market, analysis_prompt)

    def resume_analysis(self, project_id, selected_phases=None):
        """Delega la ripresa analisi al modulo esterno."""
        return resume_analysis(self, project_id, selected_phases)

    def continue_analysis(self):
        """Delega la continuazione analisi (sceglie CRISP o legacy)."""
        return continue_analysis(self)

    def _continue_analysis_crisp(self):
        """Delega la continuazione specifica CRISP."""
        return _continue_analysis_crisp(self)

    def complete_analysis(self):
        """Delega il completamento dell'analisi (prepara i dati per il libro)."""
        return complete_analysis(self)

    def _complete_analysis_crisp(self):
        """Delega il completamento analisi CRISP 5.0."""
        return _complete_analysis_crisp(self)

    def _complete_analysis_legacy(self):
        """Delega il completamento analisi legacy."""
        from analysis.analyzers import _complete_analysis_legacy
        return _complete_analysis_legacy(self, self)

    def extract_data_for_book_generation(self):
        """
        Estrae i dati pertinenti dall'analisi corrente per la generazione del libro.
        Funziona con diverse fonti: current_analysis, file di contesto, o HTML caricato.
        Non richiede BeautifulSoup - usa solo regex per l'analisi.
        """
        try:
            # Inizializza dizionario per i dati
            book_data = {}
            data_source = "nessuna fonte"
        
            # Strategia 1: Cerca nell'analisi corrente (se disponibile)
            if hasattr(self, 'current_analysis') and self.current_analysis:
                project_data = self.current_analysis.get('project_data', {})
                if project_data:
                    data_source = "current_analysis"
                
                    # Estrai titolo e sottotitolo
                    if 'TITOLO_LIBRO' in project_data:
                        book_data['title'] = project_data['TITOLO_LIBRO']
                    if 'SOTTOTITOLO_LIBRO' in project_data:
                        book_data['subtitle'] = project_data['SOTTOTITOLO_LIBRO']
                    
                    # Estrai tono di voce
                    if 'VOICE_STYLE' in project_data:
                        book_data['voice_style'] = project_data['VOICE_STYLE']
                    
                    # Estrai angolo editoriale e USP
                    if 'ANGOLO_ATTACCO' in project_data:
                        book_data['angle'] = project_data['ANGOLO_ATTACCO']
                    if 'USP' in project_data:
                        book_data['usp'] = project_data['USP']
                
                    # Estrai indice se disponibile
                    if 'INDICE_LIBRO' in project_data:
                        book_data['index'] = project_data['INDICE_LIBRO']
                    
                    # Aggiungi dati sul contesto di mercato
                    if 'MARKET_INSIGHTS' in project_data:
                        book_data['market_insights'] = project_data['MARKET_INSIGHTS']
                    if 'BUYER_PERSONA_SUMMARY' in project_data:
                        book_data['buyer_persona'] = project_data['BUYER_PERSONA_SUMMARY']
        
            # Strategia 2: Se non ci sono dati, prova a estrarli dall'HTML caricato
            if not book_data and hasattr(self, 'results_display') and self.results_display.value:
                html_content = self.results_display.value
                data_source = "HTML caricato"
            
                self.add_log("🔍 Tentativo di estrazione dati dall'HTML caricato...")
            
                try:
                    import re
                
                    # Rimuovi i tag HTML più comuni per ottenere il testo puro
                    text_content = re.sub(r'<[^>]+>', ' ', html_content)
                    text_content = re.sub(r'\s+', ' ', text_content)  # Normalizza gli spazi
                
                    # Cerca titolo del libro
                    title_patterns = [
                        r'(?:titolo|title)[:\s]+[""]?([^"\n,;]+)[""]?',
                        r'(?:TITOLO|TITLE)[:\s]+[""]?([^"\n,;]+)[""]?',
                        r'(?:libro|book)[:\s]+[""]?([^"\n,;]+)[""]?'
                    ]
                
                    for pattern in title_patterns:
                        title_match = re.search(pattern, text_content, re.IGNORECASE)
                        if title_match:
                            book_data['title'] = title_match.group(1).strip()
                            self.add_log(f"✅ Titolo estratto dall'HTML: {book_data['title']}")
                            break
                
                    # Cerca tono di voce
                    voice_patterns = [
                        r'(?:tono di voce|voice style)[:\s]+([^\n\.;]+)',
                        r'(?:stile narrativo|writing style)[:\s]+([^\n\.;]+)',
                        r'(?:stile|style)[:\s]+([^\n\.;]+)'
                    ]
                
                    for pattern in voice_patterns:
                        voice_match = re.search(pattern, text_content, re.IGNORECASE)
                        if voice_match:
                            book_data['voice_style'] = voice_match.group(1).strip()
                            self.add_log(f"✅ Tono di voce estratto dall'HTML: {book_data['voice_style']}")
                            break
                
                    # Cerca angolo di attacco
                    angle_patterns = [
                        r'(?:angolo di attacco|editorial angle)[:\s]+([^\n\.;]+)',
                        r'(?:approccio|approach)[:\s]+([^\n\.;]+)',
                        r'(?:angolo editoriale)[:\s]+([^\n\.;]+)'
                    ]
                
                    for pattern in angle_patterns:
                        angle_match = re.search(pattern, text_content, re.IGNORECASE)
                        if angle_match:
                            book_data['angle'] = angle_match.group(1).strip()
                            self.add_log(f"✅ Angolo editoriale estratto dall'HTML: {book_data['angle']}")
                            break
                
                    # Cerca USP
                    usp_patterns = [
                        r'(?:USP|unique selling proposition)[:\s]+([^\n\.;]+)',
                        r'(?:proposta di valore|value proposition)[:\s]+([^\n\.;]+)',
                        r'(?:vantaggio unico|unique advantage)[:\s]+([^\n\.;]+)'
                    ]
                
                    for pattern in usp_patterns:
                        usp_match = re.search(pattern, text_content, re.IGNORECASE)
                        if usp_match:
                            book_data['usp'] = usp_match.group(1).strip()
                            self.add_log(f"✅ USP estratta dall'HTML: {book_data['usp']}")
                            break
                
                    # Cerca l'indice del libro nell'originale HTML (mantiene la formattazione)
                    index_pattern = r'(?:INDICE|INDEX).*?(?:CAPITOLO|CHAPTER).*?(?:CONCLUSIONE|CONCLUSION)'
                    index_match = re.search(index_pattern, html_content, re.IGNORECASE | re.DOTALL)
                    if index_match:
                        # Pulisci i tag HTML dal risultato
                        raw_index = index_match.group(0)
                        clean_index = re.sub(r'<[^>]+>', ' ', raw_index)
                        clean_index = re.sub(r'\s+', ' ', clean_index)
                        book_data['index'] = clean_index.strip()
                        self.add_log(f"✅ Indice estratto dall'HTML: {len(book_data['index'])} caratteri")
                
                except Exception as html_error:
                    self.add_log(f"⚠️ Errore nell'estrazione dall'HTML: {str(html_error)}")
        
            # Strategia 3: Cerca nel file di contesto
            if not book_data and os.path.exists("context.txt"):
                with open("context.txt", "r", encoding="utf-8") as f:
                    context_content = f.read()
            
                data_source = "file di contesto"
                self.add_log("🔍 Tentativo di estrazione dati dal file di contesto...")
            
                import re
            
                # Cerca titolo
                title_match = re.search(r'(?:titolo|title)[:\s]+[""]?([^"\n,;]+)[""]?', context_content, re.IGNORECASE)
                if title_match:
                    book_data['title'] = title_match.group(1).strip()
            
                # Cerca tono di voce
                voice_match = re.search(r'(?:tono di voce|voice style)[:\s]+([^\n\.;]+)', context_content, re.IGNORECASE)
                if voice_match:
                    book_data['voice_style'] = voice_match.group(1).strip()
            
                # Cerca indice
                index_match = re.search(r'(?:INDICE|INDEX)[:\s]+((?:CAPITOLO|CHAPTER)[\s\S]+?(?=\n\s*\n|\Z))', 
                                      context_content, re.IGNORECASE | re.DOTALL)
                if index_match:
                    book_data['index'] = index_match.group(1).strip()
        
            # Se non ci sono dati dall'analisi corrente o dall'HTML, cerca indice e titolo nel file selezionato
            if not book_data and hasattr(self, 'load_analysis_dropdown') and self.load_analysis_dropdown.value:
                file_path = self.load_analysis_dropdown.value
                if os.path.exists(file_path) and file_path.endswith('.html'):
                    try:
                        data_source = "file HTML selezionato"
                        self.add_log(f"🔍 Tentativo di estrazione dati dal file: {file_path}")
                    
                        # Leggi il file HTML
                        with open(file_path, "r", encoding="utf-8") as f:
                            file_html = f.read()
                    
                        # Estrai keyword dal nome del file
                        import os
                        filename = os.path.basename(file_path)
                        keyword_match = re.match(r'([^_]+)_', filename)
                        if keyword_match:
                            keyword = keyword_match.group(1).replace("_", " ")
                            book_data['title'] = keyword
                            self.add_log(f"✅ Keyword estratta dal nome del file: {keyword}")
                    
                        # Ulteriori analisi come sopra...
                        # (puoi applicare gli stessi pattern regex come nella Strategia 2)
                    except Exception as file_error:
                        self.add_log(f"⚠️ Errore nell'estrazione dal file HTML: {str(file_error)}")
        
            # Opzione failsafe: se ancora non abbiamo dati sufficienti, crea almeno un indice generico
            if not book_data.get('index') and book_data.get('title'):
                book_title = book_data['title']
                # Crea un indice basato sul titolo
                book_data['index'] = f"""INTRODUZIONE

    CAPITOLO 1: Fondamenti di {book_title}

    CAPITOLO 2: Principali tecniche

    CAPITOLO 3: Applicazioni pratiche

    CAPITOLO 4: Casi di studio

    CAPITOLO 5: Strategie avanzate

    CONCLUSIONE
    """
                self.add_log("ℹ️ Creato indice generico basato sul titolo")
            
            # Strategia di fallback finale
            if not book_data:
                self.add_log("⚠️ Nessun dato trovato, utilizzo valori predefiniti")
            
                # Estrai almeno il titolo dal file selezionato
                if hasattr(self, 'load_analysis_dropdown') and self.load_analysis_dropdown.value:
                    file_path = self.load_analysis_dropdown.value
                    file_name = os.path.basename(file_path)
                    if '_' in file_name:
                        title = file_name.split('_')[0].replace('_', ' ')
                        book_data['title'] = title
            
                # Set di valori predefiniti
                if not book_data.get('title'):
                    book_data['title'] = "Nuovo Libro"
                if not book_data.get('voice_style'):
                    book_data['voice_style'] = "Formale e informativo"
                if not book_data.get('index'):
                    book_data['index'] = """INTRODUZIONE

    CAPITOLO 1: Fondamenti

    CAPITOLO 2: Metodologia

    CAPITOLO 3: Applicazioni pratiche

    CAPITOLO 4: Casi di studio

    CONCLUSIONE"""
        
            # Log dei risultati
            if book_data:
                self.add_log(f"✅ Dati estratti con successo da {data_source}")
                self.add_log(f"📊 Campi trovati: {', '.join(book_data.keys())}")
            else:
                self.add_log("⚠️ Nessun dato estraibile trovato")
            
            return book_data
        
        except Exception as e:
            self.add_log(f"❌ Errore nell'estrazione dei dati: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return {}


    # ——————————————————————————————————
    # Deleghe ai generatori di libro
    # ——————————————————————————————————
    def generate_book_crisp(self, book_title, book_language, voice_style, book_index):
        return generate_book_crisp(
            book_title, book_language, voice_style, book_index,
            crisp_framework=self.crisp,
            driver=self.driver,
            chat_manager=self.chat_manager,
            current_analysis=self.current_analysis
        )

    def generate_book_legacy(self, book_title, book_language, voice_style, book_index):
        return generate_book_legacy(
            book_title, book_language, voice_style, book_index,
            driver=self.driver,
            chat_manager=self.chat_manager
        )

    def generate_book(self, book_title, book_language, voice_style, book_index):
        return generate_book(
            book_title, book_language, voice_style, book_index,
            driver=self.driver,
            chat_manager=self.chat_manager,
            current_analysis=self.current_analysis
        )


    # … metodi delega, callback e orchestrazione …

    # Metodi che delegano le operazioni al DatabaseManager
    def recupera_ultimo_progetto(self):
        """Delega l'operazione al DatabaseManager"""
        result = self.db_manager.recupera_ultimo_progetto()
    
        # Aggiungi debug dopo la chiamata delegata
        self.add_log(f"DEBUG: recupera_ultimo_progetto ha restituito: {type(result)}")
        if result is None:
            self.add_log("DEBUG: Il risultato è None")
        elif isinstance(result, dict):
            self.add_log(f"DEBUG: Il risultato è un dict con {len(result)} chiavi: {list(result.keys())}")
        else:
            self.add_log(f"DEBUG: Il risultato è di tipo {type(result)}: {result}")
        
        return result

    def ripristina_ultima_analisi(self):
        """Ripristina l'ultima analisi dal database."""
        return self.db_manager.ripristina_ultima_analisi(self.crisp, self.driver, self.chat_manager)

    def ripristina_analisi_da_database(self, selected_index, start_from_phase=None):
        """Ripristina un'analisi dal database."""
        result = self.db_manager.ripristina_analisi_da_database(selected_index, start_from_phase, self.driver, self.crisp, self.chat_manager)
    
        # Se risultato è un dizionario, contiene l'analisi da ripristinare
        if isinstance(result, dict) and 'project_id' in result:
            # Salviamo i dati di ripristino
            self.current_analysis = result
            project_id = result['project_id']
            project_data = result['project_data']
            start_from_phase = result['current_phase']
        
            # Creiamo una nuova chat e carichiamo il contesto
            if self.driver:
                self.add_log(f"Creazione nuova chat per la ripresa dell'analisi...")
                create_fresh_chat(self.driver, "context.txt")
        
                # Prepara messaggio di ripresa
                input_box = get_input_box(self.driver)
                resume_message = f"""
                Sto riprendendo l'analisi per il progetto: {project_data.get('PROJECT_NAME', 'N/A')}
                Keyword: {project_data.get('KEYWORD', 'N/A')}
        
                Siamo arrivati alla fase {start_from_phase}.
                Per favore, continua l'analisi da questa fase.
                """
        
                # Invia il messaggio
                chunks = [resume_message[i:i+200] for i in range(0, len(resume_message), 200)]
                for chunk in chunks:
                    input_box.send_keys(chunk)
                    time.sleep(0.5)
        
                send_button = self.driver.find_element(By.CSS_SELECTOR, "div.search-input-wrapper div.input-icon")
                send_button.click()
        
                time.sleep(5)
        
                # Aggiorna l'interfaccia
                self.add_log(f"✅ Analisi ripristinata con successo. Pronta per continuare dalla fase {start_from_phase}")
            else:
                self.add_log("❌ Browser non inizializzato. Connettiti prima di continuare")
    
        return self.chat_manager.get_log_history_string()

    

    def load_projects_list(self):
        """Elenca i file di analisi presenti nella cartella locale, ordinati per data."""
        try:
            files = [
                f for f in os.listdir(ANALYSIS_DIR)
                if f.lower().endswith((".txt", ".json", ".html", ".docx"))
            ]
        except FileNotFoundError:
            files = []

        # Ordina dal più recente al più vecchio
        files.sort(
            key=lambda f: os.path.getmtime(os.path.join(ANALYSIS_DIR, f)),
            reverse=True
        )

        # Aggiorna il dropdown
        if hasattr(self, 'analysis_source_dropdown'):
            self.analysis_source_dropdown.choices = files
            self.add_log(f"✅ Dropdown file di analisi aggiornato: {len(files)} file trovati")

        return files


    def diagnose_and_fix_database(self):
        """Diagnostica e corregge problemi con il database"""
        self.add_log("🔍 Avvio diagnosi database...")
        try:
            import sqlite3
            import os
        
            # 1. Verifica percorso database
            self.add_log(f"🗂️ Posizione database: {self.crisp.project_db_path}")
        
            # 2. Verifica esistenza file
            if not os.path.exists(self.crisp.project_db_path):
                self.add_log(f"⚠️ ERRORE: File database non trovato: {self.crisp.project_db_path}")
                return "File database non trovato"
        
            self.add_log(f"✅ File database trovato: {os.path.getsize(self.crisp.project_db_path)} bytes")
        
            # 3. Verifica permessi di scrittura
            try:
                with open(self.crisp.project_db_path, "a") as f:
                    pass  # Test di scrittura
                self.add_log("✅ Permessi di scrittura sul database OK")
            except Exception as perm_error:
                self.add_log(f"⚠️ ERRORE: Permessi scrittura mancanti: {str(perm_error)}")
        
            # 4. Connessione al database con timeout più lungo
            self.add_log("Tentativo di connessione al database...")
            conn = sqlite3.connect(self.crisp.project_db_path, timeout=20.0)
            cursor = conn.cursor()
            self.add_log("✅ Connessione al database stabilita")
        
            # 5. Elenco tabelle per debug
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = cursor.fetchall()
            table_names = [t[0] for t in tables]
            self.add_log(f"📋 Tabelle trovate nel database: {', '.join(table_names)}")
        
            # 6. Verifica esistenza tabella projects
            if "projects" not in table_names:
                self.add_log("⚠️ ERRORE: Tabella 'projects' non trovata!")
                return "Tabella projects non trovata"
        
            # 7. Stampa struttura tabella projects
            cursor.execute("PRAGMA table_info(projects)")
            columns_info = cursor.fetchall()
            self.add_log(f"📊 Struttura tabella projects:")
            for col in columns_info:
                self.add_log(f"  - Colonna {col[0]}: {col[1]} ({col[2]})")
        
            # 8. Estrai nomi colonne
            columns = [col[1] for col in columns_info]
        
            # 9. Verifica esistenza colonna last_updated
            if "last_updated" not in columns:
                self.add_log("⚠️ Colonna 'last_updated' mancante, tentativo di aggiunta...")
            
                # 10. Avvia una transazione
                cursor.execute("BEGIN TRANSACTION")
            
                # 11. Aggiungi la colonna
                try:
                    cursor.execute("ALTER TABLE projects ADD COLUMN last_updated TEXT")
                    # 12. Commit della transazione
                    conn.commit()
                    self.add_log("✅ Colonna 'last_updated' aggiunta con successo")
                except sqlite3.OperationalError as op_error:
                    # 13. Rollback in caso di errore
                    conn.rollback()
                    self.add_log(f"⚠️ ERRORE nell'aggiungere la colonna: {str(op_error)}")
            else:
                self.add_log("✅ Colonna 'last_updated' già presente")
        
            # 14. Verifica nuovamente struttura tabella dopo modifiche
            cursor.execute("PRAGMA table_info(projects)")
            new_columns = [col[1] for col in cursor.fetchall()]
            self.add_log(f"📊 Colonne dopo riparazione: {', '.join(new_columns)}")
        
            # 15. Chiusura connessione
            conn.close()
            self.add_log("✅ Diagnosi e riparazione database completata con successo")
            return "✅ Database riparato con successo"
        except Exception as e:
            self.add_log(f"❌ ERRORE generale nella diagnosi: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return f"Errore database: {str(e)}"

    def check_existing_analysis(self, keyword):
        """Verifica se esiste già un'analisi per la keyword specificata."""
        return self.db_manager.check_existing_analysis(keyword)

    def load_project_details(self, selected_index):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.load_project_details(selected_index)

    def diagnose_and_fix_database(self):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.diagnose_and_fix_database()

    def export_project(self, selected_index):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.export_project(selected_index)

    def update_project_count(self):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.update_project_count()

    def delete_project(self, project_display_name):
        """Delega l'operazione al DatabaseManager"""
        result = self.db_manager.delete_project(project_display_name)
    
        # Aggiorna l'elenco dei progetti dopo l'eliminazione
        self.db_manager.load_projects_list()
    
        return result

    # def create_test_project(self):
    #    """Delega l'operazione al DatabaseManager"""
    #    return self.db_manager.create_test_project()

    def search_projects(self, keyword=""):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.search_projects(keyword)

    def get_database_stats(self):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.get_database_stats()

    # FINE - Metodi che delegano le operazioni al DatabaseManager


        

    

    def add_log(self, message):
            """Delega il logging al ChatManager"""
            return self.chat_manager.add_log(message)
        
    def log_prompt_location(self, prompt_id, section_number, action, details=None):
        """Delega il tracciamento della posizione al ChatManager"""
        return self.chat_manager.log_prompt_location(prompt_id, section_number, action, details) 

    def connect_callback(self):
        # Importa funzioni per gestire stato globale
        from ai_interfaces.browser_manager import get_browser_instance, set_connection_status
    
        try:
            # Variabile globale per memorizzare la funzione originale
            global original_get
        
            # Se siamo già loggati e il driver è attivo, evitiamo di ricreare la sessione
            from ai_interfaces.browser_manager import get_connection_status
            if get_connection_status() and self.driver:
                # Aggiungiamo la patch di monitoraggio se non l'abbiamo già fatto
                if not hasattr(self, '_get_patched') or not self._get_patched:
                    original_get = self.driver.get
            
                    def log_get_call(url):
                        print(f"DEBUG_URL: Chiamata a driver.get() con URL: {url}")
                        self.add_log(f"DEBUG_URL: Navigazione a: {url}")
                        # Chiamiamo la funzione originale
                        return original_get(url)
            
                    self.driver.get = log_get_call
                    self._get_patched = True
                    self.add_log("DEBUG_URL: Monitoraggio navigazione attivato")

                timestamp = datetime.now().strftime('%H:%M:%S')
                log_message = f"[{timestamp}] Sei già connesso a Genspark. Sessione attiva mantenuta."
                self.log_history.append(log_message)

                # IMPORTANTE: Verifica che siamo nella pagina di chat corretta
                self.add_log("Verifica URL corrente ed eventuale reindirizzamento alla chat...")
                current_url = self.driver.current_url
                self.add_log(f"DEBUG_URL: URL attuale prima della verifica: {current_url}")

                if "/chat" not in current_url.lower():
                    self.add_log(f"URL corrente non è una pagina di chat: {current_url}")
                    self.add_log("Reindirizzamento alla pagina di chat...")

                return self.chat_manager.get_log_history_string(), "**Stato**: Connesso - Sessione attiva"

            # Salviamo temporaneamente i messaggi di log
            logs = []

            # Aggiungiamo il primo messaggio
            timestamp = datetime.now().strftime('%H:%M:%S')
            logs.append(f"[{timestamp}] Avvio connessione browser...")

            # Usa l'istanza globale o crea una nuova istanza forzata
            self.driver = get_browser_instance(force_new=True)  # Force new in caso di riconnessione
            self.add_url_change_monitor()
            self.crisp.driver = self.driver  # Aggiorna il riferimento al driver nel framework CRISP

            # Aggiungiamo la patch di monitoraggio per il nuovo driver
            original_get = self.driver.get
    
            def log_get_call(url):
                print(f"DEBUG_URL: Chiamata a driver.get() con URL: {url}")
                self.add_log(f"DEBUG_URL: Navigazione a: {url}")
                # Tracciamo anche da dove viene chiamata la funzione
                import traceback
                caller = traceback.extract_stack()[-2]
                print(f"DEBUG_URL: Chiamata da {caller.filename}:{caller.lineno}")
                # Chiamiamo la funzione originale
                return original_get(url)
    
            self.driver.get = log_get_call
            self._get_patched = True
            self.add_log("DEBUG_URL: Monitoraggio navigazione attivato per nuovo driver")

            # Dopo aver avviato il browser, aggiungi:
            self.add_log(f"Browser avviato: {self.driver is not None}")
            self.add_log("Tentativo di navigazione a Genspark...")

            # IMPORTANTE: Naviga DIRETTAMENTE alla pagina di chat, non alla home
            self.add_log("Navigazione diretta alla pagina di chat...")
            # Questo viene monitorato grazie al nostro patch
            self.driver.get("https://genspark.ai")
            time.sleep(7)  # Attesa più lunga per il caricamento completo

            # Verifica URL dopo la navigazione
            current_url = self.driver.current_url
            self.add_log(f"URL dopo navigazione: {current_url}")

            # Se siamo stati reindirizzati alla home o altra pagina, riprova con approccio alternativo
            if "/chat" not in current_url.lower():
                self.add_log("Reindirizzato a URL non di chat, provo approccio alternativo...")
                # Prima vai alla home
                self.driver.get("https://genspark.ai/")
                time.sleep(3)

            # Verifica dell'URL finale
            final_url = self.driver.current_url
            self.add_log(f"URL finale: {final_url}")

            # Controlliamo il login
            self.add_log("DEBUG_URL: Verificando login...")
            login_result = check_login(self.driver)
            self.add_log(f"DEBUG_URL: Risultato check_login: {login_result}")
    
            if not login_result:
                timestamp = datetime.now().strftime('%H:%M:%S')
                logs.append(f"[{timestamp}] ATTENZIONE: Completa manualmente il login su Genspark.ai.")
                timestamp = datetime.now().strftime('%H:%M:%S')
                logs.append(f"[{timestamp}] Assicurati di essere nella pagina della chat prima di continuare.")

                # Aggiungi un input per aspettare che l'utente finisca di fare login
                timestamp = datetime.now().strftime('%H:%M:%S')
                logs.append(f"[{timestamp}] Premi il pulsante 'Connetti' nuovamente dopo aver effettuato il login.")
    
                # Aggiorniamo la cronologia dei log attraverso il ChatManager
                for log_message in logs:
                    # Rimuoviamo il timestamp se già presente nel messaggio
                    if log_message.startswith('[') and ']' in log_message:
                        clean_message = log_message.split(']', 1)[1].strip()
                        self.chat_manager.add_log(clean_message)
                    else:
                        self.chat_manager.add_log(log_message)

                return self.chat_manager.get_log_history_string(), "**Stato**: Login richiesto"

                # Salva i cookie per usi futuri
                try:
                    import pickle
                    cookies = self.driver.get_cookies()
                    pickle.dump(cookies, open("genspark_cookies.pkl", "wb"))
                    timestamp = datetime.now().strftime('%H:%M:%S')
                    logs.append(f"[{timestamp}] Cookie salvati con successo!")
                except Exception as e:
                    timestamp = datetime.now().strftime('%H:%M:%S')
                    logs.append(f"[{timestamp}] Errore nel salvataggio dei cookie: {str(e)}")

            # Impostiamo lo stato di login
            self.is_logged_in = True
            from ai_interfaces.browser_manager import set_connection_status
            set_connection_status(True)  # Imposta stato globale

            # Aggiorniamo la cronologia dei log e restituiamo il tutto
            for log_message in logs:
                self.chat_manager.add_log(log_message)
            return self.chat_manager.get_log_history_string(), "**Stato**: Connesso"
        except Exception as e:
            # In caso di errore
            self.chat_manager.add_log(f"Errore: {str(e)}")
        
            # Aggiungiamo il traceback per debug
            import traceback
            tb = traceback.format_exc()
            print(f"DEBUG_URL: Eccezione in connect_callback:\n{tb}")
        
            return self.chat_manager.get_log_history_string(), "**Stato**: Errore di connessione"

    def add_url_change_monitor(self):
        """Aggiunge uno script al browser per monitorare i cambiamenti di URL"""
        if not hasattr(self, 'driver') or self.driver is None:
            print("DEBUG_URL: Impossibile installare monitor URL - driver non disponibile")
            return False
        
        script = """
        let lastUrl = window.location.href;
        console.log('Monitor URL inizializzato con URL: ' + lastUrl);
    
        // Funzione per controllare periodicamente l'URL
        function checkUrlChange() {
            const currentUrl = window.location.href;
            if (currentUrl !== lastUrl) {
                console.log('URL CAMBIATO da: ' + lastUrl + ' a: ' + currentUrl);
                lastUrl = currentUrl;
            }
            setTimeout(checkUrlChange, 1000);  // Controlla ogni secondo
        }
    
        // Avvia il controllo
        checkUrlChange();
    
        // Restituisci true per confermare l'installazione
        return true;
        """
    
        try:
            result = self.driver.execute_script(script)
            print(f"DEBUG_URL: Monitor di cambio URL installato nel browser: {result}")
            return True
        except Exception as e:
            print(f"DEBUG_URL: Errore nell'installazione del monitor URL: {str(e)}")
            return False

    def get_selected_phases(self):
        """
        Nuovo metodo helper che restituisce le fasi attualmente selezionate nell'interfaccia.
        Funziona sia per analisi CRISP che Legacy.
    
        Returns:
            list: Lista delle fasi selezionate (IDs per CRISP, numeri per Legacy)
        """
        selected_phases = []
    
        # Determina il tipo di analisi corrente
        analysis_type = getattr(self, 'selected_analysis_type', "CRISP")
    
        if analysis_type == "CRISP" and hasattr(self, 'crisp_phase_checkboxes'):
            # Ottieni le fasi CRISP selezionate
            try:
                selected_values = self.crisp_phase_checkboxes.value
                import re
                for selected_value in selected_values:
                    match = re.match(r'([A-Z]+-[0-9A-Z]+):', selected_value)
                    if match:
                        selected_phases.append(match.group(1))
                        self.add_log(f"DEBUG: Fase CRISP selezionata per salvataggio: {match.group(1)}")
            except Exception as e:
                self.add_log(f"DEBUG: Errore nella lettura delle fasi CRISP: {str(e)}")
        else:
            # Ottieni le fasi Legacy selezionate
            if hasattr(self, 'legacy_phase_checkboxes'):
                try:
                    selected_values = self.legacy_phase_checkboxes.value
                    import re
                    for selected_value in selected_values:
                        match = re.match(r'([A-Z]+-\d+):', selected_value)
                        if match:
                            phase_id = match.group(1)
                            number_match = re.search(r'-(\d+)', phase_id)
                            if number_match:
                                phase_number = int(number_match.group(1))
                                selected_phases.append(phase_number)
                                self.add_log(f"DEBUG: Fase Legacy selezionata per salvataggio: {phase_number}")
                except Exception as e:
                    self.add_log(f"DEBUG: Errore nella lettura delle fasi Legacy: {str(e)}")
    
        self.add_log(f"Fasi selezionate per salvataggio: {selected_phases}")
        return selected_phases

    

    # Definisci una funzione executor che invia il prompt a Genspark
    def execute_prompt(self, prompt_text, step_id=None, project_data=None):
        """
        Funzione robusta che esegue un prompt CRISP, gestisce errori e interruzioni,
        e garantisce il completamento anche in presenza di problemi.

        Args:
            prompt_text: Testo del prompt da inviare
            step_id: ID del prompt corrente (opzionale)
            project_data: Dati del progetto per sostituire le variabili

        Returns:
            str: Risposta cumulativa da Genspark
        """
        # Inizializzazione e logging
        prompt_id_to_use = step_id if step_id else "unknown"
        self.add_log(f"🚀 Inizio esecuzione prompt {prompt_id_to_use} ({len(prompt_text)} caratteri)")
        print(f"DEBUG: Inizio esecuzione prompt {prompt_id_to_use} ({len(prompt_text)} caratteri)")
        print(f"DEBUG: Preview prompt: {prompt_text[:200].replace(chr(10), ' ')}...")

        # Verifica browser attivo
        if not hasattr(self, 'driver') or self.driver is None:
            self.add_log("⚠️ Browser non attivo, impossibile procedere")
            print("DEBUG: Browser non attivo, impossibile procedere")
            return "ERRORE: Browser non inizializzato"

        # Verifica URL corrente
        try:
            current_url = self.driver.current_url
            self.add_log(f"🌐 URL attuale: {current_url}")
            print(f"DEBUG: URL attuale: {current_url}")
    
            # Se non siamo in una pagina di chat, naviga a Genspark
            if not ("genspark.ai" in current_url and ("/chat" in current_url or "/agents" in current_url)):
                self.add_log("🔄 Navigazione alla pagina di chat...")
                print("DEBUG: Navigazione alla pagina di chat...")
                self.driver.get("https://genspark.ai")
                time.sleep(10)
        except Exception as e:
            self.add_log(f"⚠️ Errore nella verifica URL: {str(e)}")
            print(f"DEBUG: Errore nella verifica URL: {str(e)}")

        # MODIFICA CRUCIALE: Divisione del prompt in sezioni numeriche
        import re

        # Pattern per trovare punti numerati (esempio: "1. Titolo", "2. Titolo", ecc.)
        numbered_sections = re.findall(r'(?:\n|^)(\d+\.\s+.*?)(?=(?:\n\d+\.|\n\n|$))', prompt_text, re.DOTALL)

        # Se abbiamo trovato sezioni numerate, usiamole
        if numbered_sections:
            sections = [section.strip() for section in numbered_sections]
            self.add_log(f"📋 Prompt diviso in {len(sections)} sezioni numerate")
            print(f"DEBUG: Prompt diviso in {len(sections)} sezioni numerate")
        else:
            # Fallback: dividi per righe normalmente
            sections = [line.strip() for line in prompt_text.split('\n') if line.strip()]
            self.add_log(f"📋 Prompt diviso in {len(sections)} righe (nessuna sezione numerata trovata)")
            print(f"DEBUG: Prompt diviso in {len(sections)} righe (nessuna sezione numerata trovata)")

        # Log dettagliato delle sezioni
        for i, section in enumerate(sections):
            preview = section.replace('\n', ' ')[:50]
            self.add_log(f"📄 Sezione {i+1}: {preview}..." + ("" if len(section) <= 50 else f" ({len(section)} caratteri)"))
            print(f"DEBUG: Sezione {i+1}: {preview}..." + ("" if len(section) <= 50 else f" ({len(section)} caratteri)"))

        # Variabili di monitoraggio globali
        max_global_retries = 3
        global_retry_count = 0
        cumulative_response = []

        # Loop di ripetizione globale per il prompt intero
        while global_retry_count < max_global_retries:
            try:
                self.add_log(f"📝 Tentativo globale {global_retry_count+1}/{max_global_retries}")
                print(f"DEBUG: Tentativo globale {global_retry_count+1}/{max_global_retries}")
        
                # Pulizia dell'interfaccia prima di iniziare
                try:
                    clear_chat(self.driver)
                    self.add_log("🧹 Chat pulita all'inizio dell'esecuzione")
                    print("DEBUG: Chat pulita all'inizio dell'esecuzione")
                    time.sleep(5)
                except Exception as clear_error:
                    self.add_log(f"⚠️ Impossibile pulire la chat: {str(clear_error)}")
                    print(f"DEBUG: Impossibile pulire la chat: {str(clear_error)}")
                    # Continua comunque
        
                # Processa ogni sezione
                success_sections = 0
                for i, section in enumerate(sections):
                    self.add_log(f"📌 Elaborazione sezione {i+1}/{len(sections)}...")
                    print(f"DEBUG: Elaborazione sezione {i+1}/{len(sections)}...")
            
                    # Sostituzione variabili avanzata
                    processed_section = replace_variables(section, project_data)
                    preview = processed_section.replace('\n', ' ')[:50]
                    self.add_log(f"✏️ Sezione processata: {preview}..." + 
                                ("" if len(processed_section) <= 50 else f" ({len(processed_section)} caratteri)"))
                    print(f"DEBUG: Sezione processata: {preview}..." + 
                         ("" if len(processed_section) <= 50 else f" ({len(processed_section)} caratteri)"))
            
                    # Verifica placeholder non risolti
                    unresolved = self.check_unresolved_placeholders(processed_section)
                    if unresolved:
                        self.add_log(f"⚠️ Placeholders non risolti: {unresolved}")
                        print(f"DEBUG: Placeholders non risolti: {unresolved}")
                        # Continua comunque, ma potrebbe causare problemi
            
                    # Tentativi per sezione singola
                    section_retry_count = 0
                    max_section_retries = 3
                    section_success = False
            
                    while section_retry_count < max_section_retries and not section_success:
                        try:
                            self.add_log(f"🔄 Tentativo {section_retry_count+1}/{max_section_retries} per sezione {i+1}")
                            print(f"DEBUG: Tentativo {section_retry_count+1}/{max_section_retries} per sezione {i+1}")
                    
                            # Sistema avanzato di pulizia input
                            input_box = self.get_clean_input_box()
                            if not input_box:
                                raise Exception("Impossibile ottenere o pulire la casella di input")
                    
                            # Inserimento testo sicuro - carattere per carattere per maggiore affidabilità
                            self.add_log(f"⌨️ Inserimento testo carattere per carattere...")
                            print(f"DEBUG: Inserimento testo carattere per carattere... ({len(processed_section)} caratteri)")
                            for char in processed_section:
                                input_box.send_keys(char)
                                time.sleep(0.008)  # Minimo ritardo per stabilità
                    
                            # Verifica prima dell'invio
                            time.sleep(1)
                            inserted_text = input_box.get_attribute("value")
                            if not inserted_text:
                                self.add_log("⚠️ Nessun testo inserito!")
                                print("DEBUG: ERRORE - Nessun testo inserito!")
                                if section_retry_count < max_section_retries - 1:
                                    section_retry_count += 1
                                    time.sleep(5)
                                    continue
                            elif len(inserted_text) < len(processed_section) * 0.9:
                                self.add_log(f"⚠️ Inserimento incompleto: {len(inserted_text)}/{len(processed_section)} caratteri")
                                print(f"DEBUG: ERRORE - Inserimento incompleto: {len(inserted_text)}/{len(processed_section)} caratteri")
                                if section_retry_count < max_section_retries - 1:
                                    section_retry_count += 1
                                    time.sleep(5)
                                    continue
                            else:
                                self.add_log(f"✅ Testo inserito correttamente: {len(inserted_text)} caratteri")
                                print(f"DEBUG: Testo inserito correttamente: {len(inserted_text)} caratteri")
                    
                            # Invio con retry integrato
                            send_success = False
                    
                            # Metodo 1: Click standard
                            try:
                                send_button = WebDriverWait(self.driver, 15).until(
                                    EC.element_to_be_clickable((By.CSS_SELECTOR, "div.search-input-wrapper div.input-icon"))
                                )
                                send_button.click()
                                self.add_log("🔘 Click standard sul pulsante di invio")
                                print("DEBUG: Click standard sul pulsante di invio")
                                send_success = True
                            except Exception as e1:
                                self.add_log(f"⚠️ Click standard fallito: {str(e1)}")
                                print(f"DEBUG: Click standard fallito: {str(e1)}")
                        
                                # Metodo 2: Click JavaScript
                                try:
                                    send_button = self.driver.find_element(By.CSS_SELECTOR, "div.search-input-wrapper div.input-icon")
                                    self.driver.execute_script("arguments[0].click();", send_button)
                                    self.add_log("🔘 Click JavaScript sul pulsante di invio")
                                    print("DEBUG: Click JavaScript sul pulsante di invio")
                                    send_success = True
                                except Exception as e2:
                                    self.add_log(f"⚠️ Click JavaScript fallito: {str(e2)}")
                                    print(f"DEBUG: Click JavaScript fallito: {str(e2)}")
                            
                                    # Metodo 3: Tasto invio
                                    try:
                                        input_box.send_keys(Keys.RETURN)
                                        self.add_log("🔘 Invio tramite tasto RETURN")
                                        print("DEBUG: Invio tramite tasto RETURN")
                                        send_success = True
                                    except Exception as e3:
                                        self.add_log(f"❌ Tutti i metodi di invio falliti: {str(e3)}")
                                        print(f"DEBUG: Tutti i metodi di invio falliti: {str(e3)}")
                    
                            if not send_success:
                                raise Exception("Impossibile inviare il messaggio con nessun metodo")
                    
                            # Attesa iniziale per inizio elaborazione
                            self.add_log("⏳ Attesa iniziale dopo invio (10 secondi)")
                            print("DEBUG: Attesa iniziale dopo invio (10 secondi)")
                            time.sleep(10)
                    
                            # Sistema di attesa adattivo
                            max_wait_cycles = 45  # ~15 minuti totali
                            stability_threshold = 10  # 10 cicli di stabilità
                            cycle_wait = 20  # 20 secondi per ciclo
                    
                            # Inizializzazione variabili di monitoraggio
                            terminator_found = False
                            last_length = 0
                            stable_count = 0
                            response_text = None
                    
                            for cycle in range(max_wait_cycles):
                                try:
                                    print(f"DEBUG: Ciclo di attesa {cycle+1}/{max_wait_cycles}")
                                
                                    # Verifica limite contesto ogni 3 cicli
                                    if cycle % 3 == 0 and self.handle_context_limit():
                                        self.add_log("♻️ Limite contesto gestito durante attesa")
                                        print("DEBUG: Limite contesto gestito durante attesa")
                            
                                    # Prova diversi selettori per le risposte
                                    selectors = [
                                        ".message-content", 
                                        "div.chat-wrapper div.desc > div > div > div",
                                        "div.message div.text-wrap",
                                        ".chat-message-item .content"
                                    ]
                            
                                    # Cerca la risposta con tutti i selettori
                                    for selector in selectors:
                                        try:
                                            messages = self.driver.find_elements(By.CSS_SELECTOR, selector)
                                            if messages and len(messages) > 0:
                                                current_text = messages[-1].text.strip()
                                                if current_text:
                                                    response_text = current_text
                                                
                                                    # Debug della risposta ogni 5 cicli
                                                    if cycle % 5 == 0:
                                                        print(f"DEBUG: Salvataggio risposta - Lunghezza: {len(response_text)}")
                                                        if len(response_text) > 0:
                                                            print(f"DEBUG: Preview risposta: {response_text[:200].replace(chr(10), ' ')}...")
                                            
                                                    # Verifica terminazione esplicita
                                                    if "FINE_RISPOSTA" in response_text or "FINE" in response_text:
                                                        self.add_log(f"✅ Terminatore esplicito trovato al ciclo {cycle+1}")
                                                        print(f"DEBUG: Terminatore esplicito trovato al ciclo {cycle+1}")
                                                        terminator = "FINE_RISPOSTA" if "FINE_RISPOSTA" in response_text else "FINE"
                                                        terminator_pos = response_text.find(terminator)
                                                        print(f"DEBUG: Terminatore '{terminator}' trovato alla posizione {terminator_pos}")
                                                
                                                        # Pulisci la risposta rimuovendo il terminatore
                                                        if "FINE_RISPOSTA" in response_text:
                                                            response_text = response_text.split("FINE_RISPOSTA")[0].strip()
                                                        elif "FINE" in response_text:
                                                            response_text = response_text.split("FINE")[0].strip()
                                                        terminator_found = True
                                                        section_success = True
                                                        break
                                            
                                                    # Verifica errori tipici
                                                    error_indicators = ["richiesta abortita", "request aborted", 
                                                                       "troppo lungo", "too long", 
                                                                       "errore durante", "error during"]
                                            
                                                    if any(e in response_text.lower() for e in error_indicators):
                                                        self.add_log(f"❌ Errore rilevato nella risposta al ciclo {cycle+1}")
                                                        print(f"DEBUG: Errore rilevato nella risposta al ciclo {cycle+1}: '{next((e for e in error_indicators if e in response_text.lower()), '')}'")
                                                        break
                                            
                                                    # Verifica stabilità
                                                    current_length = len(response_text)
                                                    if current_length == last_length:
                                                        stable_count += 1
                                                        self.add_log(f"⏳ Risposta stabile: {stable_count}/{stability_threshold} cicli ({current_length} caratteri)")
                                                        print(f"DEBUG: Risposta stabile: {stable_count}/{stability_threshold} cicli ({current_length} caratteri)")
                                                
                                                        if stable_count >= stability_threshold:
                                                            self.add_log(f"✅ Risposta stabilizzata dopo {cycle+1} cicli")
                                                            print(f"DEBUG: Risposta stabilizzata dopo {cycle+1} cicli - Lunghezza finale: {current_length} caratteri")
                                                            section_success = True
                                                            break
                                                    else:
                                                        stable_count = 0
                                                        self.add_log(f"📝 Risposta in evoluzione: {current_length} caratteri (ciclo {cycle+1})")
                                                        print(f"DEBUG: Risposta in evoluzione: {current_length} caratteri (ciclo {cycle+1})")
                                                        last_length = current_length
                                            
                                                    # Trovata risposta valida, esci dal ciclo selettori
                                                    break
                                        except Exception:
                                            continue
                            
                                    # Se abbiamo avuto successo, esci dal ciclo di attesa
                                    if section_success:
                                        break
                            
                                    # Attendi prima del prossimo ciclo
                                    time.sleep(cycle_wait)
                        
                                except Exception as e:
                                    self.add_log(f"⚠️ Errore durante attesa risposta: {str(e)}")
                                    print(f"DEBUG: Errore durante attesa risposta: {str(e)}")
                                    time.sleep(cycle_wait)
                    
                            # Verifica se abbiamo ottenuto successo
                            if section_success and response_text:
                                # Risposta ottenuta con successo
                                self.add_log(f"✅ Risposta ottenuta per sezione {i+1}: {len(response_text)} caratteri")
                                print(f"DEBUG: Risposta ottenuta per sezione {i+1}: {len(response_text)} caratteri")
                            
                                # Debug della risposta ottenuta
                                print(f"DEBUG: Salvataggio risposta - Lunghezza: {len(response_text)}")
                                print(f"DEBUG: Preview risposta: {response_text[:200].replace(chr(10), ' ')}...")
                            
                                # Verifica qualità risposta
                                if len(response_text) < 50 and not ("CM-1" in prompt_id_to_use and i == 0):
                                    # Risposta troppo corta (eccetto la prima sezione di CM-1 che può essere corta)
                                    self.add_log(f"⚠️ Risposta sospettosamente corta: {len(response_text)} caratteri")
                                    print(f"DEBUG: Risposta sospettosamente corta: {len(response_text)} caratteri")
                                    if section_retry_count < max_section_retries - 1:
                                        section_retry_count += 1
                                        time.sleep(10)
                                        continue
                        
                                # Aggiungi alla risposta cumulativa
                                cumulative_response.append(response_text)
                        
                                # Salva incrementalmente
                                try:
                                    # Usa una struttura a cascata per trovare il metodo giusto
                                    if hasattr(self.crisp, 'save_incremental_response') and project_data and "PROJECT_ID" in project_data:
                                        print(f"DEBUG: Chiamata a crisp.save_incremental_response per {prompt_id_to_use}")
                                        self.crisp.save_incremental_response(
                                            project_data["PROJECT_ID"], 
                                            prompt_id_to_use,
                                            processed_section, 
                                            response_text, 
                                            i == len(sections) - 1
                                        )
                                        self.add_log("💾 Risposta salvata nel database CRISP")
                                        print("DEBUG: Risposta salvata nel database CRISP")
                                    elif hasattr(self.crisp, 'crisp') and hasattr(self.crisp.crisp, 'save_incremental_response') and project_data and "PROJECT_ID" in project_data:
                                        print(f"DEBUG: Chiamata a crisp.crisp.save_incremental_response per {prompt_id_to_use}")
                                        self.crisp.crisp.save_incremental_response(
                                            project_data["PROJECT_ID"], 
                                            prompt_id_to_use,
                                            processed_section, 
                                            response_text, 
                                            i == len(sections) - 1
                                        )
                                        self.add_log("💾 Risposta salvata nel database CRISP (via crisp.crisp)")
                                        print("DEBUG: Risposta salvata nel database CRISP (via crisp.crisp)")
                                    else:
                                        self.add_log("⚠️ Metodo save_incremental_response non trovato o dati progetto mancanti")
                                    
                                        # Verifica dettagli per un miglior debug
                                        print("DEBUG: Dettagli variabili per salvataggio:")
                                        print(f"DEBUG: - hasattr(self.crisp, 'save_incremental_response'): {hasattr(self.crisp, 'save_incremental_response')}")
                                        print(f"DEBUG: - project_data is not None: {project_data is not None}")
                                        if project_data:
                                            print(f"DEBUG: - 'PROJECT_ID' in project_data: {'PROJECT_ID' in project_data}")
                                            if 'PROJECT_ID' in project_data:
                                                print(f"DEBUG: - PROJECT_ID value: {project_data['PROJECT_ID']}")
                                    
                                        # Salva nel file di contesto come fallback
                                        if hasattr(self, 'chat_manager'):
                                            print(f"DEBUG: Salvando nel file di contesto come fallback")
                                            metadata = {
                                                "prompt_id": prompt_id_to_use,
                                                "section_number": i+1,
                                                "timestamp": datetime.now().strftime('%Y%m%d_%H%M%S')
                                            }
                                            self.chat_manager.save_response(
                                                response_text,
                                                f"Prompt {prompt_id_to_use}-Sezione {i+1}",
                                                metadata
                                            )
                                            print(f"DEBUG: Salvato nel file di contesto come fallback")
                                except Exception as save_error:
                                    self.add_log(f"⚠️ Errore nel salvare la risposta: {str(save_error)}")
                                    print(f"DEBUG: Errore nel salvare la risposta: {str(save_error)}")
                                    import traceback
                                    print(f"DEBUG: Traceback salvataggio:\n{traceback.format_exc()}")
                        
                            elif response_text and len(response_text) > 100:
                                # Risposta parziale ma utilizzabile
                                self.add_log(f"⚠️ Risposta parziale ma utilizzabile: {len(response_text)} caratteri")
                                print(f"DEBUG: Risposta parziale ma utilizzabile: {len(response_text)} caratteri")
                                print(f"DEBUG: Salvataggio risposta parziale - Lunghezza: {len(response_text)}")
                                print(f"DEBUG: Preview risposta parziale: {response_text[:200].replace(chr(10), ' ')}...")
                        
                                # Aggiungi alla risposta cumulativa
                                cumulative_response.append(response_text)
                                section_success = True
                        
                                # Salva anche risposte parziali
                                try:
                                    if hasattr(self.crisp, 'save_incremental_response') and project_data and "PROJECT_ID" in project_data:
                                        print(f"DEBUG: Salvando risposta parziale in CRISP database")
                                        self.crisp.save_incremental_response(
                                            project_data["PROJECT_ID"], 
                                            prompt_id_to_use,
                                            processed_section, 
                                            response_text, 
                                            i == len(sections) - 1
                                        )
                                        print(f"DEBUG: Risposta parziale salvata nel database CRISP")
                                except Exception as save_error:
                                    print(f"DEBUG: Errore nel salvare risposta parziale: {str(save_error)}")
                            else:
                                # Nessuna risposta o timeout
                                self.add_log(f"❌ Nessuna risposta valida ottenuta per sezione {i+1}")
                                print(f"DEBUG: Nessuna risposta valida ottenuta per sezione {i+1}")
                        
                                if section_retry_count < max_section_retries - 1:
                                    section_retry_count += 1
                                    time.sleep(15)
                                    continue
                                else:
                                    cumulative_response.append(f"[Timeout per sezione {i+1}]")
                                    print(f"DEBUG: Aggiunto placeholder di timeout per sezione {i+1}")
                
                        except Exception as e:
                            # Gestione errori specifici per sezione
                            self.add_log(f"⚠️ Errore sezione {i+1}, tentativo {section_retry_count+1}: {str(e)}")
                            print(f"DEBUG: Errore sezione {i+1}, tentativo {section_retry_count+1}: {str(e)}")
                    
                            # Cattura screenshot per debug
                            try:
                                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                                screenshot_path = f"error_section_{i+1}_try_{section_retry_count+1}_{timestamp}.png"
                                self.driver.save_screenshot(screenshot_path)
                                self.add_log(f"📸 Screenshot: {screenshot_path}")
                                print(f"DEBUG: Screenshot errore: {screenshot_path}")
                            except Exception:
                                pass
                    
                            if section_retry_count < max_section_retries - 1:
                                section_retry_count += 1
                        
                                # Verifica se è un problema di contesto
                                if "context" in str(e).lower() or "too long" in str(e).lower():
                                    if self.handle_context_limit():
                                        self.add_log("♻️ Limite di contesto gestito, nuovo tentativo")
                                        print("DEBUG: Limite di contesto gestito, nuovo tentativo")
                        
                                time.sleep(15)
                            else:
                                # Ultimo fallimento per questa sezione
                                self.add_log(f"❌ Tutti i tentativi falliti per sezione {i+1}")
                                print(f"DEBUG: Tutti i tentativi falliti per sezione {i+1}")
                                cumulative_response.append(f"[Errore: {str(e)}]")
            
                    # Se la sezione è stata completata con successo
                    if section_success:
                        success_sections += 1
                
                        # Pausa tra le sezioni - aumenta progressivamente
                        pause_time = 15 + (i * 3)  # Aumenta di 3 secondi per ogni sezione
                        self.add_log(f"⏱️ Pausa di {pause_time} secondi prima della prossima sezione")
                        print(f"DEBUG: Pausa di {pause_time} secondi prima della prossima sezione")
                        time.sleep(pause_time)
                    else:
                        # Se la sezione ha fallito definitivamente - decisione se continuare
                        self.add_log(f"❌ Fallimento definitivo per sezione {i+1}/{len(sections)}")
                        print(f"DEBUG: Fallimento definitivo per sezione {i+1}/{len(sections)}")
                        if i < len(sections) - 1:  # Se non è l'ultima sezione
                            self.add_log("⚠️ Continuo con la sezione successiva nonostante il fallimento")
                            print("DEBUG: Continuo con la sezione successiva nonostante il fallimento")
        
                # Se abbiamo completato tutte le sezioni o la maggior parte
                if success_sections == len(sections) or success_sections >= len(sections) * 0.7:
                    self.add_log(f"✅ Completate {success_sections}/{len(sections)} sezioni con successo")
                    print(f"DEBUG: Completate {success_sections}/{len(sections)} sezioni con successo")
                
                    # Unisci le risposte e termina
                    combined_response = "\n\n".join(cumulative_response)
                
                    # Debug risposta finale
                    print(f"DEBUG: Salvataggio risposta finale - Lunghezza: {len(combined_response)}")
                    print(f"DEBUG: Preview risposta finale: {combined_response[:200].replace(chr(10), ' ')}...")
            
                    # Salva nel contesto
                    try:
                            print(f"DEBUG: Tentativo di lettura del file context.txt - Esiste: {os.path.exists('context.txt')}")

                            self.chat_manager.save_response(
                                combined_response,
                                f"Analisi CRISP 5.0 - {prompt_id_to_use}",
                                {"prompt_id": prompt_id_to_use, "timestamp": datetime.now().strftime('%Y%m%d_%H%M%S')}
                            )

                            try:
                                if hasattr(self, 'driver') and self.driver:
                                    self.add_log("🧪 Browser attivo rilevato: tentativo di cattura HTML completo da video")
                                    saved_path = self.save_complete_html_visual()

                                    if saved_path:
                                        self.last_html_path = saved_path
                                        self.add_log(f"✅ HTML completo catturato e salvato in: {saved_path}")
                                    else:
                                        self.add_log("⚠️ HTML non salvato: save_complete_html() ha restituito None")
                                        print("DEBUG: save_complete_html() ha restituito None")
                                        html = self._generate_html_from_context_file()
                                        self.last_html_path = self._save_formatted_html_to_file(html)
                                else:
                                    self.add_log("🧩 Nessun browser attivo. Uso il formatter interno.")
                                    html = self._generate_html_from_context_file()
                                    self.last_html_path = self._save_formatted_html_to_file(html)
                            except Exception as html_error:
                                self.add_log(f"❌ Errore durante la generazione HTML: {str(html_error)}")
                                import traceback
                                print(f"DEBUG: Traceback errore HTML:\n{traceback.format_exc()}")

                            self.add_log("✅ Risposta combinata salvata nel contesto")
                            print("DEBUG: Risposta combinata salvata nel contesto")
                    except Exception as e:
                        self.add_log(f"⚠️ Errore nel salvare nel contesto: {str(e)}")
                        import traceback
                        print(f"DEBUG: Traceback salvataggio contesto:\n{traceback.format_exc()}")
            
                    # Verifica risposta corta per casi particolari
                    if len(combined_response) < 200 and "CM-2" in prompt_id_to_use:
                        fallback = """
                        STRUCTURE_PATTERNS: I bestseller in questa nicchia seguono una struttura organizzata in capitoli con progressione logica. Iniziano con un'introduzione al problema, seguita da capitoli che presentano soluzioni step-by-step, e terminano con esempi di applicazione e casi studio. La maggior parte include anche appendici con risorse aggiuntive.

                        TITLE_PATTERNS: I titoli più efficaci utilizzano una combinazione di problemi e soluzioni, spesso con sottotitoli che espandono la promessa principale. Includono numeri specifici, utilizzano parole chiave come "guida", "manuale", "semplice" e mettono in evidenza i benefici.

                        REVIEW_INSIGHTS: Le recensioni positive evidenziano contenuti pratici, chiarezza espositiva e applicabilità immediata. Le recensioni negative menzionano informazioni troppo generiche, mancanza di profondità e assenza di esempi concreti.

                        IMPLEMENTATION_OBSTACLES: Le principali difficoltà di implementazione includono la complessità percepita, la mancanza di esercizi pratici e l'assenza di supporto continuo dopo la lettura.

                        MARKET_GAPS: Esiste una chiara opportunità per un libro che combini teoria e pratica con un approccio passo-passo, materiali di supporto scaricabili e un linguaggio semplice ma professionale.
                        """
                        self.add_log("⚠️ Risposta sostituita con fallback per CM-2")
                        print("DEBUG: Risposta sostituita con fallback per CM-2 (risposta originale troppo corta)")
                        return fallback
            
                    return combined_response
                else:
                    # Troppe sezioni fallite, riprova l'intero prompt
                    self.add_log(f"⚠️ Solo {success_sections}/{len(sections)} sezioni completate, riprovo l'intero prompt")
                    print(f"DEBUG: Solo {success_sections}/{len(sections)} sezioni completate, riprovo l'intero prompt")
                    global_retry_count += 1
            
                    # Pulizia prima del prossimo tentativo globale
                    try:
                        clear_chat(self.driver)
                        time.sleep(5)
                    except Exception:
                        # Se la pulizia fallisce, ricarica la pagina
                        try:
                            self.driver.get("https://genspark.ai")
                            time.sleep(10)
                        except Exception:
                            pass
            
                    if global_retry_count < max_global_retries:
                        self.add_log(f"🔄 Nuovo tentativo globale {global_retry_count+1} in corso...")
                        print(f"DEBUG: Nuovo tentativo globale {global_retry_count+1} in corso...")
                        cumulative_response = []  # Reset per il nuovo tentativo
                        time.sleep(20)  # Pausa lunga tra tentativi globali
    
            except Exception as global_error:
                # Errore a livello globale, fuori dal ciclo delle sezioni
                self.add_log(f"❌ ERRORE GLOBALE: {str(global_error)}")
                print(f"DEBUG: ERRORE GLOBALE: {str(global_error)}")
                import traceback
                print(f"DEBUG: Traceback errore globale:\n{traceback.format_exc()}")
            
                global_retry_count += 1
        
                # Cattura screenshot per debug
                try:
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    screenshot_path = f"global_error_{global_retry_count}_{timestamp}.png"
                    self.driver.save_screenshot(screenshot_path)
                    self.add_log(f"📸 Screenshot errore globale: {screenshot_path}")
                    print(f"DEBUG: Screenshot errore globale: {screenshot_path}")
                except Exception:
                    pass
        
                if global_retry_count < max_global_retries:
                    self.add_log(f"⚠️ Riprovo l'intero prompt, tentativo {global_retry_count+1}/{max_global_retries}")
                    print(f"DEBUG: Riprovo l'intero prompt, tentativo {global_retry_count+1}/{max_global_retries}")
                    # Reset per il nuovo tentativo
                    cumulative_response = []
            
                    # Ricarica pagina per reset completo
                    try:
                        self.driver.get("https://genspark.ai")
                        time.sleep(15)
                    except Exception:
                        time.sleep(10)  # Pausa se la navigazione fallisce
                else:
                    self.add_log("❌ Tutti i tentativi globali falliti")
                    print("DEBUG: Tutti i tentativi globali falliti dopo tutti i retry")

        # Se siamo qui, tutti i tentativi globali sono falliti
        if cumulative_response:
            # Usa le risposte parziali se disponibili
            self.add_log("⚠️ Utilizzo risultati parziali dai tentativi falliti")
            print("DEBUG: Utilizzo risultati parziali dai tentativi falliti")
            combined_response = "\n\n".join(cumulative_response)
        
            # Debug risposta parziale finale
            print(f"DEBUG: Salvataggio risposta parziale finale - Lunghezza: {len(combined_response)}")
            print(f"DEBUG: Preview risposta parziale finale: {combined_response[:200].replace(chr(10), ' ')}...")
        
            return combined_response
        else:
            # Fallback di emergenza
            self.add_log("❌ Nessuna risposta ottenuta, utilizzo fallback di emergenza")
            print("DEBUG: Nessuna risposta ottenuta, utilizzo fallback di emergenza")
        
            if "CM-1" in prompt_id_to_use:
                keyword = project_data.get("KEYWORD", "keyword sconosciuta") if project_data else "keyword sconosciuta"
                fallback_response = f"""
                MARKET_INSIGHTS: Il mercato per {keyword} mostra un interesse crescente con volume di ricerca medio-alto. Si tratta di un mercato competitivo ma con spazio per nuovi contenuti di qualità che affrontino gap specifici.

                KEYWORD_DATA: La keyword principale {keyword} ha un volume di ricerca medio con competitività moderata. Le keyword correlate mostrano interesse per guide pratiche, soluzioni a problemi specifici e approcci step-by-step.

                BESTSELLER_OVERVIEW: I bestseller in questa nicchia tendono ad avere titoli diretti che promettono soluzioni concrete, utilizzano un linguaggio accessibile e offrono contenuti strutturati con esempi pratici e casi studio.
                """
                print(f"DEBUG: Utilizzato fallback per CM-1 con keyword '{keyword}'")
                return fallback_response
            elif "CM-2" in prompt_id_to_use:
                fallback_response = """
                STRUCTURE_PATTERNS: I bestseller in questa nicchia seguono una struttura organizzata in capitoli con progressione logica. Iniziano con un'introduzione al problema, seguita da capitoli che presentano soluzioni step-by-step, e terminano con esempi di applicazione e casi studio. La maggior parte include anche appendici con risorse aggiuntive.

                TITLE_PATTERNS: I titoli più efficaci utilizzano una combinazione di problemi e soluzioni, spesso con sottotitoli che espandono la promessa principale. Includono numeri specifici, utilizzano parole chiave come "guida", "manuale", "semplice" e mettono in evidenza i benefici.

                REVIEW_INSIGHTS: Le recensioni positive evidenziano contenuti pratici, chiarezza espositiva e applicabilità immediata. Le recensioni negative menzionano informazioni troppo generiche, mancanza di profondità e assenza di esempi concreti.

                IMPLEMENTATION_OBSTACLES: Le principali difficoltà di implementazione includono la complessità percepita, la mancanza di esercizi pratici e l'assenza di supporto continuo dopo la lettura.

                MARKET_GAPS: Esiste una chiara opportunità per un libro che combini teoria e pratica con un approccio passo-passo, materiali di supporto scaricabili e un linguaggio semplice ma professionale.
                """
                print("DEBUG: Utilizzato fallback per CM-2")
                return fallback_response
            else:
                fallback_response = f"[Risposta di fallback generica per {prompt_id_to_use}]"
                print(f"DEBUG: Utilizzato fallback generico per {prompt_id_to_use}")
                return fallback_response

    def check_unresolved_placeholders(self, text):
        """
        Verifica placeholders non risolti nel testo
        Delega alla funzione in ai_interfaces/file_text_utils.py
    
        Args:
            text: Testo da controllare
        
        Returns:
            list: Lista di placeholders non risolti o None se non ce ne sono
        """
        from ai_interfaces.file_text_utils import check_unresolved_placeholders as utils_check_unresolved_placeholders
    
        return utils_check_unresolved_placeholders(text)

    def get_clean_input_box(self):
        """
        Ottiene e pulisce completamente la casella di input
        Delega alla funzione in ai_interfaces/interaction_utils.py
        """
        from ai_interfaces.interaction_utils import get_clean_input_box as utils_get_clean_input_box
    
        return utils_get_clean_input_box(
            driver=self.driver,
            log_callback=self.add_log
        )

    # --- METODI DI SUPPORTO ---

    
    def safe_text_input(self, input_box, text):
        """
        Inserisce il testo in modo sicuro nella casella di input
        Delega alla funzione in ai_interfaces/interaction_utils.py
        """
        from ai_interfaces.interaction_utils import safe_text_input as utils_safe_text_input
    
        utils_safe_text_input(
            driver=self.driver,
            input_box=input_box,
            text=text,
            log_callback=self.add_log
        )

    def click_send_button(self):
        """
        Tenta di cliccare il pulsante di invio con metodi multipli
        Delega alla funzione in ai_interfaces/interaction_utils.py
        """
        from ai_interfaces.interaction_utils import click_send_button as utils_click_send_button
    
        return utils_click_send_button(
            driver=self.driver,
            log_callback=self.add_log
        )

    def wait_for_stable_response(self, max_wait_cycles=45, stability_threshold=10, cycle_wait=20):
        """
        Sistema avanzato di attesa per risposta stabile
        Delega alla funzione in ai_interfaces/interaction_utils.py
        """
        from ai_interfaces.interaction_utils import wait_for_stable_response as utils_wait_for_stable_response
    
        return utils_wait_for_stable_response(
            driver=self.driver,
            max_wait_cycles=max_wait_cycles,
            stability_threshold=stability_threshold,
            cycle_wait=cycle_wait,
            log_callback=self.add_log
        )

    def save_response_to_project(self, project_data, prompt_id, line, response, is_final):

            from framework.formatters import process_table_html, process_text

            # All'inizio del metodo save_response_to_project
            has_terminator = "FINE" in response
            terminator_pos = response.find("FINE") if has_terminator else -1

            if has_terminator:
                self.add_log(f"🛑 'FINE' rilevato alla posizione {terminator_pos}")
                # Imposta l'attributo per uso futuro
                self.last_response_complete = True
            else:
                self.add_log(f"⚠️ Risposta NON contiene terminatore FINE")
                # Imposta l'attributo per uso futuro
                self.last_response_complete = False

            """Salva la risposta nel framework CRISP in modo affidabile"""
            if not project_data or "PROJECT_ID" not in project_data:
                self.add_log("⚠️ Impossibile salvare: dati progetto mancanti")
                return False

            try:
                # Ottieni la keyword corrente per il nome del file di contesto
                keyword = project_data.get("KEYWORD", "unknown").strip()
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

                # Determina in quale fase dell'analisi siamo basato sul prompt_id
                analysis_phase = "Analisi"
                if prompt_id:
                    if isinstance(prompt_id, str):
                        if "market" in prompt_id.lower():
                            analysis_phase = "Analisi di Mercato"
                        elif "buyer" in prompt_id.lower():
                            analysis_phase = "Buyer Persona"
                        elif "gap" in prompt_id.lower():
                            analysis_phase = "Gap Analysis"

                # Anche se keyword è vuota, usa comunque nomi file significativi
                safe_keyword = "unknown"
                if keyword:
                    # Sanitizza la keyword per usarla come parte del nome file
                    import re
                    safe_keyword = re.sub(r'[\\/*?:"<>|]', "", keyword).replace(" ", "_")[:30]
    
                # Crea nomi file specifici per questa sessione di analisi
                context_filename = os.path.join(self.context_dir, f"context_{safe_keyword}.txt")
                html_filename = f"analisi_{safe_keyword}_{timestamp}.html"
    
                # Crea la directory output se non esiste
                os.makedirs("output", exist_ok=True)
                html_filepath = os.path.join("output", html_filename)
    
                # Se è il primo salvataggio per questa keyword, imposta il file di contesto
                if hasattr(self.chat_manager, 'context_file'):
                    self.chat_manager.context_file = context_filename
                    self.add_log(f"📄 File di contesto impostato a: {context_filename}")
        
                    # Crea una copia del contesto per questa keyword se non esiste
                    if not os.path.exists(context_filename) and os.path.exists("context.txt"):
                        import shutil
                        shutil.copy2("context.txt", context_filename)
                        self.add_log(f"📄 Creata copia del contesto per keyword: {keyword}")

                # Salva nel database CRISP
                success = False

                # Usa una struttura a cascata per trovare il metodo giusto
                if hasattr(self.crisp, 'save_incremental_response'):
                    self.crisp.save_incremental_response(
                        project_data["PROJECT_ID"], 
                        prompt_id,
                        line, 
                        response, 
                        is_final
                    )
                    success = True
                elif hasattr(self.crisp, 'crisp') and hasattr(self.crisp.crisp, 'save_incremental_response'):
                    self.crisp.crisp.save_incremental_response(
                        project_data["PROJECT_ID"], 
                        prompt_id,
                        line, 
                        response, 
                        is_final
                    )
                    success = True
                else:
                    self.add_log("⚠️ Metodo save_incremental_response non trovato")

                # Salva nel file di testo (sempre)
                try:
                    # Salva nel file di testo (append per mantenere la storia)
                    with open(context_filename, "a", encoding="utf-8") as f:
                        # Intestazione che identifica la sezione
                        f.write(f"\n=== {prompt_id or 'Analisi Legacy'}: {keyword} - {safe_keyword} - {timestamp} ===\n")
    
                        # Se la risposta inizia con il numero della domanda (es. "5) Identifica..."), rimuovilo
                        if re.match(r'^\s*\d+[\)\.]\s+', response):
                            # Prova a trovare una riga vuota dopo la domanda
                            lines = response.split('\n')
                            start_line = 0
                            first_content_line = 1  # Default: inizia dalla seconda riga
        
                            # Cerca una riga vuota dopo la prima riga (la domanda)
                            for i, line in enumerate(lines):
                                if i > 0 and not line.strip():  # Riga vuota trovata dopo la prima riga
                                    start_line = i + 1
                                    break
        
                            if start_line > 0 and start_line < len(lines):
                                # Prendi tutto dopo la riga vuota
                                response_only = '\n'.join(lines[start_line:])
                                self.add_log(f"✂️ Domanda rimossa dalla risposta (inizio dalla riga {start_line+1})")
                                f.write(response_only)
                            else:
                                # Se non troviamo una riga vuota, usa il pattern regex per rimuovere la prima riga
                                response_only = re.sub(r'^\s*\d+[\)\.]\s+.*?[\n\r]+', '', response)
                                self.add_log("✂️ Prima riga della domanda rimossa dalla risposta")
                                f.write(response_only)
                        else:
                            # Se non sembra iniziare con una domanda numerata, lascia invariato
                            f.write(response)
    
                        f.write("\n\n")

                    self.add_log(f"💾 Risposta salvata in {context_filename}")
        
                    # Salva in un file HTML separato (nuovo file per ogni risposta completa)
                    if is_final:
                        # Crea metadati per il file HTML
                        metadata = {
                            "Keyword": keyword,
                            "Data": datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
                            "Tipo di analisi": analysis_phase,
                            "Prompt ID": prompt_id or "N/A",
                            "Progetto ID": project_data.get("PROJECT_ID", "N/A")
                        }
            
                        # Crea CSS avanzato per tabelle e formattazione
                        css_styles = """
                        body { 
                            font-family: Arial, sans-serif; 
                            line-height: 1.6; 
                            max-width: 1000px; 
                            margin: 0 auto; 
                            padding: 20px;
                            background-color: #f9f9f9;
                        }
                        .header {
                            background-color: #2563eb;
                            color: white;
                            padding: 20px;
                            text-align: center;
                            border-radius: 8px;
                            margin-bottom: 20px;
                            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
                        }
                        h1, h2, h3, h4 { 
                            color: #2563eb; 
                            margin-top: 24px; 
                            margin-bottom: 16px; 
                        }
                        .metadata {
                            background-color: #f0f4f8;
                            padding: 15px;
                            border-radius: 8px;
                            margin-bottom: 20px;
                            border-left: 4px solid #2563eb;
                        }
                        table { 
                            border-collapse: collapse; 
                            width: 100%; 
                            margin: 20px 0;
                            background-color: white;
                            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                        }
                        th, td { 
                            border: 1px solid #ddd; 
                            padding: 12px; 
                            text-align: left; 
                        }
                        th { 
                            background-color: #f2f2f2; 
                            font-weight: bold;
                        }
                        tr:nth-child(even) { 
                            background-color: #f9f9f9; 
                        }
                        .content {
                            background-color: white;
                            border-radius: 8px;
                            padding: 20px;
                            margin-top: 20px;
                            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                        }
                        ul, ol {
                            margin-top: 15px;
                            margin-bottom: 15px;
                        }
                        li {
                            margin-bottom: 8px;
                        }
                        """
                
                        # Formatta il testo
                        formatted_response = response
                
                        # Conversione migliorata di tabelle
                        import re
                
                        # Rileva tabelle in formato markdown
                        table_pattern = r'(\|[^\n]+\|\n\|[\-\|: ]+\|\n(?:\|[^\n]+\|\n?)+)'
                
                        # Importa la funzione process_table_html dal modulo formatters
                        from framework.formatters import process_table_html, process_text

                        # Funzione wrapper che utilizza process_table_html
                        def format_markdown_table(match):
                            table_text = match.group(1)
                            return process_table_html(table_text)

                        # Applica la formattazione delle tabelle
                        formatted_response = re.sub(table_pattern, format_markdown_table, formatted_response, flags=re.MULTILINE)

                        # Converti elenchi puntati e numerati
                        formatted_response = self.convert_lists_to_html(formatted_response)

                        # Converti markdown di base in HTML
                        formatted_response = formatted_response.replace('\n\n', '</p><p>')
                        formatted_response = f"<p>{formatted_response}</p>"
                        formatted_response = formatted_response.replace('**', '<strong>').replace('**', '</strong>')
                        formatted_response = formatted_response.replace('*', '<em>').replace('*', '</em>')

                        # Aggiungi classi alle tabelle esistenti
                        formatted_response = formatted_response.replace('<table>', '<table class="data-table">')

                        # Creiamo l'HTML completo
                        html_content = f"""<!DOCTYPE html>
                        <html>
                        <head>
                            <meta charset="UTF-8">
                            <meta name="viewport" content="width=device-width, initial-scale=1.0">
                            <title>Analisi {keyword} - {analysis_phase}</title>
                            <style>
                            {css_styles}
                            </style>
                        </head>
                        <body>
                            <div class="header">
                                <h1>Analisi: {keyword}</h1>
                                <p>{analysis_phase}</p>
                            </div>

                            <div class="metadata">
                        """

                        # Aggiungi metadati
                        for key, value in metadata.items():
                            html_content += f"<p><strong>{key}:</strong> {value}</p>\n"

                        html_content += """
                            </div>

                            <div class="content">
                        """

                        html_content += formatted_response

                        html_content += """
                            </div>
                        </body>
                        </html>
                        """

                        # Salva il file HTML
                        with open(html_filepath, "w", encoding="utf-8") as f:
                            f.write(html_content)

                        self.add_log(f"✅ Report HTML completo salvato in: {html_filepath}")

                        # Apri automaticamente il file nel browser
                        import webbrowser
                        try:
                            file_path = os.path.abspath(html_file).replace('\\', '/')
                            webbrowser.open("file:///" + file_path)    
                            self.add_log("✅ File HTML aperto nel browser")
                        except Exception as browser_error:
                            self.add_log(f"⚠️ Impossibile aprire il browser: {str(browser_error)}")
                
                        # Crea anche una directory per i file formattati Genspark
                        os.makedirs("output/genspark_formatted", exist_ok=True)
                        genspark_html_path = os.path.join("output/genspark_formatted", f"{safe_keyword}_genspark_{timestamp}.html")
                
                        # Salva una copia del contesto in HTML stilizzato di alta qualità
                        # try:
                            # Se esiste il metodo per generare HTML formattato, usalo
                        #    if hasattr(self, '_generate_html_from_context_file'):
                                # Assicurati che l'attributo _analysis_session_id esista
                        #        if not hasattr(self, '_analysis_session_id'):
                        #            import uuid
                        #            self._analysis_session_id = str(uuid.uuid4())[:8]
                            
                        #        formatted_html = self._generate_html_from_context_file()
                        #        if formatted_html:
                        #            with open(genspark_html_path, "w", encoding="utf-8") as f:
                        #                f.write(formatted_html)
                        #            self.add_log(f"✅ HTML formattato Genspark salvato in: {genspark_html_path}")
                                   
                                     # Memorizza il percorso per uso futuro
                        #             self.last_html_path = genspark_html_path
                        # except Exception as html_error:
                        #    self.add_log(f"⚠️ Errore nella generazione HTML formattato: {str(html_error)}")
                
                        # Aggiorna il display dei risultati con l'HTML completo
                        if hasattr(self, 'results_display') and self.results_display is not None:
                            try:
                                if hasattr(self.results_display, 'update'):
                                    self.results_display.update(value=html_content)
                                else:
                                    self.results_display.value = html_content
                                self.add_log("✅ Visualizzazione risultati aggiornata")
                            except Exception as ui_error:
                                self.add_log(f"❌ Errore nell'aggiornamento UI: {str(ui_error)}")
                        
                        # Se vogliamo aprire il file nel browser
                        # import webbrowser
                        # webbrowser.open(f"file:///{os.path.abspath(html_filepath)}")

                except Exception as file_error:
                    self.add_log(f"⚠️ Errore nel salvare nel file: {str(file_error)}")

                return success
    
            except Exception as e:
                self.add_log(f"⚠️ Errore nel salvare la risposta: {str(e)}")
                import traceback
                self.add_log(traceback.format_exc())
                return False
    
    def convert_lists_to_html(self, text):
        """Converte liste testuali in liste HTML"""
        import re
    
        # Converte liste numeriche
        pattern = r'(\d+\.\s+.+?)(?=\n\d+\.|$)'
        if re.search(pattern, text, re.DOTALL):
            matches = re.finditer(pattern, text, re.DOTALL)
            html_list = "<ol>\n"
            for match in matches:
                item = match.group(1).strip()
                # Rimuovi il numero iniziale
                item = re.sub(r'^\d+\.\s+', '', item)
                html_list += f"  <li>{item}</li>\n"
            html_list += "</ol>"
            text = re.sub(pattern, '', text, flags=re.DOTALL)
            text += html_list
    
        # Converte liste puntate
        pattern = r'(•\s+.+?)(?=\n•|$)'
        if re.search(pattern, text, re.DOTALL):
            matches = re.finditer(pattern, text, re.DOTALL)
            html_list = "<ul>\n"
            for match in matches:
                item = match.group(1).strip()
                # Rimuovi il punto elenco iniziale
                item = re.sub(r'^•\s+', '', item)
                html_list += f"  <li>{item}</li>\n"
            html_list += "</ul>"
            text = re.sub(pattern, '', text, flags=re.DOTALL)
            text += html_list
    
        return text

    def convert_tables_to_html(self, text):
        """Converte tabelle testuali in tabelle HTML"""
        import re
    
        # Cerca pattern di tabelle (righe con | che si ripetono)
        # Esempio: | Colonna1 | Colonna2 | Colonna3 |
        table_pattern = r'(\|[\s\S]+?\|[\s\S]+?\|[\s\S]*?)(?=\n\s*\n|\Z)'
    
        matches = re.finditer(table_pattern, text, re.MULTILINE)
        for match in matches:
            table_text = match.group(1)
            rows = table_text.strip().split('\n')
        
            # Crea la tabella HTML
            html_table = "<table>\n"
        
            # La prima riga è l'intestazione
            if rows:
                html_table += "  <thead>\n    <tr>\n"
                headers = rows[0].strip().split('|')
                for header in headers:
                    if header.strip():
                        html_table += f"      <th>{header.strip()}</th>\n"
                html_table += "    </tr>\n  </thead>\n"
        
            # Le righe successive sono dati
            if len(rows) > 1:
                html_table += "  <tbody>\n"
                for i in range(1, len(rows)):
                    if rows[i].strip():
                        html_table += "    <tr>\n"
                        cells = rows[i].strip().split('|')
                        for cell in cells:
                            if cell.strip():
                                html_table += f"      <td>{cell.strip()}</td>\n"
                        html_table += "    </tr>\n"
                html_table += "  </tbody>\n"
        
            html_table += "</table>"
        
            # Sostituisci la tabella testuale con quella HTML
            text = text.replace(table_text, html_table)
    
        return text

    def process_combined_response(self, response_list, prompt_id, project_data):
        """Processa e salva la risposta combinata"""
        # Unisci le risposte
        combined = "\n\n".join(response_list)
        self.add_log(f"📋 Risposta combinata: {len(combined)} caratteri")
    
        # Salva nel contesto
        try:
            self.chat_manager.save_response(
                combined,
                f"Analisi CRISP 5.0 - {prompt_id}",
                {"prompt_id": prompt_id, "timestamp": datetime.now().strftime('%Y%m%d_%H%M%S')}
            )
            self.add_log("✅ Risposta salvata nel contesto")
        except Exception as e:
            self.add_log(f"⚠️ Errore nel salvare nel contesto: {str(e)}")
    
        # Verifica risposta corta e applica fallback se necessario
        if len(combined) < 200 and "CM-2" in prompt_id:
            fallback = self.get_fallback_response(prompt_id, project_data)
            self.add_log("⚠️ Applicato fallback per risposta troppo corta")
            return fallback
    
        return combined

    def get_fallback_response(self, prompt_id, project_data):
        """Fornisce risposte di fallback per casi specifici"""
        if "CM-2" in prompt_id:
            return """
            STRUCTURE_PATTERNS: I bestseller in questa nicchia seguono una struttura organizzata in capitoli con progressione logica. Iniziano con un'introduzione al problema, seguita da capitoli che presentano soluzioni step-by-step, e terminano con esempi di applicazione e casi studio. La maggior parte include anche appendici con risorse aggiuntive.

            TITLE_PATTERNS: I titoli più efficaci utilizzano una combinazione di problemi e soluzioni, spesso con sottotitoli che espandono la promessa principale. Includono numeri specifici, utilizzano parole chiave come "guida", "manuale", "semplice" e mettono in evidenza i benefici.

            REVIEW_INSIGHTS: Le recensioni positive evidenziano contenuti pratici, chiarezza espositiva e applicabilità immediata. Le recensioni negative menzionano informazioni troppo generiche, mancanza di profondità e assenza di esempi concreti.

            IMPLEMENTATION_OBSTACLES: Le principali difficoltà di implementazione includono la complessità percepita, la mancanza di esercizi pratici e l'assenza di supporto continuo dopo la lettura.

            MARKET_GAPS: Esiste una chiara opportunità per un libro che combini teoria e pratica con un approccio passo-passo, materiali di supporto scaricabili e un linguaggio semplice ma professionale.
            """
        elif "CM-1" in prompt_id:
            keyword = project_data.get("KEYWORD", "keyword sconosciuta") if project_data else "keyword sconosciuta"
            return f"""
            MARKET_INSIGHTS: Il mercato per {keyword} mostra un interesse crescente con volume di ricerca medio-alto. Si tratta di un mercato competitivo ma con spazio per nuovi contenuti di qualità che affrontino gap specifici.

            KEYWORD_DATA: La keyword principale {keyword} ha un volume di ricerca medio con competitività moderata. Le keyword correlate mostrano interesse per guide pratiche, soluzioni a problemi specifici e approcci step-by-step.

            BESTSELLER_OVERVIEW: I bestseller in questa nicchia tendono ad avere titoli diretti che promettono soluzioni concrete, utilizzano un linguaggio accessibile e offrono contenuti strutturati con esempi pratici e casi studio.
            """
        else:
            return f"[Risposta di fallback generica per {prompt_id}]"
                
    def execute_with_updates(self, func, *args, **kwargs):
        """
        Esegue una funzione aggiornando l'interfaccia periodicamente.
        Da usare per operazioni lunghe come l'analisi di mercato.
        Delega alla funzione in framework/utils.py

        Args:
            func: La funzione da eseguire
            *args, **kwargs: Argomenti per la funzione
    
        Returns:
            Il risultato finale della funzione
        """
        from framework.utils import execute_with_updates as utils_execute_with_updates
    
        result = utils_execute_with_updates(
            func=func,
            log_callback=self.add_log,
            *args,
            **kwargs
        )
    
        # Se la funzione nel modulo utils ha restituito None a causa di un errore,
        # restituisci il log delle chat
        if result is None and hasattr(self, 'chat_manager'):
            return self.chat_manager.get_log_history_string()
    
        return result

    


    def load_project_list(self):
        """Carica la lista dei progetti dal database"""
        try:
            conn = sqlite3.connect(self.crisp.project_db_path)
            cursor = conn.cursor()
        
            # Query migliorata che include anche informazioni sulla fase massima raggiunta
            query = """
            SELECT p.id, p.name, p.creation_date, 
                   (SELECT keyword FROM project_variables WHERE project_id = p.id AND name = 'KEYWORD' LIMIT 1) as keyword,
                   (SELECT COUNT(*) FROM project_results WHERE project_id = p.id) as results_count,
                   (SELECT prompt_id FROM project_results 
                    WHERE project_id = p.id 
                    ORDER BY id DESC LIMIT 1) as last_phase
            FROM projects p
            ORDER BY p.creation_date DESC
            """
        
            cursor.execute(query)
            projects = cursor.fetchall()
            conn.close()
        
            # Formatta i risultati
            formatted_projects = []
            for proj in projects:
                proj_id, name, date, keyword, results, last_phase = proj
                date_formatted = datetime.fromisoformat(date).strftime('%d/%m/%Y %H:%M')
            
                # Crea un nome display informativo
                if keyword:
                    display_name = f"{keyword} - {date_formatted} ({results} risultati)"
                else:
                    display_name = f"{name} - {date_formatted} ({results} risultati)"
                
                formatted_projects.append({
                    "id": proj_id,
                    "name": name,
                    "date": date_formatted,
                    "keyword": keyword or "N/A",
                    "results_count": results,
                    "last_phase": last_phase or "N/A",
                    "display": display_name
                })
        
            return formatted_projects
    
        except Exception as e:
            self.add_log(f"❌ Errore nel caricamento progetti: {str(e)}")
            return []

    def restart_current_analysis(self):
        """
        Riavvia l'analisi corrente usando i parametri salvati.
        """
        if not hasattr(self, 'current_analysis_params'):
            self.add_log("⚠️ Impossibile riavviare: nessuna analisi in corso")
            return False
    
        params = self.current_analysis_params
        self.add_log("🔄 Riavvio analisi con i seguenti parametri:")
        for key, value in params.items():
            if key != 'analysis_prompt':  # Evitiamo di loggare il prompt completo
                self.add_log(f"  - {key}: {value}")
    
        # Reset variabili di stato
        if hasattr(self, 'question_status'):
            self.question_status = {}
    
        # Riavvia l'analisi chiamando il metodo originale
        return self.analyze_market(
            book_type=params['book_type'],
            keyword=params['keyword'],
            language=params['language'],
            market=params['market'],
            analysis_prompt=params['analysis_prompt'],
            use_crisp=params['use_crisp']
        )

    



    def select_all_phases(self, analysis_type):
        """Seleziona tutte le fasi del tipo di analisi specificato"""
        try:
            # Preparare gli aggiornamenti per TUTTI i checkbox
            all_updates = []
        
            # Aggiornamenti per checkbox CRISP
            for _ in range(len(self.crisp_phase_checkboxes)):
                # Seleziona solo se è selezionato CRISP
                all_updates.append(gr.update(value=(analysis_type == "CRISP")))
        
            # Aggiornamenti per checkbox Legacy
            for _ in range(len(self.legacy_phase_checkboxes)):
                # Seleziona solo se è selezionato Legacy
                all_updates.append(gr.update(value=(analysis_type == "Legacy")))
        
            return all_updates
        except Exception as e:
            self.add_log(f"Errore nella selezione di tutte le fasi: {str(e)}")
            # Restituisci aggiornamenti vuoti per tutti i checkbox
            return [gr.update() for _ in range(len(self.crisp_phase_checkboxes) + len(self.legacy_phase_checkboxes))]

    def deselect_all_phases(self, analysis_type):
        """Deseleziona tutte le fasi del tipo di analisi specificato"""
        try:
            # Preparare gli aggiornamenti per TUTTI i checkbox
            all_updates = []
        
            # Aggiornamenti per checkbox CRISP
            for _ in range(len(self.crisp_phase_checkboxes)):
                # Imposta tutti a False
                all_updates.append(gr.update(value=False))
        
            # Aggiornamenti per checkbox Legacy
            for _ in range(len(self.legacy_phase_checkboxes)):
                # Imposta tutti a False
                all_updates.append(gr.update(value=False))
        
            return all_updates
        except Exception as e:
            self.add_log(f"Errore nella deselezione di tutte le fasi: {str(e)}")
            # Restituisci aggiornamenti vuoti per tutti i checkbox
            return [gr.update() for _ in range(len(self.crisp_phase_checkboxes) + len(self.legacy_phase_checkboxes))]

    
    def _filter_legacy_prompt_sections(self, analysis_prompt, selected_phases=None):
        """
        Filtra il prompt legacy per includere solo le sezioni selezionate.
        Ora delega alla funzione in framework/analysis/market_analysis.py
    
        Args:
            analysis_prompt: Prompt di analisi completo
            selected_phases: Lista dei numeri di fase da includere (se None, legge dai checkbox)
    
        Returns:
            str: Prompt filtrato contenente solo le sezioni selezionate
        """
        # Se selected_phases non è fornito, leggi lo stato attuale dei checkbox
        if selected_phases is None:
            selected_phases = []
            for phase_id, checkbox in self.legacy_phase_checkboxes.items():
                # Controlla esplicitamente lo stato attuale del checkbox
                try:
                    # Prova prima ad accedere al valore come proprietà
                    if hasattr(checkbox, "value") and checkbox.value:
                        selected_phases.append(phase_id)
                    # Poi prova ad accedere come elemento DOM
                    elif hasattr(checkbox, "get_value") and checkbox.get_value():
                        selected_phases.append(phase_id)
                except Exception as e:
                    self.add_log(f"Errore nel leggere stato checkbox {phase_id}: {str(e)}")

            # Log per debugging
            self.add_log(f"DEBUG: Fasi selezionate dai checkbox: {selected_phases}")
        
            # Assicurati che ci sia almeno una fase selezionata
            if not selected_phases:
                self.add_log("⚠️ Nessuna fase selezionata, utilizzo fase 1 come default")
                selected_phases = [1]  # Default alla prima fase se niente è selezionato
    
        from framework.analysis.market_analysis import filter_legacy_prompt_sections
    
        # Passa le fasi selezionate alla funzione delegata
        return filter_legacy_prompt_sections(
            analysis_prompt, 
            selected_phases, 
            log_callback=self.add_log
        )
    
    def take_debug_screenshot(self, prefix):
        """
        Scatta uno screenshot per debugging
        Delega alla funzione in ai_interfaces/browser_manager.py
    
        Args:
            prefix: Prefisso per il nome del file
        
        Returns:
            str: Nome del file screenshot o None in caso di errore
        """
        from ai_interfaces.browser_manager import take_debug_screenshot as manager_take_debug_screenshot
    
        return manager_take_debug_screenshot(
            driver=self.driver,
            prefix=prefix,
            log_callback=self.add_log
        )


    def get_last_response(self):
        """
        Recupera l'ultima risposta dalla chat con controlli migliorati per terminazione.
        Delega alla funzione in ai_interfaces/genspark_driver.py
        """
        from ai_interfaces.genspark_driver import get_last_response as driver_get_last_response
    
        return driver_get_last_response(self.driver, self.add_log)

    def handle_context_limit(self, driver=None):
        """
        Gestisce il limite di contesto in Genspark.
    
        Args:
            driver: WebDriver di Selenium (opzionale, usa self.driver se non specificato)
        """
        from ai_interfaces.browser_manager import reset_context_manual as manager_reset_context_manual

        # Definisci la funzione che riavvierà l'analisi se necessario
        def restart_analysis_if_needed():
            self.add_log("🔄 Riavvio dell'analisi dalla domanda #1 dopo nuova sessione")
    
            # Verifica che abbiamo i parametri dell'analisi corrente
            if not hasattr(self, 'current_analysis_params'):
                self.add_log("⚠️ Impossibile riavviare l'analisi: parametri non disponibili")
                return False
    
            # Impostiamo il flag per il riavvio
            self.restart_analysis_needed = True
            return True

        # Usa il driver fornito o quello dell'istanza
        browser_driver = driver if driver is not None else self.driver
    
        success, result = manager_reset_context_manual(
            driver=browser_driver,
            log_callback=self.add_log,
            restart_analysis_callback=restart_analysis_if_needed
        )

        return success
    def reset_context_manual(self, driver):
        """
        Reset completo del contesto: chiude la chat corrente, apre una nuova sessione,
        e ricarica il contesto se necessario.
        Delega alla funzione in ai_interfaces/browser_manager.py
    
        Returns:
            bool: True se il reset è riuscito, False altrimenti
        """
        from ai_interfaces.browser_manager import reset_context_manual as manager_reset_context_manual
    
        return manager_reset_context_manual(
            driver=driver,  # Per mantenere la compatibilità con la firma originale
            log_callback=self.add_log
        )

    def update_analysis_status(self, status_text, progress_percentage=None):
        """
        Aggiorna lo stato dell'analisi nell'interfaccia.
        Delega alla funzione in ui/interface_utils.py

        Args:
            status_text: Testo dello stato
            progress_percentage: Percentuale di completamento (opzionale)
        """
        from ui.interface_utils import update_analysis_status as utils_update_analysis_status
    
        utils_update_analysis_status(
            analysis_status=self.analysis_status if hasattr(self, 'analysis_status') else None,
            status_text=status_text,
            progress_percentage=progress_percentage,
            log_callback=self.add_log
        )

    def show_feedback(self, title, message, type="info"):
        """
        Mostra un messaggio di feedback all'utente.
        Delega alla funzione in ui/interface_utils.py

        Args:
            title: Titolo del messaggio
            message: Testo del messaggio
            type: Tipo di messaggio (info, success, warning, error)
        """
        from ui.interface_utils import show_feedback as utils_show_feedback
    
        utils_show_feedback(
            results_display=self.results_display if hasattr(self, 'results_display') else None,
            title=title,
            message=message,
            type=type,
            log_callback=self.add_log
        )


    def create_focused_context(self, original_context_file, max_size=8000):
        """
        Crea una versione condensata del contesto focalizzata sulle informazioni essenziali.
    
        Args:
            original_context_file: Percorso del file di contesto originale
            max_size: Dimensione massima in caratteri
        
        Returns:
            str: Percorso del file di contesto condensato o None in caso di errore
        """
        try:
            import re
        
            self.add_log(f"Creazione contesto condensato da {original_context_file}...")
        
            with open(original_context_file, 'r', encoding='utf-8') as f:
                full_content = f.read()
        
            # Estrai le sezioni più importanti
            important_sections = []
            condensed_content = ["CONTESTO CONDENSATO PER CONTINUAZIONE ANALISI:\n\n"]
        
            # 1. Estrai informazioni sul progetto corrente
            project_info = re.search(r'===\s+Analisi CRISP 5\.0[^=]+(.*?)(?=\n===|$)', full_content, re.DOTALL)
            if project_info:
                condensed_content.append(f"[INFORMAZIONI PROGETTO]\n{project_info.group(1).strip()}\n\n")
        
            # 2. Estrai variabili di mercato importanti
            key_vars = ["MARKET_INSIGHTS", "KEYWORD_DATA", "BESTSELLER_OVERVIEW", 
                       "STRUCTURE_PATTERNS", "TITLE_PATTERNS", "REVIEW_INSIGHTS"]
        
            for var in key_vars:
                var_pattern = re.compile(f"{var}[\\s]*:[\\s]*(.*?)(?=\\n[A-Z_]+:|$)", re.DOTALL)
                var_match = var_pattern.search(full_content)
                if var_match:
                    content = var_match.group(1).strip()
                    # Tronca a 500 caratteri se più lungo
                    if len(content) > 500:
                        content = content[:497] + "..."
                    condensed_content.append(f"{var}:\n{content}\n\n")
        
            # 3. Estrai l'ultima parte della conversazione (ultime 2-3 sezioni)
            sections = re.findall(r'===\s+([^=\n]+)\s+-\s+([^=\n]+)\s+===\n([\s\S]*?)(?=\n===|$)', full_content)
        
            # Prendi solo le ultime 3 sezioni
            recent_sections = sections[-3:] if len(sections) > 3 else sections
        
            for section_name, timestamp, content in recent_sections:
                # Tronca il contenuto se troppo lungo
                if len(content) > 1000:
                    content = content[:997] + "..."
                condensed_content.append(f"=== {section_name} - {timestamp} ===\n{content.strip()}\n\n")
        
            # Combina e verifica la dimensione
            final_content = "".join(condensed_content)
        
            # Se ancora troppo grande, tronca ulteriormente
            if len(final_content) > max_size:
                half_size = max_size // 2
                final_content = (
                    final_content[:half_size] + 
                    "\n\n[...CONTENUTO OMESSO PER LIMITI DI DIMENSIONE...]\n\n" + 
                    final_content[-half_size:]
                )
        
            # Salva il contesto condensato
            condensed_file = original_context_file.replace('.txt', '_condensed.txt')
            with open(condensed_file, 'w', encoding='utf-8') as f:
                f.write(final_content)
        
            self.add_log(f"Contesto condensato creato: {len(final_content)} caratteri")
            return condensed_file
    
        except Exception as e:
            self.add_log(f"❌ Errore nella creazione del contesto condensato: {str(e)}")
            return None

    def send_essential_context(self, driver, context_file):
        """
        Invia un riepilogo essenziale del contesto come messaggio diretto.
    
        Args:
            driver: WebDriver di Selenium
            context_file: Percorso del file di contesto
    
        Returns:
            bool: True se l'invio è riuscito, False altrimenti
        """
        try:
            self.add_log("Invio riepilogo essenziale del contesto...")
        
            # Estrai le informazioni chiave dal file di contesto
            with open(context_file, 'r', encoding='utf-8') as f:
                content = f.read()
        
            # Estrai il progetto e la keyword
            project_match = re.search(r'Progetto:\s*([^\n]+)', content)
            keyword_match = re.search(r'Keyword:\s*([^\n]+)', content)
            market_match = re.search(r'Mercato:\s*([^\n]+)', content)
        
            project = project_match.group(1).strip() if project_match else "Progetto sconosciuto"
            keyword = keyword_match.group(1).strip() if keyword_match else "Keyword sconosciuta"
            market = market_match.group(1).strip() if market_match else "Mercato sconosciuto"
        
            # Crea un messaggio di riepilogo conciso
            summary = f"""
            RIPRISTINO ANALISI CRISP:
        
            Stavo analizzando il mercato di "{keyword}" su {market}.
            ID progetto: {project}
        
            Ho dovuto aprire una nuova sessione per limiti di contesto.
            Sto continuando l'analisi dal punto in cui era stata interrotta.
        
            Per favore, continua l'analisi precedente.
            """
        
            # Trova l'input box
            input_box = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "div.search-input-wrapper textarea"))
            )
        
            # Pulisci l'input box
            input_box.clear()
            time.sleep(0.5)
        
            # Invia il riepilogo per piccoli blocchi
            for chunk in [summary[i:i+100] for i in range(0, len(summary), 100)]:
                input_box.send_keys(chunk)
                time.sleep(0.1)
        
            time.sleep(1)
        
            # Trova e clicca il pulsante di invio
            send_button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "div.search-input-wrapper div.input-icon"))
            )
            send_button.click()
        
            # Attendi l'elaborazione
            time.sleep(10)
        
            self.add_log("Riepilogo essenziale inviato con successo")
            return True
        
        except Exception as e:
            self.add_log(f"❌ Errore nell'invio del riepilogo essenziale: {str(e)}")
            return False


    


    


    

    def create_selection_dialog(self, titoli_options, indici_options, voice_style):
        """Crea una finestra di dialogo per selezionare titolo, indice e stile di voce"""
        import gradio as gr

        # Verifica che ci siano opzioni da mostrare
        if not titoli_options:
            titoli_options = [{"id": 1, "titolo": "Il tuo nuovo libro", "sottotitolo": "", "display": "Titolo predefinito"}]
        if not indici_options:
            indici_options = [{"id": 1, "content": "INTRODUZIONE\n\nCAPITOLO 1: Fondamenti\n\nCAPITOLO 2: Metodologia\n\nCAPITOLO 3: Applicazione\n\nCONCLUSIONE", "display": "Indice predefinito"}]

        # Creazione interfaccia di selezione
        with gr.Blocks(title="Selezione Opzioni") as selection_interface:
            with gr.Row():
                gr.Markdown("## Seleziona le opzioni per il tuo libro")
        
            # Titolo
            with gr.Row():
                with gr.Column(scale=3):
                    gr.Markdown("### Seleziona un titolo")
                    title_radio = gr.Radio(
                        choices=[t["display"] for t in titoli_options],
                        label="Titoli disponibili",
                        value=titoli_options[0]["display"] if titoli_options else None
                    )
            
                with gr.Column(scale=1):
                    gr.Markdown("### Anteprima")
                    title_preview = gr.Textbox(
                        label="Titolo selezionato",
                        value=titoli_options[0]["titolo"] if titoli_options else "",
                        interactive=False
                    )
                    subtitle_preview = gr.Textbox(
                        label="Sottotitolo",
                        value=titoli_options[0]["sottotitolo"] if titoli_options else "",
                        interactive=False
                    )
        
            # Indice
            with gr.Row():
                with gr.Column(scale=1):
                    gr.Markdown("### Seleziona un indice")
                    index_radio = gr.Radio(
                        choices=[idx["display"] for idx in indici_options],
                        label="Indici disponibili",
                        value=indici_options[0]["display"] if indici_options else None
                    )
            
                with gr.Column(scale=2):
                    gr.Markdown("### Anteprima indice")
                    index_preview = gr.TextArea(
                        label="Indice selezionato",
                        value=indici_options[0]["content"] if indici_options else "",
                        interactive=False,
                        lines=10
                    )
        
            # Stile di voce
            with gr.Row():
                voice_style_input = gr.Textbox(
                    label="Stile di voce",
                    value=voice_style,
                    interactive=True
                )
        
            # Pulsanti azione
            with gr.Row():
                confirm_btn = gr.Button("Conferma selezione", variant="primary")
                cancel_btn = gr.Button("Annulla", variant="secondary")
        
            # Funzioni di callback
            def update_title_preview(title_display):
                for t in titoli_options:
                    if t["display"] == title_display:
                        return t["titolo"], t["sottotitolo"]
                return "", ""
        
            def update_index_preview(index_display):
                for idx in indici_options:
                    if idx["display"] == index_display:
                        return idx["content"]
                return ""
        
            def confirm_selection(title_display, index_display, voice_style):
                # Trova il titolo selezionato
                selected_title = ""
                selected_subtitle = ""
                for t in titoli_options:
                    if t["display"] == title_display:
                        selected_title = t["titolo"]
                        selected_subtitle = t["sottotitolo"]
                        break
            
                # Trova l'indice selezionato
                selected_index = ""
                for idx in indici_options:
                    if idx["display"] == index_display:
                        selected_index = idx["content"]
                        break
            
                # Aggiorna i campi nella scheda di generazione del libro
                full_title = f"{selected_title}" + (f" - {selected_subtitle}" if selected_subtitle else "")
            
                # Aggiorna i campi dell'interfaccia principale
                if hasattr(self, 'book_title'):
                    self.book_title.update(value=full_title)
            
                if hasattr(self, 'book_index'):
                    self.book_index.update(value=selected_index)
            
                if hasattr(self, 'voice_style'):
                    self.voice_style.update(value=voice_style)
            
                self.add_log(f"✅ Selezione confermata: Titolo='{full_title}', Stile='{voice_style}'")
            
                # Chiudi la finestra di dialogo
                return gr.update(visible=False)
        
            def cancel_selection():
                # Usa i valori predefiniti
                default_title = titoli_options[0]["titolo"] + " - " + titoli_options[0]["sottotitolo"] if titoli_options else "Il tuo nuovo libro"
                default_index = indici_options[0]["content"] if indici_options else "INTRODUZIONE\n\nCAPITOLO 1: Fondamenti\n\nCAPITOLO 2: Metodologia\n\nCAPITOLO 3: Applicazione\n\nCONCLUSIONE"
            
                # Aggiorna i campi nella scheda di generazione del libro
                if hasattr(self, 'book_title'):
                    self.book_title.update(value=default_title)
            
                if hasattr(self, 'book_index'):
                    self.book_index.update(value=default_index)
            
                if hasattr(self, 'voice_style'):
                    self.voice_style.update(value=voice_style)
            
                self.add_log("⚠️ Selezione annullata, usati valori predefiniti")
            
                # Chiudi la finestra di dialogo
                return gr.update(visible=False)
        
            # Connessione dei callback
            title_radio.change(
                fn=update_title_preview,
                inputs=[title_radio],
                outputs=[title_preview, subtitle_preview]
            )
        
            index_radio.change(
                fn=update_index_preview,
                inputs=[index_radio],
                outputs=[index_preview]
            )
        
            confirm_btn.click(
                fn=confirm_selection,
                inputs=[title_radio, index_radio, voice_style_input],
                outputs=[selection_interface]
            )
        
            cancel_btn.click(
                fn=cancel_selection,
                outputs=[selection_interface]
            )
    
        # Lancia l'interfaccia in una nuova finestra
        selection_interface.launch(inbrowser=True, prevent_thread_lock=True)
        self.add_log("🔍 Finestra di selezione aperta. Scegli le opzioni desiderate.")

    def handle_selection_result(self, title_display, index_display, voice_style):
        """Gestisce il risultato della selezione dalla finestra di dialogo"""
        try:
            title_id = None
            index_id = None
        
            # Trova l'ID del titolo selezionato
            for t in self.temp_titles:
                if t["display"] == title_display:
                    title_id = t["id"]
                    break
        
            # Trova l'ID dell'indice selezionato
            for idx in self.temp_indices:
                if idx["display"] == index_display:
                    index_id = idx["id"]
                    break
        
            # Aggiorna i campi nella scheda di generazione del libro
            if title_id is not None:
                selected_title = next((t for t in self.temp_titles if t["id"] == title_id), None)
                if selected_title:
                    full_title = f"{selected_title['titolo']}" + (f" - {selected_title['sottotitolo']}" if selected_title['sottotitolo'] else "")
                    self.book_title.update(value=full_title)
        
            if index_id is not None:
                selected_index = next((idx for idx in self.temp_indices if idx["id"] == index_id), None)
                if selected_index:
                    self.book_index.update(value=selected_index["content"])
        
            # Aggiorna lo stile di voce
            self.voice_style.update(value=voice_style)
        
            self.add_log(f"✅ Selezione applicata alla scheda di generazione del libro")
            return True
        
        except Exception as e:
            self.add_log(f"❌ Errore nell'applicazione della selezione: {str(e)}")
            return False
    
    def load_from_analysis(self):
        """Carica titolo, indice e tono di voce dall'analisi"""
        try:
            self.add_log("🔄 Tentativo di caricare dati dall'analisi...")
    
            # Verifica se l'analisi corrente esiste
            if not hasattr(self, 'current_analysis') or not self.current_analysis:
                # Prova a recuperare l'ultimo progetto dal database come fallback
                self.add_log("⚠️ Nessuna analisi corrente trovata, tentativo di recupero dall'ultimo progetto...")
        
                try:
                    project_id = self.recupera_ultimo_progetto()
                    self.add_log(f"DEBUG: recupera_ultimo_progetto ha restituito: {type(project_id)}")

                    # Verifica se è un ID progetto (stringa)
                    if project_id and isinstance(project_id, str):
                        self.add_log(f"✅ Trovato ID ultimo progetto: {project_id}")
                    
                        # Usa load_project_details per caricare i dettagli del progetto
                        project_details = self.load_project_details(project_id)
                    
                        if project_details:
                            self.add_log(f"✅ Dettagli progetto caricati con successo")
                        
                            # IMPORTANTE: Assegna i dettagli a current_analysis
                            self.current_analysis = {
                                'project_id': project_id,
                                'project_data': {}  # Inizializza un dizionario vuoto per i dati
                            }
                        
                            # Se i dettagli sono in formato HTML, estrai i dati
                            if isinstance(project_details, str) and project_details.strip().startswith("<"):
                                try:
                                    from bs4 import BeautifulSoup
                                    soup = BeautifulSoup(project_details, 'html.parser')
                                
                                    # Cerca tabelle o definizioni
                                    project_data = {}
                                
                                    # Cerca nelle tabelle
                                    for table in soup.find_all('table'):
                                        for row in table.find_all('tr'):
                                            cells = row.find_all(['td', 'th'])
                                            if len(cells) >= 2:
                                                key = cells[0].text.strip()
                                                value = cells[1].text.strip()
                                                project_data[key] = value
                                
                                    # Cerca nelle definizioni
                                    for dt in soup.find_all('dt'):
                                        if dt.next_sibling and dt.next_sibling.name == 'dd':
                                            key = dt.text.strip()
                                            value = dt.next_sibling.text.strip()
                                            project_data[key] = value
                                
                                    # Cerca nei paragrafi
                                    for p in soup.find_all('p'):
                                        text = p.text.strip()
                                        if ':' in text:
                                            parts = text.split(':', 1)
                                            key = parts[0].strip()
                                            value = parts[1].strip()
                                            project_data[key] = value
                                
                                    if project_data:
                                        self.current_analysis['project_data'] = project_data
                                        self.add_log(f"✅ Estratti {len(project_data)} campi dall'HTML")
                                    else:
                                        # Salva il contenuto originale come fallback
                                        self.current_analysis['html_content'] = project_details
                                        self.add_log("⚠️ Nessun dato strutturato estratto dall'HTML")
                                except ImportError:
                                    self.add_log("⚠️ BeautifulSoup non disponibile, impossibile analizzare HTML")
                                    # Salva il contenuto originale
                                    self.current_analysis['html_content'] = project_details
                                except Exception as html_error:
                                    self.add_log(f"⚠️ Errore nell'estrazione dall'HTML: {str(html_error)}")
                                    # Salva il contenuto originale
                                    self.current_analysis['html_content'] = project_details
                        
                            # Se è già un dizionario, usalo direttamente
                            elif isinstance(project_details, dict):
                                self.current_analysis['project_data'] = project_details
                                self.add_log(f"✅ Dizionario caricato direttamente con {len(project_details)} campi")
                        
                            # Prova anche a caricare con CRISP per sicurezza
                            if hasattr(self, 'crisp') and hasattr(self.crisp, '_load_project_data'):
                                try:
                                    crisp_data = self.crisp._load_project_data(project_id)
                                    if crisp_data:
                                        self.current_analysis['project_data'] = crisp_data
                                        self.add_log(f"✅ Dati progetto aggiornati con CRISP")
                                except Exception as crisp_error:
                                    self.add_log(f"⚠️ Errore caricamento CRISP: {str(crisp_error)}")
                        else:
                            self.add_log("⚠️ Nessun dettaglio progetto trovato per questo ID")
                            return "Dettagli progetto non trovati", "", "", "ID esistente ma dettagli mancanti"
                    elif isinstance(project_id, dict):
                        # Nel caso in cui restituisca un dizionario completo
                        self.current_analysis = project_id
                        self.add_log(f"✅ Ultimo progetto caricato in current_analysis come dizionario")
                    else:
                        self.add_log("⚠️ Nessun progetto recente trovato nel database")
                        return "Nessuna analisi trovata", "", "", "Usa l'analisi per generare questi dati"

                except Exception as db_error:
                    self.add_log(f"⚠️ Errore nel recupero dell'ultimo progetto: {str(db_error)}")
                    import traceback
                    self.add_log(traceback.format_exc())
                    return "Errore database", "", "", "Verifica connessione al database"
    
            # Stampiamo info di debug sull'analisi corrente
            if isinstance(self.current_analysis, dict):
                self.add_log(f"🔍 Analisi trovata con {len(self.current_analysis)} elementi")
                for key in list(self.current_analysis.keys())[:5]:  # Mostriamo solo le prime 5 chiavi
                    self.add_log(f"- Chiave: {key}")
            
                # Se c'è project_data, mostriamo anche quelle chiavi
                if 'project_data' in self.current_analysis and isinstance(self.current_analysis['project_data'], dict):
                    self.add_log(f"🔍 Project data con {len(self.current_analysis['project_data'])} variabili")
                    for key in list(self.current_analysis['project_data'].keys())[:5]:
                        self.add_log(f"- Variabile: {key}")
            else:
                self.add_log(f"⚠️ current_analysis è di tipo {type(self.current_analysis)}, non un dizionario")
                return "Formato analisi non valido", "", "", "Errore nel formato dati"
        
            # Estrazione del titolo
            title = ""
            project_data = self.current_analysis.get('project_data', {})
        
            if 'TITOLO_LIBRO' in project_data:
                title = project_data['TITOLO_LIBRO']
            else:
                # Cerca in vari campi possibili
                for key in ['TITOLO', 'BOOK_TITLE', 'SELECTED_TITLE']:
                    if key in project_data and project_data[key]:
                        title = project_data[key]
                        self.add_log(f"✅ Titolo trovato in campo {key}")
                        break
                    
            # Se ancora non abbiamo trovato un titolo, cerchiamo nella root dell'analisi
            if not title:
                for key in ['TITOLO_LIBRO', 'TITOLO', 'BOOK_TITLE']:
                    if key in self.current_analysis and self.current_analysis[key]:
                        title = self.current_analysis[key]
                        self.add_log(f"✅ Titolo trovato nella root dell'analisi: {key}")
                        break
        
            # Estrazione dell'indice
            book_index = ""
            if 'INDICE_LIBRO' in project_data:
                book_index = project_data['INDICE_LIBRO']
            else:
                # Cerca in vari campi possibili
                for key in ['INDICE', 'BOOK_INDEX', 'SELECTED_INDEX']:
                    if key in project_data and project_data[key]:
                        book_index = project_data[key]
                        self.add_log(f"✅ Indice trovato in campo {key}")
                        break
                    
            # Se ancora non abbiamo trovato un indice, cerchiamo nella root dell'analisi
            if not book_index:
                for key in ['INDICE_LIBRO', 'INDICE', 'BOOK_INDEX']:
                    if key in self.current_analysis and self.current_analysis[key]:
                        book_index = self.current_analysis[key]
                        self.add_log(f"✅ Indice trovato nella root dell'analisi: {key}")
                        break
        
            # Estrazione del tono di voce
            voice_style = ""
            if 'VOICE_STYLE' in project_data:
                voice_style = project_data['VOICE_STYLE']
            else:
                # Cerca in vari campi possibili
                for key in ['TONE', 'STYLE', 'WRITING_STYLE']:
                    if key in project_data and project_data[key]:
                        voice_style = project_data[key]
                        self.add_log(f"✅ Stile di voce trovato in campo {key}")
                        break
                    
            # Se ancora non abbiamo trovato uno stile, cerchiamo nella root dell'analisi
            if not voice_style:
                for key in ['VOICE_STYLE', 'TONE', 'STYLE']:
                    if key in self.current_analysis and self.current_analysis[key]:
                        voice_style = self.current_analysis[key]
                        self.add_log(f"✅ Stile di voce trovato nella root dell'analisi: {key}")
                        break
        
            # Estrazione del tipo di libro
            book_type = ""
            if 'LIBRO_TIPO' in project_data:
                book_type = project_data['LIBRO_TIPO']
            else:
                # Cerca in vari campi possibili
                for key in ['BOOK_TYPE', 'TIPO', 'GENRE']:
                    if key in project_data and project_data[key]:
                        book_type = project_data[key]
                        self.add_log(f"✅ Tipo di libro trovato in campo {key}")
                        break
                    
            # Se ancora non abbiamo trovato un tipo, cerchiamo nella root dell'analisi
            if not book_type:
                for key in ['LIBRO_TIPO', 'BOOK_TYPE', 'TIPO']:
                    if key in self.current_analysis and self.current_analysis[key]:
                        book_type = self.current_analysis[key]
                        self.add_log(f"✅ Tipo di libro trovato nella root dell'analisi: {key}")
                        break
        
            # Prova a estrarre dati dal campo ANALISI se disponibile
            if (not title or not book_index or not voice_style or not book_type) and 'project_data' in self.current_analysis:
                project_data = self.current_analysis['project_data']
            
                # Se abbiamo un campo ANALISI, proviamo a estrarre i dati da lì
                if 'ANALISI' in project_data and project_data['ANALISI']:
                    analysis_text = project_data['ANALISI']
                    self.add_log(f"🔍 Tentativo di estrazione dai dati di analisi ({len(analysis_text)} caratteri)")
                
                    # Cerca titoli nel testo dell'analisi se non lo abbiamo ancora
                    if not title:
                        import re
                        title_patterns = [
                            r'(?:titolo|title)[:\s]+[""]?([^"\n,;]+)[""]?',
                            r'(?:TITOLO|TITLE)[:\s]+[""]?([^"\n,;]+)[""]?',
                            r'(?:libro)[:\s]+[""]?([^"\n,;]+)[""]?'
                        ]
                    
                        for pattern in title_patterns:
                            title_match = re.search(pattern, analysis_text, re.IGNORECASE)
                            if title_match:
                                title = title_match.group(1).strip()
                                self.add_log(f"✅ Titolo estratto dal testo dell'analisi: {title}")
                                break
                
                    # Estrai indice se non lo abbiamo ancora
                    if not book_index:
                        import re
                        index_patterns = [
                            r'(?:INDICE|INDEX)[:\s]+((?:CAPITOLO|CHAPTER)[\s\S]+?(?=\n\s*\n|\Z))',
                            r'(INTRODUZIONE\s*\n+\s*CAPITOLO[\s\S]+?(?=\n\s*\n|\Z))',
                            r'((?:CAPITOLO|CHAPTER)\s+\d+[:\s]+[^\n]+(?:\n+(?:CAPITOLO|CHAPTER)\s+\d+[:\s]+[^\n]+)+)'
                        ]
                    
                        for pattern in index_patterns:
                            index_match = re.search(pattern, analysis_text, re.IGNORECASE)
                            if index_match:
                                book_index = index_match.group(1).strip()
                                self.add_log(f"✅ Indice estratto dal testo dell'analisi: {len(book_index)} caratteri")
                                break
                
                    # Estrai tono di voce se non lo abbiamo ancora
                    if not voice_style:
                        import re
                        voice_patterns = [
                            r'(?:tono di voce|voice style|stile)[:\s]+([^\n\.;]+)',
                            r'(?:VOICE_STYLE|TONE)[:\s]+([^\n\.;]+)'
                        ]
                    
                        for pattern in voice_patterns:
                            voice_match = re.search(pattern, analysis_text, re.IGNORECASE)
                            if voice_match:
                                voice_style = voice_match.group(1).strip()
                                self.add_log(f"✅ Tono di voce estratto dal testo dell'analisi: {voice_style}")
                                break
                
                    # Estrai tipo di libro se non lo abbiamo ancora
                    if not book_type:
                        import re
                        type_patterns = [
                            r'(?:tipo di libro|book type)[:\s]+([^\n\.;]+)',
                            r'(?:genere|genre)[:\s]+([^\n\.;]+)'
                        ]
                    
                        for pattern in type_patterns:
                            type_match = re.search(pattern, analysis_text, re.IGNORECASE)
                            if type_match:
                                book_type = type_match.group(1).strip()
                                self.add_log(f"✅ Tipo di libro estratto dal testo dell'analisi: {book_type}")
                                break

            # Usa valori predefiniti se necessario
            if not title:
                title = "Titolo del tuo libro"  # Valore predefinito
                self.add_log("⚠️ Usando titolo predefinito")

            if not book_index:
                book_index = """INTRODUZIONE

    CAPITOLO 1: Fondamenti

    CAPITOLO 2: Metodologia

    CAPITOLO 3: Applicazione

    CAPITOLO 4: Esempi Pratici

    CONCLUSIONE"""
                self.add_log("⚠️ Usando indice predefinito")

            if not voice_style:
                voice_style = "Conversazionale e informativo"
                self.add_log("⚠️ Usando stile di voce predefinito")

            if not book_type:
                book_type = "Manuale (Non-Fiction)"
                self.add_log("⚠️ Usando tipo libro predefinito")
        
            # Messaggi di log sui risultati finali
            self.add_log(f"✅ Titolo finale: {title}")
            self.add_log(f"✅ Indice finale: {len(book_index)} caratteri")
            self.add_log(f"✅ Stile di voce finale: {voice_style}")
            self.add_log(f"✅ Tipo di libro finale: {book_type}")
        
            return title, book_index, voice_style, book_type
        
        except Exception as e:
            self.add_log(f"❌ Errore nel caricamento dati dall'analisi: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return "", "", "", "Errore nel caricamento"

    def update_image_settings(self, book_type, content_type, step_by_step):
        """Aggiorna automaticamente le impostazioni delle immagini in base alle scelte"""
        recommended_images = "No"
        notes = ""
    
        # Regole per impostare automaticamente l'uso di immagini
        if book_type == "Ricettario" or content_type == "Ricette":
            recommended_images = "Sì"
            if step_by_step == "Sì":
                notes = "Configurazione raccomandata: 3-4 immagini per ricetta, mostrando le fasi principali."
            else:
                notes = "Configurazione raccomandata: 1 immagine per ricetta, mostrando il piatto finito."
    
        elif book_type == "Craft & Hobby" or content_type == "Progetti":
            recommended_images = "Sì"
            if step_by_step == "Sì":
                notes = "Configurazione raccomandata: 4-6 immagini per progetto, illustrando ogni passaggio importante."
            else:
                notes = "Configurazione raccomandata: 1-2 immagini per progetto."
    
        elif content_type == "Esercizi" or content_type == "Guide":
            recommended_images = "Sì"
            if step_by_step == "Sì":
                notes = "Configurazione raccomandata: Immagini per ogni passo fondamentale degli esercizi."
            else:
                notes = "Configurazione raccomandata: 1 immagine per esercizio/guida."
    
        elif book_type == "Manuale (Non-Fiction)" or book_type == "Test Study":
            if step_by_step == "Sì":
                recommended_images = "Sì"
                notes = "Configurazione raccomandata: Immagini per concetti chiave e diagrammi."
            else:
                notes = "Configurazione raccomandata: 1 immagine introduttiva per capitolo."
    
        return recommended_images, notes

    def transfer_to_book_tab(self):
        """
        Trasferisce i dati dall'analisi alla tab di generazione del libro.
        Include la pulizia dei dati estratti dall'HTML.
        """
        try:
            # Estrai i dati dall'analisi
            book_data = self.extract_data_for_book_generation()
            if not book_data:
                self.add_log("⚠️ Nessun dato disponibile da trasferire")
                return "Nessun dato disponibile", "", "", "", "", ""
        
            # Funzione per pulire i valori dai tag HTML e codifiche
            def clean_value(value):
                if not value:
                    return ""
            
                # Rimuovi HTML entities e caratteri speciali
                import re
                import html
            
                # Decodifica HTML entities come &quot;
                value = html.unescape(value)
            
                # Rimuovi tag HTML rimanenti
                value = re.sub(r'<[^>]+>', '', value)
            
                # Rimuovi prefissi come "Percepita: " o simili
                prefixes = ["Percepita:", "Percepita", "strutturato:", "strutturato", "Persona Narrativa"]
                for prefix in prefixes:
                    if value.startswith(prefix):
                        value = value[len(prefix):].strip()
            
                return value.strip()
        
            # Prepara i dati per l'aggiornamento dei componenti
            title = clean_value(book_data.get('title', ''))
            subtitle = clean_value(book_data.get('subtitle', ''))
            voice_style = clean_value(book_data.get('voice_style', ''))
            angle = clean_value(book_data.get('angle', ''))
            usp = clean_value(book_data.get('usp', ''))
            book_index = book_data.get('index', '')
        
            # Pulisci anche l'indice, ma preserva la formattazione delle righe
            if book_index:
                # Decodifica HTML entities nell'indice
                import html
                book_index = html.unescape(book_index)
            
                # Rimuovi tag HTML preservando interruzioni di riga
                import re
                book_index = re.sub(r'<br\s*/?>|<p>|</p>', '\n', book_index)
                book_index = re.sub(r'<[^>]+>', '', book_index)
            
                # Pulisci spazi e righe vuote multiple
                book_index = re.sub(r'\n\s*\n', '\n\n', book_index)
                book_index = book_index.strip()
        
            # Log delle informazioni pulite
            self.add_log(f"📚 Preparazione dati per generazione libro:")
            self.add_log(f"✓ Titolo: {title}")
            if subtitle:
                self.add_log(f"✓ Sottotitolo: {subtitle}")
            self.add_log(f"✓ Tono di voce: {voice_style}")
            self.add_log(f"✓ Angolo editoriale: {angle}")
            self.add_log(f"✓ USP: {usp}")
            self.add_log(f"✓ Indice: {len(book_index)} caratteri")
        
            # Prepara titolo completo
            full_title = f"{title}" + (f" - {subtitle}" if subtitle else "")
        
            self.add_log("✅ Dati trasferiti e puliti per la tab di generazione libro")
        
            # Restituisci i valori puliti per aggiornare tutti i componenti
            return (
                "Dati trasferiti con successo. Ora puoi passare alla tab 'Generazione Libro'.",
                full_title,  # book_title 
                voice_style,  # voice_style
                angle,  # editorial_angle
                usp,  # book_usp
                book_index  # book_index
            )
        
        except Exception as e:
            self.add_log(f"❌ Errore nel trasferimento dei dati: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return (
                f"Errore: {str(e)}",
                "",  # book_title
                "",  # voice_style 
                "",  # editorial_angle
                "",  # book_usp
                ""   # book_index
            )

    def refresh_projects(self):
        """Aggiorna la lista dei progetti disponibili e aggiorna il dropdown"""
        try:
            self.add_log("🔄 Aggiornamento lista progetti...")
            projects = self.load_projects_list() if hasattr(self, 'load_projects_list') else []
            return gr.update(choices=projects)
        except Exception as e:
            self.add_log(f"❌ Errore aggiornamento progetti: {str(e)}")
            return gr.update()

    def generate_book_enhanced(self, book_title, book_language, voice_style, book_index, 
                               book_type, content_type, use_images, step_by_step, include_analysis, custom_notes):
        """Versione migliorata del metodo generate_book che considera i nuovi parametri"""
        try:
            self.add_log(f"📚 Avvio generazione libro: '{book_title}'")
            self.add_log(f"📋 Tipo: {book_type}, Contenuto: {content_type}, Immagini: {use_images}, Step-by-step: {step_by_step}")
        
            # Preparazione dei parametri per il prompt
            params = {
                "book_title": book_title,
                "book_language": book_language,
                "voice_style": voice_style,
                "book_index": book_index,
                "book_type": book_type,
                "content_type": content_type,
                "use_images": use_images,
                "step_by_step": step_by_step,
                "custom_notes": custom_notes
            }
        
            # Aggiungi contesto dell'analisi se richiesto
            analysis_context = ""
            if include_analysis:
                analysis_context = self.prepare_analysis_context()
                self.add_log(f"✅ Contesto dell'analisi incluso: {len(analysis_context)} caratteri")
        
            # Determina quale metodo di generazione utilizzare
            if hasattr(self, 'use_crisp') and self.use_crisp:
                result = self._generate_book_crisp(
                    book_title, book_language, voice_style, book_index, 
                    params, analysis_context
                )
            else:
                result = self._generate_book_legacy(
                    book_title, book_language, voice_style, book_index,
                    params, analysis_context
                )
            
            return self.chat_manager.get_log_history_string()
        
        except Exception as e:
            self.add_log(f"❌ Errore nella generazione del libro: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return self.chat_manager.get_log_history_string()

    def prepare_analysis_context(self):
        """Prepara e formatta il contesto dell'analisi per la generazione del libro"""
        try:
            context = ""
        
            # Prova a caricare dal file HTML formattato più recente
            keyword = self.get_current_keyword()
            if keyword:
                # Cerca il file HTML più recente
                html_path = f"output/genspark_formatted/{keyword}_genspark_current.html"
                if os.path.exists(html_path):
                    self.add_log(f"📄 Trovato file HTML formattato per '{keyword}'")
                
                    # Estrai il contenuto utile dall'HTML
                    import re
                    from bs4 import BeautifulSoup
                
                    with open(html_path, "r", encoding="utf-8") as f:
                        html_content = f.read()
                
                    soup = BeautifulSoup(html_content, 'html.parser')
                
                    # Estrai solo il testo rilevante
                    content_div = soup.find('div', {'class': 'genspark-response'})
                    if content_div:
                        text_content = content_div.get_text(separator="\n\n")
                        context = f"=== ANALISI DI MERCATO PER '{keyword}' ===\n\n{text_content}\n\n"
                    else:
                        # Fallback all'estrazione basica
                        text_content = re.sub(r'<[^>]+>', ' ', html_content)
                        text_content = re.sub(r'\s+', ' ', text_content).strip()
                        context = f"=== ANALISI DI MERCATO PER '{keyword}' ===\n\n{text_content}\n\n"
                else:
                    # Fallback al file di contesto testuale
                    context_file = f"context_{keyword}.txt"
                    if os.path.exists(context_file):
                        with open(context_file, "r", encoding="utf-8") as f:
                            context = f.read()
                    
                        self.add_log(f"📄 Utilizzato file di contesto testuale per '{keyword}'")
                    else:
                        self.add_log("⚠️ Nessun file di analisi trovato")
        
            return context
        except Exception as e:
            self.add_log(f"⚠️ Errore nella preparazione del contesto dell'analisi: {str(e)}")
            return ""
    
    
    
    def split_prompt(self, text, prompt_id=None, section_number=None):
        """
        Divide il prompt in sezioni numeriche mantenendo l'integrità.
        Aggiunto tracciamento della posizione nel flusso.
        Delega alla funzione in ai_interfaces/interaction_utils.py
        """
        from ai_interfaces.interaction_utils import split_prompt as utils_split_prompt
    
        return utils_split_prompt(
            text=text,
            prompt_id=prompt_id,
            section_number=section_number,
            log_callback=self.add_log
        )

    def send_to_genspark(self, text, prompt_id=None, section_number=None):
        """
        Invia un messaggio a Genspark e attende la risposta.
        Versione completamente rivista per garantire stabilità e affidabilità,
        con aggiornamento dell'interfaccia utente integrato.
        """
        from ai_interfaces.genspark_driver import send_to_genspark as genspark_send
    
        return genspark_send(
            driver=self.driver,
            text=text,
            log_callback=self.add_log,
            prompt_id=prompt_id,
            section_number=section_number,
            cooldown_manager=getattr(self, 'cooldown_manager', None),
            chat_manager=getattr(self, 'chat_manager', None),
            results_display=getattr(self, 'results_display', None)
        )

    # Altre funzioni di supporto da aggiungere a genspark_driver.py:

    def get_last_response(driver, log_callback=None):
        """
        Ottiene l'ultima risposta da Genspark.
    
        Args:
            driver: WebDriver di Selenium
            log_callback: Funzione di callback per il logging
        
        Returns:
            str: Ultima risposta di Genspark
        """
        try:
            selectors = [
                ".message-content", 
                "div.chat-wrapper div.desc > div > div > div",
                "div.message div.text-wrap",
                ".chat-message-item .content"
            ]
        
            for selector in selectors:
                try:
                    messages = driver.find_elements(By.CSS_SELECTOR, selector)
                    if messages and len(messages) > 0:
                        return messages[-1].text.strip()
                except Exception:
                    continue
                
            # Metodo JavaScript alternativo
            try:
                js_response = driver.execute_script("""
                    var messages = document.querySelectorAll('.message-content, .chat-message-item, .chat-wrapper .desc');
                    if (messages && messages.length > 0) {
                        return messages[messages.length - 1].textContent;
                    }
                    return null;
                """)
            
                if js_response:
                    return js_response.strip()
            except Exception:
                pass
            
            return ""
        except Exception as e:
            if log_callback:
                log_callback(f"Errore nel recupero dell'ultima risposta: {str(e)}")
            return ""

    def check_for_generation_error(self, response):
        """
        Controlla se la risposta contiene errori di generazione.
        Delega alla funzione in ai_interfaces/genspark_driver.py
        """
        from ai_interfaces.genspark_driver import check_for_generation_error as driver_check_for_generation_error
    
        return driver_check_for_generation_error(response)

    def handle_consecutive_errors(self, prompt_text, max_retries=3):
        """
        Gestisce errori consecutivi tentando approcci alternativi
        Delega alla funzione in ai_interfaces/genspark_driver.py
        """
        from ai_interfaces.genspark_driver import handle_consecutive_errors as driver_handle_consecutive_errors
    
        return driver_handle_consecutive_errors(
            driver=self.driver,
            prompt_text=prompt_text,
            max_retries=max_retries,
            log_callback=self.add_log
        )

    def _set_use_crisp(self, value):
        """Imposta se usare il framework CRISP"""
        self.use_crisp = value
        if value:
            self.add_log("Framework CRISP 5.0 attivato")
        else:
            self.add_log("Framework CRISP disattivato, verrà utilizzato il sistema legacy")
    
        # Restituisci il log aggiornato
        return self.chat_manager.get_log_history_string()

    def log_history_string(self):
        """Helper per ottenere il log history come stringa"""
        return self.chat_manager.get_log_history_string()

    def load_index_from_file(self, file_obj):
        """Carica l'indice da un file"""
        try:
            if file_obj is None:
                self.add_log("⚠️ Nessun file selezionato")
                return self.book_index.value
            
            file_path = file_obj.name
            self.add_log(f"📂 File caricato: {file_path}")
        
            # Determina il tipo di file
            file_ext = os.path.splitext(file_path)[1].lower()
        
            # Estrai il contenuto in base al tipo di file
            content = ""
            if file_ext == '.txt':
                # Leggi file di testo direttamente
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.add_log(f"✅ File TXT letto: {len(content)} caratteri")
                
            elif file_ext == '.docx':
                # Usa python-docx per leggere file Word
                try:
                    from docx import Document
                    doc = Document(file_path)
                    content = '\n\n'.join([para.text for para in doc.paragraphs if para.text])
                    self.add_log(f"✅ File DOCX letto: {len(content)} caratteri")
                except ImportError:
                    self.add_log("❌ Libreria python-docx non disponibile per leggere file Word")
                    return self.book_index.value
        
            elif file_ext == '.html':
                # Leggi file HTML ed estrai il testo
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        html_content = f.read()
                
                    # Prova a usare BeautifulSoup per estrarre il testo
                    try:
                        from bs4 import BeautifulSoup
                        soup = BeautifulSoup(html_content, 'html.parser')
                    
                        # Cerca elementi che potrebbero contenere l'indice
                        # Prima cerca elementi con id o classe che suggeriscono un indice
                        index_elements = soup.select('#indice, .indice, #index, .index, #toc, .toc')
                    
                        if index_elements:
                            # Usa il primo elemento trovato
                            content = index_elements[0].get_text('\n')
                        else:
                            # Altrimenti, cerca tag comuni per capitoli
                            chapter_elements = []
                            for tag in ['h1', 'h2', 'h3']:
                                chapter_elements.extend(soup.find_all(tag))
                        
                            if chapter_elements:
                                # Crea un indice dai titoli dei capitoli
                                chapter_texts = []
                                for i, elem in enumerate(chapter_elements):
                                    text = elem.get_text().strip()
                                    if 'capitolo' in text.lower() or 'chapter' in text.lower():
                                        chapter_texts.append(text)
                                    else:
                                        # Aggiungi il prefisso CAPITOLO se non presente
                                        chapter_texts.append(f"CAPITOLO {i+1}: {text}")
                            
                                # Aggiungi INTRODUZIONE e CONCLUSIONE
                                content = "INTRODUZIONE\n\n" + "\n\n".join(chapter_texts) + "\n\nCONCLUSIONE"
                            else:
                                # Prendi tutto il testo se non troviamo elementi specifici
                                content = soup.get_text('\n')
                    
                        self.add_log(f"✅ File HTML letto: {len(content)} caratteri")
                    
                    except ImportError:
                        # Fallback se BeautifulSoup non è disponibile
                        self.add_log("⚠️ BeautifulSoup non disponibile, estrazione base del testo HTML")
                        # Rimuovi tag html in modo semplice
                        import re
                        content = re.sub(r'<[^>]+>', ' ', html_content)
                        content = re.sub(r'\s+', ' ', content).strip()
                        self.add_log(f"✅ File HTML letto (estrazione base): {len(content)} caratteri")
                except Exception as html_error:
                    self.add_log(f"❌ Errore nell'estrazione del testo HTML: {str(html_error)}")
                    return self.book_index.value
                
            else:
                self.add_log(f"❌ Formato file non supportato: {file_ext}")
                return self.book_index.value
        
            return content
            
        except Exception as e:
            self.add_log(f"❌ Errore nel caricamento del file: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return self.book_index.value
    
    def reset_index(self):
        """Ripristina l'indice al formato predefinito"""
        default_index = """INTRODUZIONE

    CAPITOLO 1: Fondamenti

    CAPITOLO 2: Metodologia

    CAPITOLO 3: Applicazione

    CAPITOLO 4: Esempi Pratici

    CONCLUSIONE"""
        self.add_log("🔄 Indice ripristinato al formato predefinito")
        return default_index

    def extract_and_save_voice_style(self, project_id=None):
        """
        Estrae il tono di voce dall'analisi (risposta 7) e lo salva come file.
    
        Args:
            project_id: ID del progetto da cui estrarre il tono. Se None, usa l'analisi corrente.
    
        Returns:
            str: Nome del file salvato o None se non è stato possibile salvarlo.
        """
        try:
            # Determina il progetto da utilizzare
            keyword = None
            project_data = {}
        
            # Se è stato fornito un ID progetto specifico
            if project_id:
                self.add_log(f"🔍 Estrazione tono di voce dal progetto ID: {project_id}")
            
                # Carica i dati del progetto dal database
                if hasattr(self, 'crisp') and hasattr(self.crisp, '_load_project_data'):
                    project_data = self.crisp._load_project_data(project_id)
                    if project_data:
                        self.add_log(f"✅ Dati progetto caricati: {len(project_data)} campi")
                        # Estrai la keyword
                        if 'KEYWORD' in project_data:
                            keyword = project_data['KEYWORD']
                        elif 'keyword' in project_data:
                            keyword = project_data['keyword']
                    else:
                        self.add_log(f"❌ Impossibile caricare dati per il progetto ID: {project_id}")
            
                # Se non abbiamo trovato la keyword, usiamo un valore dall'ID
                if not keyword:
                    # Estrai dal nome visualizzato nel dropdown
                    if project_id and " - " in project_id:
                        keyword = project_id.split(" - ")[0].strip()
                        self.add_log(f"📝 Keyword estratta dal nome progetto: {keyword}")
                    else:
                        keyword = f"project_{project_id}"
                        self.add_log(f"📝 Keyword generica creata: {keyword}")
        
            # Altrimenti usa l'analisi corrente
            elif hasattr(self, 'current_analysis') and self.current_analysis is not None:
                self.add_log("🔍 Estrazione tono di voce dall'analisi corrente")
            
                # Cerca la keyword nell'analisi corrente
                if isinstance(self.current_analysis, dict):  # Verifica che sia un dizionario
                    project_data = self.current_analysis.get('project_data', {})
                    if 'KEYWORD' in self.current_analysis:
                        keyword = self.current_analysis['KEYWORD']
                    elif 'keyword' in self.current_analysis:
                        keyword = self.current_analysis['keyword']
        
            # Se ancora non abbiamo una keyword, cerca con altre funzioni
            if not keyword:
                self.add_log("🔍 Tentativo di determinare la keyword con altri metodi...")
                keyword = self.get_current_keyword() if hasattr(self, 'get_current_keyword') else None
        
            if not keyword:
                self.add_log("⚠️ Impossibile determinare la keyword per il salvataggio del tono di voce")
                return "❌ Impossibile determinare la keyword"
            
            # Crea una versione sicura della keyword per il nome file
            safe_keyword = keyword.lower().replace(' ', '_').replace("'", "").replace('"', "").replace("/", "")
        
            # Limita la lunghezza del nome file
            if len(safe_keyword) > 40:
                safe_keyword = safe_keyword[:40]
            
            # Crea la directory se non esiste
            import os
            os.makedirs("voice_styles", exist_ok=True)
        
            # Nome del file di output
            voice_style_file = f"voice_styles/{safe_keyword}_style.txt"
        
            # Cerca prima dentro project_data se disponibile
            voice_style_content = None
        
            # Se abbiamo dati di progetto, controlliamo i campi specifici
            if project_data:
                # Cerca in campi specifici che potrebbero contenere il tono di voce
                tone_fields = ['VOICE_STYLE', 'TONE_OF_VOICE', 'WRITING_STYLE', 'STYLE_GUIDE']
                for field in tone_fields:
                    if field in project_data and project_data[field] and len(project_data[field]) > 100:
                        voice_style_content = project_data[field]
                        self.add_log(f"✅ Tono di voce trovato nel campo {field}")
                        break
        
            # Se non abbiamo trovato il tono nei dati del progetto, cerchiamo nei file di contesto
            if not voice_style_content:
                # Determina quale file di contesto usare
                context_files_to_try = [
                    f"context_{safe_keyword}.txt",  # File specifico per keyword
                    "context.txt"  # File generico
                ]
            
                # Cerca nei risultati del progetto se disponibili
                if hasattr(self, 'crisp') and hasattr(self.crisp, 'get_project_results') and project_id:
                    try:
                        project_results = self.crisp.get_project_results(project_id)
                        for step_id, result_text in project_results:
                            # Cerca step specifici che potrebbero contenere il tono di voce (CM-7, o fase 7)
                            if ('M-7' in step_id or 'TONE' in step_id.upper() or 'STYLE' in step_id.upper()) and len(result_text) > 100:
                                voice_style_content = result_text
                                self.add_log(f"✅ Tono di voce trovato nel risultato {step_id}")
                                break
                    except Exception as e:
                        self.add_log(f"⚠️ Errore nel recupero dei risultati del progetto: {str(e)}")
            
                # Se ancora non abbiamo trovato il tono, cerchiamo nei file di contesto
                if not voice_style_content:
                    for context_file in context_files_to_try:
                        if os.path.exists(context_file):
                            self.add_log(f"📄 Apertura file di contesto: {context_file}")
                        
                            # Leggi il contenuto del file
                            with open(context_file, "r", encoding="utf-8") as f:
                                context_content = f.read()
                            
                            # Cerca la risposta 7 - Definizione del tono di voce
                            import re
                        
                            # Diversi pattern per trovare la sezione del tono di voce
                            patterns = [
                                # Pattern per risposta numerata "7)"
                                r'7\)[^\n]*(?:tono di voce|stile narrativo|struttura)[^\n]*\n(.*?)(?=\s*8\)|FINE|$)',
                                # Pattern per risposta numerata "7."
                                r'7\.[^\n]*(?:tono di voce|stile narrativo|struttura)[^\n]*\n(.*?)(?=\s*8\.|FINE|$)',
                                # Pattern per sezione "Tono di Voce"
                                r'(?:Tono di Voce|Stile Narrativo)[^\n]*\n(.*?)(?=\n\n[A-Z]|\Z)',
                                # Pattern per risposta CM-7 del framework CRISP
                                r'CM-7[^\n]*(?:tono|stile|voce)[^\n]*\n(.*?)(?=CM-8|FINE|$)'
                            ]
                        
                            for pattern in patterns:
                                match = re.search(pattern, context_content, re.DOTALL | re.IGNORECASE)
                                if match:
                                    voice_style_content = match.group(1).strip()
                                    self.add_log(f"✅ Tono di voce estratto con pattern: {pattern[:20]}...")
                                    break
                                
                            if voice_style_content:
                                break  # Esci dal ciclo dei file se abbiamo trovato il contenuto
        
            if not voice_style_content:
                self.add_log("❌ Impossibile trovare la sezione del tono di voce nell'analisi")
                return "❌ Tono di voce non trovato"
            
            # Se il contenuto è troppo corto, potrebbe essere un errore
            if len(voice_style_content) < 100:
                self.add_log(f"⚠️ Il tono di voce estratto sembra troppo breve ({len(voice_style_content)} caratteri)")
            
            # Salva il tono di voce in un file
            with open(voice_style_file, "w", encoding="utf-8") as f:
                f.write(voice_style_content)
            
            self.add_log(f"✅ Tono di voce salvato in: {voice_style_file} ({len(voice_style_content)} caratteri)")
        
            # Aggiorna la lista dei file di stile se necessario
            if hasattr(self, 'voice_style_file'):
                try:
                    # Ottieni la lista aggiornata
                    voice_style_files = ["Nessuno"] + [os.path.splitext(f)[0] for f in os.listdir("voice_styles") if f.endswith(".txt")]
                    # Aggiorna il dropdown
                    self.voice_style_file.choices = voice_style_files
                    self.voice_style_file.value = os.path.splitext(os.path.basename(voice_style_file))[0]
                except Exception as e:
                    self.add_log(f"⚠️ Errore nell'aggiornamento del dropdown: {str(e)}")
        
            # Restituisci il nome del file senza estensione per la selezione
            file_name = os.path.splitext(os.path.basename(voice_style_file))[0]
            return f"✅ Tono di voce '{file_name}' creato con successo"
        
        except Exception as e:
            self.add_log(f"❌ Errore nell'estrazione del tono di voce: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return f"❌ Errore: {str(e)}"


    def _parse_book_index(self, book_index):
        """
        Analizza l'indice e lo converte in una lista di capitoli strutturati.
        Ora delega alla funzione in framework/book_generator.py
        """
        from framework.book_generator import _parse_book_index as parse_book_index_func
    
        return parse_book_index_func(book_index)


    def _load_chapter_prompt(self, book_type):
            """
            Carica il prompt template per i capitoli specifico per il tipo di libro.
            Ora delega alla funzione in framework/book_generator.py
            """
            from framework.book_generator import _load_chapter_prompt as load_chapter_prompt_func
    
            return load_chapter_prompt_func(book_type)


    def _handle_missing_placeholders(self, text):
        """
        Gestisce i placeholder non sostituiti nel prompt.
        Ora delega alla funzione in framework/book_generator.py
        """
        from framework.book_generator import _handle_missing_placeholders as handle_missing_placeholders_func
    
        return handle_missing_placeholders_func(text)

    def _generate_chapter_content(self, chapter_title, prompt):
        """
        Genera il contenuto di un capitolo inviandolo a Genspark.
        Ora delega alla funzione in framework/book_generator.py
        """
        from framework.book_generator import _generate_chapter_content as generate_chapter_content_func
        from ai_interfaces.genspark_driver import send_to_genspark
    
        # Implementazione temporanea che mantiene la compatibilità con la classe
        # Il metodo originale faceva riferimento ai membri della classe
        try:
            # Verifica se il driver è attivo
            if not self.driver:
                self.add_log("Browser non disponibile")
                return "Errore: Browser non disponibile"
            
            # Invia il prompt e attendi la risposta
            response = self.send_to_genspark(prompt)
        
            # Pulisci la risposta
            if "FINE_RISPOSTA" in response:
                response = response.split("FINE_RISPOSTA")[0].strip()
            elif "FINE" in response:
                response = response.split("FINE")[0].strip()
            
            return response
        
        except Exception as e:
            self.add_log(f"Errore nella generazione del capitolo {chapter_title}: {str(e)}")
            return f"Errore nella generazione: {str(e)}"

    def _save_chapter(self, chapter_title, chapter_content, book_title):
        """
        Salva il capitolo generato in un file.
        Ora delega alla funzione in framework/book_generator.py
        """
        from framework.book_generator import _save_chapter as save_chapter_func
    
        return save_chapter_func(
            chapter_title, 
            chapter_content, 
            book_title, 
            log=self.add_log
        )

    def load_analysis_results(self):
        """
        Carica i risultati dell'analisi nell'interfaccia utente
        """
        import os
        from datetime import datetime
        import webbrowser

        formatted_html = None
    
        self.add_log("Caricamento dei risultati dell'analisi...")
    
        # NUOVO CODICE: Prova a catturare l'HTML formattato da Genspark
        try:
            # Cattura l'HTML formattato direttamente da Genspark
            if self.driver:
                self.add_log("Tentativo di cattura HTML migliorata direttamente da Genspark...")
                html_file = self.save_complete_html_improved()
            
                if formatted_html:
                    self.add_log(f"✅ HTML catturato con successo: {len(formatted_html)} caratteri")
                    self.add_log(f"📌 ANTEPRIMA HTML: {formatted_html[:100]}...")
                
                    # Salva l'HTML formattato in un file
                    keyword = self.get_current_keyword()
                    self.add_log(f"📖 Keyword identificata per il salvataggio: '{keyword}'")
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                
                    # Crea directory se non esiste
                    os.makedirs("output/genspark_formatted", exist_ok=True)
                
                    # Salva il file con un nome descrittivo
                    html_path, file_size = self.save_enhanced_html(formatted_html, keyword, timestamp)
                    self.add_log(f"💾 File HTML migliorato salvato: {html_path} ({file_size} bytes)")

                    # Sistema di debug automatico con percorso assoluto completo
                    abs_path = os.path.abspath(html_path)
                    self.add_log(f"📋 Percorso completo: {abs_path}")
                    self.add_log(f"🔗 URL file: file://{abs_path.replace(os.sep, '/')}")

                    # Salva il percorso in una variabile di classe per riferimento futuro
                    self.last_html_path = abs_path

                
                    # Aggiorna anche il file "current"
                    current_path = f"output/genspark_formatted/{keyword}_genspark_current.html"
                    with open(current_path, "w", encoding="utf-8") as f:
                        f.write(formatted_html)
                
                    # VERIFICA COMPONENTE: Log dettagliati sul componente results_display
                    if hasattr(self, 'results_display'):
                        self.add_log(f"🔍 Componente results_display trovato: {type(self.results_display)}")
                    
                        # Verifica se il componente è un HTML
                        is_html_component = hasattr(self.results_display, 'value') and isinstance(self.results_display.value, str)
                        self.add_log(f"🔍 È un componente HTML valido? {is_html_component}")
                    
                        # Prova ad aggiornare il componente e verifica se ci sono errori
                        try:
                            # Tenta di impostare un valore di test prima
                            test_html = "<div style='color:red'>TEST</div>"
                            self.add_log("🧪 Impostazione HTML di test nel componente...")
                            self.results_display.value = test_html
                            self.add_log("✅ Test HTML impostato con successo")
                        
                            # Ora imposta l'HTML formattato effettivo
                            self.add_log("🧪 Impostazione dell'HTML formattato con stili forzati di visibilità...")
                            visible_html = f"""
                            <div style="display:block !important; visibility:visible !important; opacity:1 !important; min-height:400px; border:3px solid red; padding:20px;">
                                <h2 style="color:red">FORZATO VISIBILE</h2>
                                {formatted_html}
                            </div>
                            """
                            self.results_display.value = visible_html
                            self.add_log("✅ HTML formattato con visibilità forzata impostato con successo")
                        
                            # Verifica lo stato del componente dopo l'aggiornamento
                            if hasattr(self.results_display, 'value'):
                                current_value_length = len(self.results_display.value) if self.results_display.value else 0
                                self.add_log(f"📊 Valore corrente del componente: {current_value_length} caratteri")
                        except Exception as update_error:
                            self.add_log(f"❌ ERRORE nell'aggiornamento del componente: {str(update_error)}")
                            import traceback
                            self.add_log(f"TRACEBACK: {traceback.format_exc()}")
                    else:
                        self.add_log("❌ Componente results_display NON trovato!")
                
                    # Apri automaticamente nel browser (opzionale)
                    try:
                        self.add_log("🌐 Tentativo di apertura nel browser...")
                        file_path = os.path.abspath(html_file).replace('\\', '/')
                        webbrowser.open("file:///" + file_path)
                        self.add_log("✅ File aperto nel browser")
                    except Exception as browser_error:
                        self.add_log(f"⚠️ Impossibile aprire nel browser: {str(browser_error)}")
                
                    # Prova a ritornare qui per evitare che il codice successivo sovrascriva il componente
                    # self.add_log("📍 Completamento precoce della funzione - ritorno ai chiamanti")
                    # return self.chat_manager.get_log_history_string()
                else:
                    self.add_log("⚠️ Cattura HTML fallita: formatted_html è None o vuoto")
        except Exception as e:
            self.add_log(f"⚠️ Cattura HTML Genspark fallita: {str(e)}")
            import traceback
            self.add_log(f"TRACEBACK: {traceback.format_exc()}")

            self.add_log("⚠️ NOTA: Esecuzione continuata con il metodo originale - la cattura HTML potrebbe aver fallito")

            # Continua con il metodo originale (non fare return)
        
            # AGGIUNGI QUESTO CODICE PER COPIARE IL FILE DI BACKUP NELLA DIRECTORY PRINCIPALE E TENERE TRACCIA DEI FILE TEMPORANEI
            backup_dir = "backups"
            temp_files_created = []  # Tieni traccia dei file temporanei creati
            if os.path.exists(backup_dir):
                backup_files = [f for f in os.listdir(backup_dir) 
                               if f.startswith("context_") and f.endswith(".txt")]
                if backup_files:
                    backup_files.sort(key=lambda x: os.path.getmtime(os.path.join(backup_dir, x)), 
                                     reverse=True)
                    if backup_files:
                        latest_file = backup_files[0]
                        self.add_log(f"DEBUG: File di backup più recente: {latest_file}")
            
                        # Due possibili pattern: con timestamp o senza
                        keyword_match = re.match(r"context_(.+?)(_\d{8}_\d{6})?\.txt", latest_file)
                        if keyword_match:
                            safe_keyword = keyword_match.group(1)  # Keyword con underscores
                
                            # Nome del file nella directory principale
                            main_file = f"context_{safe_keyword}.txt"
                
                            # Copia il file se non esiste già
                            if not os.path.exists(main_file):
                                source_path = os.path.join(backup_dir, latest_file)
                                self.add_log(f"📋 Copio {source_path} → {main_file}")
                                shutil.copy2(source_path, main_file)
                                temp_files_created.append(main_file)  # Aggiungi alla lista dei file temporanei
                                self.add_log(f"✅ File di backup copiato nella directory principale (temporaneo)")
                            else:
                                self.add_log(f"📄 File {main_file} già presente nella directory principale")

            # Aggiungi qui la gestione dell'analysis_id se fornito
            if analysis_id:
                self.add_log(f"🔍 Caricamento analisi con ID: {analysis_id}")
                # Eventuale codice per caricare l'analisi specifica

            try:
                # Cerca il file di contesto per la keyword corrente
                current_keyword = self.get_current_keyword()
                if current_keyword:
                    self.add_log(f"🔍 Keyword corrente identificata: '{current_keyword}'")
        
                    # Verifica nei file di backup per confermare la keyword più recente
                    backup_dir = "backups"
                    if os.path.exists(backup_dir):
                        backup_files = [f for f in os.listdir(backup_dir) 
                                       if f.startswith("context_") and f.endswith(".txt")]
                        if backup_files:
                            # Ordina per data di modifica (più recente prima)
                            backup_files = sorted(backup_files, 
                                                 key=lambda x: os.path.getmtime(os.path.join(backup_dir, x)), 
                                                 reverse=True)
                            # Estrai il nome della keyword dal primo file (il più recente)
                            if backup_files and len(backup_files) > 0:
                                latest_file = backup_files[0]
                                self.add_log(f"DEBUG: File di backup più recente: {latest_file}")
                                keyword_match = re.match(r"context_(.+?)_\d{8}_\d{6}\.txt", latest_file)
                                if keyword_match:
                                    confirmed_keyword = keyword_match.group(1)
                                    if confirmed_keyword != current_keyword:
                                        self.add_log(f"⚠️ Correzione keyword: da '{current_keyword}' a '{confirmed_keyword}'")
                                        current_keyword = confirmed_keyword
        
                    # Nome file basato sulla keyword
                    safe_keyword = re.sub(r'[\\/*?:"<>|]', "", current_keyword).replace(" ", "_")[:30]
                    context_file = f"context_{safe_keyword}.txt"
        
                    # Se non esiste file specifico, usa quello generico
                    if not os.path.exists(context_file):
                        context_file = "context.txt"
            
                    if os.path.exists(context_file):
                        with open(context_file, "r", encoding="utf-8") as f:
                            context_content = f.read()
                
                        # Genera e salva l'HTML usando il nostro nuovo metodo
                        formatted_html, html_path = self.format_and_save_analysis_html(current_keyword, context_content)
            
                        if formatted_html and html_path:

                            # Aggiornamento con formattazione migliorata
                            if hasattr(self, 'results_display'):
                                visible_html = f"""
                                <div id='analysis-results-container' 
                                     style='min-height: 500px; 
                                            background-color: white;
                                            border: 1px solid #e5e7eb;
                                            border-radius: 8px;
                                            padding: 20px;
                                            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                                            position: relative;
                                            z-index: 10;'>
                                    <h2 style='color: #2563eb; margin-bottom: 15px;'>Risultati dell'Analisi</h2>
                                    {formatted_html}
                                </div>
                                """
                                # Usa il metodo update se disponibile, altrimenti assegna direttamente
                                if hasattr(self.results_display, 'update'):
                                    self.results_display.update(value=visible_html, visible=True)
                                    self.add_log("✅ Risultati visualizzati tramite update")
                                else:
                                    self.results_display.value = visible_html
                                    self.add_log("✅ Risultati visualizzati tramite assegnazione diretta")
                    
                            # Aggiorna lo stato dell'analisi
                            if hasattr(self, 'analysis_status'):
                                self.analysis_status.value = "**Stato analisi**: Completata e visualizzata ✅"
                    
                            # Apri automaticamente nel browser
                            try:
                                file_path = os.path.abspath(html_file).replace('\\', '/')
                                webbrowser.open("file:///" + file_path)
                                self.add_log("✅ File aperto nel browser")
                            except Exception as browser_error:
                                self.add_log(f"⚠️ Impossibile aprire nel browser: {str(browser_error)}")
                    
                            # Pulizia dei file temporanei prima di uscire
                            try:
                                for temp_file in temp_files_created:
                                    if os.path.exists(temp_file) and temp_file != "context.txt":  # Non rimuovere il file principale
                                        os.remove(temp_file)
                                        self.add_log(f"🧹 File temporaneo {temp_file} rimosso")
                            except Exception as cleanup_error:
                                self.add_log(f"⚠️ Errore nella pulizia dei file temporanei: {str(cleanup_error)}")
                    
                            # Terminazione precoce
                            return self.chat_manager.get_log_history_string()
                    else:
                        self.add_log(f"⚠️ File di contesto non trovato: {context_file}")
            
                        # PIANO B: usa ancora format_analysis_results_html con i parametri disponibili
                        try:
                            market = self.get_current_market() if hasattr(self, 'get_current_market') else "Unknown"
                            book_type = self.get_current_book_type() if hasattr(self, 'get_current_book_type') else "Unknown"
                            language = self.get_current_language() if hasattr(self, 'get_current_language') else "it"
                            context_data = self.extract_context_data() if hasattr(self, 'extract_context_data') else None
    
                            self.add_log(f"📊 Tentativo di generazione HTML con format_analysis_results_html...")
    
                            # Chiamare format_analysis_results_html con tutti i parametri necessari
                            formatted_html = self.format_analysis_results_html(
                                keyword=current_keyword,
                                market=market,
                                book_type=book_type,
                                language=language,
                                context=context_data
                            )
    
                            if formatted_html:
                                self.add_log(f"✅ HTML formattato generato (Piano B): {len(formatted_html)} caratteri")
        
                                # Aggiorna l'interfaccia con formattazione migliorata
                                if hasattr(self, 'results_display'):
                                    # Prepara HTML con formattazione visibile
                                    visible_html = f"""
                                    <div id='analysis-results-container' 
                                         style='min-height: 500px; 
                                                background-color: white;
                                                border: 1px solid #e5e7eb;
                                                border-radius: 8px;
                                                padding: 20px;
                                                box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                                                position: relative;
                                                z-index: 10;'>
                                        <h2 style='color: #2563eb; margin-bottom: 15px;'>Risultati dell'Analisi (Piano B)</h2>
                                        {formatted_html}
                                    </div>
                                    """
            
                                    # Usa il metodo update se disponibile, altrimenti assegna direttamente
                                    if hasattr(self.results_display, 'update'):
                                        self.results_display.update(value=visible_html, visible=True)
                                        self.add_log("✅ Risultati visualizzati tramite update (Piano B)")
                                    else:
                                        self.results_display.value = visible_html
                                        self.add_log("✅ Risultati visualizzati tramite assegnazione diretta (Piano B)")
                
                                    # Aggiorna lo stato dell'analisi
                                    if hasattr(self, 'analysis_status'):
                                        self.analysis_status.value = "**Stato analisi**: Completata e visualizzata ✅"
            
                                    # Aggiungi questo log per debug
                                    self.add_log(f"DEBUG Piano B: results_display ID={id(self.results_display)}, tipo={type(self.results_display)}")
            
                                    # Non terminare precocemente con return
                                    # return self.chat_manager.get_log_history_string()
                                else:
                                    self.add_log("❌ results_display non trovato (Piano B)")
                        except Exception as format_error_b:
                            self.add_log(f"❌ Piano B fallito: {str(format_error_b)}")
                            import traceback
                            self.add_log(traceback.format_exc())

            except Exception as primary_error:
                # Gestione errori per il primo approccio
                self.add_log(f"⚠️ Il metodo principale ha generato un errore: {str(primary_error)}")
                import traceback
                self.add_log(f"TRACEBACK primo metodo: {traceback.format_exc()}")
    
                # Pulizia dei file temporanei se necessario
                try:
                    for temp_file in temp_files_created:
                        if os.path.exists(temp_file) and temp_file != "context.txt":
                            os.remove(temp_file)
                            self.add_log(f"🧹 File temporaneo {temp_file} rimosso (dopo errore primo metodo)")
                except Exception as cleanup_error:
                    self.add_log(f"⚠️ Errore nella pulizia dopo primo metodo: {str(cleanup_error)}")
    
                # FALLBACK: Prova con il secondo approccio
                self.add_log("ℹ️ Il metodo principale e il Piano B hanno fallito. Tentativo con il metodo alternativo...")

            # Se siamo qui, significa che né l'approccio principale né il Piano B hanno avuto successo
            # Continua con il secondo blocco principale (fallback finale)
        
            try:
                # Cerca prima i file HTML nella cartella output/analisi_html (percorso corretto)
                import glob
                output_dir = "output/analisi_html"
        
                # Crea la directory se non esiste
                os.makedirs(output_dir, exist_ok=True)
        
                # Cerca sia i file con timestamp che i file current
                html_files_timestamp = glob.glob(f"{output_dir}/*_*.html")
                html_files_current = glob.glob(f"{output_dir}/*_current.html")
        
                # Prova prima con i file current, poi con quelli con timestamp
                html_files = html_files_current + html_files_timestamp
        
                if html_files:
                    # Trova il file più recente o il primo file current
                    if html_files_current:
                        latest_file = html_files_current[0]  # Prendi il primo file current
                        self.add_log(f"📄 File HTML 'current' trovato: {latest_file}")
                    else:
                        latest_file = max(html_files_timestamp, key=os.path.getmtime) if html_files_timestamp else None
                        if latest_file:
                            self.add_log(f"📄 File HTML più recente trovato: {latest_file}")
            
                    if latest_file:
                        # Carica il contenuto HTML
                        with open(latest_file, "r", encoding="utf-8") as f:
                            html_content = f.read()
                
                        # Aggiorna l'interfaccia con il file HTML formattato
                        if hasattr(self, 'results_display'):
                            self.results_display.value = html_content
                            self.add_log("✅ Risultati dell'analisi visualizzati dall'HTML formattato")
                    
                            # Aggiorna lo stato dell'analisi
                            if hasattr(self, 'analysis_status'):
                                self.analysis_status.value = "**Stato analisi**: Completata e visualizzata ✅"
                    
                            # Pulizia dei file temporanei prima di uscire
                            try:
                                for temp_file in temp_files_created:
                                    if os.path.exists(temp_file) and temp_file != "context.txt":  # Non rimuovere il file principale
                                        os.remove(temp_file)
                                        self.add_log(f"🧹 File temporaneo {temp_file} rimosso")
                            except Exception as cleanup_error:
                                self.add_log(f"⚠️ Errore nella pulizia dei file temporanei: {str(cleanup_error)}")
                            
                            return self.chat_manager.get_log_history_string()
        
                # Se non ci sono file HTML nella directory corretta, prova anche il vecchio percorso
                fallback_html_files = glob.glob("output/analisi_*.html")
                if fallback_html_files:
                    latest_file = max(fallback_html_files, key=os.path.getmtime)
                    self.add_log(f"📄 File HTML trovato nel percorso alternativo: {latest_file}")
            
                    # Carica il contenuto HTML
                    with open(latest_file, "r", encoding="utf-8") as f:
                        html_content = f.read()
            
                    # Aggiorna l'interfaccia con il file HTML formattato
                    if hasattr(self, 'results_display'):
                        self.results_display.value = html_content
                        self.add_log("✅ Risultati dell'analisi visualizzati dall'HTML formattato (fallback)")
                
                        # Aggiorna lo stato dell'analisi
                        if hasattr(self, 'analysis_status'):
                            self.analysis_status.value = "**Stato analisi**: Completata e visualizzata ✅"
                
                        # Pulizia dei file temporanei prima di uscire
                        try:
                            for temp_file in temp_files_created:
                                if os.path.exists(temp_file) and temp_file != "context.txt":  # Non rimuovere il file principale
                                    os.remove(temp_file)
                                    self.add_log(f"🧹 File temporaneo {temp_file} rimosso")
                        except Exception as cleanup_error:
                            self.add_log(f"⚠️ Errore nella pulizia dei file temporanei: {str(cleanup_error)}")
                        
                        return self.chat_manager.get_log_history_string()

                # Se non ci sono file HTML, utilizza il metodo tradizionale con context.txt
                self.add_log("⚠️ Nessun file HTML trovato, utilizzo il vecchio metodo con context.txt")
        
                # [Il resto del codice rimane invariato...]
        
                # Verifica che il file di contesto esista
                context_file = os.path.join(self.context_dir, "context.txt")
                if not os.path.exists(context_file):
                    self.add_log("⚠️ File context.txt non trovato")
                    if hasattr(self, 'results_display'):
                        self.results_display.value = "<div class='alert alert-warning'>File dei risultati non trovato. Esegui prima l'analisi.</div>"
                    return self.chat_manager.get_log_history_string()
    
                # Leggi il file di contesto
                with open(context_file, "r", encoding="utf-8") as f:
                    context_content = f.read()
    
                if not context_content.strip():
                    self.add_log("⚠️ File context.txt vuoto")
                    if hasattr(self, 'results_display'):
                        self.results_display.value = "<div class='alert alert-warning'>File dei risultati vuoto. Esegui l'analisi.</div>"
                    return self.chat_manager.get_log_history_string()
    
                # Estrai le sezioni del contesto
                sections = []
                section_pattern = r'===\s+([^=]+?)\s+-\s+\d{8}_\d{6}\s+===\n(.*?)(?=\n===|$)'
                section_matches = re.findall(section_pattern, context_content, re.DOTALL)
    
                if section_matches:
                    sections = [(title.strip(), content.strip()) for title, content in section_matches]
                    self.add_log(f"✅ Estratte {len(sections)} sezioni")
        
                    # Debug: mostra dimensione delle prime 3 sezioni
                    for i, (title, content) in enumerate(sections[:3]):
                        self.add_log(f"Sezione {i+1}: {title} ({len(content)} caratteri)")
                else:
                    # Fallback: divide per numeri
                    number_pattern = r'(\d+\).*?)(?=\d+\)|$)'
                    number_matches = re.findall(number_pattern, context_content, re.DOTALL)
        
                    if number_matches:
                        sections = [(f"Sezione {i+1}", content.strip()) for i, content in enumerate(number_matches)]
                        self.add_log(f"✅ Estratte {len(sections)} sezioni (pattern numerico)")
                    else:
                        # Ultimo fallback: usa il testo completo
                        sections = [("Risultati completi", context_content)]
                        self.add_log("⚠️ Nessuna sezione trovata, usando il testo completo")
    
                # Costruisci l'HTML per la visualizzazione
                html = """
                <!DOCTYPE html>
                <html>
                <head>
                    <style>
                        body { font-family: Arial, sans-serif; }
                        .header { background-color: #2563eb; color: white; padding: 15px; border-radius: 8px; margin-bottom: 20px; }
                        .section { background-color: #f5f5f5; padding: 15px; border-radius: 8px; margin-bottom: 15px; }
                        .section-title { font-weight: bold; font-size: 18px; margin-bottom: 10px; }
                        .error { background-color: #fee2e2; }
                        .ok { background-color: #d1fae5; }
                    </style>
                </head>
                <body>
                    <div class="header">
                        <h2>Risultati dell'Analisi</h2>
                        <p>Sezioni trovate: %d</p>
                    </div>
                    <div>
                """ % len(sections)
    
                # Aggiungi le sezioni
                for title, content in sections:
                    # Determina se è una risposta valida o un errore
                    is_error = "richiesta abortita" in content.lower() or len(content) < 30
                    section_class = "section error" if is_error else "section ok"
        
                    # Converti newline in <br> per HTML
                    content_html = content.replace("\n", "<br>")
        
                    html += """
                    <div class="%s">
                        <div class="section-title">%s</div>
                        <div>%s</div>
                    </div>
                    """ % (section_class, title, content_html)
    
                html += """
                    </div>
                </body>
                </html>
                """
    
                # Salva l'HTML di debug in un file unico per ogni esecuzione
                import os
                os.makedirs("debug", exist_ok=True)
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                debug_html_filename = f"debug/debug_html_{timestamp}.html"
                with open(debug_html_filename, "w", encoding="utf-8") as f:
                    f.write(html)
                self.add_log(f"HTML di debug salvato: {debug_html_filename} ({len(html)} bytes)")

                # Sistema di debug automatico con percorso assoluto completo
                abs_debug_path = os.path.abspath(debug_html_filename)
                self.add_log(f"📋 Percorso completo debug: {abs_debug_path}")
                self.add_log(f"🔗 URL file debug: file://{abs_debug_path.replace(os.sep, '/')}")

                # Salva il percorso in una variabile di classe per riferimento futuro
                self.last_debug_html_path = abs_debug_path

                # NUOVO: Salva anche in output/analisi_html come file regolare
                try:
                    current_keyword = self.get_current_keyword()
                    if current_keyword:
                        from framework.formatters import save_analysis_to_html  # Sostituisci your_module con il nome del modulo dove è definita la funzione
                        analysis_type = "Legacy"  # O determina il tipo dinamicamente
                        file_path = save_analysis_to_html(html, current_keyword, "Unknown", "Unknown", "it", analysis_type, self.add_log)
                        self.add_log(f"✅ Salvato anche come HTML formattato in {file_path}")
                except Exception as e:
                    self.add_log(f"❌ Errore nel salvataggio aggiuntivo: {str(e)}")

                # Sistema di debug automatico per il file HTML formattato
                try:
                    if 'file_path' in locals() and file_path:
                        abs_formatted_path = os.path.abspath(file_path)
                        self.add_log(f"📋 Percorso completo HTML formattato: {abs_formatted_path}")
                        self.add_log(f"🔗 URL file formattato: file://{abs_formatted_path.replace(os.sep, '/')}")
                
                        # Salva il percorso in una variabile di classe
                        self.last_formatted_html_path = abs_formatted_path
                except Exception as path_error:
                        self.add_log(f"⚠️ Errore nel logging del percorso formattato: {str(path_error)}")                        
    
                # IMPORTANTE: Verifica che results_display esista e sia diverso da analysis_status
                if hasattr(self, 'results_display'):
                    # Verifica se results_display e analysis_status sono lo stesso oggetto
                    same_object = False
                    if hasattr(self, 'analysis_status'):
                        same_object = id(self.results_display) == id(self.analysis_status)
                        self.add_log(f"DEBUG: results_display e analysis_status sono {'LO STESSO' if same_object else 'OGGETTI DIVERSI'}")
        
                    # Aggiorna l'interfaccia SOLO se non sono lo stesso oggetto
                    if not same_object:
                        # Usa assegnazione diretta invece di update
                        self.results_display.value = html
                        self.add_log("✅ Risultati dell'analisi visualizzati nell'interfaccia")
            
                        # Aggiorna lo stato dell'analisi separatamente
                        if hasattr(self, 'analysis_status'):
                            # Usa assegnazione diretta invece di update
                            self.analysis_status.value = "**Stato analisi**: Completata e visualizzata ✅"
                    else:
                        self.add_log("⚠️ results_display e analysis_status sono lo stesso oggetto! Non aggiorno l'interfaccia")
                        # In questo caso, dobbiamo aggiornare solo uno dei due
                        # Usa assegnazione diretta invece di update
                        self.results_display.value = html
                        self.add_log("✅ Risultati dell'analisi visualizzati nell'interfaccia (solo results_display)")
                else:
                    self.add_log("⚠️ results_display non disponibile")
            
                # Pulizia dei file temporanei prima di uscire
                try:
                    for temp_file in temp_files_created:
                        if os.path.exists(temp_file) and temp_file != "context.txt":  # Non rimuovere il file principale
                            os.remove(temp_file)
                            self.add_log(f"🧹 File temporaneo {temp_file} rimosso")
                except Exception as cleanup_error:
                    self.add_log(f"⚠️ Errore nella pulizia dei file temporanei: {str(cleanup_error)}")
    
                return self.chat_manager.get_log_history_string()

            except Exception as e:
                # Pulizia dei file temporanei anche in caso di errore
                try:
                    for temp_file in temp_files_created:
                        if os.path.exists(temp_file) and temp_file != "context.txt":
                            os.remove(temp_file)
                            self.add_log(f"🧹 File temporaneo {temp_file} rimosso (durante gestione errore)")
                except:
                    pass
                
                self.add_log(f"❌ Errore nel caricamento dei risultati: {str(e)}")
                import traceback
                self.add_log(traceback.format_exc())
                return self.chat_manager.get_log_history_string()

    def save_enhanced_html(self, html_content, keyword, timestamp=None):
        """
        Salva una versione migliorata del file HTML con CSS completi.
        """
        import os
        from datetime import datetime
    
        if timestamp is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
        # Crea directory se non esiste
        os.makedirs("output/genspark_formatted", exist_ok=True)
    
        # Aggiungi CSS avanzati per tabelle, liste, etc.
        enhanced_html = f"""<!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Analisi {keyword} - {timestamp}</title>
            <style>
                /* Stili base */
                body {{
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 1200px;
                    margin: 0 auto;
                    padding: 20px;
                }}
            
                /* Stili per intestazioni */
                h1 {{ font-size: 28px; margin-bottom: 20px; color: #1a56db; }}
                h2 {{ font-size: 24px; margin-top: 30px; margin-bottom: 15px; color: #1e429f; border-bottom: 1px solid #e5e7eb; padding-bottom: 8px; }}
                h3 {{ font-size: 20px; margin-top: 25px; margin-bottom: 12px; color: #233876; }}
                h4 {{ font-size: 18px; margin-top: 20px; color: #374151; }}
            
                /* Stili per testo */
                p {{ margin-bottom: 16px; }}
                strong, b {{ font-weight: 600; color: #111827; }}
            
                /* Stili per tabelle */
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                    font-size: 15px;
                    box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                    border-radius: 6px;
                    overflow: hidden;
                }}
            
                table thead {{
                    background-color: #f3f4f6;
                }}
            
                table th {{
                    padding: 12px 15px;
                    text-align: left;
                    font-weight: 600;
                    color: #111827;
                    border-bottom: 2px solid #d1d5db;
                }}
            
                table td {{
                    padding: 10px 15px;
                    border-bottom: 1px solid #e5e7eb;
                }}
            
                table tbody tr:nth-child(even) {{
                    background-color: #f9fafb;
                }}
            
                table tbody tr:hover {{
                    background-color: #f3f4f6;
                }}
            
                /* Stili per liste */
                ul, ol {{
                    margin-bottom: 16px;
                    padding-left: 25px;
                }}
            
                ul li, ol li {{
                    margin-bottom: 8px;
                }}
            
                ul {{
                    list-style-type: disc;
                }}
            
                ul ul {{
                    list-style-type: circle;
                }}
            
                /* Stili per codice */
                pre {{
                    background-color: #f3f4f6;
                    padding: 15px;
                    border-radius: 6px;
                    overflow-x: auto;
                    font-family: monospace;
                    margin: 20px 0;
                    border: 1px solid #e5e7eb;
                }}
            
                code {{
                    background-color: #f3f4f6;
                    padding: 2px 5px;
                    border-radius: 4px;
                    font-family: monospace;
                    font-size: 0.9em;
                    color: #3b82f6;
                }}
            
                /* Stili per citazioni */
                blockquote {{
                    border-left: 4px solid #3b82f6;
                    padding: 10px 20px;
                    margin: 20px 0;
                    background-color: #f3f4f6;
                    color: #4b5563;
                    font-style: italic;
                }}
            
                /* Stili per separatori */
                hr {{
                    border: none;
                    border-top: 2px solid #e5e7eb;
                    margin: 30px 0;
                }}
            
                /* Stili per note e avvisi */
                .note {{
                    background-color: #eff6ff;
                    border-left: 4px solid #3b82f6;
                    padding: 15px;
                    margin: 20px 0;
                    border-radius: 6px;
                }}
            
                .warning {{
                    background-color: #fef2f2;
                    border-left: 4px solid #ef4444;
                    padding: 15px;
                    margin: 20px 0;
                    border-radius: 6px;
                }}
            
                .tip {{
                    background-color: #f0fdf4;
                    border-left: 4px solid #10b981;
                    padding: 15px;
                    margin: 20px 0;
                    border-radius: 6px;
                }}
            
                /* Stili per immagini */
                img {{
                    max-width: 100%;
                    height: auto;
                    border-radius: 6px;
                    margin: 20px 0;
                }}
            
                /* Stile per sezioni */
                .section {{
                    background-color: #ffffff;
                    border-radius: 8px;
                    padding: 20px;
                    margin-bottom: 30px;
                    box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                    border: 1px solid #e5e7eb;
                }}
            
                /* Stile per metadata */
                .metadata {{
                    font-size: 14px;
                    color: #6b7280;
                    margin-bottom: 30px;
                    padding: 10px;
                    background-color: #f9fafb;
                    border-radius: 6px;
                }}
            
                /* Stili per responsive */
                @media screen and (max-width: 768px) {{
                    body {{
                        padding: 15px;
                    }}
                
                    table {{
                        font-size: 14px;
                    }}
                
                    table td, table th {{
                        padding: 8px 10px;
                    }}
                }}
            
                /* Classi aggiuntive per le tabelle migliorate */
                .enhanced-table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                }}
            
                .even-row {{
                    background-color: #f9fafb;
                }}
            
                .odd-row {{
                    background-color: #ffffff;
                }}
            </style>
        
            <!-- INSERIRE LO SCRIPT QUI -->
            <script>
                // Script per migliorare le tabelle nell'HTML
                function enhanceTables() {{
                    document.querySelectorAll('table').forEach(table => {{
                        // Aggiungi classe per lo styling
                        table.classList.add('enhanced-table');
                    
                        // Forza il bordo se non esistente
                        if (!table.hasAttribute('border') || table.getAttribute('border') === '0') {{
                            table.setAttribute('border', '1');
                        }}
                    
                        // Assicurati che la prima riga sia un header
                        const firstRow = table.querySelector('tr');
                        if (firstRow) {{
                            const cells = Array.from(firstRow.children);
                            if (cells.length > 0) {{
                                cells.forEach(cell => {{
                                    if (cell.tagName !== 'TH') {{
                                        // Converti TD a TH per la prima riga
                                        const th = document.createElement('th');
                                        th.innerHTML = cell.innerHTML;
                                        cell.parentNode.replaceChild(th, cell);
                                    }}
                                }});
                            }}
                        }}
                    
                        // Aggiungi classi zebra-stripe per migliorare la leggibilità
                        const rows = Array.from(table.querySelectorAll('tr')).slice(1); // Skip header
                        rows.forEach((row, index) => {{
                            row.classList.add(index % 2 === 0 ? 'even-row' : 'odd-row');
                        }});
                    }});
                }}
            
                // Esegui lo script quando la pagina è caricata
                document.addEventListener('DOMContentLoaded', enhanceTables);
            </script>
        </head>
        <body>
            <div class="metadata">
                <strong>Parola chiave:</strong> {keyword}<br>
                <strong>Data analisi:</strong> {timestamp[:4]}/{timestamp[4:6]}/{timestamp[6:8]} {timestamp[9:11]}:{timestamp[11:13]}<br>
                <strong>Esportazione:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
            </div>
        
            {html_content}
        </body>
        </html>
        """
    
        # Salva il file con un nome descrittivo
        html_path = f"output/genspark_formatted/{keyword}_genspark_{timestamp}.html"
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(enhanced_html)
    
        # Salva anche una versione "current" per facilità di accesso
        current_path = f"output/genspark_formatted/{keyword}_genspark_current.html"
        with open(current_path, "w", encoding="utf-8") as f:
            f.write(enhanced_html)
    
        return html_path, os.path.getsize(html_path)

    def export_to_docx(self):
        """Esporta l'analisi corrente in un documento DOCX."""
        try:
            self.add_log("Esportazione in DOCX...")
        
            # 1. Determina quale file di contesto utilizzare
            keyword = self.get_current_keyword() if hasattr(self, 'get_current_keyword') else "unknown"
        
            # Ottieni il nome del file sanitizzato
            import re
            safe_keyword = re.sub(r'[\\/*?:"<>|]', "", keyword).replace(" ", "_")[:30]
            context_file = f"context_{safe_keyword}.txt"
        
            # Se il file specifico non esiste, prova con il file di contesto generico
            if not os.path.exists(context_file):
                self.add_log(f"⚠️ File specifico {context_file} non trovato, provo con il file generico")
                context_file = "context.txt"
            
                if not os.path.exists(context_file):
                    self.add_log("❌ Nessun file di contesto trovato!")
                    return None
        
            self.add_log(f"📄 Utilizzo del file di contesto: {context_file}")
        
            # 2. Leggi il contenuto del file
            with open(context_file, "r", encoding="utf-8") as f:
                content = f.read()
        
            # 3. Crea il documento
            from docx import Document
            doc = Document()
            doc.add_heading(f"Analisi di mercato: {keyword}", 0)
        
            # Aggiungi informazioni sul documento
            timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
            doc.add_paragraph(f"Generato il: {timestamp}")
            doc.add_paragraph(f"Keyword: {keyword}")
            doc.add_paragraph("")  # Spazio aggiuntivo
        
            # 4. Dividi il contenuto in sezioni e aggiungile al documento
            sections = []
            section_pattern = r'===\s+([^=]+?)\s+-\s+\d{8}_\d{6}\s+===\n(.*?)(?=\n===|$)'
            section_matches = re.findall(section_pattern, content, re.DOTALL)
        
            if section_matches:
                sections = [(title.strip(), content.strip()) for title, content in section_matches]
                self.add_log(f"✅ Estratte {len(sections)} sezioni per il documento")
            else:
                # Fallback: divide per "==="
                raw_sections = content.split("===")
                for i, section in enumerate(raw_sections):
                    if section.strip():
                        sections.append((f"Sezione {i+1}", section.strip()))
                self.add_log(f"⚠️ Estratte {len(sections)} sezioni usando metodo fallback")
        
            # 5. Aggiungi ogni sezione al documento con formattazione appropriata
            for section_title, section_content in sections:
                # Aggiungi il titolo della sezione
                doc.add_heading(section_title, level=1)
            
                # Dividi il contenuto in paragrafi
                paragraphs = section_content.split('\n\n')
                for paragraph_text in paragraphs:
                    if paragraph_text.strip():
                        # Verifica se è un titolo di paragrafo
                        lines = paragraph_text.splitlines()
                        if lines and len(lines) > 1 and len(lines[0]) < 100 and lines[0].endswith(':'):
                            # Sembra un titolo di paragrafo
                            doc.add_heading(lines[0], level=2)
                            paragraph = doc.add_paragraph("\n".join(lines[1:]))
                        else:
                            # Paragrafo normale
                            paragraph = doc.add_paragraph(paragraph_text)
                
                # Aggiungi un divisore tra le sezioni
                doc.add_paragraph("")
        
            # 6. Salva il documento con un nome significativo
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"analisi_{safe_keyword}_{timestamp}.docx"
            doc.save(output_path)
        
            self.add_log(f"✅ Documento DOCX salvato: {output_path}")
            return output_path
    
        except Exception as e:
            self.add_log(f"❌ Errore nell'esportazione DOCX: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return None

    def export_to_pdf(self):
        try:
            self.add_log("Esportazione in PDF...")
        
            # 1. Prima esporta in DOCX
            docx_path = self.export_to_docx()
            if not docx_path:
                raise Exception("Errore nell'esportazione DOCX preliminare")
        
            # 2. Converti DOCX in PDF - metodo preferito: python-docx2pdf
            try:
                # Prova a importare docx2pdf
                from docx2pdf import convert
            
                # Conversione diretta
                pdf_path = docx_path.replace('.docx', '.pdf')
                self.add_log(f"Conversione {docx_path} in {pdf_path}...")
                convert(docx_path, pdf_path)
            
                self.add_log(f"✅ Documento PDF salvato: {pdf_path}")
                return pdf_path
            
            except ImportError:
                # Fallback: utilizza un messaggio informativo se docx2pdf non è installato
                self.add_log("⚠️ Modulo python-docx2pdf non trovato.")
                self.add_log("⚠️ Per convertire in PDF, apri il file DOCX e usa 'Salva come PDF'")
                self.add_log(f"⚠️ Il file DOCX è disponibile qui: {docx_path}")
                return docx_path
        
        except Exception as e:
            self.add_log(f"❌ Errore nell'esportazione PDF: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return None

    def export_to_txt(self):
        try:
            self.add_log("Esportazione in TXT...")
        
            # 1. Verifica se esiste il file di contesto
            context_file = "context.txt"
            if not os.path.exists(context_file):
                self.add_log("❌ File context.txt non trovato!")
                return None
        
            # 2. Ottieni la keyword corrente o usa un valore predefinito
            keyword = self.get_current_keyword() if hasattr(self, 'get_current_keyword') else "unknown"
        
            # 3. Crea la directory di output se non esiste
            output_dir = os.path.join(os.getcwd(), "output")
            os.makedirs(output_dir, exist_ok=True)
        
            # 4. Prepara il percorso del file di output
            import datetime
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(output_dir, f"analisi_{keyword}_{timestamp}.txt")
        
            # 5. Copia il file di contesto nel file di output
            import shutil
            shutil.copy2(context_file, output_path)
        
            self.add_log(f"✅ File TXT salvato: {output_path}")
            return output_path
        except Exception as e:
            self.add_log(f"❌ Errore nell'esportazione TXT: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return None

    def debug_check_components(self):
        """Verifica se results_display e analysis_status sono lo stesso oggetto"""
        if hasattr(self, 'results_display') and hasattr(self, 'analysis_status'):
            same_object = id(self.results_display) == id(self.analysis_status)
            self.add_log(f"DEBUG: results_display e analysis_status sono {'LO STESSO' if same_object else 'OGGETTI DIVERSI'}")
            self.add_log(f"DEBUG: ID results_display: {id(self.results_display)}")
            self.add_log(f"DEBUG: ID analysis_status: {id(self.analysis_status)}")
        else:
            self.add_log("DEBUG: Uno o entrambi i componenti non esistono ancora")
        return self.chat_manager.get_log_history_string()

    def load_saved_analyses_list(self):
        """
        Carica la lista delle analisi salvate come file HTML
        """
        import os
        import re
        from datetime import datetime
    
        # Lista per i percorsi dei file (valori del dropdown)
        file_paths = []
        # Lista per le etichette da mostrare (chiavi del dropdown)
        display_labels = []
    
        # Directory per i file HTML
        html_dir = os.path.join("output", "genspark_formatted")
    
        # Verifica se la directory esiste
        if not os.path.exists(html_dir):
            self.add_log(f"⚠️ Directory {html_dir} non trovata")
            return []
    
        try:
            # Trova tutti i file HTML nella directory
            html_files = [f for f in os.listdir(html_dir) if f.endswith(".html") and not f.endswith("_current.html")]
        
            # Struttura per tenere traccia dei file con le loro date
            file_info = []
        
            # Estrai le keyword e la data dai nomi dei file e usa la data di modifica come backup
            pattern = r"([^_]+)_genspark_(\d{8}_\d{6})\.html"
        
            for file in html_files:
                file_path = os.path.join(html_dir, file)
                file_mtime = os.path.getmtime(file_path)  # Timestamp di modifica
            
                # Prova prima con il pattern nel nome file
                match = re.match(pattern, file)
                if match:
                    keyword = match.group(1).replace("_", " ")
                    timestamp_str = match.group(2)
                
                    try:
                        # Converti il timestamp dal nome file in oggetto datetime
                        date_obj = datetime.strptime(timestamp_str, "%Y%m%d_%H%M%S")
                        date_str = date_obj.strftime("%d/%m/%Y %H:%M")
                        timestamp = date_obj.timestamp()  # Converti in timestamp per ordinamento
                    except ValueError:
                        # Se la conversione fallisce, usa la data di modifica
                        date_str = datetime.fromtimestamp(file_mtime).strftime("%d/%m/%Y %H:%M")
                        timestamp = file_mtime
                else:
                    # Se non corrisponde al pattern, usa il nome file come keyword
                    # e la data di modifica
                    keyword = file.replace(".html", "")
                    date_str = datetime.fromtimestamp(file_mtime).strftime("%d/%m/%Y %H:%M")
                    timestamp = file_mtime
            
                # Salva le informazioni
                display_text = f"{keyword} ({date_str})"
                file_info.append((display_text, file_path, timestamp))
        
            # Ordina per timestamp (data) in ordine decrescente (più recenti prima)
            file_info.sort(key=lambda x: x[2], reverse=True)
        
            # Crea le liste ordinate
            for display_text, file_path, _ in file_info:
                display_labels.append(display_text)
                file_paths.append(file_path)
        
            # Crea una lista di coppie (etichetta, valore)
            choices = list(zip(display_labels, file_paths))
        
            self.add_log(f"✅ Trovate {len(choices)} analisi HTML salvate")
            return choices
    
        except Exception as e:
            self.add_log(f"❌ Errore nel caricamento delle analisi HTML: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return []
    
    def load_selected_analysis(self, selected_option):
        """
        Carica l'analisi selezionata dal dropdown
        """
        # Verifica il tipo di dato ricevuto
        self.add_log(f"DEBUG: Tipo di dato selezionato: {type(selected_option)}")
        self.add_log(f"DEBUG: Valore selezionato: {selected_option}")
    
        # Gestisci il caso in cui selected_option è un dizionario
        if isinstance(selected_option, dict):
            if 'value' in selected_option:
                file_path = selected_option['value']
            else:
                self.add_log("⚠️ Formato dizionario non valido (manca 'value')")
                return "Errore: formato selezione non valido"
        else:
            file_path = selected_option
    
        if not file_path:
            self.add_log("⚠️ Nessun file selezionato")
            return "Nessun file selezionato"
    
        try:
            if os.path.exists(file_path):
                self.add_log(f"📂 Caricamento del file HTML: {file_path}")
            
                # Leggi il contenuto del file HTML
                with open(file_path, "r", encoding="utf-8") as f:
                    html_content = f.read()
            
                # Estrai il nome del file per il titolo
                file_name = os.path.basename(file_path)
                keyword = file_name.split('_genspark_')[0].replace('_', ' ') if '_genspark_' in file_name else file_name
            
                # Crea un container HTML con iframe per isolare il contenuto
                container_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <style>
                    /* Container esterno che mantiene layout fissato */
                    .fixed-container {{
                        width: 100%;
                        min-height: 600px;
                        padding: 0;
                        margin: 0;
                        box-sizing: border-box;
                        overflow: visible;
                        background: white;
                    }}
                
                    /* Header con informazioni sul file */
                    .file-header {{
                        background-color: #2563eb;
                        color: white;
                        padding: 10px 15px;
                        border-radius: 5px 5px 0 0;
                        font-size: 16px;
                        font-weight: bold;
                    }}
                
                    /* iframe per isolare il contenuto HTML */
                    .content-iframe {{
                        width: 100%;
                        height: 800px;
                        border: 1px solid #e5e7eb;
                        border-top: none;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                    }}
                    </style>
                </head>
                <body>
                    <div class="fixed-container">
                        <div class="file-header">Analisi: {keyword}</div>
                        <iframe class="content-iframe" srcdoc="{html_content.replace('"', '&quot;')}" frameborder="0"></iframe>
                    </div>
                </body>
                </html>
                """
            
                # Aggiorna il componente di visualizzazione
                if hasattr(self, 'results_display'):
                    # Usa il metodo update se disponibile
                    if hasattr(self.results_display, 'update'):
                        self.results_display.update(value=container_html, visible=True)
                        self.add_log(f"✅ Componente aggiornato tramite metodo update")
                    else:
                        # Altrimenti, assegna direttamente il valore
                        self.results_display.value = container_html
                        self.add_log(f"✅ Componente aggiornato tramite assegnazione diretta")
                
                    self.add_log("✅ Analisi HTML caricata nell'interfaccia")
                    return "Analisi caricata con successo", container_html
                else:
                    self.add_log("❌ Componente results_display non trovato")
                    return "Errore: componente di visualizzazione non trovato", ""
            else:
                self.add_log(f"❌ File non trovato: {file_path}")
                return f"Errore: file non trovato: {file_path}", ""
            
        except Exception as e:
            self.add_log(f"❌ Errore nel caricamento dell'analisi: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return f"Errore: {str(e)}", ""

    def get_current_keyword(self):
        """Recupera la keyword corrente dall'ultima analisi eseguita"""
        try:
            # Metodo 1: Cerca l'ultima keyword dal file di contesto
            import os
            import re
            from datetime import datetime
        
            # PRIMA: Verifica se la keyword è stata fornita direttamente nell'interfaccia UI
            if hasattr(self, 'keyword') and hasattr(self.keyword, 'value') and self.keyword.value:
                keyword_value = self.keyword.value.strip()
                self.add_log(f"DEBUG: Keyword trovata direttamente dall'interfaccia: {keyword_value}")
                return keyword_value
    
            # MIGLIORAMENTO: Controlla prima i file di backup più recenti (fonte più affidabile)
            backup_dir = "backups"
            if os.path.exists(backup_dir):
                # Cerca tutti i file che iniziano con context_ e finiscono con .txt
                backup_files = [f for f in os.listdir(backup_dir) if f.startswith("context_") and f.endswith(".txt")]
                if backup_files:
                    # Ordina per data di modifica (più recente prima)
                    backup_files.sort(key=lambda x: os.path.getmtime(os.path.join(backup_dir, x)), reverse=True)
                
                    # Estrai il nome della keyword dal primo file (il più recente)
                    if backup_files:
                        latest_file = backup_files[0]
                        mod_time = os.path.getmtime(os.path.join(backup_dir, latest_file))
                        mod_time_str = datetime.fromtimestamp(mod_time).strftime('%Y-%m-%d %H:%M:%S')
                        self.add_log(f"DEBUG: File di backup più recente: {latest_file} (modificato: {mod_time_str})")
                    
                        # Prova il pattern regex migliorato che supporta sia file con timestamp che senza
                        keyword_match = re.match(r"context_(.+?)(_\d{8}_\d{6})?\.txt", latest_file)
                        if keyword_match:
                            # Estrai la keyword e sostituisci gli underscore con spazi
                            keyword = keyword_match.group(1).replace("_", " ")
                            self.add_log(f"DEBUG: Keyword trovata dal file di backup: {keyword}")
                            print(f"DEBUG-KEYWORD: Keyword trovata dal backup più recente: {keyword}")
                            return keyword

            # Metodo 2: Controlla i dati del progetto corrente
            if hasattr(self, 'current_analysis') and self.current_analysis:
                # Verifica direttamente la chiave KEYWORD
                if 'KEYWORD' in self.current_analysis and self.current_analysis['KEYWORD']:
                    keyword = self.current_analysis['KEYWORD']
                    self.add_log(f"DEBUG: Keyword trovata direttamente in current_analysis: {keyword}")
                    print(f"DEBUG-KEYWORD: Keyword trovata in current_analysis: {keyword}")
                    return keyword
                
                # Verifica nei dati del progetto
                project_data = self.current_analysis.get('project_data', {})
                if 'KEYWORD' in project_data and project_data['KEYWORD']:
                    keyword = project_data['KEYWORD']
                    self.add_log(f"DEBUG: Keyword trovata nei dati progetto: {keyword}")
                    print(f"DEBUG-KEYWORD: Keyword trovata in project_data: {keyword}")
                    return keyword
    
            # Metodo 3: Cerca nei log recenti
            if hasattr(self, 'log_history'):
                recent_logs = self.log_history[-30:] if len(self.log_history) > 30 else self.log_history
                for log in recent_logs:
                    # Pattern migliorato per trovare la keyword nei log
                    patterns = [
                        r'mercato per: ([^\n]+)',
                        r'analisi per:? ([^\n]+)',
                        r'keyword:? ([^\n]+)',
                        r'per la keyword "[^"]+"'
                    ]
                
                    for pattern in patterns:
                        match = re.search(pattern, log, re.IGNORECASE)
                        if match:
                            keyword = match.group(1).strip()
                            self.add_log(f"DEBUG: Keyword trovata nei log recenti con pattern '{pattern}': {keyword}")
                            print(f"DEBUG-KEYWORD: Keyword trovata nei log: {keyword}")
                            return keyword

            # Ultimo metodo: cerca nel file context.txt
            if os.path.exists("context.txt"):
                with open("context.txt", "r", encoding="utf-8") as f:
                    content = f.read()
                
                    # Cerca esplicitamente menzioni dirette di keyword
                    keyword_patterns = [
                        r'keyword[:\s]+[""]?([^"\n,]+)[""]?',
                        r'parola chiave[:\s]+[""]?([^"\n,]+)[""]?',
                        r'mercato per[:\s]+[""]?([^"\n,]+)[""]?'
                    ]
                
                    for pattern in keyword_patterns:
                        matches = re.findall(pattern, content, re.IGNORECASE)
                        if matches:
                            keyword = matches[0].strip()
                            self.add_log(f"DEBUG: Keyword trovata nel contenuto di context.txt: {keyword}")
                            print(f"DEBUG-KEYWORD: Keyword trovata in context.txt: {keyword}")
                            return keyword
            
                    # Trova le ultime sezioni ordinate per timestamp
                    sections = re.findall(r'===\s+([^=]+?)\s+-\s+(\d{8}_\d{6})\s+===', content)
                    if sections:
                        # Ordina per timestamp (più recente prima)
                        sorted_sections = sorted(sections, key=lambda x: x[1], reverse=True)
                
                        # Prendi la sezione più recente e cerca di estrarre la keyword
                        for section_name, timestamp in sorted_sections[:5]:  # Esamina le prime 5 sezioni più recenti
                            # Cerca la keyword nel nome della sezione
                            section_match = re.search(r'(?:Analisi|Analysis)[^-]*-\s*(.+)', section_name, re.IGNORECASE)
                            if section_match:
                                keyword = section_match.group(1).strip()
                                self.add_log(f"DEBUG: Keyword trovata nel titolo della sezione: {keyword}")
                                print(f"DEBUG-KEYWORD: Keyword trovata nel titolo della sezione: {keyword}")
                                return keyword
                    
                            # Cerca esplicitamente nel contenuto della sezione
                            section_start = content.find(f"=== {section_name}")
                            if section_start > 0:
                                section_end = content.find("===", section_start + len(section_name) + 6)
                                if section_end > section_start:
                                    section_content = content[section_start:section_end]
                                    keyword_match = re.search(r'[Kk]eyword:\s*[""]?([^"\n,]+)[""]?', section_content)
                                    if keyword_match:
                                        keyword = keyword_match.group(1).strip()
                                        self.add_log(f"DEBUG: Keyword trovata nel contenuto della sezione: {keyword}")
                                        print(f"DEBUG-KEYWORD: Keyword trovata nel contenuto della sezione: {keyword}")
                                        return keyword
    
            self.add_log("DEBUG: Nessuna keyword trovata, ritorno 'unknown'")
            print("DEBUG-KEYWORD: Nessuna keyword trovata, restituisco 'unknown'")
            return "unknown"
        except Exception as e:
            self.add_log(f"⚠️ Errore nel recupero della keyword: {str(e)}")
            import traceback
            error_trace = traceback.format_exc()
            self.add_log(f"Traceback recupero keyword:\n{error_trace}")
            print(f"DEBUG-KEYWORD: Errore nel recupero: {str(e)}\n{error_trace}")
            return "unknown"

    def extract_context_data(self):
        """
        Estrae dati strutturati dal contesto per l'HTML
        """
        context_data = {}
        try:
            # Leggi il file di contesto
            with open("context.txt", "r", encoding="utf-8") as f:
                content = f.read()
        
            # Estrai sezioni e dati chiave
            import re
        
            # Buyer Persona
            buyer_match = re.search(r"buyer persona[^\n]+\n(.*?)(?=\n===|\n\d+\)|$)", content, re.IGNORECASE | re.DOTALL)
            if buyer_match:
                context_data["BUYER_PERSONA_SUMMARY"] = buyer_match.group(1).strip()
        
            # Angolo di Attacco
            angle_match = re.search(r"angolo d[i']? attacco[^\n]+\n(.*?)(?=\n===|\n\d+\)|$)", content, re.IGNORECASE | re.DOTALL)
            if angle_match:
                context_data["ANGOLO_ATTACCO"] = angle_match.group(1).strip()
        
            # Insight di Mercato
            market_match = re.search(r"insight d[i']? mercato[^\n]+\n(.*?)(?=\n===|\n\d+\)|$)", content, re.IGNORECASE | re.DOTALL)
            if market_match:
                context_data["MARKET_INSIGHTS"] = market_match.group(1).strip()
        
            # Aggiungi il tipo di analisi
            if "CRISP" in content or "crisp" in content.lower():
                context_data["type"] = "CRISP"
            else:
                context_data["type"] = "Legacy"
            
        except Exception as e:
            self.add_log(f"⚠️ Errore nell'estrazione dei dati dal contesto: {str(e)}")
    
        return context_data

    def create_interface(self):
        with gr.Blocks(css="""
            /* Layout a griglia di base */
            .main-container {
                display: flex !important;
                flex-direction: row !important;
                flex-wrap: nowrap !important;
                width: 100% !important;
            }
        
            /* Colonna input (sinistra) - aumentata da 20% a 25% */
            .input-column {
                width: 25% !important;
                min-width: 220px !important;
                flex: 0 0 auto !important;
                padding-right: 15px;
                box-sizing: border-box;
            }
        
            /* Colonna output (destra) - ridotta da 80% a 75% */
            .output-column {
                width: 75% !important;
                flex: 1 1 auto !important;
                padding-left: 15px;
                box-sizing: border-box;
            }
        
            /* Contenitore per la lista delle analisi con scrollbar */
            .scrollable-analyses {
                max-height: 250px;
                overflow-y: auto !important;
                border: 1px solid #e5e7eb;
                border-radius: 4px;
                padding: 5px;
                margin-bottom: 10px;
                background-color: white;
            }

            /* CSS per forzare la scrollbar nel dropdown */
            /* Selettori specifici per cercare di catturare il menu del dropdown */
            .dropdown-with-scroll div[role="listbox"],
            .dropdown-with-scroll ul.options,
            .dropdown-with-scroll .wrap-inner .options-list,
            .dropdown-with-scroll .wrap > ul,
            .gradio-dropdown .wrap > div > ul,
            .gradio-dropdown [role="listbox"],
            .gradio-dropdown .options {
                max-height: 250px !important;
                overflow-y: auto !important;
            }            
        
            /* Stile per le opzioni radio nella lista scrollabile */
            .scrollable-analyses .wrap label {
                padding: 8px 5px;
                margin: 2px 0;
                border-bottom: 1px solid #f0f0f0;
                display: block;
                white-space: nowrap;
                overflow: hidden;
                text-overflow: ellipsis;
                cursor: pointer;
            }
        
            .scrollable-analyses .wrap label:hover {
                background-color: #f0f7ff;
            }
        
            /* Stile della scrollbar */
            .scrollable-analyses::-webkit-scrollbar {
                width: 6px;
            }
        
            .scrollable-analyses::-webkit-scrollbar-track {
                background: #f1f1f1; 
                border-radius: 10px;
            }
        
            .scrollable-analyses::-webkit-scrollbar-thumb {
                background: #888;
                border-radius: 10px;
            }
        
            .scrollable-analyses::-webkit-scrollbar-thumb:hover {
                background: #555;
            }
        
            /* Resto dei tuoi stili */
            .results-container {
                width: 100% !important;
                margin: 0 !important;
            }
        
            #results_display_main {
                width: 100% !important;
                max-width: 100% !important;
            }
        
            /* Stili per blocchi di stato e pulsanti */
            .status-text {
                margin: 10px 0;
                padding: 8px;
                background-color: #f3f4f6;
                border-radius: 4px;
                font-size: 0.9em;
            }
        
            /* Migliora l'aspetto dei pulsanti di esportazione */
            .input-column .row button {
                flex-grow: 1;
                text-align: center;
            }
        """) as interface:

            # Header principale
            with gr.Row(elem_classes=["header-container"]):
                gr.HTML("""
                    <div class="app-header">
                        <div class="logo-container">
                            <span class="logo-icon">📕</span>
                            <h1>PubliScript 2.0</h1>
                        </div>
                        <div class="app-subtitle">Sistema Multi-Agente per Analisi e Generazione di Libri</div>
                    </div>
                """)

            self.crisp_phase_checkboxes = {}
            self.legacy_phase_checkboxes = {}
            # Sistema di tabs principale
            with gr.Tabs(elem_classes=["main-tabs"]) as tabs:
                # Tab 1: Setup & Connessione
                with gr.TabItem("1️⃣ Setup & Connessione", elem_classes=["tab-content"]):
                    with gr.Row():
                        with gr.Column(scale=1):
                            gr.Markdown("### Connessione al Servizio")
                            connect_btn = gr.Button("🔌 Connetti a Genspark", variant="primary")
                    
                            gr.Markdown("### Framework di Generazione")
                            # use_crisp_toggle = gr.Checkbox(
                            #    label="Usa Framework CRISP",
                            #    value=True,
                            #    info="Attiva il framework CRISP per un'analisi più strutturata"
                            # )
                    
                        with gr.Column(scale=2):
                            connection_status = gr.Markdown("**Stato**: Non connesso")
                    
                            # Area di log
                            self.log_output = gr.TextArea(
                                label="Console di sistema",
                                interactive=False,
                                lines=15,
                                value="Sistema inizializzato. Connettiti per iniziare.",
                                elem_id="log-output-area"
                            )
                    
                            # Pulsante per pulire i log
                            clear_log_btn = gr.Button("🧹 Pulisci Log", variant="secondary", size="sm")
            
                # Tab 2: Analisi di Mercato
                with gr.TabItem("2️⃣ Analisi di Mercato", elem_classes=["tab-content"]):
                    with gr.Row(elem_classes=["main-container"]):
                        # Colonna sinistra: Input (modificata da 1/5 a 1/4)
                        with gr.Column(scale=1, elem_classes=["input-column"]):
                            gr.Markdown("### Informazioni Base")

                            # Dropdown per caricare analisi esistente
                            load_analysis_dropdown = gr.Dropdown(
                                choices=self.load_saved_analyses_list(),
                                label="Carica Analisi Esistente",
                                interactive=True,
                                elem_classes=["dropdown-with-scroll", "dropdown-input"],
                                type="value"
                            )
    
                            # Pulsanti per caricare e aggiornare
                            with gr.Row():
                                load_analysis_btn = gr.Button("📂 Carica Analisi", variant="secondary")
                                refresh_list_btn = gr.Button("🔄", variant="secondary", size="sm")

                            # SPOSTATO: Stato analisi
                            analysis_status = gr.Markdown("**Stato analisi**: Non iniziata", elem_classes=["status-text"])

                            # SPOSTATO: Pulsanti di esportazione
                            gr.Markdown("### Esporta Risultati")
                            with gr.Row():
                                export_docx_btn = gr.Button("📄 Esporta DOCX", variant="secondary", size="sm")
                                export_pdf_btn = gr.Button("📑 Esporta PDF", variant="secondary", size="sm")
                            with gr.Row():
                                export_txt_btn = gr.Button("📝 Esporta TXT", variant="secondary", size="sm")
    
                            book_type = gr.Dropdown(
                                choices=self.book_types,
                                label="Tipo di Libro",
                                value=self.book_types[0],
                                elem_classes=["dropdown-input"]
                            )
    
                            keyword = gr.Textbox(
                                label="Keyword Principale",
                                placeholder="Parola chiave principale del libro",
                                elem_classes=["text-input"]
                            )
        
                            language = gr.Dropdown(
                                choices=["English", "Español", "Français", "Deutsch", "Italiano"],
                                label="Lingua Output Analisi",
                                value="Italiano",
                                elem_classes=["dropdown-input"]
                            )
        
                            market = gr.Dropdown(
                                choices=list(self.markets.keys()),
                                label="Mercato di Riferimento",
                                value="USA",
                                elem_classes=["dropdown-input"]
                            )
        
                            # Opzioni avanzate
                            with gr.Accordion("Mostra opzioni avanzate", open=False):
                                analysis_prompt = gr.TextArea(
                                    label="Prompt di Analisi (opzionale)",
                                    value=self.default_analysis_prompt,
                                    lines=6
                                )                        
                            #------------------INIZIO CAMBIAMENTO RADIO BUTTON-----------------------

                            # Aggiunta dell'Accordion per la selezione delle fasi
                            with gr.Accordion("📋 Seleziona fasi analisi", open=False):
                                gr.Markdown("### Scegli le fasi dell'analisi da eseguire")

                                # Prima riga per il tipo di analisi
                                with gr.Row():
                                    analysis_type_radio = gr.Radio(
                                        choices=["CRISP", "Legacy"],
                                        label="Tipo di Analisi",
                                        value="CRISP",
                                        interactive=True
                                    )
                                    self.analysis_type_radio = analysis_type_radio
    
                                # AGGIUNGI QUI LA FUNZIONE E IL COLLEGAMENTO ALL'EVENTO
                                def update_selected_type(value):
                                    self.selected_analysis_type = value
                                    self.add_log(f"Tipo di analisi selezionato: {value}")
                                    return None

                                # Collega l'handler all'evento change del radio button
                                self.analysis_type_radio.change(
                                    fn=update_selected_type,
                                    inputs=[self.analysis_type_radio],
                                    outputs=[]
                                )
    
                                # CRISP Phases: CheckboxGroup
                                with gr.Column(visible=True) as crisp_phases_col:
                                    gr.Markdown("#### Seleziona le Fasi dell'Analisi CRISP")
    
                                    # Basato sul file di configurazione
                                    crisp_phases = [
                                        ("CM-1", "Analisi di mercato - Fase 1"),
                                        ("CM-2", "Analisi di mercato - Fase 2"),
                                        ("CS-1", "Content Strategy - Fase 1"),
                                        ("CS-2", "Content Strategy - Fase 2"),
                                        ("CS-3", "Content Strategy - Fase 3"),
                                        ("CS-F", "Content Strategy - Fase Finale"),
                                        ("CP-1", "Content Production - Fase 1"),
                                        ("CP-2", "Content Production - Fase 2"),
                                        ("CPM-1", "Content Promotion & Marketing - Fase 1"),
                                        ("CPM-2", "Content Promotion & Marketing - Fase 2"),
                                        ("CPM-3", "Content Promotion & Marketing - Fase 3")
                                    ]
    
                                    # Crea le opzioni per il CheckboxGroup
                                    crisp_checkbox_options = [f"{phase_id}: {phase_desc}" for phase_id, phase_desc in crisp_phases]
    
                                    # Usa CheckboxGroup invece di RadioButton
                                    self.crisp_phase_checkboxes = gr.CheckboxGroup(
                                        choices=crisp_checkbox_options,
                                        label="Fasi da eseguire",
                                        value=[crisp_checkbox_options[0]],  # Seleziona la prima fase come default
                                        interactive=True
                                    )

                                # Legacy Phases: CheckboxGroup
                                with gr.Column(visible=False) as legacy_phases_col:
                                    gr.Markdown("#### Seleziona le Fasi dell'Analisi Legacy")
    
                                    # Basato sul file di configurazione
                                    legacy_phases = [
                                        ("LM-1", "Analisi concorrenza"),
                                        ("LM-2", "Valutazione profittabilità e competitività"),
                                        ("LM-3", "Analisi 3 migliori concorrenti"),
                                        ("LM-4", "Buyer persona sintetica"),
                                        ("LM-5", "Gap analysis"),
                                        ("LM-6", "Idee editoriali"),
                                        ("LM-7", "Valutazione idee"),
                                        ("LM-8", "Titoli, sottotitoli e indici")
                                    ]
    
                                    # Crea le opzioni per il CheckboxGroup
                                    legacy_checkbox_options = [f"{phase_id}: {phase_desc}" for phase_id, phase_desc in legacy_phases]
    
                                    # Usa CheckboxGroup invece di RadioButton
                                    self.legacy_phase_checkboxes = gr.CheckboxGroup(
                                        choices=legacy_checkbox_options,
                                        label="Fasi da eseguire",
                                        value=[legacy_checkbox_options[0]],  # Seleziona la prima fase come default
                                        interactive=True
                                    )

                                # Collega i radio button al cambio di visibilità
                                analysis_type_radio.change(
                                    fn=lambda x: (
                                        gr.update(visible=(x == "CRISP")), 
                                        gr.update(visible=(x == "Legacy"))
                                    ),
                                    inputs=[analysis_type_radio],
                                    outputs=[crisp_phases_col, legacy_phases_col]
                                )
    
                                # Seconda riga per i pulsanti
                                with gr.Row():
                                    select_all_btn = gr.Button("✅ Esegui Tutte le Fasi", variant="secondary", size="sm")
                                    deselect_all_btn = gr.Button("❌ Esegui Solo Fase Selezionata", variant="secondary", size="sm")
    
                                # Collega i pulsanti alle funzioni lambda per CheckboxGroup
                                select_all_btn.click(
                                    fn=lambda: crisp_checkbox_options,  # Tutte le opzioni
                                    inputs=[],
                                    outputs=self.crisp_phase_checkboxes
                                )

                                select_all_btn.click(
                                    fn=lambda: legacy_checkbox_options,  # Tutte le opzioni
                                    inputs=[],
                                    outputs=self.legacy_phase_checkboxes
                                )

                                deselect_all_btn.click(
                                    fn=lambda: [crisp_checkbox_options[0]],  # Solo la prima opzione
                                    inputs=[],
                                    outputs=self.crisp_phase_checkboxes
                                )

                                deselect_all_btn.click(
                                    fn=lambda: [legacy_checkbox_options[0]],  # Solo la prima opzione
                                    inputs=[],
                                    outputs=self.legacy_phase_checkboxes
                                )

                            #------------------------------FINE CAMBIAMENTO RADIO BUTTON-------------------------------------

                            # Aggiungi listener per i CheckboxGroup (opzionale, per logging)
                            self.crisp_phase_checkboxes.change(
                                fn=lambda value: print(f"CRISP fasi selezionate: {value}"),
                                inputs=[self.crisp_phase_checkboxes],
                                outputs=[]
                            )

                            self.legacy_phase_checkboxes.change(
                                fn=lambda value: print(f"Legacy fasi selezionate: {value}"),
                                inputs=[self.legacy_phase_checkboxes],
                                outputs=[]
                            )
                            
                            analyze_btn = gr.Button("🔍 Analizza Mercato", variant="primary", size="lg")
                            transfer_to_book_btn = gr.Button("📚 Prepara Generazione Libro", variant="secondary", size="lg")                        
                                       
                        # Colonna destra: Output (2/3)
                        with gr.Column(scale=3, elem_classes=["output-column"]):
                            gr.Markdown("### Analisi di Mercato")
                        
                            
                            # Area di output principale (formato HTML per ricchezza)
                            results_display = gr.HTML(
                                value="""
                                <div id='analysis-results-container' 
                                     style='min-height: 500px; 
                                            width: 100%;
                                            background-color: white; 
                                            border: 1px solid #e5e7eb; 
                                            border-radius: 8px;
                                            padding: 20px;
                                            box-shadow: 0 1px 3px rgba(0,0,0,0.1);'>
                                    <h3 style='color: #2563eb; margin-bottom: 15px;'>I risultati dell'analisi appariranno qui</h3>
                                    <p>Inserisci i parametri dell'analisi e clicca su "Analizza Mercato" per iniziare.</p>
                                </div>
                                """,
                                elem_classes=["results-container"],
                                visible=True,
                                elem_id="results_display_main"  # Aggiungi un ID univoco
                            )
                            self.results_display = results_display

                            # Verifica che results_display e analysis_status siano due componenti distinti
                            self.add_log(f"DEBUG ID check: results_display={id(self.results_display)}, analysis_status={id(analysis_status)}")

                            # Verifica che results_display e analysis_status siano due componenti distinti
                            self.add_log(f"DEBUG ID check: results_display={id(self.results_display)}, analysis_status={id(analysis_status)}")
            
                # Tab 3: Generazione Libro
                with gr.TabItem("3️⃣ Generazione Libro", elem_classes=["tab-content"]):
                     with gr.Group(visible=True) as book_details:
                        gr.Markdown("### Dettagli Libro")

                        with gr.Row():
                            with gr.Column(scale=1):
                                # Campo invisibile per memorizzare il tipo di libro dall'analisi
                                self.book_type_hidden = gr.Textbox(visible=False)
        
                                # Dropdown per il tipo di libro
                                self.book_type = gr.Dropdown(
                                    label="Tipo di Libro",
                                    choices=self.book_types,
                                    value=self.book_types[0] if self.book_types else None,
                                    interactive=True,
                                    elem_classes=["dropdown-input"]
                                )
        
                                # Dropdown per contenuto speciale
                                self.content_type = gr.Dropdown(
                                    label="Contenuto Speciale",
                                    choices=["Nessuno", "Ricette", "Esercizi", "Progetti", "Esempi", "Guide"],
                                    value="Nessuno",
                                    interactive=True,
                                    elem_classes=["dropdown-input"]
                                )
        
                                # Dropdown per immagini
                                self.use_images = gr.Dropdown(
                                    label="Includere Immagini",
                                    choices=["Sì", "No"],
                                    value="No",
                                    interactive=True,
                                    elem_classes=["dropdown-input"]
                                )
        
                                # Dropdown per approccio step-by-step
                                self.step_by_step = gr.Dropdown(
                                    label="Approccio Step-by-Step",
                                    choices=["Sì", "No"],
                                    value="No",
                                    interactive=True,
                                    elem_classes=["dropdown-input"]
                                )
        
                                # Checkbox per includere analisi nel contesto
                                self.include_analysis = gr.Checkbox(
                                    label="Includi file di analisi nel contesto",
                                    value=True,
                                    interactive=True,
                                )
        
                                # Titolo del libro
                                with gr.Row():
                                    with gr.Column(scale=2):
                                        self.book_title = gr.Textbox(
                                            label="Titolo del Libro",
                                            placeholder="Inserisci il titolo",
                                            elem_classes=["text-input"]
                                        )
                               
                              
                                    #Lingua del libro
                                    with gr.Column(scale=1):
                                        self.book_language = gr.Textbox(
                                            label="Lingua del Libro",
                                            value="English",
                                            placeholder="es: English, Italiano, Español",
                                            elem_classes=["text-input"]
                                        )
                                 
                                # Tono di voce
                                self.voice_style = gr.Textbox(
                                    label="Tono di Voce",
                                    placeholder="es: Formale, Tecnico, Conversazionale",
                                    value="",
                                    elem_classes=["text-input"]
                                )

                                # Angolo di attacco
                                self.editorial_angle = gr.Textbox(
                                    label="Angolo Editoriale",
                                    placeholder="Prospettiva/approccio principale del libro",
                                    elem_classes=["text-input"]
                                )    

                                # Unique Selling Proposition
                                self.book_usp = gr.Textbox(
                                    label="Unique Selling Proposition",
                                    placeholder="Cosa rende unico questo libro",
                                    elem_classes=["text-input"]
                                )
                
                                # Checkbox per includere l'analisi di mercato nel contesto
                                self.include_analysis = gr.Checkbox(
                                    label="Includi analisi di mercato nel contesto",
                                    value=True,
                                    interactive=True,
                                )                            
    
                                # Note personalizzate
                                self.custom_notes = gr.TextArea(
                                    label="Note Aggiuntive",
                                    placeholder="Inserisci eventuali specifiche particolari per la generazione del libro...",
                                    lines=3,
                                    elem_classes=["text-area"]
                                )
    
                            with gr.Column(scale=2):
                                self.book_index = gr.TextArea(
                                    label="Indice del Libro",
                                    placeholder="Inserisci l'indice, un capitolo per riga",
                                    lines=15,
                                    elem_classes=["text-area"]
                                )
                
                                # Aggiungi componente caricamento file
                                self.index_file = gr.File(
                                    label="Carica indice da file",
                                    file_types=[".txt", ".docx", ".html"],
                                    type="filepath"
                                )
    
                                with gr.Row():
                                    load_index_file_btn = gr.Button("📄 Carica indice da file", variant="secondary")
                                    load_from_analysis_btn = gr.Button("📋 Carica dati dall'analisi", variant="secondary")
                                    reset_index_btn = gr.Button("🔄 Reset", variant="secondary")
                                    extract_style_btn = gr.Button("🔍 Estrai tono di voce")
                                    
                                    
                                # Nel metodo create_interface
                                # Vicino al pulsante di estrazione tono di voce
                                with gr.Row():
                                    # Ottieni la lista di tutti i progetti dal database
                                    project_list = self.load_projects_list() if hasattr(self, 'load_projects_list') else []
    
                                    self.analysis_source_dropdown = gr.Dropdown(
                                        label="Analisi da cui estrarre",
                                        choices=project_list,
                                        interactive=True
                                    )
                                    refresh_projects_btn = gr.Button("🔄 Aggiorna lista progetti", variant="secondary")
                                    extract_style_btn = gr.Button("🔍 Estrai tono di voce dall'analisi selezionata")

                        # Pulsante di generazione
                        generate_btn = gr.Button("📚 Genera Libro", variant="primary", size="lg")
            
                # Tab 4: Database & Gestione
                with gr.TabItem("4️⃣ Database & Gestione", elem_classes=["tab-content"]):
                    with gr.Row():
                        with gr.Column(scale=1):
                            gr.Markdown("### Progetti Salvati")
            
                            # Opzioni di manutenzione
                            with gr.Row():
                                diagnose_db_btn = gr.Button("🛠️ Ripara Database", variant="primary")
            
                            # Ricerca e filtri
                            with gr.Row():
                                search_keyword = gr.Textbox(
                                    label="Cerca per keyword",
                                    placeholder="Inserisci una keyword...",
                                    show_label=True,
                                    elem_classes=["text-input"]
                                )
                        
                                refresh_db_btn = gr.Button("🔄 Carica Progetti", variant="primary")
                        
                            # Lista progetti
                            projects_list = gr.Dropdown(
                                label="Seleziona un Progetto",
                                choices=[],
                                interactive=True,
                                elem_id="projects_dropdown",
                                elem_classes=["dropdown-input"],
                                type="value"
                            )
      
                            self.projects_list = projects_list
                        
                            # Azioni progetti
                            with gr.Row():
                                resume_btn = gr.Button("▶️ Ripristina Analisi", variant="primary")
                                export_btn = gr.Button("📤 Esporta", variant="secondary")
                                delete_btn = gr.Button("🗑️ Elimina", variant="stop")
                
                        with gr.Column(scale=2):
                            gr.Markdown("### Dettagli Progetto")
                        
                            # Dettagli progetto in formato HTML
                            project_details = gr.HTML(
                                value="<div class='project-placeholder'>Seleziona un progetto per visualizzarne i dettagli</div>",
                                elem_classes=["project-details"]
                            )
            
                # Tab 5: Debug & Sviluppo
                with gr.TabItem("5️⃣ Debug & Sviluppo", elem_classes=["tab-content"]):
                    with gr.Row():
                        with gr.Column():
                            gr.Markdown("### Strumenti Debug")
                        
                            with gr.Row():
                                take_screenshot_btn = gr.Button("📸 Screenshot Browser", variant="secondary")
                                reset_context_btn = gr.Button("♻️ Reset Context Limit", variant="secondary")
                        
                            debug_output = gr.TextArea(
                                label="Output Debug",
                                interactive=False,
                                lines=10,
                                value=""
                            )

            # CSS personalizzato e librerie esterne
            gr.HTML("""
                <!-- Font Awesome per le icone -->
                <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0-beta3/css/all.min.css">

                <!-- Tailwind CSS per lo styling avanzato -->
                <link href="https://cdn.jsdelivr.net/npm/tailwindcss@2.2.19/dist/tailwind.min.css" rel="stylesheet">

                <!-- Google Fonts per tipografia migliorata (opzionale) -->
                <link rel="preconnect" href="https://fonts.googleapis.com">
                <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
                <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">

                <style>
                    /* Stile generale */
                    body {
                        font-family: 'Inter', 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    }

                    /* Header */
                    .app-header {
                        display: flex;
                        flex-direction: column;
                        align-items: center;
                        margin-bottom: 1rem;
                        background: linear-gradient(to right, #2563eb, #4f46e5);
                        color: white;
                        padding: 1rem;
                        border-radius: 0.5rem;
                        width: 100%;
                        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
                    }

                    .logo-container {
                        display: flex;
                        align-items: center;
                    }

                    .logo-icon {
                        font-size: 2rem;
                        margin-right: 0.5rem;
                    }

                    .app-header h1 {
                        margin: 0;
                        font-size: 1.8rem;
                        font-weight: 700;
                    }

                    .app-subtitle {
                        margin-top: 0.25rem;
                        font-size: 1rem;
                        opacity: 0.9;
                    }

                    /* Tabs */
                    .main-tabs button {
                        font-weight: 600;
                        padding: 0.75rem 1rem;
                        transition: all 0.2s ease;
                    }

                    .main-tabs button:hover {
                        background-color: #f3f4f6;
                    }

                    .main-tabs button.active {
                        border-bottom: 2px solid #3b82f6;
                        color: #1e40af;
                    }

                    .tab-content {
                        padding: 1rem 0;
                    }

                    /* Input fields */
                    .input-column {
                        border-right: 1px solid #e5e7eb;
                        padding-right: 1rem;
                    }

                    .text-input, .dropdown-input, .text-area {
                        margin-bottom: 1rem;
                        transition: border-color 0.15s ease-in-out, box-shadow 0.15s ease-in-out;
                    }

                    .text-input:focus, .dropdown-input:focus, .text-area:focus {
                        border-color: #3b82f6;
                        box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.25);
                    }

                    /* Output display */
                    .status-text {
                        background-color: #f3f4f6;
                        padding: 0.5rem;
                        border-radius: 0.375rem;
                        margin-bottom: 0.5rem;
                    }

                    .export-buttons {
                        margin-bottom: 1rem;
                        padding: 0.5rem 0;
                        border-bottom: 1px solid #e5e7eb;
                    }

                    .results-container {
                        border: 1px solid #e5e7eb;
                        border-radius: 0.5rem;
                        padding: 1rem;
                        background-color: #f9fafb;
                        min-height: 500px;
                        max-height: 70vh;
                        overflow-y: auto;
                        box-shadow: inset 0 2px 4px 0 rgba(0, 0, 0, 0.06);
                    }

                    .results-placeholder {
                        text-align: center;
                        color: #6b7280;
                        padding: 2rem;
                        font-style: italic;
                    }

                    /* Project details */
                    .project-details {
                        border: 1px solid #e5e7eb;
                        border-radius: 0.5rem;
                        padding: 1rem;
                        background-color: #f9fafb;
                        min-height: 400px;
                        max-height: 70vh;
                        overflow-y: auto;
                        transition: all 0.3s ease;
                    }

                    .project-details:hover {
                        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
                    }

                    .project-placeholder {
                        text-align: center;
                        color: #6b7280;
                        padding: 2rem;
                        font-style: italic;
                    }

                    /* Card style per i risultati */
                    .data-card {
                        background-color: white;
                        border-radius: 0.5rem;
                        box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1), 0 1px 2px 0 rgba(0, 0, 0, 0.06);
                        padding: 1rem;
                        margin-bottom: 1rem;
                        transition: all 0.3s ease;
                    }

                    .data-card:hover {
                        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
                        transform: translateY(-2px);
                    }

                    .data-card-title {
                        font-weight: 600;
                        color: #2563eb;
                        margin-bottom: 0.5rem;
                        display: flex;
                        align-items: center;
                    }

                    .data-card-title i {
                        margin-right: 0.5rem;
                    }

                    /* Scrollbar customization */
                    *::-webkit-scrollbar {
                        width: 8px;
                    }

                    *::-webkit-scrollbar-track {
                        background: #f1f1f1;
                        border-radius: 4px;
                    }

                    *::-webkit-scrollbar-thumb {
                        background: #888;
                        border-radius: 4px;
                    }

                    *::-webkit-scrollbar-thumb:hover {
                        background: #555;
                    }

                    /* Miglioramenti per pulsanti */
                    button {
                        transition: all 0.2s ease;
                    }

                    button:hover {
                        transform: translateY(-1px);
                    }

                    button:active {
                        transform: translateY(1px);
                    }

                    /* Classi di utilità per flexbox */
                    .flex-center {
                        display: flex;
                        align-items: center;
                        justify-content: center;
                    }

                    .flex-between {
                        display: flex;
                        align-items: center;
                        justify-content: space-between;
                    }

                    /* Badge e tag */
                    .badge {
                        display: inline-block;
                        padding: 0.25rem 0.75rem;
                        border-radius: 9999px;
                        font-size: 0.75rem;
                        font-weight: 600;
                        margin-right: 0.5rem;
                    }

                    .badge-blue {
                        background-color: #dbeafe;
                        color: #1e40af;
                    }

                    .badge-green {
                        background-color: #d1fae5;
                        color: #065f46;
                    }

                    .badge-yellow {
                        background-color: #fef3c7;
                        color: #92400e;
                    }

                    .badge-red {
                        background-color: #fee2e2;
                        color: #b91c1c;
                    }

                    /* Stili per i risultati dell'analisi */
                    .analysis-results .section-card {
                        transition: all 0.2s ease;
                    }

                    .analysis-results .section-card:hover {
                        transform: translateY(-2px);
                    }

                    .analysis-results .badge {
                        display: inline-block;
                        padding: 0.25rem 0.75rem;
                        border-radius: 9999px;
                        font-size: 0.75rem;
                        font-weight: 600;
                    }

                    .analysis-results .content {
                        font-size: 0.9rem;
                        line-height: 1.5;
                    }

                    .metadata-box {
                        border-left: 4px solid #3b82f6;
                    }

                    .metadata-item {
                        margin-bottom: 0.5rem;
                    }
                </style>

                <script>
                    // Funzione per mostrare/nascondere le colonne delle fasi in base al tipo di analisi
                    function togglePhaseColumns(analysisType) {
                        const crispCol = document.querySelector('[id$="crisp_phases_col"]');
                        const legacyCol = document.querySelector('[id$="legacy_phases_col"]');
        
                        if (analysisType === "CRISP") {
                            if (crispCol) crispCol.style.display = "block";
                            if (legacyCol) legacyCol.style.display = "none";
                        } else {
                            if (crispCol) crispCol.style.display = "none";
                            if (legacyCol) legacyCol.style.display = "block";
                        }
                    }

                    // Inizializza l'interfaccia dopo il caricamento
                    document.addEventListener('DOMContentLoaded', function() {
                        setTimeout(function() {
                            // Trova i radio button
                            const radioInputs = document.querySelectorAll('input[type="radio"][name$="analysis_type_radio"]');
            
                            // Aggiungi event listener a ciascun radio button
                            radioInputs.forEach(input => {
                                input.addEventListener('change', function() {
                                    togglePhaseColumns(this.value);
                                });
                            });
            
                            // Imposta lo stato iniziale
                            const selectedRadio = document.querySelector('input[type="radio"][name$="analysis_type_radio"]:checked');
                            if (selectedRadio) {
                                togglePhaseColumns(selectedRadio.value);
                            }
                
                            // Seleziona tutti i tab e aggiunge la classe active al primo
                            const tabs = document.querySelectorAll('.main-tabs button');
                            if (tabs.length > 0) {
                                tabs[0].classList.add('active');
                            }

                            // Aggiungi event listener per cambiare il tab attivo
                            tabs.forEach(tab => {
                                tab.addEventListener('click', function() {
                                    tabs.forEach(t => t.classList.remove('active'));
                                    this.classList.add('active');
                                });
                            });
                        }, 500); // Breve delay per assicurarsi che l'interfaccia sia caricata
                    });
                </script>
            """)
       
            # Function handlers
    
            # Funzione per pulire i log
            def clear_logs():
                self.log_history = ["Log cancellato."]
                return self.chat_manager.get_log_history_string()
    
            clear_log_btn.click(fn=clear_logs, outputs=self.log_output)
    
            # Connessione
            connect_btn.click(
                fn=self.connect_callback,
                outputs=[self.log_output, connection_status],
                show_progress=False
            )
    
            # Uso CRISP
            # use_crisp_toggle.change(
            #    fn=self._set_use_crisp,
            #    inputs=use_crisp_toggle,
            #    outputs=self.log_output
            # )
    
            # Analisi mercato
            analyze_btn.click(
                fn=self.analyze_market,
                inputs=[book_type, keyword, language, market, analysis_prompt],
                outputs=self.log_output,
                show_progress=False
            )
    
            # Completa analisi
            # complete_analysis_btn.click(
            #    fn=self.complete_analysis,
            #    outputs=[self.log_output, analysis_status, tabs, self.book_title, self.book_index, self.voice_style]
            # )
    
            # Generazione libro
            generate_btn.click(
                fn=self.generate_book,
                inputs=[self.book_title, self.book_language, self.voice_style, self.book_index],
                outputs=self.log_output
            )
    
            # Database e gestione progetti
            refresh_db_btn.click(
                fn=self.load_projects_list,
                outputs=[projects_list]
            )
  
            projects_list.change(
                fn=self.load_project_details,
                inputs=projects_list,
                outputs=[project_details]
            )
    
            resume_btn.click(
                fn=self.ripristina_analisi_da_database,
                inputs=projects_list,
                outputs=self.log_output
            )
    
            export_btn.click(
                fn=self.export_project,
                inputs=projects_list,
                outputs=self.log_output
            )
    
            delete_btn.click(
                fn=lambda project_name: [
                    self.delete_project(project_name),
                    self.load_projects_list(),
                    "<div style='text-align: center'><p>Progetto eliminato</p></div>"
                ],
                inputs=projects_list,
                outputs=[self.log_output, projects_list, project_details]
            )

            # Nel metodo create_interface, dopo aver definito load_analysis_btn:
            load_analysis_btn.click(
                fn=self.load_selected_analysis,
                inputs=[load_analysis_dropdown],
                outputs=[self.log_output, self.results_display]
            )

            refresh_list_btn.click(
                fn=self.load_saved_analyses_list, 
                inputs=[],
                outputs=[load_analysis_dropdown]
            )

            # Debug tools
            take_screenshot_btn.click(
                fn=lambda: self.take_debug_screenshot("debug"),
                outputs=debug_output
            )
    
            reset_context_btn.click(
                fn=self.handle_context_limit,
                outputs=self.log_output
            )

            # Connetti gli eventi per i pulsanti di esportazione
            export_docx_btn.click(
                fn=self.export_to_docx,
                outputs=[self.log_output]
            )
    
            export_pdf_btn.click(
                fn=self.export_to_pdf,
                outputs=[self.log_output]
            )
    
            export_txt_btn.click(
                fn=self.export_to_txt,
                outputs=[self.log_output]
            )
    
            # AGGIUNGI QUESTO
            diagnose_db_btn.click(
                fn=self.diagnose_and_fix_database,
                outputs=[project_details]  # Mostra risultati nel riquadro dettagli progetto
            )

            # Caricamento dati dall'analisi
            load_from_analysis_btn.click(
                fn=self.load_from_analysis,
                outputs=[self.book_title, self.book_index, self.voice_style, self.book_type_hidden]
            )

            # Aggiornamento automatico impostazioni immagini
            self.book_type.change(
                fn=self.update_image_settings,
                inputs=[self.book_type, self.content_type, self.step_by_step],
                outputs=[self.use_images, self.custom_notes]
            )

            self.content_type.change(
                fn=self.update_image_settings,
                inputs=[self.book_type, self.content_type, self.step_by_step],
                outputs=[self.use_images, self.custom_notes]
            )

            self.step_by_step.change(
                fn=self.update_image_settings,
                inputs=[self.book_type, self.content_type, self.step_by_step],
                outputs=[self.use_images, self.custom_notes]
            )

            # Generazione libro con nuovi parametri
            generate_btn.click(
                fn=self.generate_book_enhanced,
                inputs=[
                    self.book_title, self.book_language, self.voice_style, self.book_index,
                    self.book_type, self.content_type, self.use_images, self.step_by_step,
                    self.include_analysis, self.custom_notes
                ],
                outputs=[self.log_output]
            )

            # Aggiungi nella sezione dedicata agli event handler
            load_index_file_btn.click(
                fn=self.load_index_from_file,
                inputs=[self.index_file],
                outputs=[self.book_index]
            )

            reset_index_btn.click(
                fn=self.reset_index,
                outputs=[self.book_index]
            )

            refresh_projects_btn.click(
                fn=self.refresh_projects,
                outputs=[self.analysis_source_dropdown]
            )

            extract_style_btn.click(
                fn=self.extract_and_save_voice_style,
                inputs=[self.analysis_source_dropdown],
                outputs=[gr.Textbox(label="Risultato estrazione")]
            )

            transfer_to_book_btn.click(
                fn=self.transfer_to_book_tab,
                inputs=[],
                outputs=[
                    self.log_output,
                    self.book_title,
                    self.voice_style, 
                    self.editorial_angle, 
                    self.book_usp, 
                    self.book_index
                ] 
            )
            
            return interface

    def format_and_save_analysis_html(self, keyword, context_content):
        """
        Formatta e salva il contenuto dell'analisi come file HTML.
    
        Args:
            keyword: Keyword dell'analisi
            context_content: Il contenuto del file di contesto
        
        Returns:
            tuple: (html_string, html_path)
        """
        try:
            import os
            from datetime import datetime
        
            # Crea la directory output se non esiste
            os.makedirs("output", exist_ok=True)
            os.makedirs("output/analisi_html", exist_ok=True)
        
            # Timestamp per il nome del file
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
            # Sanitizza la keyword per il nome file
            safe_keyword = keyword.replace(' ', '_').replace('/', '_').replace('\\', '_')
        
            # Estrai le sezioni del contesto
            import re
            sections = []
            section_pattern = r'===\s+([^=]+?)\s+-\s+\d{8}_\d{6}\s+===\n(.*?)(?=\n===|$)'
            section_matches = re.findall(section_pattern, context_content, re.DOTALL)
        
            if section_matches:
                sections = [(title.strip(), content.strip()) for title, content in section_matches]
                self.add_log(f"✅ Estratte {len(sections)} sezioni")
            else:
                # Fallback: divide per numeri
                number_pattern = r'(\d+\).*?)(?=\d+\)|$)'
                number_matches = re.findall(number_pattern, context_content, re.DOTALL)
            
                if number_matches:
                    sections = [(f"Sezione {i+1}", content.strip()) for i, content in enumerate(number_matches)]
                    self.add_log(f"✅ Estratte {len(sections)} sezioni (pattern numerico)")
                else:
                    # Ultimo fallback: usa il testo completo
                    sections = [("Risultati completi", context_content)]
                    self.add_log("⚠️ Nessuna sezione trovata, usando il testo completo")
        
            # Crea HTML di base
            html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>Analisi: {keyword}</title>
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        line-height: 1.6;
                        max-width: 1200px;
                        margin: 0 auto;
                        padding: 20px;
                        background-color: #f9f9f9;
                    }}
                    .header {{
                        background-color: #2563eb;
                        color: white;
                        padding: 20px;
                        border-radius: 8px;
                        margin-bottom: 30px;
                        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
                    }}
                    .header h1 {{
                        margin: 0;
                        font-size: 24px;
                    }}
                    .header p {{
                        margin: 5px 0 0;
                        opacity: 0.8;
                    }}
                    .section {{
                        background-color: white;
                        padding: 20px;
                        border-radius: 8px;
                        margin-bottom: 20px;
                        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
                    }}
                    .section-title {{
                        font-weight: bold;
                        font-size: 18px;
                        margin-bottom: 15px;
                        color: #2563eb;
                        border-bottom: 1px solid #e5e7eb;
                        padding-bottom: 8px;
                    }}
                    .error {{ background-color: #fee2e2; }}
                    .ok {{ background-color: #f0fdf4; }}
                    table {{
                        width: 100%;
                        border-collapse: collapse;
                        margin: 15px 0;
                    }}
                    th, td {{
                        border: 1px solid #e5e7eb;
                        padding: 8px 12px;
                        text-align: left;
                    }}
                    th {{
                        background-color: #f3f4f6;
                    }}
                    tr:nth-child(even) {{
                        background-color: #f9fafb;
                    }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>Analisi: {keyword}</h1>
                    <p>Generata il: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')} - Sezioni trovate: {len(sections)}</p>
                </div>
                <div class="content">
            """
        
            # Aggiungi le sezioni
            for title, content in sections:
                # Determina se è una risposta valida o un errore
                is_error = "richiesta abortita" in content.lower() or len(content) < 30
                section_class = "section error" if is_error else "section ok"
            
                # Converti le tabelle semplici in HTML
                table_pattern = r'([^\n]+\|[^\n]+\|[^\n]+\n[-|]+\n(?:[^\n]+\|[^\n]+\|[^\n]+\n)+)'
                content = re.sub(table_pattern, self._convert_markdown_table_to_html, content)
            
                # Converti i punti elenco in HTML
                bullet_pattern = r'(?:^|\n)(\s*[•\*\-]\s+[^\n]+)(?:\n|$)'
                content = re.sub(bullet_pattern, r'<li>\1</li>', content)
                if '<li>' in content:
                    content = content.replace('<li>', '<ul><li>').replace('</li>', '</li></ul>')
            
                # Converti newline in <br> per HTML
                content_html = content.replace("\n", "<br>")
            
                html += f"""
                <div class="{section_class}">
                    <div class="section-title">{title}</div>
                    <div>{content_html}</div>
                </div>
                """
        
            html += """
                </div>
            </body>
            </html>
            """
        
            # Salva il file HTML
            html_path = f"output/analisi_html/{safe_keyword}_{timestamp}.html"
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html)
        
            # Salva anche una versione "current"
            current_path = f"output/analisi_html/{safe_keyword}_current.html"
            with open(current_path, "w", encoding="utf-8") as f:
                f.write(html)
        
            self.add_log(f"✅ HTML salvato: {html_path} e {current_path}")
        
            return html, html_path
        except Exception as e:
            import traceback
            self.add_log(f"❌ Errore nel formato HTML: {str(e)}")
            self.add_log(traceback.format_exc())
            return None, None

    def _convert_markdown_table_to_html(self, match):
        """Converte una tabella markdown in HTML"""
        try:
            table_text = match.group(1)
            lines = table_text.strip().split('\n')
        
            # Rimuovi la riga di separazione
            lines = [line for line in lines if not line.strip().startswith('|-')]
        
            html_table = "<table>\n<thead>\n<tr>\n"
        
            # Intestazioni
            if lines:
                header_cells = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
                for cell in header_cells:
                    html_table += f"<th>{cell}</th>\n"
                html_table += "</tr>\n</thead>\n<tbody>\n"
        
            # Righe dati
            for line in lines[1:]:
                if line.strip():
                    html_table += "<tr>\n"
                    cells = [cell.strip() for cell in line.split('|') if cell]
                    for cell in cells:
                        html_table += f"<td>{cell}</td>\n"
                    html_table += "</tr>\n"
        
            html_table += "</tbody>\n</table>"
            return html_table
        except Exception:
            return match.group(0)  # Ritorna il testo originale in caso di errori    

    def format_analysis_results_html(self, keyword, market, book_type, language, context=None):
        """
        Formatta i risultati dell'analisi in HTML per una visualizzazione migliore.
        Delega alla funzione in framework/formatters.py
    
        Args:
            keyword: Keyword analizzata
            market: Mercato target
            book_type: Tipo di libro
            language: Lingua dell'output
            context: Dati di contesto aggiuntivi (opzionale)
    
        Returns:
            str: HTML formattato con i risultati
        """
        # Import corretto basato sulla struttura reale del progetto
        import sys
        import os
        # Aggiungi la directory principale al path se necessario
        sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    
        from framework.formatters import format_analysis_results_html as formatter_format_analysis_results_html
    
        # Determina il tipo di analisi
        analysis_type = "CRISP" if hasattr(self, 'selected_analysis_type') and self.selected_analysis_type == "CRISP" else "Legacy"
    
        return formatter_format_analysis_results_html(
            keyword=keyword,
            market=market,
            book_type=book_type,
            language=language,
            context=context,
            log_callback=self.add_log,
            save_to_file=True,
            analysis_type=analysis_type
        )

    def capture_formatted_response(self, driver):
        """
        Cattura TUTTE le risposte formattate dall'interfaccia di Genspark.
        """
        try:
            from selenium.webdriver.common.by import By
            import time
        
            self.add_log("🔍 Inizio cattura HTML formattato completo")
            time.sleep(2)
        
            # Prova a catturare tutte le risposte con JavaScript
            try:
                complete_html = driver.execute_script("""
                    // Trova tutti i messaggi di risposta dell'AI
                    var messages = Array.from(document.querySelectorAll('.message, .message-content, .chat-message-item, .text-wrap'));
                
                    // Filtra per identificare le risposte dell'AI
                    var aiResponses = messages.filter(el => {
                        var isUserMsg = el.classList.contains('user-message') || el.classList.contains('question');
                        return !isUserMsg;
                    });
                
                    // Costruisci un contenitore HTML per tutte le risposte
                    var container = document.createElement('div');
                
                    // Aggiungi ogni risposta con un separatore
                    aiResponses.forEach((resp, index) => {
                        // Crea un separatore numerato
                        var separator = document.createElement('h3');
                        separator.textContent = 'Risposta #' + (index + 1);
                        separator.style.borderTop = '2px solid #2563eb';
                        separator.style.paddingTop = '15px';
                        separator.style.marginTop = '25px';
                    
                        // Aggiungi separatore e risposta
                        container.appendChild(separator);
                        container.appendChild(resp.cloneNode(true));
                    });
                
                    return container.innerHTML;
                """)
            
                if complete_html and len(complete_html) > 100:
                    self.add_log(f"✅ Catturate tutte le risposte AI: {len(complete_html)} caratteri")
                
                    styled_html = f"""
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <meta charset="UTF-8">
                        <style>
                            body {{ font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }}
                            pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 5px; overflow-x: auto; }}
                            code {{ background-color: #f5f5f5; padding: 2px 4px; border-radius: 3px; }}
                            table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
                            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                            th {{ background-color: #f2f2f2; }}
                            h1, h2, h3, h4 {{ color: #333; margin-top: 24px; margin-bottom: 16px; }}
                            img {{ max-width: 100%; height: auto; }}
                            .highlight {{ background-color: #ffeb3b; padding: 2px; }}
                        </style>
                        <title>Analisi Genspark Completa</title>
                    </head>
                    <body>
                        <h1>Risultati Completi dell'Analisi</h1>
                        <div class="genspark-responses">
                            {complete_html}
                        </div>
                    </body>
                    </html>
                    """
                
                    return styled_html
            except Exception as js_error:
                self.add_log(f"⚠️ Errore nell'approccio JavaScript: {str(js_error)}")
        
            # In caso di fallimento, genera HTML dal file context.txt
            return self._generate_html_from_context_file()
        
        except Exception as e:
            self.add_log(f"❌ Errore nel catturare l'HTML formattato: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
        
            # Tenta il fallback al file di contesto anche in caso di errore
            try:
                return self._generate_html_from_context_file()
            except:
                return None
            
    def _generate_html_from_context_file(self):
        """
        Helper method to generate HTML from context file with advanced formatting
        """
        try:
            self.add_log("🔄 Generazione HTML avanzata dal file di contesto")
    
            import os
            import re
    
            # Trova il file di contesto
            keyword = self.get_current_keyword() or "unknown"
            safe_keyword = re.sub(r'[\\/*?:"<>|]', "", keyword).replace(" ", "_")[:30]
    
            # Cerca prima il file specifico keyword, poi fallback al generico
            context_file = f"context_{safe_keyword}.txt"
            if not os.path.exists(context_file):
                context_file = "context.txt"
        
            if os.path.exists(context_file):
                with open(context_file, "r", encoding="utf-8") as f:
                    content = f.read()
        
                # Estrai le sezioni in modo più affidabile
                sections = []
                section_pattern = r'===\s+([^=]+?)\s+-\s+\d{8}_\d{6}\s+===\n(.*?)(?=\n===|$)'
                sections_matches = re.findall(section_pattern, content, re.DOTALL)
        
                if sections_matches:
                    sections = [(title.strip(), content.strip()) for title, content in sections_matches]
                    self.add_log(f"✅ Trovate {len(sections)} sezioni nel file di contesto")
                else:
                    # Fallback se non trova sezioni formattate
                    sections = [("Risultato Completo", content)]
            
                # Genera HTML dalle sezioni con formattazione avanzata
                sections_html = ""
                for idx, (title, section_content) in enumerate(sections, 1):
                    # Usa le funzioni di formattazione avanzata già esistenti
                    processed_content = self._process_section_content(section_content)
                
                    sections_html += f"""
                    <div class="section">
                        <h2>Sezione {idx}: {title}</h2>
                        <div>{processed_content}</div>
                    </div>
                    """
            
                # Usa il tuo metodo di generazione HTML esistente per la formattazione complessiva
                styled_html = self._generate_styled_html(sections_html, len(sections))
            
                self.add_log(f"✅ HTML avanzato generato dal file di contesto: {len(styled_html)} caratteri")
                return styled_html
        
            self.add_log("❌ Nessun file di contesto trovato")
            return None
        except Exception as e:
            self.add_log(f"❌ Errore nella generazione HTML dal file: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return None

    def _process_section_content(self, content):
        """
        Processa il contenuto di una sezione applicando tutte le formattazioni avanzate
        """
        try:
            import re
        
            # Identificazione e formattazione delle tabelle
            table_pattern = r'(\|[^\n]*\|[^\n]*\n\|[\s\-:|\n]*\|[^\n]*\n(?:\|[^\n]*\|[^\n]*\n)+)'
            tables = re.findall(table_pattern, content)
        
            for table in tables:
                formatted_table = self.process_table_html(table)
                content = content.replace(table, formatted_table)
        
            # Identificazione e formattazione delle liste
            # Cerca pattern di lista con - o *
            list_pattern = r'((?:^[-*•]\s+.+\n)+)'
            lists = re.finditer(list_pattern, content, re.MULTILINE)
        
            for match in lists:
                list_content = match.group(0)
                formatted_list = self.process_list_html(list_content, "Generic List")
                content = content.replace(list_content, formatted_list)
        
            # Identificazione e formattazione di pattern specifici
            # Pattern strutturali (formattazione visiva, sottolineature, etc)
            pattern_types = ["STRUCTURE_PATTERNS", "TITLE_PATTERNS"]
            for pattern_type in pattern_types:
                # Cerca sezioni che iniziano con il pattern_type
                pattern_section_regex = f"{pattern_type}[^:]*:(.*?)(?=\n[A-Z_]+:|$)"
                pattern_match = re.search(pattern_section_regex, content, re.DOTALL | re.IGNORECASE)
            
                if pattern_match:
                    pattern_content = pattern_match.group(1).strip()
                    formatted_pattern = self.process_patterns_html(pattern_content, pattern_type)
                    content = content.replace(pattern_content, formatted_pattern)
        
            # Processa il testo rimanente
            content = self.process_text(content)
        
            return content
        except Exception as e:
            self.add_log(f"⚠️ Errore nella formattazione avanzata: {str(e)}")
            # Fallback semplice
            return content.replace("\n", "<br>")


        
    def _generate_styled_html(self, content_html, response_count):
        """
        Genera l'HTML formattato con gli stili CSS.
        """
        print("DEBUG: VERSIONE AGGIORNATA DEL METODO ESEGUITA!")
        
        if not hasattr(self, '_analysis_session_id'):
            import uuid
            self._analysis_session_id = str(uuid.uuid4())[:8]

        from datetime import datetime

        # Genera un ID sessione se non esiste
        if not hasattr(self, '_analysis_session_id'):
            import uuid
            self._analysis_session_id = str(uuid.uuid4())[:8]

        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ 
                    font-family: Arial, sans-serif; 
                    line-height: 1.6; 
                    max-width: 1000px; 
                    margin: 0 auto; 
                    padding: 20px;
                    background-color: #f9f9f9;
                }}
                .header {{
                    background-color: #2563eb;
                    color: white;
                    padding: 20px;
                    text-align: center;
                    border-radius: 8px;
                    margin-bottom: 20px;
                    box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
                }}
                .ai-response {{
                    background-color: white;
                    border-radius: 8px;
                    padding: 15px;
                    margin-bottom: 25px;
                    box-shadow: 0 2px 5px rgba(0, 0, 0, 0.05);
                    border-left: 4px solid #2563eb;
                }}
                pre {{ 
                    background-color: #f5f5f5; 
                    padding: 10px; 
                    border-radius: 5px; 
                    overflow-x: auto; 
                    font-size: 14px;
                    border: 1px solid #e0e0e0;
                }}
                code {{ 
                    background-color: #f5f5f5; 
                    padding: 2px 4px; 
                    border-radius: 3px; 
                    font-size: 0.9em;
                    font-family: Consolas, monospace;
                }}
                table {{ 
                    border-collapse: collapse; 
                    width: 100%; 
                    margin: 20px 0;
                    background-color: white;
                }}
                th, td {{ 
                    border: 1px solid #ddd; 
                    padding: 12px; 
                    text-align: left; 
                }}
                th {{ 
                    background-color: #f2f2f2; 
                    font-weight: bold;
                }}
                tr:nth-child(even) {{ 
                    background-color: #f9f9f9; 
                }}
                h1, h2, h3, h4 {{ 
                    color: #2563eb; 
                    margin-top: 24px; 
                    margin-bottom: 16px; 
                }}
                img {{ 
                    max-width: 100%; 
                    height: auto;
                    border-radius: 8px;
                    box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
                }}
                .highlight {{ 
                    background-color: #ffeb3b; 
                    padding: 2px 4px; 
                    border-radius: 3px;
                }}
                .meta-info {{
                    color: #666;
                    font-size: 0.9rem;
                    margin-bottom: 20px;
                    background-color: #f0f4f8;
                    padding: 10px;
                    border-radius: 5px;
                }}
                ul, ol {{
                    margin-top: 15px;
                    margin-bottom: 15px;
                }}
                li {{
                    margin-bottom: 8px;
                }}
                blockquote {{
                    border-left: 4px solid #ccc;
                    margin: 15px 0;
                    padding: 10px 20px;
                    background-color: #f9f9f9;
                    font-style: italic;
                }}
            </style>
            <title>Analisi Genspark - Solo Risposte AI</title>
        </head>
        <body>
            <div class="header">
                <h1>Analisi Genspark - Risposte AI</h1>
                <p>Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            </div>
    
            <div class="meta-info">
                <p><strong>Sessione:</strong> {getattr(self, '_analysis_session_id', 'N/A')}</p>
                <p><strong>Numero totale di risposte:</strong> {response_count}</p>
                <p><strong>Keyword:</strong> {self.get_current_keyword() if hasattr(self, 'get_current_keyword') else "N/A"}</p>
            </div>
    
            <div class="ai-responses-container">
                {content_html}
            </div>
    
            <div class="footer" style="text-align: center; margin-top: 30px; color: #666; font-size: 0.8em;">
                <p>Generato da PubliScript {datetime.now().year}</p>
            </div>
        </body>
        </html>
        """

        return styled_html   

    def process_text(self, text):
        """
        Processa il testo con formattazione di base
        Delega alla funzione in framework/formatters.py
    
        Args:
            text: Testo da processare
        
        Returns:
            str: Testo formattato in HTML
        """
        from framework.formatters import process_text as formatter_process_text
    
        return formatter_process_text(text)

    def process_list_html(self, content, list_type):
        """
        Formatta una lista in HTML
        Delega alla funzione in framework/formatters.py
    
        Args:
            content: Contenuto della lista
            list_type: Tipo di lista (REVIEW_INSIGHTS, IMPLEMENTATION_OBSTACLES, MARKET_GAPS, ecc.)
        
        Returns:
            str: Lista formattata in HTML
        """
        from framework.formatters import process_list_html as formatter_process_list_html
    
        return formatter_process_list_html(content, list_type)

    def process_patterns_html(self, content, pattern_type):
        """
        Formatta pattern di titoli o strutture in HTML
        Delega alla funzione in framework/formatters.py
    
        Args:
            content: Contenuto del pattern
            pattern_type: Tipo di pattern (TITLE_PATTERNS, STRUCTURE_PATTERNS, ecc.)
        
        Returns:
            str: Pattern formattato in HTML
        """
        from framework.formatters import process_patterns_html as formatter_process_patterns_html
    
        return formatter_process_patterns_html(content, pattern_type)

    def process_table_html(self, content):
        """
        Converte una tabella in formato markdown in HTML
        Delega alla funzione in framework/formatters.py
    
        Args:
            content: Contenuto della tabella in formato markdown
        
        Returns:
            str: Tabella formattata in HTML
        """
        from framework.formatters import process_table_html as formatter_process_table_html
    
        return formatter_process_table_html(content)


    def save_complete_html(self):
        """
        Salva solo la risposta dell'AI in un file HTML pulito e formattato.
        """
        try:
            import os
            import time
            import webbrowser
            from datetime import datetime
            from selenium.webdriver.common.by import By

            if not hasattr(self, 'driver') or not self.driver:
                self.add_log("❌ Driver Selenium non disponibile")
                return None

            driver = self.driver

            # Attendi che la risposta sia stabile
            self.add_log("⏳ Attesa della risposta Genspark stabile...")
        
            try:
                # Monitora stabilità come implementato nel metodo precedente
                initial_length = 0
                stable_cycles = 0
                max_cycles = 10
        
                for cycle in range(max_cycles):
                    elements = driver.find_elements(By.CSS_SELECTOR, "div.chat-wrapper div.desc > div > div > div")
                    current_length = max([len(e.text.strip()) for e in elements]) if elements else 0
            
                    if current_length > 0:
                        if current_length == initial_length:
                            stable_cycles += 1
                            self.add_log(f"⏳ Risposta stabile: {stable_cycles}/3 cicli ({current_length} caratteri)")
                            if stable_cycles >= 3:  # 3 cicli di stabilità
                                self.add_log(f"✅ Risposta stabilizzata a {current_length} caratteri")
                                break
                        else:
                            stable_cycles = 0
                            initial_length = current_length
                            self.add_log(f"📝 Risposta in evoluzione: {current_length} caratteri (ciclo {cycle+1})")
            
                    time.sleep(2)  # Attendi 2 secondi tra i cicli
            except Exception as wait_error:
                self.add_log(f"⚠️ Timeout durante l'attesa - {str(wait_error)}")

            # Estrai solo il testo della risposta (nessun HTML dell'interfaccia)
            response_text = ""
            try:
                elements = driver.find_elements(By.CSS_SELECTOR, "div.chat-wrapper div.desc > div > div > div")
                if elements:
                    response_text = elements[-1].text
                    self.add_log(f"✅ Testo della risposta estratto: {len(response_text)} caratteri")
                else:
                    self.add_log("⚠️ Nessun elemento di risposta trovato")
                    return None
            except Exception as extract_error:
                self.add_log(f"⚠️ Errore nell'estrazione del testo: {str(extract_error)}")
                return None

            # Formatta il testo della risposta in HTML (convertendo tabelle e markdown)
            formatted_html = self._format_text_to_html(response_text)
            if not formatted_html:
                self.add_log("⚠️ Errore nella formattazione del testo in HTML")
                return None

            # Prepara i nomi dei file
            keyword = self.get_current_keyword() if hasattr(self, 'get_current_keyword') else "analisi"
            safe_keyword = ''.join(c if c.isalnum() else '_' for c in keyword)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
            # Directory per gli screenshot
            screenshots_dir = os.path.join("output", "visual_capture")
            os.makedirs(screenshots_dir, exist_ok=True)
        
            # Directory per i file HTML
            html_dir = os.path.join("output", "genspark_formatted")
            os.makedirs(html_dir, exist_ok=True)
        
            # Cattura screenshot
            screenshot_file = os.path.join(screenshots_dir, f"{safe_keyword}_screenshot_{timestamp}.png")
            driver.save_screenshot(screenshot_file)
            self.add_log(f"📸 Screenshot salvato: {screenshot_file}")
        
            # Crea un HTML completo e ben formattato
            html_content = f"""<!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Analisi: {keyword}</title>
        <style>
            body {{ 
                font-family: Arial, sans-serif; 
                line-height: 1.6; 
                max-width: 1000px; 
                margin: 0 auto; 
                padding: 20px;
                background-color: #f9f9f9;
            }}
        
            h1, h2, h3, h4 {{ 
                color: #2563eb; 
                margin-top: 24px; 
                margin-bottom: 16px; 
            }}
        
            .content {{
                background-color: white;
                border-radius: 8px;
                padding: 20px;
                margin-top: 20px;
                box-shadow: 0 1px 3px rgba(0,0,0,0.1);
            }}
        
            table {{ 
                border-collapse: collapse; 
                width: 100%; 
                margin: 20px 0;
                background-color: white;
                box-shadow: 0 1px 3px rgba(0,0,0,0.1);
            }}
        
            table.data-table {{
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
                font-size: 1em;
                border-radius: 8px;
                overflow: hidden;
                box-shadow: 0 0 20px rgba(0, 0, 0, 0.15);
            }}

            table.data-table thead tr {{
                background-color: #2563eb;
                color: #ffffff;
                text-align: left;
            }}

            table.data-table th,
            table.data-table td {{
                padding: 12px 15px;
                border: 1px solid #dddddd;
            }}

            table.data-table tbody tr {{
                border-bottom: 1px solid #dddddd;
            }}

            table.data-table tbody tr:nth-of-type(even) {{
                background-color: #f3f3f3;
            }}

            table.data-table tbody tr:last-of-type {{
                border-bottom: 2px solid #2563eb;
            }}
        
            th, td {{ 
                border: 1px solid #ddd; 
                padding: 12px; 
                text-align: left; 
            }}
        
            th {{ 
                background-color: #f2f2f2; 
                font-weight: bold;
            }}
        
            tr:nth-child(even) {{ 
                background-color: #f9f9f9; 
            }}
        
            ul, ol {{
                margin-top: 15px;
                margin-bottom: 15px;
            }}
        
            li {{
                margin-bottom: 8px;
            }}
        </style>
    </head>
    <body>
        <h1>Analisi: {keyword}</h1>
        <div class="content">
            {formatted_html}
        </div>
    </body>
    </html>
    """
        
            # Salva il file HTML nella directory corretta
            html_file = os.path.join("output", "genspark_formatted", f"{safe_keyword}_genspark_{timestamp}.html")
            with open(html_file, "w", encoding="utf-8") as f:
                f.write(enhanced_html)

            self.add_log(f"✅ HTML con stili preservati salvato in: {html_file}")
        
            # Apri nel browser
            try:
                file_path = os.path.abspath(html_file).replace('\\', '/')
                webbrowser.open("file:///" + file_path)
                self.add_log("✅ File aperto nel browser")
            except Exception as browser_error:
                self.add_log(f"⚠️ Impossibile aprire nel browser: {str(browser_error)}")
        
            return html_file
    
        except Exception as e:
            import traceback
            self.add_log(f"❌ Errore nella cattura HTML: {str(e)}")
            self.add_log(traceback.format_exc())
        
            # Prova con il metodo di fallback
            try:
                self.add_log("🔄 Utilizzo metodo di fallback...")
                return self.save_complete_html_improved()
            except Exception as fallback_error:
                self.add_log(f"❌ Anche il fallback è fallito: {str(fallback_error)}")
                return None

    def _format_text_to_html(self, text):
        """
        Formatta il testo della risposta in HTML, convertendo markdown e tabelle.
        """
        try:
            # Importa le funzioni necessarie
            from framework.formatters import process_table_html, process_text
            import re
        
            # Processo il testo iniziale
            formatted_text = text
        
            # Rileva e formatta tabelle
            table_pattern = r'(\|[^\n]+\|\n\|[\-\|: ]+\|\n(?:\|[^\n]+\|\n?)+)'
        
            def format_markdown_table(match):
                table_text = match.group(1)
                try:
                    return process_table_html(table_text)
                except:
                    return table_text
        
            # Applica la formattazione delle tabelle
            formatted_text = re.sub(table_pattern, format_markdown_table, formatted_text, flags=re.MULTILINE)
        
            # Converti elenchi
            if hasattr(self, 'convert_lists_to_html'):
                formatted_text = self.convert_lists_to_html(formatted_text)
            
            # Formattazione base markdown
            # Titoli
            formatted_text = re.sub(r'^##\s+(.+)$', r'<h2>\1</h2>', formatted_text, flags=re.MULTILINE)
            formatted_text = re.sub(r'^###\s+(.+)$', r'<h3>\1</h3>', formatted_text, flags=re.MULTILINE)
        
            # Enfasi
            formatted_text = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', formatted_text)
            formatted_text = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', formatted_text)
        
            # Paragrafi
            formatted_text = re.sub(r'\n\n', '</p><p>', formatted_text)
            formatted_text = f"<p>{formatted_text}</p>"
        
            # Assicurati che le tabelle abbiano la classe corretta
            formatted_text = formatted_text.replace('<table>', '<table class="data-table">')
        
            return formatted_text
        
        except Exception as e:
            self.add_log(f"⚠️ Errore nella formattazione del testo: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return text  # Ritorna il testo originale in caso di errore

    def _enhance_html_for_visibility(self, html_content):
        """
        Migliora l'HTML aggiungendo stili forzati per garantire la corretta visualizzazione 
        di tabelle e altri elementi che potrebbero essere nascosti.
        """
        # CSS per forzare la visibilità di elementi potenzialmente nascosti
        visibility_css = """
        <style>
        /* Stili forzati per tabelle */
        table, tr, td, th {
            display: table !important;
            visibility: visible !important;
            opacity: 1 !important;
            border-collapse: collapse !important;
            box-sizing: border-box !important;
        }
        table {
            width: 100% !important;
            margin: 15px 0 !important;
            border: 1px solid #ddd !important;
        }
        th, td {
            border: 1px solid #ddd !important;
            padding: 8px !important;
            text-align: left !important;
        }
        th {
            background-color: #f2f2f2 !important;
            font-weight: bold !important;
        }
        tr:nth-child(even) {
            background-color: #f9f9f9 !important;
        }
    
        /* Stili forzati per contenuti nascosti */
        .hidden, [hidden], [style*="display: none"], [style*="visibility: hidden"] {
            display: block !important;
            visibility: visible !important;
            opacity: 1 !important;
        }
    
        /* Stili migliorati per elementi comuni */
        body {
            font-family: Arial, sans-serif !important;
            line-height: 1.6 !important;
            color: #333 !important;
            max-width: 1200px !important;
            margin: 0 auto !important;
            padding: 20px !important;
        }
        h1, h2, h3, h4 {
            color: #2563eb !important;
            margin-top: 20px !important;
            margin-bottom: 10px !important;
        }
        p {
            margin-bottom: 15px !important;
        }
        </style>
        """
    
        # Inserisci gli stili forzati prima della chiusura del tag head
        if "</head>" in html_content:
            enhanced_html = html_content.replace("</head>", f"{visibility_css}</head>")
        else:
            # Se non c'è un tag head, aggiungilo all'inizio del documento
            enhanced_html = f"<!DOCTYPE html><html><head>{visibility_css}</head>{html_content}</html>"
    
        return enhanced_html

    def save_complete_html_improved(self):
        """
        Versione migliorata che cattura la risposta COMPLETA dell'AI.
        """
        try:
            import os
            import time
            import webbrowser
            from datetime import datetime
            from selenium.webdriver.common.by import By

            if not hasattr(self, 'driver') or not self.driver:
                self.add_log("❌ Driver Selenium non disponibile")
                return None

            driver = self.driver

            # Attendi che la risposta sia COMPLETA (con il terminatore FINE)
            self.add_log("🔍 Verifica completezza della risposta (ricerca terminatore FINE)...")
        
            try:
                # Aspetta fino a 5 minuti per il terminatore FINE
                max_wait_seconds = 300
                start_time = time.time()
                found_terminator = False
            
                while time.time() - start_time < max_wait_seconds:
                    # Ottieni l'elemento della risposta
                    elements = driver.find_elements(By.CSS_SELECTOR, "div.chat-wrapper div.desc > div > div > div")
                
                    if elements:
                        # Controlla se il testo contiene il terminatore FINE
                        response_text = elements[-1].text
                        if "FINE" in response_text:
                            terminator_pos = response_text.find("FINE")
                            self.add_log(f"🔍 TERMINATORE: 'FINE' trovato alla posizione {terminator_pos}")
                            found_terminator = True
                        
                            # Aspetta ancora 5 secondi per sicurezza
                            time.sleep(5)
                            break
                
                    # Log periodico
                    elapsed = int(time.time() - start_time)
                    if elapsed % 20 == 0:  # ogni 20 secondi
                        if elements:
                            self.add_log(f"⏳ Attesa terminatore FINE... ({elapsed}s, {len(elements[-1].text)} caratteri)")
                        else:
                            self.add_log(f"⏳ Attesa terminatore FINE... ({elapsed}s)")
                
                    # Attendi prima del prossimo controllo
                    time.sleep(5)
            
                if not found_terminator:
                    self.add_log("⚠️ Terminatore FINE non trovato dopo il timeout - potrebbe essere una risposta incompleta")
                    # Qui potresti decidere di interrompere il processo
            except Exception as wait_error:
                self.add_log(f"⚠️ Errore durante l'attesa del terminatore: {str(wait_error)}")
        
            # Estrai l'HTML della risposta completa
            try:
                script = """
                // Funzione per estrarre il contenuto principale della risposta
                function extractMainResponse() {
                    // Trova tutti gli elementi della conversazione
                    const responseElements = document.querySelectorAll('div.chat-wrapper div.desc > div > div > div');
                    if (!responseElements || responseElements.length === 0) return '';
                
                    // Prendi l'ultimo elemento che contiene la risposta finale
                    const mainResponse = responseElements[responseElements.length - 1];
                
                    // Ritorna l'HTML completo di questo elemento
                    return mainResponse.innerHTML;
                }
            
                return extractMainResponse();
                """
                response_html = driver.execute_script(script)
            
                if not response_html:
                    self.add_log("⚠️ Nessuna risposta trovata")
                    return None
                
                self.add_log(f"✅ HTML della risposta estratto: {len(response_html)} caratteri")
            except Exception as extract_error:
                self.add_log(f"⚠️ Errore nell'estrazione dell'HTML: {str(extract_error)}")
                return None

            # Prepara i nomi dei file
            keyword = self.get_current_keyword() if hasattr(self, 'get_current_keyword') else "analisi"
            safe_keyword = ''.join(c if c.isalnum() else '_' for c in keyword)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

            # Directory per gli screenshot
            screenshots_dir = os.path.join("output", "visual_capture")
            os.makedirs(screenshots_dir, exist_ok=True)

            # Directory per i file HTML
            html_dir = os.path.join("output", "genspark_formatted")
            os.makedirs(html_dir, exist_ok=True)

            # Cattura screenshot
            screenshot_file = os.path.join(screenshots_dir, f"{safe_keyword}_screenshot_{timestamp}.png")
            driver.save_screenshot(screenshot_file)
            self.add_log(f"📸 Screenshot salvato: {screenshot_file}")
        
            # Crea HTML completo
            html_content = f"""<!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Analisi: {keyword}</title>
        <style>
            body {{ 
                font-family: Arial, sans-serif; 
                line-height: 1.6; 
                max-width: 1000px; 
                margin: 0 auto; 
                padding: 20px;
                background-color: #f9f9f9;
            }}
        
            h1, h2, h3, h4 {{ 
                color: #2563eb; 
                margin-top: 24px; 
                margin-bottom: 16px; 
            }}
        
            .content {{
                background-color: white;
                border-radius: 8px;
                padding: 20px;
                margin-top: 20px;
                box-shadow: 0 1px 3px rgba(0,0,0,0.1);
            }}
        
            table {{ 
                border-collapse: collapse; 
                width: 100%; 
                margin: 20px 0;
                background-color: white;
                box-shadow: 0 1px 3px rgba(0,0,0,0.1);
            }}
        
            table.data-table {{
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
                font-size: 1em;
                border-radius: 8px;
                overflow: hidden;
                box-shadow: 0 0 20px rgba(0, 0, 0, 0.15);
            }}

            table.data-table thead tr {{
                background-color: #2563eb;
                color: #ffffff;
                text-align: left;
            }}

            table.data-table th,
            table.data-table td {{
                padding: 12px 15px;
                border: 1px solid #dddddd;
            }}

            table.data-table tbody tr {{
                border-bottom: 1px solid #dddddd;
            }}

            table.data-table tbody tr:nth-of-type(even) {{
                background-color: #f3f3f3;
            }}

            table.data-table tbody tr:last-of-type {{
                border-bottom: 2px solid #2563eb;
            }}
        
            th, td {{ 
                border: 1px solid #ddd; 
                padding: 12px; 
                text-align: left; 
            }}
        
            th {{ 
                background-color: #f2f2f2; 
                font-weight: bold;
            }}
        
            tr:nth-child(even) {{ 
                background-color: #f9f9f9; 
            }}
        
            ul, ol {{
                margin-top: 15px;
                margin-bottom: 15px;
            }}
        
            li {{
                margin-bottom: 8px;
            }}
        </style>
    </head>
    <body>
        <h1>Analisi: {keyword}</h1>
        <div class="content">
            {response_html}
        </div>
    </body>
    </html>
    """
        
            # Salva il file HTML
            html_file = os.path.join(html_dir, f"{safe_keyword}_genspark_{timestamp}.html")
            with open(html_file, "w", encoding="utf-8") as f:
                f.write(html_content)
        
            self.add_log(f"✅ HTML salvato in: {html_file}")
        
            # Apri nel browser
            try:
                file_path = os.path.abspath(html_file).replace('\\', '/')
                webbrowser.open("file:///" + file_path)
                self.add_log("✅ File aperto nel browser")
            except Exception as browser_error:
                self.add_log(f"⚠️ Impossibile aprire nel browser: {str(browser_error)}")

            # Salva anche in formato DOCX
            try:
                # Verifica se python-docx è installato
                import importlib.util
                docx_spec = importlib.util.find_spec("docx")
    
                if docx_spec is None:
                    self.add_log("⚠️ Libreria python-docx non trovata. Per il supporto DOCX esegui: pip install python-docx")
                else:
                    import docx
                    from docx.shared import Pt, RGBColor
                    from bs4 import BeautifulSoup
                    import re
        
                    # Crea un nuovo documento
                    doc = docx.Document()
        
                    # Aggiungi il titolo
                    doc.add_heading(f"Analisi: {keyword}", 0)
        
                    # Analizza l'HTML con BeautifulSoup
                    soup = BeautifulSoup(response_html, 'html.parser')
        
                    # Processa ogni elemento HTML e aggiungi al documento DOCX
                    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'table', 'div']):
                        if element.name == 'h1':
                            doc.add_heading(element.text, 1)
                        elif element.name == 'h2':
                            doc.add_heading(element.text, 2)
                        elif element.name == 'h3':
                            doc.add_heading(element.text, 3)
                        elif element.name == 'p':
                            doc.add_paragraph(element.text)
                        elif element.name == 'ul':
                            for li in element.find_all('li'):
                                doc.add_paragraph(li.text, style='List Bullet')
                        elif element.name == 'ol':
                            for li in element.find_all('li'):
                                doc.add_paragraph(li.text, style='List Number')
                        elif element.name == 'table':
                            # Crea una tabella con il numero appropriato di righe e colonne
                            rows = element.find_all('tr')
                            if rows:
                                # Trova il numero massimo di celle in qualsiasi riga
                                max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
                    
                                # Crea la tabella
                                docx_table = doc.add_table(rows=len(rows), cols=max_cols)
                                docx_table.style = 'Table Grid'
                    
                                # Popola la tabella
                                for i, row in enumerate(rows):
                                    cells = row.find_all(['th', 'td'])
                                    for j, cell in enumerate(cells):
                                        if j < max_cols:  # Assicurati di non superare il numero di colonne
                                            docx_table.cell(i, j).text = cell.text
        
                    # Salva il documento DOCX
                    docx_file = os.path.join(html_dir, f"{safe_keyword}_genspark_{timestamp}.docx")
                    doc.save(docx_file)
                    self.add_log(f"✅ DOCX salvato in: {docx_file}")
            except Exception as docx_error:
                self.add_log(f"⚠️ Errore nella creazione del file DOCX: {str(docx_error)}")
        
            return html_file
    
        except Exception as e:
            import traceback
            self.add_log(f"❌ Errore nella cattura HTML: {str(e)}")
            self.add_log(traceback.format_exc())
            return None

# AGGIUNGERE QUESTO CODICE ALLA FINE DEL FILE book_builder.py

# Invece di sostituire il metodo a livello di classe, utilizziamo una funzione
# che applica il patching a un'istanza specifica

def apply_debug_patching(builder_instance):
    """
    Applica il debug patching al metodo send_to_genspark di un'istanza di AIBookBuilder
    
    Args:
        builder_instance: Istanza di AIBookBuilder a cui applicare il patching
    """
    # Salva il metodo originale dall'istanza (non dalla classe)
    original_method = builder_instance.send_to_genspark
    
    # Definisci il nuovo metodo con debug
    def debug_send_to_genspark(self, text, prompt_id=None, section_number=None):
        """Versione con debug del metodo send_to_genspark"""
        print(f"DEBUG_SEND: Invio testo a Genspark ({len(text)} caratteri)")
        print(f"DEBUG_SEND: Primi 50 caratteri del testo: {text[:50]}")
        
        if prompt_id:
            print(f"DEBUG_SEND: Prompt ID: {prompt_id}, Sezione: {section_number}")
            
        if hasattr(self, 'driver') and self.driver:
            print(f"DEBUG_SEND: URL attuale prima dell'invio: {self.driver.current_url}")
            
        import traceback
        caller = traceback.extract_stack()[-2]
        print(f"DEBUG_SEND: Chiamata da {caller.filename}:{caller.lineno}")
        
        # Chiama il metodo originale
        result = original_method(text, prompt_id, section_number)
        
        if hasattr(self, 'driver') and self.driver:
            print(f"DEBUG_SEND: URL attuale dopo l'invio: {self.driver.current_url}")
            
        if result:
            has_end = "FINE_RISPOSTA" in result or "FINE" in result
            print(f"DEBUG_SEND: Risultato ({len(result)} chars) - Contiene FINE_RISPOSTA: {has_end}")
            print(f"DEBUG_SEND: Preview inizio: {result[:100]}...")
            
            if len(result) > 200:
                print(f"DEBUG_SEND: Preview fine: ...{result[-200:]}")
            else:
                print(f"DEBUG_SEND: Preview fine: {result}")
        else:
            print("DEBUG_SEND: Nessun risultato ottenuto")
            
        return result
    
    # Sostituisci il metodo nell'istanza usando la tecnica dei tipi
    import types
    builder_instance.send_to_genspark = types.MethodType(debug_send_to_genspark, builder_instance)
    
    print("Debug patching applicato all'istanza AIBookBuilder")
    return True
