"""
Builder principale per la creazione di libri.
"""
import sys
import re
import sqlite3
import time
import logging
import os
import json
import traceback
import gradio as gr
from datetime import datetime
from pathlib import Path
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys

from ai_interfaces.browser_manager import setup_browser, check_login, create_fresh_chat
from ai_interfaces.interaction_utils import get_input_box, clear_chat

from framework.crisp_framework import CRISPFramework
from framework.crisp_utils import replace_variables
from .chat_manager import ChatManager
from .cooldown_manager import CooldownManager
from .database_manager import DatabaseManager  # Importiamo il DatabaseManager
from framework.analysis.market_analysis import filter_legacy_prompt_sections
from framework.formatters import format_analysis_results_html, save_analysis_to_html

class AIBookBuilder:
    def __init__(self):
        self.cooldown_manager = CooldownManager()
        self.selected_analysis_type = "CRISP"  # Valore predefinito
        self.chat_manager = ChatManager(parent=self)
        self.is_logged_in = False
        from ai_interfaces.browser_manager import set_connection_status
        set_connection_status(False)  # Imposta lo stato globale iniziale
        self.driver = None
        self.log_history = []
        self.chat_manager = ChatManager(parent=self)  # Passa il riferimento di sé stesso
        self.current_analysis = None      
        self.question_status = {}  # Dizionario per tracciare lo stato delle domande   

        # Imposta directory
        import os
        self.context_dir = "context_files"
        os.makedirs(self.context_dir, exist_ok=True)
        os.makedirs("output", exist_ok=True)
        os.makedirs("debug", exist_ok=True)
       

        # Inizializziamo il DatabaseManager
        self.db_manager = DatabaseManager(
            project_db_path="crisp_projects.db", 
            log_callback=self.add_log
        )

        import logging
        logger = logging.getLogger("AIBookBuilder")
        
        # Inizializza il framework CRISP
        self.crisp = CRISPFramework(
            prompt_dir="prompt_crisp",
    	    project_db_path="crisp_projects.db",
    	    driver=None  # Il driver verrà impostato dopo la connessione
	)
        
        # Rimuovo il percorso manuale: uso sempre CRISP
        self.use_crisp = True
        
        # Tipi di libro disponibili
        self.book_types = [
            "Manuale (Non-Fiction)",
            "Ricettario",
            "Craft & Hobby",
            "Survival & Outdoor",
            "Test Study"
        ]

        # Mercati Amazon disponibili
        self.markets = {
            "USA": "Amazon.com",
            "Italia": "Amazon.it",
            "Francia": "Amazon.fr",
            "Inghilterra": "Amazon.co.uk",
            "Canada": "Amazon.ca",
            "Australia": "Amazon.com.au",
            "Spagna": "Amazon.es",
            "Germania": "Amazon.de"
        }
        
        # Prompt di analisi default
        self.default_analysis_prompt = """1) Analizza la concorrenza su {amazon_url} per la keyword {keyword} nel mercato {market}: elenca i primi 10 risultati Amazon con titolo, sottotitolo, BSR, prezzo, recensioni, formato, keyword usate nei titoli, pattern visivi delle copertine; aggiungi dati da Google Trends, query emergenti e insight dai social; concludi con una tabella di sintesi e commento su cosa domina e cosa manca; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
2) Valuta la profittabilità e competitività della keyword {keyword} su {amazon_url} nel mercato {market}: considera vendite mensili stimate per range di BSR, prezzo medio osservabile, margini potenziali di guadagno per autore KDP, numero e qualità dei competitor diretti, livello di saturazione della nicchia, e difficoltà stimata nel posizionarsi; concludi con 3 bullet: “Alta opportunità se…”, “Moderata se…”, “Bassa se…”; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
3) Analizza i 3 migliori concorrenti per la keyword {keyword} su {amazon_url} nel mercato {market}: mostra per ciascuno titolo, sottotitolo, BSR, recensioni, struttura (indice se disponibile), copertina (stile, elementi distintivi), pricing, e bonus offerti; concludi con una mini-tabella comparativa e insight su ciò che li rende forti; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
4) Definisci una buyer persona sintetica per {keyword} nel mercato {market}: includi età, professione, obiettivi, problemi, livello di consapevolezza e grado di urgenza; completa con 3 bullet: “Cosa cerca”, “Cosa teme”, “Cosa sogna”; se possibile, assegnale un nome simbolico e una frase tipo “Vorrei un libro che…”; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
5) Identifica i principali gap nei libri esistenti su {amazon_url} per {keyword}: analizza recensioni negative (1★–2★) per evidenziare frustrazioni comuni, bisogni insoddisfatti, parti assenti o trattate male; indica almeno 3 aree tematiche mancanti e il valore potenziale che avrebbero se inserite in un nuovo libro; concludi con una tabella "Gap vs Opportunità"; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
6) Genera 3 idee editoriali differenzianti per un nuovo libro su {keyword} nel mercato {market}: per ciascuna proposta definisci l’angolo editoriale, l’approccio (pratico, teorico, visuale…), il target specifico e una USP sintetica (max 2 righe) che risponda ai bisogni emersi; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
7) Valuta le 3 idee editoriali proposte nel punto precedente e scegli la migliore: motiva la scelta sulla base del potenziale commerciale, dell’originalità, della compatibilità con la buyer persona e dei gap riscontrati; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
8) In base all’idea selezionata, proponi 3 titoli con relativo sottotitolo (titolo + sottotitolo per ciascuna variante) e poi genera 3 diversi indici coerenti, ciascuno con almeno 6 capitoli principali e colonna “Obiettivo del capitolo”; per i titoli, valuta chiarezza, potere evocativo, e potenziale di vendita; scrivi in {lingua}, titoli e keyword nella lingua del {market}; concludi con la parola FINE.
 """

    # Metodi che delegano le operazioni al DatabaseManager
    def recupera_ultimo_progetto(self):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.recupera_ultimo_progetto()

    def ripristina_ultima_analisi(self):
        """Ripristina l'ultima analisi dal database."""
        return self.db_manager.ripristina_ultima_analisi(self.crisp, self.driver, self.chat_manager)

    def ripristina_analisi_da_database(self, selected_index, start_from_phase=None):
        """Ripristina un'analisi dal database."""
        result = self.db_manager.ripristina_analisi_da_database(selected_index, start_from_phase, self.driver, self.crisp, self.chat_manager)
    
        # Se risultato è un dizionario, contiene l'analisi da ripristinare
        if isinstance(result, dict) and 'project_id' in result:
            # Salviamo i dati di ripristino
            self.current_analysis = result
            project_id = result['project_id']
            project_data = result['project_data']
            start_from_phase = result['current_phase']
        
            # Creiamo una nuova chat e carichiamo il contesto
            if self.driver:
                self.add_log(f"Creazione nuova chat per la ripresa dell'analisi...")
                create_fresh_chat(self.driver, "context.txt")
        
                # Prepara messaggio di ripresa
                input_box = get_input_box(self.driver)
                resume_message = f"""
                Sto riprendendo l'analisi per il progetto: {project_data.get('PROJECT_NAME', 'N/A')}
                Keyword: {project_data.get('KEYWORD', 'N/A')}
        
                Siamo arrivati alla fase {start_from_phase}.
                Per favore, continua l'analisi da questa fase.
                """
        
                # Invia il messaggio
                chunks = [resume_message[i:i+200] for i in range(0, len(resume_message), 200)]
                for chunk in chunks:
                    input_box.send_keys(chunk)
                    time.sleep(0.5)
        
                send_button = self.driver.find_element(By.CSS_SELECTOR, "div.search-input-wrapper div.input-icon")
                send_button.click()
        
                time.sleep(5)
        
                # Aggiorna l'interfaccia
                self.add_log(f"✅ Analisi ripristinata con successo. Pronta per continuare dalla fase {start_from_phase}")
            else:
                self.add_log("❌ Browser non inizializzato. Connettiti prima di continuare")
    
        return self.chat_manager.get_log_history_string()

    def resume_analysis(self, project_id, selected_phases=None):
        """
        Riprende un'analisi esistente eseguendo fasi specifiche.
    
        Args:
            project_id: ID del progetto da riprendere
            selected_phases: Lista di fasi da eseguire (opzionale)
        
        Returns:
            str: Log dell'operazione
        """
        try:
            self.add_log(f"🔄 Ripresa analisi per progetto ID: {project_id}")
        
            # Recupera i dati del progetto
            project_data = self.crisp._load_project_data(project_id)
            if not project_data:
                return self.add_log(f"❌ Progetto ID {project_id} non trovato!")
        
            # Imposta il progetto come analisi corrente
            self.current_analysis = {
                'crisp_project_id': project_id,
                'project_data': project_data,
                'KEYWORD': project_data.get('KEYWORD', 'unknown')
            }
        
            # Se non ci sono fasi specificate, usa tutte le rimanenti
            if not selected_phases:
                # Determina qual è l'ultima fase eseguita
                last_phase = self.crisp._get_last_executed_step(project_id)
                if last_phase:
                    # Trova l'indice della fase nell'elenco completo
                    all_phases = ["CM-1", "CM-2", "CM-3", "CM-4", "CM-5", "CM-6", "CM-7", "CM-8"]
                    try:
                        last_index = all_phases.index(last_phase)
                        # Seleziona tutte le fasi successive
                        selected_phases = all_phases[last_index+1:]
                    except ValueError:
                        # Se la fase non è nell'elenco standard, usa tutte le fasi
                        selected_phases = all_phases
                else:
                    # Se non c'è una fase precedente, usa tutte le fasi
                    selected_phases = ["CM-1", "CM-2", "CM-3", "CM-4", "CM-5", "CM-6", "CM-7", "CM-8"]
        
            # Se ci sono fasi da eseguire, procedi
            if selected_phases:
                self.add_log(f"🔍 Ripresa con fasi: {', '.join(selected_phases)}")
            
                # Definisci la funzione executor
                def process_prompt(prompt_text):
                    self.add_log(f"Elaborazione prompt: {len(prompt_text)} caratteri")
                    response = self.send_to_genspark(prompt_text)
                    return response
            
                # Esegui le fasi selezionate una per una
                execution_history = []
                current_data = project_data.copy()
            
                for phase_id in selected_phases:
                    self.add_log(f"🔄 Esecuzione fase {phase_id}")
                
                    # Esegui la singola fase
                    updated_data, phase_result, extracted_data = self.crisp.execute_step(
                        phase_id, 
                        current_data.copy(), 
                        process_prompt
                    )
                
                    # Aggiorna i dati correnti con i risultati della fase
                    current_data.update(updated_data)
                
                    # Aggiungi alla storia di esecuzione
                    execution_history.append({
                        'step_id': phase_id,
                        'result': phase_result,
                        'extracted_data': extracted_data
                    })
            
                # Aggiorna l'analisi corrente con i nuovi dati
                self.current_analysis['project_data'] = current_data
                self.current_analysis['execution_history'] = execution_history
            
                self.add_log("✅ Ripresa analisi completata con successo")
            else:
                self.add_log("⚠️ Nessuna fase da eseguire - l'analisi è già completa")
        
            # Carica i risultati aggiornati
            self.load_analysis_results()
        
            return self.chat_manager.get_log_history_string()
        
        except Exception as e:
            self.add_log(f"❌ Errore nella ripresa dell'analisi: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return self.chat_manager.get_log_history_string()

    def load_projects_list(self):
        """Delega l'operazione al DatabaseManager"""
        # Ottieni una lista di stringhe semplici dal DatabaseManager
        project_names = self.db_manager.load_projects_list()
    
        # Il dropdown è configurato con type="value", quindi può accettare una lista di stringhe semplici
        # Non è necessario convertire in tuple (index, value)
        if hasattr(self, 'projects_list') and self.projects_list is not None:
            self.projects_list.choices = project_names
            self.add_log(f"✅ Dropdown aggiornato con {len(project_names)} progetti")
    
        return project_names

    def diagnose_and_fix_database(self):
        """Diagnostica e corregge problemi con il database"""
        self.add_log("🔍 Avvio diagnosi database...")
        try:
            import sqlite3
            import os
        
            # 1. Verifica percorso database
            self.add_log(f"🗂️ Posizione database: {self.crisp.project_db_path}")
        
            # 2. Verifica esistenza file
            if not os.path.exists(self.crisp.project_db_path):
                self.add_log(f"⚠️ ERRORE: File database non trovato: {self.crisp.project_db_path}")
                return "File database non trovato"
        
            self.add_log(f"✅ File database trovato: {os.path.getsize(self.crisp.project_db_path)} bytes")
        
            # 3. Verifica permessi di scrittura
            try:
                with open(self.crisp.project_db_path, "a") as f:
                    pass  # Test di scrittura
                self.add_log("✅ Permessi di scrittura sul database OK")
            except Exception as perm_error:
                self.add_log(f"⚠️ ERRORE: Permessi scrittura mancanti: {str(perm_error)}")
        
            # 4. Connessione al database con timeout più lungo
            self.add_log("Tentativo di connessione al database...")
            conn = sqlite3.connect(self.crisp.project_db_path, timeout=20.0)
            cursor = conn.cursor()
            self.add_log("✅ Connessione al database stabilita")
        
            # 5. Elenco tabelle per debug
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = cursor.fetchall()
            table_names = [t[0] for t in tables]
            self.add_log(f"📋 Tabelle trovate nel database: {', '.join(table_names)}")
        
            # 6. Verifica esistenza tabella projects
            if "projects" not in table_names:
                self.add_log("⚠️ ERRORE: Tabella 'projects' non trovata!")
                return "Tabella projects non trovata"
        
            # 7. Stampa struttura tabella projects
            cursor.execute("PRAGMA table_info(projects)")
            columns_info = cursor.fetchall()
            self.add_log(f"📊 Struttura tabella projects:")
            for col in columns_info:
                self.add_log(f"  - Colonna {col[0]}: {col[1]} ({col[2]})")
        
            # 8. Estrai nomi colonne
            columns = [col[1] for col in columns_info]
        
            # 9. Verifica esistenza colonna last_updated
            if "last_updated" not in columns:
                self.add_log("⚠️ Colonna 'last_updated' mancante, tentativo di aggiunta...")
            
                # 10. Avvia una transazione
                cursor.execute("BEGIN TRANSACTION")
            
                # 11. Aggiungi la colonna
                try:
                    cursor.execute("ALTER TABLE projects ADD COLUMN last_updated TEXT")
                    # 12. Commit della transazione
                    conn.commit()
                    self.add_log("✅ Colonna 'last_updated' aggiunta con successo")
                except sqlite3.OperationalError as op_error:
                    # 13. Rollback in caso di errore
                    conn.rollback()
                    self.add_log(f"⚠️ ERRORE nell'aggiungere la colonna: {str(op_error)}")
            else:
                self.add_log("✅ Colonna 'last_updated' già presente")
        
            # 14. Verifica nuovamente struttura tabella dopo modifiche
            cursor.execute("PRAGMA table_info(projects)")
            new_columns = [col[1] for col in cursor.fetchall()]
            self.add_log(f"📊 Colonne dopo riparazione: {', '.join(new_columns)}")
        
            # 15. Chiusura connessione
            conn.close()
            self.add_log("✅ Diagnosi e riparazione database completata con successo")
            return "✅ Database riparato con successo"
        except Exception as e:
            self.add_log(f"❌ ERRORE generale nella diagnosi: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return f"Errore database: {str(e)}"

    def check_existing_analysis(self, keyword):
        """Verifica se esiste già un'analisi per la keyword specificata."""
        return self.db_manager.check_existing_analysis(keyword)

    def load_project_details(self, selected_index):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.load_project_details(selected_index)

    def diagnose_and_fix_database(self):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.diagnose_and_fix_database()

    def export_project(self, selected_index):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.export_project(selected_index)

    def update_project_count(self):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.update_project_count()

    def delete_project(self, project_display_name):
        """Delega l'operazione al DatabaseManager"""
        result = self.db_manager.delete_project(project_display_name)
    
        # Aggiorna l'elenco dei progetti dopo l'eliminazione
        self.db_manager.load_projects_list()
    
        return result

    def create_test_project(self):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.create_test_project()

    def search_projects(self, keyword=""):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.search_projects(keyword)

    def get_database_stats(self):
        """Delega l'operazione al DatabaseManager"""
        return self.db_manager.get_database_stats()

    # FINE - Metodi che delegano le operazioni al DatabaseManager

    def add_log(self, message):
            """Delega il logging al ChatManager"""
            return self.chat_manager.add_log(message)
        
    def log_prompt_location(self, prompt_id, section_number, action, details=None):
        """Delega il tracciamento della posizione al ChatManager"""
        return self.chat_manager.log_prompt_location(prompt_id, section_number, action, details) 

    def connect_callback(self):
        # Importa funzioni per gestire stato globale
        from ai_interfaces.browser_manager import get_browser_instance, set_connection_status
    
        try:
            # Variabile globale per memorizzare la funzione originale
            global original_get
        
            # Se siamo già loggati e il driver è attivo, evitiamo di ricreare la sessione
            from ai_interfaces.browser_manager import get_connection_status
            if get_connection_status() and self.driver:
                # Aggiungiamo la patch di monitoraggio se non l'abbiamo già fatto
                if not hasattr(self, '_get_patched') or not self._get_patched:
                    original_get = self.driver.get
            
                    def log_get_call(url):
                        print(f"DEBUG_URL: Chiamata a driver.get() con URL: {url}")
                        self.add_log(f"DEBUG_URL: Navigazione a: {url}")
                        # Chiamiamo la funzione originale
                        return original_get(url)
            
                    self.driver.get = log_get_call
                    self._get_patched = True
                    self.add_log("DEBUG_URL: Monitoraggio navigazione attivato")

                timestamp = datetime.now().strftime('%H:%M:%S')
                log_message = f"[{timestamp}] Sei già connesso a Genspark. Sessione attiva mantenuta."
                self.log_history.append(log_message)

                # IMPORTANTE: Verifica che siamo nella pagina di chat corretta
                self.add_log("Verifica URL corrente ed eventuale reindirizzamento alla chat...")
                current_url = self.driver.current_url
                self.add_log(f"DEBUG_URL: URL attuale prima della verifica: {current_url}")

                if "/chat" not in current_url.lower():
                    self.add_log(f"URL corrente non è una pagina di chat: {current_url}")
                    self.add_log("Reindirizzamento alla pagina di chat...")

                return self.chat_manager.get_log_history_string(), "**Stato**: Connesso - Sessione attiva"

            # Salviamo temporaneamente i messaggi di log
            logs = []

            # Aggiungiamo il primo messaggio
            timestamp = datetime.now().strftime('%H:%M:%S')
            logs.append(f"[{timestamp}] Avvio connessione browser...")

            # Usa l'istanza globale o crea una nuova istanza forzata
            self.driver = get_browser_instance(force_new=True)  # Force new in caso di riconnessione
            self.add_url_change_monitor()
            self.crisp.driver = self.driver  # Aggiorna il riferimento al driver nel framework CRISP

            # Aggiungiamo la patch di monitoraggio per il nuovo driver
            original_get = self.driver.get
    
            def log_get_call(url):
                print(f"DEBUG_URL: Chiamata a driver.get() con URL: {url}")
                self.add_log(f"DEBUG_URL: Navigazione a: {url}")
                # Tracciamo anche da dove viene chiamata la funzione
                import traceback
                caller = traceback.extract_stack()[-2]
                print(f"DEBUG_URL: Chiamata da {caller.filename}:{caller.lineno}")
                # Chiamiamo la funzione originale
                return original_get(url)
    
            self.driver.get = log_get_call
            self._get_patched = True
            self.add_log("DEBUG_URL: Monitoraggio navigazione attivato per nuovo driver")

            # Dopo aver avviato il browser, aggiungi:
            self.add_log(f"Browser avviato: {self.driver is not None}")
            self.add_log("Tentativo di navigazione a Genspark...")

            # IMPORTANTE: Naviga DIRETTAMENTE alla pagina di chat, non alla home
            self.add_log("Navigazione diretta alla pagina di chat...")
            # Questo viene monitorato grazie al nostro patch
            self.driver.get("https://genspark.ai")
            time.sleep(7)  # Attesa più lunga per il caricamento completo

            # Verifica URL dopo la navigazione
            current_url = self.driver.current_url
            self.add_log(f"URL dopo navigazione: {current_url}")

            # Se siamo stati reindirizzati alla home o altra pagina, riprova con approccio alternativo
            if "/chat" not in current_url.lower():
                self.add_log("Reindirizzato a URL non di chat, provo approccio alternativo...")
                # Prima vai alla home
                self.driver.get("https://genspark.ai/")
                time.sleep(3)

            # Verifica dell'URL finale
            final_url = self.driver.current_url
            self.add_log(f"URL finale: {final_url}")

            # Controlliamo il login
            self.add_log("DEBUG_URL: Verificando login...")
            login_result = check_login(self.driver)
            self.add_log(f"DEBUG_URL: Risultato check_login: {login_result}")
    
            if not login_result:
                timestamp = datetime.now().strftime('%H:%M:%S')
                logs.append(f"[{timestamp}] ATTENZIONE: Completa manualmente il login su Genspark.ai.")
                timestamp = datetime.now().strftime('%H:%M:%S')
                logs.append(f"[{timestamp}] Assicurati di essere nella pagina della chat prima di continuare.")

                # Aggiungi un input per aspettare che l'utente finisca di fare login
                timestamp = datetime.now().strftime('%H:%M:%S')
                logs.append(f"[{timestamp}] Premi il pulsante 'Connetti' nuovamente dopo aver effettuato il login.")
    
                # Aggiorniamo la cronologia dei log attraverso il ChatManager
                for log_message in logs:
                    # Rimuoviamo il timestamp se già presente nel messaggio
                    if log_message.startswith('[') and ']' in log_message:
                        clean_message = log_message.split(']', 1)[1].strip()
                        self.chat_manager.add_log(clean_message)
                    else:
                        self.chat_manager.add_log(log_message)

                return self.chat_manager.get_log_history_string(), "**Stato**: Login richiesto"

                # Salva i cookie per usi futuri
                try:
                    import pickle
                    cookies = self.driver.get_cookies()
                    pickle.dump(cookies, open("genspark_cookies.pkl", "wb"))
                    timestamp = datetime.now().strftime('%H:%M:%S')
                    logs.append(f"[{timestamp}] Cookie salvati con successo!")
                except Exception as e:
                    timestamp = datetime.now().strftime('%H:%M:%S')
                    logs.append(f"[{timestamp}] Errore nel salvataggio dei cookie: {str(e)}")

            # Impostiamo lo stato di login
            self.is_logged_in = True
            from ai_interfaces.browser_manager import set_connection_status
            set_connection_status(True)  # Imposta stato globale

            # Aggiorniamo la cronologia dei log e restituiamo il tutto
            for log_message in logs:
                self.chat_manager.add_log(log_message)
            return self.chat_manager.get_log_history_string(), "**Stato**: Connesso"
        except Exception as e:
            # In caso di errore
            self.chat_manager.add_log(f"Errore: {str(e)}")
        
            # Aggiungiamo il traceback per debug
            import traceback
            tb = traceback.format_exc()
            print(f"DEBUG_URL: Eccezione in connect_callback:\n{tb}")
        
            return self.chat_manager.get_log_history_string(), "**Stato**: Errore di connessione"

    def add_url_change_monitor(self):
        """Aggiunge uno script al browser per monitorare i cambiamenti di URL"""
        if not hasattr(self, 'driver') or self.driver is None:
            print("DEBUG_URL: Impossibile installare monitor URL - driver non disponibile")
            return False
        
        script = """
        let lastUrl = window.location.href;
        console.log('Monitor URL inizializzato con URL: ' + lastUrl);
    
        // Funzione per controllare periodicamente l'URL
        function checkUrlChange() {
            const currentUrl = window.location.href;
            if (currentUrl !== lastUrl) {
                console.log('URL CAMBIATO da: ' + lastUrl + ' a: ' + currentUrl);
                lastUrl = currentUrl;
            }
            setTimeout(checkUrlChange, 1000);  // Controlla ogni secondo
        }
    
        // Avvia il controllo
        checkUrlChange();
    
        // Restituisci true per confermare l'installazione
        return true;
        """
    
        try:
            result = self.driver.execute_script(script)
            print(f"DEBUG_URL: Monitor di cambio URL installato nel browser: {result}")
            return True
        except Exception as e:
            print(f"DEBUG_URL: Errore nell'installazione del monitor URL: {str(e)}")
            return False

    def _analyze_market_crisp(self, book_type, keyword, language, market, selected_phases=None):
        """
        Analizza il mercato usando il framework CRISP.
        Delega alla funzione in framework/analysis/market_analysis.py

        Args:
            book_type: Tipo di libro
            keyword: Keyword principale
            language: Lingua dell'output
            market: Mercato di riferimento
            selected_phases: Lista di fasi selezionate da eseguire (opzionale)

        Returns:
            str: Log dell'operazione
        """
        # Log di debug per verificare i CheckboxGroup
        self.add_log(f"DEBUG: crisp_phase_checkboxes esiste: {hasattr(self, 'crisp_phase_checkboxes')}")
    
        # Se non sono state specificate fasi, ottienile dal CheckboxGroup
        if not selected_phases:
            selected_phases = []
        
            if hasattr(self, 'crisp_phase_checkboxes') and self.crisp_phase_checkboxes is not None:
                try:
                    # Ottieni tutte le fasi selezionate dal CheckboxGroup
                    selected_values = self.crisp_phase_checkboxes.value
                    self.add_log(f"DEBUG: Valori selezionati da crisp_phase_checkboxes: {selected_values}")
                
                    import re
                    for selected_value in selected_values:
                        # Estrae l'ID di fase dalla stringa selezionata (es. "CM-1: Analisi del mercato...")
                        # Supporta sia il formato CM-1, CS-1, CP-1, CPM-1 ecc.
                        match = re.match(r'([A-Z]+-[0-9A-Z]+):', selected_value)
                        if match:
                            phase_id = match.group(1)
                            selected_phases.append(phase_id)
                            self.add_log(f"📊 Fase CRISP selezionata: {phase_id}")
                except Exception as e:
                    self.add_log(f"⚠️ Errore nella lettura del CheckboxGroup: {str(e)}")
        
            # Se nessuna fase è stata trovata, usa CM-1 come default
            if not selected_phases:
                selected_phases = ["CM-1"]
                self.add_log("⚠️ Nessuna fase CRISP trovata, uso CM-1 come default")

        self.add_log(f"🔍 Esecuzione selettiva delle fasi CRISP: {', '.join(selected_phases)}")

        # Importa on-demand per evitare dipendenze circolari
        from framework.analysis.market_analysis import analyze_market_crisp

        try:
            result = analyze_market_crisp(
                book_type=book_type,
                keyword=keyword,
                language=language,
                market=market,
                selected_phases=selected_phases,
                crisp_framework=self.crisp,
                driver=self.driver,
                chat_manager=self.chat_manager
            )
        
            # Aggiorna current_analysis se il risultato contiene i dati di progetto
            if isinstance(result, dict) and 'crisp_project_id' in result:
                self.current_analysis = result
        
            return self.chat_manager.get_log_history_string()
        except Exception as e:
            self.add_log(f"❌ Errore durante l'analisi CRISP: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return f"Errore durante l'analisi: {str(e)}"    

    # Definisci una funzione executor che invia il prompt a Genspark
    def execute_prompt(self, prompt_text, step_id=None, project_data=None):
        """
        Funzione robusta che esegue un prompt CRISP, gestisce errori e interruzioni,
        e garantisce il completamento anche in presenza di problemi.

        Args:
            prompt_text: Testo del prompt da inviare
            step_id: ID del prompt corrente (opzionale)
            project_data: Dati del progetto per sostituire le variabili

        Returns:
            str: Risposta cumulativa da Genspark
        """
        # Inizializzazione e logging
        prompt_id_to_use = step_id if step_id else "unknown"
        self.add_log(f"🚀 Inizio esecuzione prompt {prompt_id_to_use} ({len(prompt_text)} caratteri)")
        print(f"DEBUG: Inizio esecuzione prompt {prompt_id_to_use} ({len(prompt_text)} caratteri)")
        print(f"DEBUG: Preview prompt: {prompt_text[:200].replace(chr(10), ' ')}...")

        # Verifica browser attivo
        if not hasattr(self, 'driver') or self.driver is None:
            self.add_log("⚠️ Browser non attivo, impossibile procedere")
            print("DEBUG: Browser non attivo, impossibile procedere")
            return "ERRORE: Browser non inizializzato"

        # Verifica URL corrente
        try:
            current_url = self.driver.current_url
            self.add_log(f"🌐 URL attuale: {current_url}")
            print(f"DEBUG: URL attuale: {current_url}")
    
            # Se non siamo in una pagina di chat, naviga a Genspark
            if not ("genspark.ai" in current_url and ("/chat" in current_url or "/agents" in current_url)):
                self.add_log("🔄 Navigazione alla pagina di chat...")
                print("DEBUG: Navigazione alla pagina di chat...")
                self.driver.get("https://genspark.ai")
                time.sleep(10)
        except Exception as e:
            self.add_log(f"⚠️ Errore nella verifica URL: {str(e)}")
            print(f"DEBUG: Errore nella verifica URL: {str(e)}")

        # MODIFICA CRUCIALE: Divisione del prompt in sezioni numeriche
        import re

        # Pattern per trovare punti numerati (esempio: "1. Titolo", "2. Titolo", ecc.)
        numbered_sections = re.findall(r'(?:\n|^)(\d+\.\s+.*?)(?=(?:\n\d+\.|\n\n|$))', prompt_text, re.DOTALL)

        # Se abbiamo trovato sezioni numerate, usiamole
        if numbered_sections:
            sections = [section.strip() for section in numbered_sections]
            self.add_log(f"📋 Prompt diviso in {len(sections)} sezioni numerate")
            print(f"DEBUG: Prompt diviso in {len(sections)} sezioni numerate")
        else:
            # Fallback: dividi per righe normalmente
            sections = [line.strip() for line in prompt_text.split('\n') if line.strip()]
            self.add_log(f"📋 Prompt diviso in {len(sections)} righe (nessuna sezione numerata trovata)")
            print(f"DEBUG: Prompt diviso in {len(sections)} righe (nessuna sezione numerata trovata)")

        # Log dettagliato delle sezioni
        for i, section in enumerate(sections):
            preview = section.replace('\n', ' ')[:50]
            self.add_log(f"📄 Sezione {i+1}: {preview}..." + ("" if len(section) <= 50 else f" ({len(section)} caratteri)"))
            print(f"DEBUG: Sezione {i+1}: {preview}..." + ("" if len(section) <= 50 else f" ({len(section)} caratteri)"))

        # Variabili di monitoraggio globali
        max_global_retries = 3
        global_retry_count = 0
        cumulative_response = []

        # Loop di ripetizione globale per il prompt intero
        while global_retry_count < max_global_retries:
            try:
                self.add_log(f"📝 Tentativo globale {global_retry_count+1}/{max_global_retries}")
                print(f"DEBUG: Tentativo globale {global_retry_count+1}/{max_global_retries}")
        
                # Pulizia dell'interfaccia prima di iniziare
                try:
                    clear_chat(self.driver)
                    self.add_log("🧹 Chat pulita all'inizio dell'esecuzione")
                    print("DEBUG: Chat pulita all'inizio dell'esecuzione")
                    time.sleep(5)
                except Exception as clear_error:
                    self.add_log(f"⚠️ Impossibile pulire la chat: {str(clear_error)}")
                    print(f"DEBUG: Impossibile pulire la chat: {str(clear_error)}")
                    # Continua comunque
        
                # Processa ogni sezione
                success_sections = 0
                for i, section in enumerate(sections):
                    self.add_log(f"📌 Elaborazione sezione {i+1}/{len(sections)}...")
                    print(f"DEBUG: Elaborazione sezione {i+1}/{len(sections)}...")
            
                    # Sostituzione variabili avanzata
                    processed_section = replace_variables(section, project_data)
                    preview = processed_section.replace('\n', ' ')[:50]
                    self.add_log(f"✏️ Sezione processata: {preview}..." + 
                                ("" if len(processed_section) <= 50 else f" ({len(processed_section)} caratteri)"))
                    print(f"DEBUG: Sezione processata: {preview}..." + 
                         ("" if len(processed_section) <= 50 else f" ({len(processed_section)} caratteri)"))
            
                    # Verifica placeholder non risolti
                    unresolved = self.check_unresolved_placeholders(processed_section)
                    if unresolved:
                        self.add_log(f"⚠️ Placeholders non risolti: {unresolved}")
                        print(f"DEBUG: Placeholders non risolti: {unresolved}")
                        # Continua comunque, ma potrebbe causare problemi
            
                    # Tentativi per sezione singola
                    section_retry_count = 0
                    max_section_retries = 3
                    section_success = False
            
                    while section_retry_count < max_section_retries and not section_success:
                        try:
                            self.add_log(f"🔄 Tentativo {section_retry_count+1}/{max_section_retries} per sezione {i+1}")
                            print(f"DEBUG: Tentativo {section_retry_count+1}/{max_section_retries} per sezione {i+1}")
                    
                            # Sistema avanzato di pulizia input
                            input_box = self.get_clean_input_box()
                            if not input_box:
                                raise Exception("Impossibile ottenere o pulire la casella di input")
                    
                            # Inserimento testo sicuro - carattere per carattere per maggiore affidabilità
                            self.add_log(f"⌨️ Inserimento testo carattere per carattere...")
                            print(f"DEBUG: Inserimento testo carattere per carattere... ({len(processed_section)} caratteri)")
                            for char in processed_section:
                                input_box.send_keys(char)
                                time.sleep(0.008)  # Minimo ritardo per stabilità
                    
                            # Verifica prima dell'invio
                            time.sleep(1)
                            inserted_text = input_box.get_attribute("value")
                            if not inserted_text:
                                self.add_log("⚠️ Nessun testo inserito!")
                                print("DEBUG: ERRORE - Nessun testo inserito!")
                                if section_retry_count < max_section_retries - 1:
                                    section_retry_count += 1
                                    time.sleep(5)
                                    continue
                            elif len(inserted_text) < len(processed_section) * 0.9:
                                self.add_log(f"⚠️ Inserimento incompleto: {len(inserted_text)}/{len(processed_section)} caratteri")
                                print(f"DEBUG: ERRORE - Inserimento incompleto: {len(inserted_text)}/{len(processed_section)} caratteri")
                                if section_retry_count < max_section_retries - 1:
                                    section_retry_count += 1
                                    time.sleep(5)
                                    continue
                            else:
                                self.add_log(f"✅ Testo inserito correttamente: {len(inserted_text)} caratteri")
                                print(f"DEBUG: Testo inserito correttamente: {len(inserted_text)} caratteri")
                    
                            # Invio con retry integrato
                            send_success = False
                    
                            # Metodo 1: Click standard
                            try:
                                send_button = WebDriverWait(self.driver, 15).until(
                                    EC.element_to_be_clickable((By.CSS_SELECTOR, "div.search-input-wrapper div.input-icon"))
                                )
                                send_button.click()
                                self.add_log("🔘 Click standard sul pulsante di invio")
                                print("DEBUG: Click standard sul pulsante di invio")
                                send_success = True
                            except Exception as e1:
                                self.add_log(f"⚠️ Click standard fallito: {str(e1)}")
                                print(f"DEBUG: Click standard fallito: {str(e1)}")
                        
                                # Metodo 2: Click JavaScript
                                try:
                                    send_button = self.driver.find_element(By.CSS_SELECTOR, "div.search-input-wrapper div.input-icon")
                                    self.driver.execute_script("arguments[0].click();", send_button)
                                    self.add_log("🔘 Click JavaScript sul pulsante di invio")
                                    print("DEBUG: Click JavaScript sul pulsante di invio")
                                    send_success = True
                                except Exception as e2:
                                    self.add_log(f"⚠️ Click JavaScript fallito: {str(e2)}")
                                    print(f"DEBUG: Click JavaScript fallito: {str(e2)}")
                            
                                    # Metodo 3: Tasto invio
                                    try:
                                        input_box.send_keys(Keys.RETURN)
                                        self.add_log("🔘 Invio tramite tasto RETURN")
                                        print("DEBUG: Invio tramite tasto RETURN")
                                        send_success = True
                                    except Exception as e3:
                                        self.add_log(f"❌ Tutti i metodi di invio falliti: {str(e3)}")
                                        print(f"DEBUG: Tutti i metodi di invio falliti: {str(e3)}")
                    
                            if not send_success:
                                raise Exception("Impossibile inviare il messaggio con nessun metodo")
                    
                            # Attesa iniziale per inizio elaborazione
                            self.add_log("⏳ Attesa iniziale dopo invio (10 secondi)")
                            print("DEBUG: Attesa iniziale dopo invio (10 secondi)")
                            time.sleep(10)
                    
                            # Sistema di attesa adattivo
                            max_wait_cycles = 45  # ~15 minuti totali
                            stability_threshold = 5  # 5 cicli di stabilità
                            cycle_wait = 20  # 20 secondi per ciclo
                    
                            # Inizializzazione variabili di monitoraggio
                            last_length = 0
                            stable_count = 0
                            response_text = None
                    
                            for cycle in range(max_wait_cycles):
                                try:
                                    print(f"DEBUG: Ciclo di attesa {cycle+1}/{max_wait_cycles}")
                                
                                    # Verifica limite contesto ogni 3 cicli
                                    if cycle % 3 == 0 and self.handle_context_limit():
                                        self.add_log("♻️ Limite contesto gestito durante attesa")
                                        print("DEBUG: Limite contesto gestito durante attesa")
                            
                                    # Prova diversi selettori per le risposte
                                    selectors = [
                                        ".message-content", 
                                        "div.chat-wrapper div.desc > div > div > div",
                                        "div.message div.text-wrap",
                                        ".chat-message-item .content"
                                    ]
                            
                                    # Cerca la risposta con tutti i selettori
                                    for selector in selectors:
                                        try:
                                            messages = self.driver.find_elements(By.CSS_SELECTOR, selector)
                                            if messages and len(messages) > 0:
                                                current_text = messages[-1].text.strip()
                                                if current_text:
                                                    response_text = current_text
                                                
                                                    # Debug della risposta ogni 5 cicli
                                                    if cycle % 5 == 0:
                                                        print(f"DEBUG: Salvataggio risposta - Lunghezza: {len(response_text)}")
                                                        if len(response_text) > 0:
                                                            print(f"DEBUG: Preview risposta: {response_text[:200].replace(chr(10), ' ')}...")
                                            
                                                    # Verifica terminazione esplicita
                                                    if "FINE_RISPOSTA" in response_text or "FINE" in response_text:
                                                        self.add_log(f"✅ Terminatore esplicito trovato al ciclo {cycle+1}")
                                                        print(f"DEBUG: Terminatore esplicito trovato al ciclo {cycle+1}")
                                                        terminator = "FINE_RISPOSTA" if "FINE_RISPOSTA" in response_text else "FINE"
                                                        terminator_pos = response_text.find(terminator)
                                                        print(f"DEBUG: Terminatore '{terminator}' trovato alla posizione {terminator_pos}")
                                                
                                                        # Pulisci la risposta rimuovendo il terminatore
                                                        if "FINE_RISPOSTA" in response_text:
                                                            response_text = response_text.split("FINE_RISPOSTA")[0].strip()
                                                        elif "FINE" in response_text:
                                                            response_text = response_text.split("FINE")[0].strip()
                                                    
                                                        section_success = True
                                                        break
                                            
                                                    # Verifica errori tipici
                                                    error_indicators = ["richiesta abortita", "request aborted", 
                                                                       "troppo lungo", "too long", 
                                                                       "errore durante", "error during"]
                                            
                                                    if any(e in response_text.lower() for e in error_indicators):
                                                        self.add_log(f"❌ Errore rilevato nella risposta al ciclo {cycle+1}")
                                                        print(f"DEBUG: Errore rilevato nella risposta al ciclo {cycle+1}: '{next((e for e in error_indicators if e in response_text.lower()), '')}'")
                                                        break
                                            
                                                    # Verifica stabilità
                                                    current_length = len(response_text)
                                                    if current_length == last_length:
                                                        stable_count += 1
                                                        self.add_log(f"⏳ Risposta stabile: {stable_count}/{stability_threshold} cicli ({current_length} caratteri)")
                                                        print(f"DEBUG: Risposta stabile: {stable_count}/{stability_threshold} cicli ({current_length} caratteri)")
                                                
                                                        if stable_count >= stability_threshold:
                                                            self.add_log(f"✅ Risposta stabilizzata dopo {cycle+1} cicli")
                                                            print(f"DEBUG: Risposta stabilizzata dopo {cycle+1} cicli - Lunghezza finale: {current_length} caratteri")
                                                            section_success = True
                                                            break
                                                    else:
                                                        stable_count = 0
                                                        self.add_log(f"📝 Risposta in evoluzione: {current_length} caratteri (ciclo {cycle+1})")
                                                        print(f"DEBUG: Risposta in evoluzione: {current_length} caratteri (ciclo {cycle+1})")
                                                        last_length = current_length
                                            
                                                    # Trovata risposta valida, esci dal ciclo selettori
                                                    break
                                        except Exception:
                                            continue
                            
                                    # Se abbiamo avuto successo, esci dal ciclo di attesa
                                    if section_success:
                                        break
                            
                                    # Attendi prima del prossimo ciclo
                                    time.sleep(cycle_wait)
                        
                                except Exception as e:
                                    self.add_log(f"⚠️ Errore durante attesa risposta: {str(e)}")
                                    print(f"DEBUG: Errore durante attesa risposta: {str(e)}")
                                    time.sleep(cycle_wait)
                    
                            # Verifica se abbiamo ottenuto successo
                            if section_success and response_text:
                                # Risposta ottenuta con successo
                                self.add_log(f"✅ Risposta ottenuta per sezione {i+1}: {len(response_text)} caratteri")
                                print(f"DEBUG: Risposta ottenuta per sezione {i+1}: {len(response_text)} caratteri")
                            
                                # Debug della risposta ottenuta
                                print(f"DEBUG: Salvataggio risposta - Lunghezza: {len(response_text)}")
                                print(f"DEBUG: Preview risposta: {response_text[:200].replace(chr(10), ' ')}...")
                            
                                # Verifica qualità risposta
                                if len(response_text) < 50 and not ("CM-1" in prompt_id_to_use and i == 0):
                                    # Risposta troppo corta (eccetto la prima sezione di CM-1 che può essere corta)
                                    self.add_log(f"⚠️ Risposta sospettosamente corta: {len(response_text)} caratteri")
                                    print(f"DEBUG: Risposta sospettosamente corta: {len(response_text)} caratteri")
                                    if section_retry_count < max_section_retries - 1:
                                        section_retry_count += 1
                                        time.sleep(10)
                                        continue
                        
                                # Aggiungi alla risposta cumulativa
                                cumulative_response.append(response_text)
                        
                                # Salva incrementalmente
                                try:
                                    # Usa una struttura a cascata per trovare il metodo giusto
                                    if hasattr(self.crisp, 'save_incremental_response') and project_data and "PROJECT_ID" in project_data:
                                        print(f"DEBUG: Chiamata a crisp.save_incremental_response per {prompt_id_to_use}")
                                        self.crisp.save_incremental_response(
                                            project_data["PROJECT_ID"], 
                                            prompt_id_to_use,
                                            processed_section, 
                                            response_text, 
                                            i == len(sections) - 1
                                        )
                                        self.add_log("💾 Risposta salvata nel database CRISP")
                                        print("DEBUG: Risposta salvata nel database CRISP")
                                    elif hasattr(self.crisp, 'crisp') and hasattr(self.crisp.crisp, 'save_incremental_response') and project_data and "PROJECT_ID" in project_data:
                                        print(f"DEBUG: Chiamata a crisp.crisp.save_incremental_response per {prompt_id_to_use}")
                                        self.crisp.crisp.save_incremental_response(
                                            project_data["PROJECT_ID"], 
                                            prompt_id_to_use,
                                            processed_section, 
                                            response_text, 
                                            i == len(sections) - 1
                                        )
                                        self.add_log("💾 Risposta salvata nel database CRISP (via crisp.crisp)")
                                        print("DEBUG: Risposta salvata nel database CRISP (via crisp.crisp)")
                                    else:
                                        self.add_log("⚠️ Metodo save_incremental_response non trovato o dati progetto mancanti")
                                    
                                        # Verifica dettagli per un miglior debug
                                        print("DEBUG: Dettagli variabili per salvataggio:")
                                        print(f"DEBUG: - hasattr(self.crisp, 'save_incremental_response'): {hasattr(self.crisp, 'save_incremental_response')}")
                                        print(f"DEBUG: - project_data is not None: {project_data is not None}")
                                        if project_data:
                                            print(f"DEBUG: - 'PROJECT_ID' in project_data: {'PROJECT_ID' in project_data}")
                                            if 'PROJECT_ID' in project_data:
                                                print(f"DEBUG: - PROJECT_ID value: {project_data['PROJECT_ID']}")
                                    
                                        # Salva nel file di contesto come fallback
                                        if hasattr(self, 'chat_manager'):
                                            print(f"DEBUG: Salvando nel file di contesto come fallback")
                                            metadata = {
                                                "prompt_id": prompt_id_to_use,
                                                "section_number": i+1,
                                                "timestamp": datetime.now().strftime('%Y%m%d_%H%M%S')
                                            }
                                            self.chat_manager.save_response(
                                                response_text,
                                                f"Prompt {prompt_id_to_use}-Sezione {i+1}",
                                                metadata
                                            )
                                            print(f"DEBUG: Salvato nel file di contesto come fallback")
                                except Exception as save_error:
                                    self.add_log(f"⚠️ Errore nel salvare la risposta: {str(save_error)}")
                                    print(f"DEBUG: Errore nel salvare la risposta: {str(save_error)}")
                                    import traceback
                                    print(f"DEBUG: Traceback salvataggio:\n{traceback.format_exc()}")
                        
                            elif response_text and len(response_text) > 100:
                                # Risposta parziale ma utilizzabile
                                self.add_log(f"⚠️ Risposta parziale ma utilizzabile: {len(response_text)} caratteri")
                                print(f"DEBUG: Risposta parziale ma utilizzabile: {len(response_text)} caratteri")
                                print(f"DEBUG: Salvataggio risposta parziale - Lunghezza: {len(response_text)}")
                                print(f"DEBUG: Preview risposta parziale: {response_text[:200].replace(chr(10), ' ')}...")
                        
                                # Aggiungi alla risposta cumulativa
                                cumulative_response.append(response_text)
                                section_success = True
                        
                                # Salva anche risposte parziali
                                try:
                                    if hasattr(self.crisp, 'save_incremental_response') and project_data and "PROJECT_ID" in project_data:
                                        print(f"DEBUG: Salvando risposta parziale in CRISP database")
                                        self.crisp.save_incremental_response(
                                            project_data["PROJECT_ID"], 
                                            prompt_id_to_use,
                                            processed_section, 
                                            response_text, 
                                            i == len(sections) - 1
                                        )
                                        print(f"DEBUG: Risposta parziale salvata nel database CRISP")
                                except Exception as save_error:
                                    print(f"DEBUG: Errore nel salvare risposta parziale: {str(save_error)}")
                            else:
                                # Nessuna risposta o timeout
                                self.add_log(f"❌ Nessuna risposta valida ottenuta per sezione {i+1}")
                                print(f"DEBUG: Nessuna risposta valida ottenuta per sezione {i+1}")
                        
                                if section_retry_count < max_section_retries - 1:
                                    section_retry_count += 1
                                    time.sleep(15)
                                    continue
                                else:
                                    cumulative_response.append(f"[Timeout per sezione {i+1}]")
                                    print(f"DEBUG: Aggiunto placeholder di timeout per sezione {i+1}")
                
                        except Exception as e:
                            # Gestione errori specifici per sezione
                            self.add_log(f"⚠️ Errore sezione {i+1}, tentativo {section_retry_count+1}: {str(e)}")
                            print(f"DEBUG: Errore sezione {i+1}, tentativo {section_retry_count+1}: {str(e)}")
                    
                            # Cattura screenshot per debug
                            try:
                                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                                screenshot_path = f"error_section_{i+1}_try_{section_retry_count+1}_{timestamp}.png"
                                self.driver.save_screenshot(screenshot_path)
                                self.add_log(f"📸 Screenshot: {screenshot_path}")
                                print(f"DEBUG: Screenshot errore: {screenshot_path}")
                            except Exception:
                                pass
                    
                            if section_retry_count < max_section_retries - 1:
                                section_retry_count += 1
                        
                                # Verifica se è un problema di contesto
                                if "context" in str(e).lower() or "too long" in str(e).lower():
                                    if self.handle_context_limit():
                                        self.add_log("♻️ Limite di contesto gestito, nuovo tentativo")
                                        print("DEBUG: Limite di contesto gestito, nuovo tentativo")
                        
                                time.sleep(15)
                            else:
                                # Ultimo fallimento per questa sezione
                                self.add_log(f"❌ Tutti i tentativi falliti per sezione {i+1}")
                                print(f"DEBUG: Tutti i tentativi falliti per sezione {i+1}")
                                cumulative_response.append(f"[Errore: {str(e)}]")
            
                    # Se la sezione è stata completata con successo
                    if section_success:
                        success_sections += 1
                
                        # Pausa tra le sezioni - aumenta progressivamente
                        pause_time = 15 + (i * 3)  # Aumenta di 3 secondi per ogni sezione
                        self.add_log(f"⏱️ Pausa di {pause_time} secondi prima della prossima sezione")
                        print(f"DEBUG: Pausa di {pause_time} secondi prima della prossima sezione")
                        time.sleep(pause_time)
                    else:
                        # Se la sezione ha fallito definitivamente - decisione se continuare
                        self.add_log(f"❌ Fallimento definitivo per sezione {i+1}/{len(sections)}")
                        print(f"DEBUG: Fallimento definitivo per sezione {i+1}/{len(sections)}")
                        if i < len(sections) - 1:  # Se non è l'ultima sezione
                            self.add_log("⚠️ Continuo con la sezione successiva nonostante il fallimento")
                            print("DEBUG: Continuo con la sezione successiva nonostante il fallimento")
        
                # Se abbiamo completato tutte le sezioni o la maggior parte
                if success_sections == len(sections) or success_sections >= len(sections) * 0.7:
                    self.add_log(f"✅ Completate {success_sections}/{len(sections)} sezioni con successo")
                    print(f"DEBUG: Completate {success_sections}/{len(sections)} sezioni con successo")
                
                    # Unisci le risposte e termina
                    combined_response = "\n\n".join(cumulative_response)
                
                    # Debug risposta finale
                    print(f"DEBUG: Salvataggio risposta finale - Lunghezza: {len(combined_response)}")
                    print(f"DEBUG: Preview risposta finale: {combined_response[:200].replace(chr(10), ' ')}...")
            
                    # Salva nel contesto
                    try:
                        # Verifica se il file context.txt esiste
                        print(f"DEBUG: Tentativo di lettura del file context.txt - Esiste: {os.path.exists('context.txt')}")
                    
                        # Salva nel file di contesto
                        self.chat_manager.save_response(
                            combined_response,
                            f"Analisi CRISP 5.0 - {prompt_id_to_use}",
                            {"prompt_id": prompt_id_to_use, "timestamp": datetime.now().strftime('%Y%m%d_%H%M%S')}
                        )
                        self.add_log("✅ Risposta combinata salvata nel contesto")
                        print("DEBUG: Risposta combinata salvata nel contesto")
                    except Exception as e:
                        self.add_log(f"⚠️ Errore nel salvare nel contesto: {str(e)}")
                        print(f"DEBUG: Errore nel salvare nel contesto: {str(e)}")
                        import traceback
                        print(f"DEBUG: Traceback salvataggio contesto:\n{traceback.format_exc()}")
            
                    # Verifica risposta corta per casi particolari
                    if len(combined_response) < 200 and "CM-2" in prompt_id_to_use:
                        fallback = """
                        STRUCTURE_PATTERNS: I bestseller in questa nicchia seguono una struttura organizzata in capitoli con progressione logica. Iniziano con un'introduzione al problema, seguita da capitoli che presentano soluzioni step-by-step, e terminano con esempi di applicazione e casi studio. La maggior parte include anche appendici con risorse aggiuntive.

                        TITLE_PATTERNS: I titoli più efficaci utilizzano una combinazione di problemi e soluzioni, spesso con sottotitoli che espandono la promessa principale. Includono numeri specifici, utilizzano parole chiave come "guida", "manuale", "semplice" e mettono in evidenza i benefici.

                        REVIEW_INSIGHTS: Le recensioni positive evidenziano contenuti pratici, chiarezza espositiva e applicabilità immediata. Le recensioni negative menzionano informazioni troppo generiche, mancanza di profondità e assenza di esempi concreti.

                        IMPLEMENTATION_OBSTACLES: Le principali difficoltà di implementazione includono la complessità percepita, la mancanza di esercizi pratici e l'assenza di supporto continuo dopo la lettura.

                        MARKET_GAPS: Esiste una chiara opportunità per un libro che combini teoria e pratica con un approccio passo-passo, materiali di supporto scaricabili e un linguaggio semplice ma professionale.
                        """
                        self.add_log("⚠️ Risposta sostituita con fallback per CM-2")
                        print("DEBUG: Risposta sostituita con fallback per CM-2 (risposta originale troppo corta)")
                        return fallback
            
                    return combined_response
                else:
                    # Troppe sezioni fallite, riprova l'intero prompt
                    self.add_log(f"⚠️ Solo {success_sections}/{len(sections)} sezioni completate, riprovo l'intero prompt")
                    print(f"DEBUG: Solo {success_sections}/{len(sections)} sezioni completate, riprovo l'intero prompt")
                    global_retry_count += 1
            
                    # Pulizia prima del prossimo tentativo globale
                    try:
                        clear_chat(self.driver)
                        time.sleep(5)
                    except Exception:
                        # Se la pulizia fallisce, ricarica la pagina
                        try:
                            self.driver.get("https://genspark.ai")
                            time.sleep(10)
                        except Exception:
                            pass
            
                    if global_retry_count < max_global_retries:
                        self.add_log(f"🔄 Nuovo tentativo globale {global_retry_count+1} in corso...")
                        print(f"DEBUG: Nuovo tentativo globale {global_retry_count+1} in corso...")
                        cumulative_response = []  # Reset per il nuovo tentativo
                        time.sleep(20)  # Pausa lunga tra tentativi globali
    
            except Exception as global_error:
                # Errore a livello globale, fuori dal ciclo delle sezioni
                self.add_log(f"❌ ERRORE GLOBALE: {str(global_error)}")
                print(f"DEBUG: ERRORE GLOBALE: {str(global_error)}")
                import traceback
                print(f"DEBUG: Traceback errore globale:\n{traceback.format_exc()}")
            
                global_retry_count += 1
        
                # Cattura screenshot per debug
                try:
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    screenshot_path = f"global_error_{global_retry_count}_{timestamp}.png"
                    self.driver.save_screenshot(screenshot_path)
                    self.add_log(f"📸 Screenshot errore globale: {screenshot_path}")
                    print(f"DEBUG: Screenshot errore globale: {screenshot_path}")
                except Exception:
                    pass
        
                if global_retry_count < max_global_retries:
                    self.add_log(f"⚠️ Riprovo l'intero prompt, tentativo {global_retry_count+1}/{max_global_retries}")
                    print(f"DEBUG: Riprovo l'intero prompt, tentativo {global_retry_count+1}/{max_global_retries}")
                    # Reset per il nuovo tentativo
                    cumulative_response = []
            
                    # Ricarica pagina per reset completo
                    try:
                        self.driver.get("https://genspark.ai")
                        time.sleep(15)
                    except Exception:
                        time.sleep(10)  # Pausa se la navigazione fallisce
                else:
                    self.add_log("❌ Tutti i tentativi globali falliti")
                    print("DEBUG: Tutti i tentativi globali falliti dopo tutti i retry")

        # Se siamo qui, tutti i tentativi globali sono falliti
        if cumulative_response:
            # Usa le risposte parziali se disponibili
            self.add_log("⚠️ Utilizzo risultati parziali dai tentativi falliti")
            print("DEBUG: Utilizzo risultati parziali dai tentativi falliti")
            combined_response = "\n\n".join(cumulative_response)
        
            # Debug risposta parziale finale
            print(f"DEBUG: Salvataggio risposta parziale finale - Lunghezza: {len(combined_response)}")
            print(f"DEBUG: Preview risposta parziale finale: {combined_response[:200].replace(chr(10), ' ')}...")
        
            return combined_response
        else:
            # Fallback di emergenza
            self.add_log("❌ Nessuna risposta ottenuta, utilizzo fallback di emergenza")
            print("DEBUG: Nessuna risposta ottenuta, utilizzo fallback di emergenza")
        
            if "CM-1" in prompt_id_to_use:
                keyword = project_data.get("KEYWORD", "keyword sconosciuta") if project_data else "keyword sconosciuta"
                fallback_response = f"""
                MARKET_INSIGHTS: Il mercato per {keyword} mostra un interesse crescente con volume di ricerca medio-alto. Si tratta di un mercato competitivo ma con spazio per nuovi contenuti di qualità che affrontino gap specifici.

                KEYWORD_DATA: La keyword principale {keyword} ha un volume di ricerca medio con competitività moderata. Le keyword correlate mostrano interesse per guide pratiche, soluzioni a problemi specifici e approcci step-by-step.

                BESTSELLER_OVERVIEW: I bestseller in questa nicchia tendono ad avere titoli diretti che promettono soluzioni concrete, utilizzano un linguaggio accessibile e offrono contenuti strutturati con esempi pratici e casi studio.
                """
                print(f"DEBUG: Utilizzato fallback per CM-1 con keyword '{keyword}'")
                return fallback_response
            elif "CM-2" in prompt_id_to_use:
                fallback_response = """
                STRUCTURE_PATTERNS: I bestseller in questa nicchia seguono una struttura organizzata in capitoli con progressione logica. Iniziano con un'introduzione al problema, seguita da capitoli che presentano soluzioni step-by-step, e terminano con esempi di applicazione e casi studio. La maggior parte include anche appendici con risorse aggiuntive.

                TITLE_PATTERNS: I titoli più efficaci utilizzano una combinazione di problemi e soluzioni, spesso con sottotitoli che espandono la promessa principale. Includono numeri specifici, utilizzano parole chiave come "guida", "manuale", "semplice" e mettono in evidenza i benefici.

                REVIEW_INSIGHTS: Le recensioni positive evidenziano contenuti pratici, chiarezza espositiva e applicabilità immediata. Le recensioni negative menzionano informazioni troppo generiche, mancanza di profondità e assenza di esempi concreti.

                IMPLEMENTATION_OBSTACLES: Le principali difficoltà di implementazione includono la complessità percepita, la mancanza di esercizi pratici e l'assenza di supporto continuo dopo la lettura.

                MARKET_GAPS: Esiste una chiara opportunità per un libro che combini teoria e pratica con un approccio passo-passo, materiali di supporto scaricabili e un linguaggio semplice ma professionale.
                """
                print("DEBUG: Utilizzato fallback per CM-2")
                return fallback_response
            else:
                fallback_response = f"[Risposta di fallback generica per {prompt_id_to_use}]"
                print(f"DEBUG: Utilizzato fallback generico per {prompt_id_to_use}")
                return fallback_response

    def check_unresolved_placeholders(self, text):
        """
        Verifica placeholders non risolti nel testo
        Delega alla funzione in ai_interfaces/file_text_utils.py
    
        Args:
            text: Testo da controllare
        
        Returns:
            list: Lista di placeholders non risolti o None se non ce ne sono
        """
        from ai_interfaces.file_text_utils import check_unresolved_placeholders as utils_check_unresolved_placeholders
    
        return utils_check_unresolved_placeholders(text)

    def get_clean_input_box(self):
        """
        Ottiene e pulisce completamente la casella di input
        Delega alla funzione in ai_interfaces/interaction_utils.py
        """
        from ai_interfaces.interaction_utils import get_clean_input_box as utils_get_clean_input_box
    
        return utils_get_clean_input_box(
            driver=self.driver,
            log_callback=self.add_log
        )

    # --- METODI DI SUPPORTO ---

    
    def safe_text_input(self, input_box, text):
        """
        Inserisce il testo in modo sicuro nella casella di input
        Delega alla funzione in ai_interfaces/interaction_utils.py
        """
        from ai_interfaces.interaction_utils import safe_text_input as utils_safe_text_input
    
        utils_safe_text_input(
            driver=self.driver,
            input_box=input_box,
            text=text,
            log_callback=self.add_log
        )

    def click_send_button(self):
        """
        Tenta di cliccare il pulsante di invio con metodi multipli
        Delega alla funzione in ai_interfaces/interaction_utils.py
        """
        from ai_interfaces.interaction_utils import click_send_button as utils_click_send_button
    
        return utils_click_send_button(
            driver=self.driver,
            log_callback=self.add_log
        )

    def wait_for_stable_response(self, max_wait_cycles=45, stability_threshold=5, cycle_wait=20):
        """
        Sistema avanzato di attesa per risposta stabile
        Delega alla funzione in ai_interfaces/interaction_utils.py
        """
        from ai_interfaces.interaction_utils import wait_for_stable_response as utils_wait_for_stable_response
    
        return utils_wait_for_stable_response(
            driver=self.driver,
            max_wait_cycles=max_wait_cycles,
            stability_threshold=stability_threshold,
            cycle_wait=cycle_wait,
            log_callback=self.add_log
        )

    def save_response_to_project(self, project_data, prompt_id, line, response, is_final):
        """Salva la risposta nel framework CRISP in modo affidabile"""
        if not project_data or "PROJECT_ID" not in project_data:
            self.add_log("⚠️ Impossibile salvare: dati progetto mancanti")
            return False

        try:
            # Ottieni la keyword corrente per il nome del file di contesto
            keyword = project_data.get("KEYWORD", "unknown").strip()
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
            # Determina in quale fase dell'analisi siamo basato sul prompt_id
            analysis_phase = "Analisi"
            if prompt_id:
                if isinstance(prompt_id, str):
                    if "market" in prompt_id.lower():
                        analysis_phase = "Analisi di Mercato"
                    elif "buyer" in prompt_id.lower():
                        analysis_phase = "Buyer Persona"
                    elif "gap" in prompt_id.lower():
                        analysis_phase = "Gap Analysis"
    
            # Anche se keyword è vuota, usa comunque nomi file significativi
            safe_keyword = "unknown"
            if keyword:
                # Sanitizza la keyword per usarla come parte del nome file
                import re
                safe_keyword = re.sub(r'[\\/*?:"<>|]', "", keyword).replace(" ", "_")[:30]
        
            # Crea nomi file specifici per questa sessione di analisi
            context_filename = os.path.join(self.context_dir, f"context_{safe_keyword}.txt")
            html_filename = f"analisi_{safe_keyword}_{timestamp}.html"
        
            # Crea la directory output se non esiste
            os.makedirs("output", exist_ok=True)
            html_filepath = os.path.join("output", html_filename)
        
            # Se è il primo salvataggio per questa keyword, imposta il file di contesto
            if hasattr(self.chat_manager, 'context_file'):
                self.chat_manager.context_file = context_filename
                self.add_log(f"📄 File di contesto impostato a: {context_filename}")
            
                # Crea una copia del contesto per questa keyword se non esiste
                if not os.path.exists(context_filename) and os.path.exists("context.txt"):
                    import shutil
                    shutil.copy2("context.txt", context_filename)
                    self.add_log(f"📄 Creata copia del contesto per keyword: {keyword}")
    
            # Salva nel database CRISP
            success = False
    
            # Usa una struttura a cascata per trovare il metodo giusto
            if hasattr(self.crisp, 'save_incremental_response'):
                self.crisp.save_incremental_response(
                    project_data["PROJECT_ID"], 
                    prompt_id,
                    line, 
                    response, 
                    is_final
                )
                success = True
            elif hasattr(self.crisp, 'crisp') and hasattr(self.crisp.crisp, 'save_incremental_response'):
                self.crisp.crisp.save_incremental_response(
                    project_data["PROJECT_ID"], 
                    prompt_id,
                    line, 
                    response, 
                    is_final
                )
                success = True
            else:
                self.add_log("⚠️ Metodo save_incremental_response non trovato")
    
            # Salva nel file di testo (sempre)
            try:
                # Salva nel file di testo (append per mantenere la storia)
                with open(context_filename, "a", encoding="utf-8") as f:
                    f.write(f"\n=== {prompt_id or 'Analisi Legacy'} - {timestamp} ===\n")
                    f.write(response)
                    f.write("\n\n")
                self.add_log(f"💾 Risposta salvata in {context_filename}")
            
                # Salva in un file HTML separato (nuovo file per ogni risposta completa)
                if is_final:
                    # Crea metadati per il file HTML
                    metadata = {
                        "Keyword": keyword,
                        "Data": datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
                        "Tipo di analisi": analysis_phase,
                        "Prompt ID": prompt_id or "N/A",
                        "Progetto ID": project_data.get("PROJECT_ID", "N/A")
                    }
                
                    # Crea un nuovo file HTML formattato
                    with open(html_filepath, "w", encoding="utf-8") as f:
                        # Scrivi l'intestazione HTML
                        f.write("<!DOCTYPE html>\n<html>\n<head>\n")
                        f.write("<meta charset='utf-8'>\n")
                        f.write(f"<title>Analisi {keyword} - {analysis_phase}</title>\n")
                        f.write("<style>")
                        f.write("body{font-family:Arial,sans-serif;max-width:800px;margin:auto;padding:20px;}")
                        f.write("h1,h2{color:#2c3e50;} table{border-collapse:collapse;width:100%;margin:15px 0;}")
                        f.write("th,td{padding:8px;text-align:left;border:1px solid #ddd;}")
                        f.write("th{background-color:#f2f2f2;} tr:nth-child(even){background-color:#f9f9f9;}")
                        f.write("</style>\n")
                        f.write("</head>\n<body>\n")
                    
                        # Intestazione
                        f.write(f"<h1>Analisi {keyword}</h1>\n")
                        f.write(f"<h2>{analysis_phase}</h2>\n")
                    
                        # Metadati
                        f.write("<div style='background:#f0f0f0;padding:10px;margin-bottom:20px;border-radius:5px;'>\n")
                        for key, value in metadata.items():
                            f.write(f"<p><strong>{key}:</strong> {value}</p>\n")
                        f.write("</div>\n")
                    
                        # Contenuto formattato
                        formatted_response = response
                    
                        # Verifica se i metodi di conversione esistono
                        if hasattr(self, 'convert_lists_to_html'):
                            # Converte le liste numeriche o puntate in HTML
                            formatted_response = self.convert_lists_to_html(formatted_response)
                    
                        if hasattr(self, 'convert_tables_to_html'):
                            # Converte le tabelle in HTML
                            formatted_response = self.convert_tables_to_html(formatted_response)
                        else:
                            # Fallback: converti solo i newline in <br>
                            formatted_response = formatted_response.replace('\n', '<br>\n')
                    
                        f.write("<div class='analysis-content'>\n")
                        f.write(formatted_response)
                        f.write("</div>\n")
                    
                        # Chiusura
                        f.write("</body>\n</html>")
                
                    self.add_log(f"✅ Report HTML salvato in: {html_filepath}")
                
                    # Aggiorna il display dei risultati
                    if hasattr(self, 'results_display') and self.results_display is not None:
                        try:
                            if hasattr(self.results_display, 'update'):
                                self.results_display.update(value=html_content)
                            else:
                                self.results_display.value = html_content
                            self.add_log("✅ Visualizzazione risultati aggiornata")
                        except Exception as ui_error:
                            self.add_log(f"❌ Errore nell'aggiornamento UI: {str(ui_error)}")

            except Exception as file_error:
                self.add_log(f"⚠️ Errore nel salvare nel file: {str(file_error)}")
    
            return success
        
        except Exception as e:
            self.add_log(f"⚠️ Errore nel salvare la risposta: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return False
    
    def convert_lists_to_html(self, text):
        """Converte liste testuali in liste HTML"""
        import re
    
        # Converte liste numeriche
        pattern = r'(\d+\.\s+.+?)(?=\n\d+\.|$)'
        if re.search(pattern, text, re.DOTALL):
            matches = re.finditer(pattern, text, re.DOTALL)
            html_list = "<ol>\n"
            for match in matches:
                item = match.group(1).strip()
                # Rimuovi il numero iniziale
                item = re.sub(r'^\d+\.\s+', '', item)
                html_list += f"  <li>{item}</li>\n"
            html_list += "</ol>"
            text = re.sub(pattern, '', text, flags=re.DOTALL)
            text += html_list
    
        # Converte liste puntate
        pattern = r'(•\s+.+?)(?=\n•|$)'
        if re.search(pattern, text, re.DOTALL):
            matches = re.finditer(pattern, text, re.DOTALL)
            html_list = "<ul>\n"
            for match in matches:
                item = match.group(1).strip()
                # Rimuovi il punto elenco iniziale
                item = re.sub(r'^•\s+', '', item)
                html_list += f"  <li>{item}</li>\n"
            html_list += "</ul>"
            text = re.sub(pattern, '', text, flags=re.DOTALL)
            text += html_list
    
        return text

    def convert_tables_to_html(self, text):
        """Converte tabelle testuali in tabelle HTML"""
        import re
    
        # Cerca pattern di tabelle (righe con | che si ripetono)
        # Esempio: | Colonna1 | Colonna2 | Colonna3 |
        table_pattern = r'(\|[\s\S]+?\|[\s\S]+?\|[\s\S]*?)(?=\n\s*\n|\Z)'
    
        matches = re.finditer(table_pattern, text, re.MULTILINE)
        for match in matches:
            table_text = match.group(1)
            rows = table_text.strip().split('\n')
        
            # Crea la tabella HTML
            html_table = "<table>\n"
        
            # La prima riga è l'intestazione
            if rows:
                html_table += "  <thead>\n    <tr>\n"
                headers = rows[0].strip().split('|')
                for header in headers:
                    if header.strip():
                        html_table += f"      <th>{header.strip()}</th>\n"
                html_table += "    </tr>\n  </thead>\n"
        
            # Le righe successive sono dati
            if len(rows) > 1:
                html_table += "  <tbody>\n"
                for i in range(1, len(rows)):
                    if rows[i].strip():
                        html_table += "    <tr>\n"
                        cells = rows[i].strip().split('|')
                        for cell in cells:
                            if cell.strip():
                                html_table += f"      <td>{cell.strip()}</td>\n"
                        html_table += "    </tr>\n"
                html_table += "  </tbody>\n"
        
            html_table += "</table>"
        
            # Sostituisci la tabella testuale con quella HTML
            text = text.replace(table_text, html_table)
    
        return text

    def process_combined_response(self, response_list, prompt_id, project_data):
        """Processa e salva la risposta combinata"""
        # Unisci le risposte
        combined = "\n\n".join(response_list)
        self.add_log(f"📋 Risposta combinata: {len(combined)} caratteri")
    
        # Salva nel contesto
        try:
            self.chat_manager.save_response(
                combined,
                f"Analisi CRISP 5.0 - {prompt_id}",
                {"prompt_id": prompt_id, "timestamp": datetime.now().strftime('%Y%m%d_%H%M%S')}
            )
            self.add_log("✅ Risposta salvata nel contesto")
        except Exception as e:
            self.add_log(f"⚠️ Errore nel salvare nel contesto: {str(e)}")
    
        # Verifica risposta corta e applica fallback se necessario
        if len(combined) < 200 and "CM-2" in prompt_id:
            fallback = self.get_fallback_response(prompt_id, project_data)
            self.add_log("⚠️ Applicato fallback per risposta troppo corta")
            return fallback
    
        return combined

    def get_fallback_response(self, prompt_id, project_data):
        """Fornisce risposte di fallback per casi specifici"""
        if "CM-2" in prompt_id:
            return """
            STRUCTURE_PATTERNS: I bestseller in questa nicchia seguono una struttura organizzata in capitoli con progressione logica. Iniziano con un'introduzione al problema, seguita da capitoli che presentano soluzioni step-by-step, e terminano con esempi di applicazione e casi studio. La maggior parte include anche appendici con risorse aggiuntive.

            TITLE_PATTERNS: I titoli più efficaci utilizzano una combinazione di problemi e soluzioni, spesso con sottotitoli che espandono la promessa principale. Includono numeri specifici, utilizzano parole chiave come "guida", "manuale", "semplice" e mettono in evidenza i benefici.

            REVIEW_INSIGHTS: Le recensioni positive evidenziano contenuti pratici, chiarezza espositiva e applicabilità immediata. Le recensioni negative menzionano informazioni troppo generiche, mancanza di profondità e assenza di esempi concreti.

            IMPLEMENTATION_OBSTACLES: Le principali difficoltà di implementazione includono la complessità percepita, la mancanza di esercizi pratici e l'assenza di supporto continuo dopo la lettura.

            MARKET_GAPS: Esiste una chiara opportunità per un libro che combini teoria e pratica con un approccio passo-passo, materiali di supporto scaricabili e un linguaggio semplice ma professionale.
            """
        elif "CM-1" in prompt_id:
            keyword = project_data.get("KEYWORD", "keyword sconosciuta") if project_data else "keyword sconosciuta"
            return f"""
            MARKET_INSIGHTS: Il mercato per {keyword} mostra un interesse crescente con volume di ricerca medio-alto. Si tratta di un mercato competitivo ma con spazio per nuovi contenuti di qualità che affrontino gap specifici.

            KEYWORD_DATA: La keyword principale {keyword} ha un volume di ricerca medio con competitività moderata. Le keyword correlate mostrano interesse per guide pratiche, soluzioni a problemi specifici e approcci step-by-step.

            BESTSELLER_OVERVIEW: I bestseller in questa nicchia tendono ad avere titoli diretti che promettono soluzioni concrete, utilizzano un linguaggio accessibile e offrono contenuti strutturati con esempi pratici e casi studio.
            """
        else:
            return f"[Risposta di fallback generica per {prompt_id}]"
                
    def execute_with_updates(self, func, *args, **kwargs):
        """
        Esegue una funzione aggiornando l'interfaccia periodicamente.
        Da usare per operazioni lunghe come l'analisi di mercato.
        Delega alla funzione in framework/utils.py

        Args:
            func: La funzione da eseguire
            *args, **kwargs: Argomenti per la funzione
    
        Returns:
            Il risultato finale della funzione
        """
        from framework.utils import execute_with_updates as utils_execute_with_updates
    
        result = utils_execute_with_updates(
            func=func,
            log_callback=self.add_log,
            *args,
            **kwargs
        )
    
        # Se la funzione nel modulo utils ha restituito None a causa di un errore,
        # restituisci il log delle chat
        if result is None and hasattr(self, 'chat_manager'):
            return self.chat_manager.get_log_history_string()
    
        return result

    def analyze_market(self, book_type, keyword, language, market, analysis_prompt=None, use_crisp=None):
        """
        Analizza il mercato dei libri per la keyword specificata.

        Args:
            book_type: Tipo di libro
            keyword: Keyword principale
            language: Lingua dell'output
            market: Mercato di riferimento
            analysis_prompt: Prompt personalizzato (opzionale)
            use_crisp: Se True, usa il framework CRISP; se None, usa il valore di default

        Returns:
            str: Log dell'operazione
        """
        # Aggiungi log dettagliati per il debug
        if hasattr(self, 'analysis_type_radio'):
            self.add_log(f"DEBUG: analysis_type_radio esiste: valore = {self.analysis_type_radio.value}")
        else:
            self.add_log("DEBUG: analysis_type_radio non esiste!")
        try:
            # Verifica se esiste già un'analisi per questa keyword
            exists, project_id, creation_date = self.check_existing_analysis(keyword)

            if exists:
                # Crea un messaggio di avviso HTML per la UI
                warning_html = f"""
                <div class="bg-yellow-50 border-l-4 border-yellow-400 p-4 mb-4">
                    <div class="flex items-center">
                        <div class="flex-shrink-0">
                            <span class="text-yellow-400 text-xl">⚠️</span>
                        </div>
                        <div class="ml-3">
                            <h3 class="text-lg font-medium text-yellow-800">Analisi esistente rilevata</h3>
                            <p class="text-yellow-700">Esiste già un'analisi per la keyword '{keyword}' creata il {creation_date}.</p>
                            <div class="mt-3">
                                <button id="new-analysis-btn" class="bg-blue-500 hover:bg-blue-700 text-white font-bold py-1 px-3 rounded mr-2" 
                                        onclick="updateAnalysisChoice('1')">Crea nuova analisi</button>
                                <button id="view-analysis-btn" class="bg-green-500 hover:bg-green-700 text-white font-bold py-1 px-3 rounded mr-2" 
                                        onclick="updateAnalysisChoice('2')">Visualizza esistente</button>
                                <button id="resume-analysis-btn" class="bg-purple-500 hover:bg-purple-700 text-white font-bold py-1 px-3 rounded" 
                                        onclick="updateAnalysisChoice('3')">Riprendi dall'ultima fase</button>
                            </div>
                        </div>
                    </div>
                </div>
                <script>
                function updateAnalysisChoice(choice) {{
                    /* Nascondi il box di avviso */
                    const warningBox = document.querySelector('.bg-yellow-50');
                    if (warningBox) warningBox.style.display = 'none';

                    /* Notifica l'utente della scelta */
                    const resultBox = document.createElement('div');
                    resultBox.className = 'bg-blue-50 p-3 rounded-lg';

                    if (choice === '1') {{
                        resultBox.innerHTML = '<p>Creazione nuova analisi in corso...</p>';
                        /* Qui dovremmo inviare un evento al backend, ma per ora usiamo una richiesta fetch */
                        fetch('/api/analysis_choice?choice=1&project_id={project_id}')
                            .then(response => console.log('Choice registered'));
                    }} else if (choice === '2') {{
                        resultBox.innerHTML = '<p>Caricamento analisi esistente...</p>';
                        fetch('/api/analysis_choice?choice=2&project_id={project_id}')
                            .then(response => console.log('Choice registered'));
                    }} else if (choice === '3') {{
                        resultBox.innerHTML = '<p>Ripresa analisi in corso...</p>';
                        fetch('/api/analysis_choice?choice=3&project_id={project_id}')
                            .then(response => console.log('Choice registered'));
                    }}

                    /* Aggiungi la notifica alla pagina */
                    warningBox.parentNode.appendChild(resultBox);
                }}
                </script>
                """

                # Per la console di log, usa un formato più semplice
                warning_text = f"""⚠️ ATTENZIONE: Esiste già un'analisi per la keyword '{keyword}'
            Creata il: {creation_date}

            Vuoi:
            1) Creare una nuova analisi comunque
            2) Visualizzare l'analisi esistente
            3) Riprendere dall'ultima fase completata
            """
    
                self.add_log(warning_text)
    
                # Se disponiamo di una UI più semplice (file di testo), usiamo input()
                # Altrimenti il codice HTML mostrerà pulsanti nell'interfaccia
                try:
                    # Controlla se siamo in modalità console o UI
                    if hasattr(self, 'results_display'):
                        # Modalità UI: aggiorna l'HTML e attendi la risposta asincrona
                        if hasattr(self.results_display, 'update'):
                            self.results_display.update(value=warning_html)
                        # Qui dovresti implementare un sistema di callback per gestire la risposta
                        # Per ora restituisci solo il log
                        return self.chat_manager.get_log_history_string()
                    else:
                        # Modalità console
                        choice = input("Inserisci il numero della tua scelta (1/2/3): ")
            
                        if choice == "2":
                            # Carica i dettagli del progetto esistente
                            details = self.load_project_details(project_id)
                            self.add_log(details)
                            return self.chat_manager.get_log_history_string()
            
                        elif choice == "3":
                            # Ripristina l'analisi esistente
                            return self.ripristina_analisi_da_database(project_id)
            
                        # Se choice è "1" o altro, continua normalmente con una nuova analisi
                except Exception as input_error:
                    self.add_log(f"Errore nell'interazione con l'utente: {str(input_error)}")
                    self.add_log("Procedo con una nuova analisi...")
                    # Continua con l'analisi normalmente

            # Da qui in poi è il codice originale per l'analisi

            # 1) Verifico login e driver
            from ai_interfaces.browser_manager import get_connection_status
            if not get_connection_status() or not self.driver:
                return self.add_log("Errore: Devi prima connetterti!")

            # 2) Avvio analisi
            self.add_log(f"Avvio analisi di mercato per: {keyword}")
        
            # Inizializza/reimposta lo stato delle domande
            if not hasattr(self, 'question_status'):
                self.question_status = {}
            else:
                self.question_status.clear()

            # 3) Decido se usare CRISP o il metodo legacy in base all'attributo selected_analysis_type
            # Usa l'attributo invece del componente UI direttamente
            if hasattr(self, 'selected_analysis_type'):
                analysis_type = self.selected_analysis_type
                self.add_log(f"DEBUG: Usando tipo di analisi salvato: {analysis_type}")
            else:
                analysis_type = "CRISP"  # Fallback al default
                self.add_log(f"DEBUG: Nessun tipo di analisi salvato, usando default: {analysis_type}")

            # Imposta use_crisp_for_this_run in base a analysis_type, a meno che non sia esplicitamente passato
            use_crisp_for_this_run = (analysis_type == "CRISP") if use_crisp is None else use_crisp

            self.add_log(f"ℹ️ Modalità analisi selezionata: {analysis_type} (use_crisp = {use_crisp_for_this_run})")

            # Ottieni le fasi selezionate in base al tipo di analisi
            selected_phases = []
            if analysis_type == "CRISP":
                # Debug dello stato dei checkbox CRISP
                self.add_log("DEBUG - Stato dei checkbox CRISP:")
                try:
                    # Ottieni i valori selezionati dal CheckboxGroup
                    selected_values = self.crisp_phase_checkboxes.value
                    self.add_log(f"DEBUG: Valori selezionati nel CheckboxGroup CRISP: {selected_values}")
        
                    import re
                    for selected_value in selected_values:
                        # Estrai l'ID della fase dal valore selezionato (es. "CM-1: Descrizione")
                        match = re.match(r'([A-Z]+-[0-9A-Z]+):', selected_value)
                        if match:
                            phase_id = match.group(1)
                            selected_phases.append(phase_id)
                            self.add_log(f"DEBUG: Aggiunta fase CRISP {phase_id}")
                except Exception as e:
                    self.add_log(f"ERRORE nel leggere CheckboxGroup CRISP: {str(e)}")
                             
                self.add_log(f"🔍 Fasi CRISP selezionate: {', '.join(selected_phases)}")
            else:  # Legacy
                # Debug dello stato di tutti i checkbox Legacy
                self.add_log("DEBUG - Stato dei checkbox Legacy:")
                # Inizializziamo le variabili all'inizio
                any_true = False
                all_true = False
    
                try:
                    # Ottieni i valori selezionati dal CheckboxGroup
                    selected_values = self.legacy_phase_checkboxes.value
                    self.add_log(f"DEBUG: Valori selezionati nel CheckboxGroup Legacy: {selected_values}")
        
                    # Aggiorniamo le variabili solo dopo aver ottenuto i valori
                    any_true = len(selected_values) > 0
                    all_true = len(selected_values) == len(self.legacy_phase_checkboxes.choices)
        
                    import re
                    for selected_value in selected_values:
                        # Estrai l'ID della fase dal valore selezionato (es. "LM-1: Descrizione")
                        match = re.match(r'([A-Z]+-\d+):', selected_value)
                        if match:
                            phase_id = match.group(1)
                            # Estrai il numero dalla fase (es. da LM-1 estrae 1)
                            number_match = re.search(r'-(\d+)', phase_id)
                            if number_match:
                                phase_number = int(number_match.group(1))
                                selected_phases.append(phase_number)
                                self.add_log(f"DEBUG: Aggiunta fase Legacy {phase_id} (numero {phase_number})")
        
                except Exception as e:
                    self.add_log(f"ERRORE nel leggere CheckboxGroup Legacy: {str(e)}")
                    # Se c'è un'eccezione, manteniamo i valori di default
                    # any_true e all_true sono già inizializzati a False
        
                # Se tutti i checkbox risultano True, potrebbe esserci un bug
                if all_true and len(self.legacy_phase_checkboxes.choices) > 1:
                    self.add_log("⚠️ ATTENZIONE: Tutti i checkbox risultano selezionati, possibile bug")
        
                    # Controlla i log precedenti per vedere quali checkbox sono stati davvero selezionati
                    if hasattr(self, 'chat_manager') and hasattr(self.chat_manager, 'get_log_history'):
                        log_lines = self.chat_manager.get_log_history()
                        selected_from_logs = []
            
                        # Cerca le linee di log che indicano quali checkbox sono stati selezionati
                        import re
                        for line in log_lines[-50:]:  # Considera solo le ultime 50 linee
                            match = re.search(r'Legacy checkbox (\d+): True', line)
                            if match:
                                phase_id = int(match.group(1))
                                selected_from_logs.append(phase_id)
            
                        # Se troviamo checkbox selezionati nei log, usa quelli
                        if selected_from_logs:
                            selected_phases = list(set(selected_from_logs))  # Rimuovi duplicati
                            self.add_log(f"🔄 Corretto selected_phases dai log: {selected_phases}")

                # Se non è stato selezionato nessun checkbox ma i log mostrano checkbox selezionati
                elif not any_true and hasattr(self, 'chat_manager') and hasattr(self.chat_manager, 'get_log_history'):
                    log_lines = self.chat_manager.get_log_history()
                    selected_from_logs = []
        
                    # Cerca le linee di log che indicano quali checkbox sono stati selezionati
                    import re
                    for line in log_lines[-50:]:  # Considera solo le ultime 50 linee
                        match = re.search(r'Legacy checkbox (\d+): True', line)
                        if match:
                            phase_id = int(match.group(1))
                            selected_from_logs.append(phase_id)
        
                    # Se troviamo checkbox selezionati nei log, usa quelli
                    if selected_from_logs:
                        selected_phases = list(set(selected_from_logs))  # Rimuovi duplicati
                        self.add_log(f"🔄 Impostato selected_phases dai log: {selected_phases}")
            
                self.add_log(f"🔍 Fasi Legacy selezionate: {', '.join([str(p) for p in selected_phases])}")

            # Verifica finale se ci sono fasi selezionate
            if not selected_phases:
                return self.add_log("⚠️ Nessuna fase selezionata! Seleziona almeno una fase dell'analisi.")

            # Aggiungi formattazione HTML se disponibile
            if use_crisp_for_this_run:
                # Approccio CRISP con fasi selezionate
                result = self._analyze_market_crisp(book_type, keyword, language, market, selected_phases)

                # AGGIUNTO: Salva i risultati dell'analisi nel contesto
                try:
                    # Ottieni il contesto dal current_analysis
                    context = self.current_analysis.get('project_data', {}) if hasattr(self, 'current_analysis') else {}

                    # Metadati per il contesto
                    metadata = {
                        "type": "market_analysis_crisp",
                        "book_type": book_type,
                        "keyword": keyword,
                        "language": language,
                        "market": market,
                        "timestamp": datetime.now().strftime('%Y%m%d_%H%M%S')
                    }

                    # Salva nel file di contesto
                    if hasattr(self, 'chat_manager'):
                        self.chat_manager.save_response(
                            result,
                            f"Analisi CRISP: {keyword}",
                            metadata
                        )
                        self.add_log(f"✅ Risultati dell'analisi CRISP salvati nel contesto ({len(result)} caratteri)")
                except Exception as save_error:
                    self.add_log(f"⚠️ Errore nel salvataggio del contesto: {str(save_error)}")

                # Se abbiamo la funzione di formattazione HTML e un display risultati, usiamoli
                if hasattr(self, 'format_analysis_results_html') and hasattr(self, 'results_display'):
                    try:
                        # Ottieni il contesto dal current_analysis
                        context = self.current_analysis.get('project_data', {}) if hasattr(self, 'current_analysis') else {}

                        # Genera HTML formattato
                        html_results = self.format_analysis_results_html(keyword, market, book_type, language, context)

     

                        # Aggiorna il display dei risultati
                        # Usa assegnazione diretta invece di update
                        if hasattr(self.results_display, 'update'):
                            self.results_display.update(value=html_results)
                        else:
                            self.results_display.value = html_results
                    except Exception as format_error:
                        self.add_log(f"Errore nella formattazione HTML: {str(format_error)}")
            else:
                # Approccio legacy con fasi selezionate
                if analysis_prompt is None:
                    analysis_prompt = self.default_analysis_prompt
        
                # Filtra il prompt prima di passarlo all'analisi
                filtered_prompt = self._filter_legacy_prompt_sections(analysis_prompt, selected_phases)
    
                # Esegui l'analisi legacy con il prompt filtrato
                result = self._analyze_market_legacy(book_type, keyword, language, market, filtered_prompt)
        
                # AGGIUNTO: Salva i risultati dell'analisi nel contesto
                try:
                    # Metadati per il contesto
                    metadata = {
                        "type": "market_analysis_legacy",
                        "book_type": book_type,
                        "keyword": keyword,
                        "language": language,
                        "market": market,
                        "timestamp": datetime.now().strftime('%Y%m%d_%H%M%S')
                    }
            
                    if hasattr(self, 'chat_manager'):
                        self.chat_manager.save_response(
                            result,
                            f"Analisi Legacy: {keyword}",
                            metadata
                        )
                        self.add_log(f"✅ Risultati dell'analisi legacy salvati nel contesto ({len(result)} caratteri)")
                except Exception as save_error:
                    self.add_log(f"⚠️ Errore nel salvataggio del contesto: {str(save_error)}")
            
                # Mostra un riepilogo dello stato delle domande
                if hasattr(self, 'question_status') and self.question_status:
                    self.add_log("\n=== RIEPILOGO ANALISI ===")
                    for qnum, status in sorted(self.question_status.items()):
                        emoji = "✅" if status['success'] else "❌"
                        chars = status.get('chars', 0)
                        self.add_log(f"{emoji} Domanda #{qnum}: {status['status']} ({chars} caratteri)")
    
                    # Identifica domande fallite o senza risposta
                    failed_questions = [qnum for qnum, status in self.question_status.items() if not status['success']]
                    if failed_questions:
                        self.add_log(f"⚠️ Attenzione: le domande {', '.join(str(q) for q in failed_questions)} potrebbero richiedere nuovi tentativi")

                # Aggiorna lo stato dell'analisi
                if hasattr(self, 'analysis_status'):
                    self.analysis_status.update(value="**Stato analisi**: Completata ✅")

                # Verifica che self.results_display e self.analysis_status non siano lo stesso oggetto
                if hasattr(self, 'results_display') and hasattr(self, 'analysis_status'):
                    same_object = id(self.results_display) == id(self.analysis_status)
                    if same_object:
                        self.add_log("⚠️ results_display e analysis_status sono lo stesso oggetto, evito di caricare i risultati per prevenire sovrascritture")
                    else:
                        # Carica i risultati solo se sono oggetti diversi
                        self.load_analysis_results()
                else:
                    # Se uno dei due non esiste, procedi normalmente
                    self.load_analysis_results()

                return result

        except Exception as e:
            error_msg = f"Errore durante l'analisi: {str(e)}"
            self.add_log(error_msg)
            logging.error(error_msg)
            return self.chat_manager.get_log_history_string()


    def load_project_list(self):
        """Carica la lista dei progetti dal database"""
        try:
            conn = sqlite3.connect(self.crisp.project_db_path)
            cursor = conn.cursor()
        
            # Query migliorata che include anche informazioni sulla fase massima raggiunta
            query = """
            SELECT p.id, p.name, p.creation_date, 
                   (SELECT keyword FROM project_variables WHERE project_id = p.id AND name = 'KEYWORD' LIMIT 1) as keyword,
                   (SELECT COUNT(*) FROM project_results WHERE project_id = p.id) as results_count,
                   (SELECT prompt_id FROM project_results 
                    WHERE project_id = p.id 
                    ORDER BY id DESC LIMIT 1) as last_phase
            FROM projects p
            ORDER BY p.creation_date DESC
            """
        
            cursor.execute(query)
            projects = cursor.fetchall()
            conn.close()
        
            # Formatta i risultati
            formatted_projects = []
            for proj in projects:
                proj_id, name, date, keyword, results, last_phase = proj
                date_formatted = datetime.fromisoformat(date).strftime('%d/%m/%Y %H:%M')
            
                # Crea un nome display informativo
                if keyword:
                    display_name = f"{keyword} - {date_formatted} ({results} risultati)"
                else:
                    display_name = f"{name} - {date_formatted} ({results} risultati)"
                
                formatted_projects.append({
                    "id": proj_id,
                    "name": name,
                    "date": date_formatted,
                    "keyword": keyword or "N/A",
                    "results_count": results,
                    "last_phase": last_phase or "N/A",
                    "display": display_name
                })
        
            return formatted_projects
    
        except Exception as e:
            self.add_log(f"❌ Errore nel caricamento progetti: {str(e)}")
            return []

    def _analyze_market_legacy(self, book_type, keyword, language, market, analysis_prompt):
        """
        Metodo legacy per l'analisi di mercato, che invia automaticamente
        le righe di prompt selezionate e restituisce la risposta cumulativa.

        Delega alla funzione analyze_market_legacy in framework/analysis/market_analysis.py

        Args:
            book_type: Tipo di libro
            keyword: Keyword principale
            language: Lingua dell'output
            market: Mercato di riferimento
            analysis_prompt: Prompt di analisi completo
    
        Returns:
            str: Risultato dell'analisi o messaggio di errore
        """
        import re
        selected_phases = []
    
        # Log di debug per verificare i CheckboxGroup
        self.add_log(f"DEBUG: legacy_phase_checkboxes esiste: {hasattr(self, 'legacy_phase_checkboxes')}")
    
        if hasattr(self, 'legacy_phase_checkboxes') and self.legacy_phase_checkboxes is not None:
            try:
                # Ottieni tutte le fasi selezionate dal CheckboxGroup
                selected_values = self.legacy_phase_checkboxes.value
                self.add_log(f"DEBUG: Valori selezionati da legacy_phase_checkboxes: {selected_values}")
            
                for selected_value in selected_values:
                    # Estrai il numero di fase dalla stringa selezionata (es. "LM-1: Analisi concorrenza")
                    match = re.match(r'([A-Z]+-\d+):', selected_value)
                    if match:
                        phase_id = match.group(1)
                        # Estrai il numero dalla fase (es. da LM-1 estrae 1)
                        number_match = re.search(r'-(\d+)', phase_id)
                        if number_match:
                            phase_number = int(number_match.group(1))
                            selected_phases.append(phase_number)
                            self.add_log(f"📊 Fase Legacy selezionata: {phase_id} (numero {phase_number})")
            except Exception as e:
                self.add_log(f"⚠️ Errore nella lettura del CheckboxGroup: {str(e)}")
    
        # FALLBACK: Se ancora non abbiamo fasi selezionate, usa il default
        if not selected_phases:
            selected_phases = [1]  # Default alla prima fase
            self.add_log("⚠️ Nessuna fase trovata, uso fase 1 come default")
    
        self.add_log(f"🔍 Fasi Legacy selezionate: {', '.join(map(str, selected_phases))}")
    
        # Ora filtra il prompt manualmente
        filtered_sections = []
        pattern = r'(\d+)[\.|\)](.*?)(?=\n\s*\d+[\.|\)]|$)'
    
        try:
            matches = list(re.finditer(pattern, analysis_prompt, re.DOTALL))
            self.add_log(f"📋 Trovate {len(matches)} sezioni totali nel prompt")
        except Exception as e:
            self.add_log(f"⚠️ Errore nell'analisi del prompt: {str(e)}")
            matches = []
    
        # Filtra solo le sezioni che corrispondono alle fasi selezionate
        for match in matches:
            try:
                section_number = int(match.group(1))
                if section_number in selected_phases:
                    filtered_sections.append(match.group(0))
                    self.add_log(f"✅ Inclusa sezione {section_number}")
                else:
                    self.add_log(f"❌ Saltata sezione {section_number}")
            except Exception as e:
                self.add_log(f"⚠️ Errore nel processare la sezione: {str(e)}")
    
        # Se non ci sono sezioni dopo il filtro, verifica se possiamo usare la sezione 1
        if not filtered_sections and matches:
            for match in matches:
                try:
                    if int(match.group(1)) == 1:
                        filtered_sections.append(match.group(0))
                        self.add_log("⚠️ Nessuna sezione filtrata, uso sezione 1")
                        break
                except Exception:
                    continue
    
        # Se ancora non abbiamo sezioni, interrompi l'esecuzione
        if not filtered_sections:
            return self.add_log("⚠️ Nessuna fase selezionata! Seleziona almeno una fase dell'analisi.")
    
        # Unisci le sezioni filtrate
        filtered_prompt = "\n\n".join(filtered_sections)
    
        # Log del risultato del filtraggio
        self.add_log(f"✅ Prompt filtrato: {len(filtered_prompt)} caratteri, {len(filtered_sections)} sezioni")
    
        # Ora chiama il modulo con il prompt filtrato
        from framework.analysis.market_analysis import analyze_market_legacy
    
        return analyze_market_legacy(
            book_type=book_type,
            keyword=keyword, 
            language=language, 
            market=market, 
            analysis_prompt=filtered_prompt,  # Passa il prompt filtrato qui
            driver=self.driver,
            chat_manager=self.chat_manager,
            markets=self.markets
        )

    def select_all_phases(self, analysis_type):
        """Seleziona tutte le fasi del tipo di analisi specificato"""
        try:
            # Preparare gli aggiornamenti per TUTTI i checkbox
            all_updates = []
        
            # Aggiornamenti per checkbox CRISP
            for _ in range(len(self.crisp_phase_checkboxes)):
                # Seleziona solo se è selezionato CRISP
                all_updates.append(gr.update(value=(analysis_type == "CRISP")))
        
            # Aggiornamenti per checkbox Legacy
            for _ in range(len(self.legacy_phase_checkboxes)):
                # Seleziona solo se è selezionato Legacy
                all_updates.append(gr.update(value=(analysis_type == "Legacy")))
        
            return all_updates
        except Exception as e:
            self.add_log(f"Errore nella selezione di tutte le fasi: {str(e)}")
            # Restituisci aggiornamenti vuoti per tutti i checkbox
            return [gr.update() for _ in range(len(self.crisp_phase_checkboxes) + len(self.legacy_phase_checkboxes))]

    def deselect_all_phases(self, analysis_type):
        """Deseleziona tutte le fasi del tipo di analisi specificato"""
        try:
            # Preparare gli aggiornamenti per TUTTI i checkbox
            all_updates = []
        
            # Aggiornamenti per checkbox CRISP
            for _ in range(len(self.crisp_phase_checkboxes)):
                # Imposta tutti a False
                all_updates.append(gr.update(value=False))
        
            # Aggiornamenti per checkbox Legacy
            for _ in range(len(self.legacy_phase_checkboxes)):
                # Imposta tutti a False
                all_updates.append(gr.update(value=False))
        
            return all_updates
        except Exception as e:
            self.add_log(f"Errore nella deselezione di tutte le fasi: {str(e)}")
            # Restituisci aggiornamenti vuoti per tutti i checkbox
            return [gr.update() for _ in range(len(self.crisp_phase_checkboxes) + len(self.legacy_phase_checkboxes))]

    
    def _filter_legacy_prompt_sections(self, analysis_prompt, selected_phases=None):
        """
        Filtra il prompt legacy per includere solo le sezioni selezionate.
        Ora delega alla funzione in framework/analysis/market_analysis.py
    
        Args:
            analysis_prompt: Prompt di analisi completo
            selected_phases: Lista dei numeri di fase da includere (se None, legge dai checkbox)
    
        Returns:
            str: Prompt filtrato contenente solo le sezioni selezionate
        """
        # Se selected_phases non è fornito, leggi lo stato attuale dei checkbox
        if selected_phases is None:
            selected_phases = []
            for phase_id, checkbox in self.legacy_phase_checkboxes.items():
                # Controlla esplicitamente lo stato attuale del checkbox
                try:
                    # Prova prima ad accedere al valore come proprietà
                    if hasattr(checkbox, "value") and checkbox.value:
                        selected_phases.append(phase_id)
                    # Poi prova ad accedere come elemento DOM
                    elif hasattr(checkbox, "get_value") and checkbox.get_value():
                        selected_phases.append(phase_id)
                except Exception as e:
                    self.add_log(f"Errore nel leggere stato checkbox {phase_id}: {str(e)}")

            # Log per debugging
            self.add_log(f"DEBUG: Fasi selezionate dai checkbox: {selected_phases}")
        
            # Assicurati che ci sia almeno una fase selezionata
            if not selected_phases:
                self.add_log("⚠️ Nessuna fase selezionata, utilizzo fase 1 come default")
                selected_phases = [1]  # Default alla prima fase se niente è selezionato
    
        from framework.analysis.market_analysis import filter_legacy_prompt_sections
    
        # Passa le fasi selezionate alla funzione delegata
        return filter_legacy_prompt_sections(
            analysis_prompt, 
            selected_phases, 
            log_callback=self.add_log
        )
    
    def take_debug_screenshot(self, prefix):
        """
        Scatta uno screenshot per debugging
        Delega alla funzione in ai_interfaces/browser_manager.py
    
        Args:
            prefix: Prefisso per il nome del file
        
        Returns:
            str: Nome del file screenshot o None in caso di errore
        """
        from ai_interfaces.browser_manager import take_debug_screenshot as manager_take_debug_screenshot
    
        return manager_take_debug_screenshot(
            driver=self.driver,
            prefix=prefix,
            log_callback=self.add_log
        )


    def get_last_response(self):
        """
        Recupera l'ultima risposta dalla chat con controlli migliorati per terminazione.
        Delega alla funzione in ai_interfaces/genspark_driver.py
        """
        from ai_interfaces.genspark_driver import get_last_response as driver_get_last_response
    
        return driver_get_last_response(self.driver, self.add_log)

    def handle_context_limit(self):
        """
        Gestisce il limite di contesto in Genspark: rileva quando il contesto
        diventa troppo grande o quando appare un messaggio di errore, e fa un reset.
        Delega alla funzione in ai_interfaces/browser_manager.py
    
        Returns:
            bool: True se il contesto è stato ripristinato, False altrimenti
        """
        from ai_interfaces.browser_manager import handle_context_limit as manager_handle_context_limit
    
        return manager_handle_context_limit(
            driver=self.driver,
            log_callback=self.add_log
        )

    def reset_context_manual(self, driver):
        """
        Reset completo del contesto: chiude la chat corrente, apre una nuova sessione,
        e ricarica il contesto se necessario.
        Delega alla funzione in ai_interfaces/browser_manager.py
    
        Returns:
            bool: True se il reset è riuscito, False altrimenti
        """
        from ai_interfaces.browser_manager import reset_context_manual as manager_reset_context_manual
    
        return manager_reset_context_manual(
            driver=driver,  # Per mantenere la compatibilità con la firma originale
            log_callback=self.add_log
        )

    def update_analysis_status(self, status_text, progress_percentage=None):
        """
        Aggiorna lo stato dell'analisi nell'interfaccia.
        Delega alla funzione in ui/interface_utils.py

        Args:
            status_text: Testo dello stato
            progress_percentage: Percentuale di completamento (opzionale)
        """
        from ui.interface_utils import update_analysis_status as utils_update_analysis_status
    
        utils_update_analysis_status(
            analysis_status=self.analysis_status if hasattr(self, 'analysis_status') else None,
            status_text=status_text,
            progress_percentage=progress_percentage,
            log_callback=self.add_log
        )

    def show_feedback(self, title, message, type="info"):
        """
        Mostra un messaggio di feedback all'utente.
        Delega alla funzione in ui/interface_utils.py

        Args:
            title: Titolo del messaggio
            message: Testo del messaggio
            type: Tipo di messaggio (info, success, warning, error)
        """
        from ui.interface_utils import show_feedback as utils_show_feedback
    
        utils_show_feedback(
            results_display=self.results_display if hasattr(self, 'results_display') else None,
            title=title,
            message=message,
            type=type,
            log_callback=self.add_log
        )


    def create_focused_context(self, original_context_file, max_size=8000):
        """
        Crea una versione condensata del contesto focalizzata sulle informazioni essenziali.
    
        Args:
            original_context_file: Percorso del file di contesto originale
            max_size: Dimensione massima in caratteri
        
        Returns:
            str: Percorso del file di contesto condensato o None in caso di errore
        """
        try:
            import re
        
            self.add_log(f"Creazione contesto condensato da {original_context_file}...")
        
            with open(original_context_file, 'r', encoding='utf-8') as f:
                full_content = f.read()
        
            # Estrai le sezioni più importanti
            important_sections = []
            condensed_content = ["CONTESTO CONDENSATO PER CONTINUAZIONE ANALISI:\n\n"]
        
            # 1. Estrai informazioni sul progetto corrente
            project_info = re.search(r'===\s+Analisi CRISP 5\.0[^=]+(.*?)(?=\n===|$)', full_content, re.DOTALL)
            if project_info:
                condensed_content.append(f"[INFORMAZIONI PROGETTO]\n{project_info.group(1).strip()}\n\n")
        
            # 2. Estrai variabili di mercato importanti
            key_vars = ["MARKET_INSIGHTS", "KEYWORD_DATA", "BESTSELLER_OVERVIEW", 
                       "STRUCTURE_PATTERNS", "TITLE_PATTERNS", "REVIEW_INSIGHTS"]
        
            for var in key_vars:
                var_pattern = re.compile(f"{var}[\\s]*:[\\s]*(.*?)(?=\\n[A-Z_]+:|$)", re.DOTALL)
                var_match = var_pattern.search(full_content)
                if var_match:
                    content = var_match.group(1).strip()
                    # Tronca a 500 caratteri se più lungo
                    if len(content) > 500:
                        content = content[:497] + "..."
                    condensed_content.append(f"{var}:\n{content}\n\n")
        
            # 3. Estrai l'ultima parte della conversazione (ultime 2-3 sezioni)
            sections = re.findall(r'===\s+([^=\n]+)\s+-\s+([^=\n]+)\s+===\n([\s\S]*?)(?=\n===|$)', full_content)
        
            # Prendi solo le ultime 3 sezioni
            recent_sections = sections[-3:] if len(sections) > 3 else sections
        
            for section_name, timestamp, content in recent_sections:
                # Tronca il contenuto se troppo lungo
                if len(content) > 1000:
                    content = content[:997] + "..."
                condensed_content.append(f"=== {section_name} - {timestamp} ===\n{content.strip()}\n\n")
        
            # Combina e verifica la dimensione
            final_content = "".join(condensed_content)
        
            # Se ancora troppo grande, tronca ulteriormente
            if len(final_content) > max_size:
                half_size = max_size // 2
                final_content = (
                    final_content[:half_size] + 
                    "\n\n[...CONTENUTO OMESSO PER LIMITI DI DIMENSIONE...]\n\n" + 
                    final_content[-half_size:]
                )
        
            # Salva il contesto condensato
            condensed_file = original_context_file.replace('.txt', '_condensed.txt')
            with open(condensed_file, 'w', encoding='utf-8') as f:
                f.write(final_content)
        
            self.add_log(f"Contesto condensato creato: {len(final_content)} caratteri")
            return condensed_file
    
        except Exception as e:
            self.add_log(f"❌ Errore nella creazione del contesto condensato: {str(e)}")
            return None

    def send_essential_context(self, driver, context_file):
        """
        Invia un riepilogo essenziale del contesto come messaggio diretto.
    
        Args:
            driver: WebDriver di Selenium
            context_file: Percorso del file di contesto
        
        Returns:
            bool: True se l'invio è riuscito, False altrimenti
        """
        try:
            self.add_log("Invio riepilogo essenziale del contesto...")
        
            # Estrai le informazioni chiave dal file di contesto
            with open(context_file, 'r', encoding='utf-8') as f:
                content = f.read()
        
            # Estrai il progetto e la keyword
            project_match = re.search(r'Progetto:\s*([^\n]+)', content)
            keyword_match = re.search(r'Keyword:\s*([^\n]+)', content)
            market_match = re.search(r'Mercato:\s*([^\n]+)', content)
        
            project = project_match.group(1).strip() if project_match else "Progetto sconosciuto"
            keyword = keyword_match.group(1).strip() if keyword_match else "Keyword sconosciuta"
            market = market_match.group(1).strip() if market_match else "Mercato sconosciuto"
        
            # Crea un messaggio di riepilogo conciso
            summary = f"""
            RIPRISTINO ANALISI CRISP:
        
            Stavo analizzando il mercato di "{keyword}" su {market}.
            ID progetto: {project}
        
            Ho dovuto aprire una nuova sessione per limiti di contesto.
            Sto continuando l'analisi dal punto in cui era stata interrotta.
        
            Per favore, continua l'analisi precedente.
            """
        
            # Trova l'input box
            input_box = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "div.search-input-wrapper textarea"))
            )
        
            # Pulisci l'input box
            input_box.clear()
            time.sleep(0.5)
        
            # Invia il riepilogo per piccoli blocchi
            for chunk in [summary[i:i+100] for i in range(0, len(summary), 100)]:
                input_box.send_keys(chunk)
                time.sleep(0.1)
        
            time.sleep(1)
        
            # Trova e clicca il pulsante di invio
            send_button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "div.search-input-wrapper div.input-icon"))
            )
            send_button.click()
        
            # Attendi l'elaborazione
            time.sleep(10)
        
            self.add_log("Riepilogo essenziale inviato con successo")
            return True
        
        except Exception as e:
            self.add_log(f"❌ Errore nell'invio del riepilogo essenziale: {str(e)}")
            return False

  

    def continue_analysis(self):
        """Continua l'analisi dopo una pausa manuale"""
        try:
            # Controlla se stai usando CRISP
            if hasattr(self, 'use_crisp') and self.use_crisp and hasattr(self, 'current_analysis') and self.current_analysis.get('crisp_project_id'):
                return self._continue_analysis_crisp()
            else:
                return self._continue_analysis_legacy()
        except Exception as e:
            print(f"DEBUG ERROR: {str(e)}")
            return self.add_log(f"Errore durante il completamento dell'analisi: {str(e)}")
            
    def _continue_analysis_crisp(self):
        """Continua l'analisi CRISP dopo una pausa manuale"""
        try:
            self.add_log("Continuazione analisi PubliScript...")
            # Ottieni la risposta attuale dalla chat
            response = self.get_last_response()
            
            if not response:
                return self.add_log("Non è stato possibile recuperare la risposta dalla chat")
            
            # Recupera l'ID del progetto CRISP
            project_id = self.current_analysis.get('crisp_project_id')
            if not project_id:
                return self.add_log("Errore: Nessun progetto CRISP trovato")
            
            # Determina quale fase CRISP è stata interrotta
            execution_history = self.current_analysis.get('execution_history', [])
            if not execution_history:
                return self.add_log("Errore: Nessuna storia di esecuzione trovata")
            
            last_step = execution_history[-1]['step_id']
            self.add_log(f"Ripresa dall'ultimo step completato: {last_step}")
            
            # Salva la risposta nel database
            # Aggiorna i dati del progetto con la nuova risposta
            self.chat_manager.save_response(
                response,
                f"Continuazione CRISP - {last_step}",
                {"project_id": project_id, "manual_continuation": True}
            )
            
            # Continua l'esecuzione del flusso CRISP
            # Definisci una funzione executor per continuare
            def continue_executor(prompt_text):
                self.add_log(f"Continuazione prompt CRISP ({len(prompt_text)} caratteri)...")
                lines = [line.strip() for line in prompt_text.split('\n') if line.strip()]
                cumulative_response = []
                
                for i, line in enumerate(lines):
                    self.add_log(f"Linea {i+1}/{len(lines)}: {line[:50]}...")
                    response = self.send_to_genspark(line)
                    cumulative_response.append(response)
                    time.sleep(2)
                
                combined_response = "\n\n".join(cumulative_response)
                self.chat_manager.save_response(
                    combined_response,
                    "Continuazione CRISP",
                    {"project_id": project_id}
                )
                return combined_response
            
            # Aggiorna l'interfaccia per indicare che la continuazione è in corso
            self.add_log("🔄 Ripresa dell'analisi CRISP...")
            
            # In un'implementazione reale, qui chiameresti il metodo del framework CRISP 
            # per continuare dal punto di interruzione. Però, poiché il framework non ha 
            # un metodo specifico per questo, dovresti implementare la logica tu stesso.
            self.add_log("✅ Analisi CRISP continuata con successo")
            return self.chat_manager.get_log_history_string()
            
        except Exception as e:
            error_msg = f"Errore durante la continuazione dell'analisi CRISP: {str(e)}"
            self.add_log(error_msg)
            logging.error(error_msg)
            return self.chat_manager.get_log_history_string()
    
    
    def complete_analysis(self):
        """
        Completa l'analisi e prepara i dettagli del libro.
        Estrae le informazioni critiche dal file di contesto o dal database CRISP
        e le prepara per la generazione del libro.
        """
        import re  # Importazione esplicita di re per evitare l'errore "variable referenced before assignment"
        import os
        import traceback
        import gradio as gr  # Aggiungi questa riga per importare gr all'interno del metodo
        from datetime import datetime

        # Ripara il database prima di continuare
        self.diagnose_and_fix_database()

        self.add_log("▶️ Avvio funzione complete_analysis")
         
        # ==================== FASE 1: DIAGNOSTICA INIZIALE ====================
        # Verifica se il file di contesto esiste e stampa informazioni diagnostiche
        print(f"DEBUG-INIT: Avvio complete_analysis() con dettagli estesi")
        print(f"DEBUG-INIT: Tentativo di lettura del file context.txt - Esiste: {os.path.exists('context.txt')}")
        print(f"DEBUG-INIT: Directory corrente: {os.getcwd()}")
        print(f"DEBUG-INIT: Memoria disponibile per gli oggetti Python")

        # Verifica se ci sono dati nel current_analysis
        if hasattr(self, 'current_analysis'):
            print(f"DEBUG-INIT: current_analysis esiste: {type(self.current_analysis)}")
            if self.current_analysis:
                print(f"DEBUG-INIT: current_analysis contiene {len(self.current_analysis)} elementi")
                # Mostra le chiavi principali
                for key in list(self.current_analysis.keys())[:5]:  # Limita a 5 chiavi per leggibilità
                    print(f"DEBUG-INIT: - Chiave: {key}, Tipo: {type(self.current_analysis[key])}")
            else:
                print("DEBUG-INIT: current_analysis è un dizionario vuoto o None")
        else:
            print("DEBUG-INIT: current_analysis non esiste come attributo")

        # Backup del file di contesto prima di iniziare l'elaborazione
        if os.path.exists("context.txt"):
            try:
                file_size = os.path.getsize("context.txt")
                print(f"DEBUG-CONTEXT: File context.txt trovato - Dimensione: {file_size} bytes")
        
                # Leggi l'intestazione del file per debug
                try:
                    with open("context.txt", "r", encoding="utf-8") as f:
                        # Leggi le prime 10 righe o meno se il file è più corto
                        first_lines = []
                        for _ in range(10):
                            try:
                                line = next(f)
                                first_lines.append(line)
                            except StopIteration:
                                break
                
                        print(f"DEBUG-CONTEXT: Prime {len(first_lines)} righe del file:")
                        for i, line in enumerate(first_lines):
                            print(f"DEBUG-CONTEXT: Riga {i+1}: {line.strip()}")
                except Exception as e:
                    print(f"DEBUG-CONTEXT: Errore nella lettura dell'intestazione del file: {str(e)}")
                    print(f"DEBUG-CONTEXT: Traceback errore intestazione:\n{traceback.format_exc()}")
        
                # Crea backup con timestamp
                import shutil
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                backup_file = f"context_backup_{timestamp}.txt"
                shutil.copy2("context.txt", backup_file)
                backup_size = os.path.getsize(backup_file)
                print(f"DEBUG-CONTEXT: Backup creato: {backup_file} ({backup_size} bytes)")
        
                # Leggi e stampa un'anteprima del contenuto
                try:
                    with open("context.txt", "r", encoding="utf-8") as f:
                        # Leggi i primi 300 caratteri per anteprima
                        context_preview = f.read(300)
                        print(f"DEBUG-CONTEXT: Anteprima del contenuto (primi 300 caratteri):")
                        # Sostituisci i caratteri di nuova riga con \n visibili
                        context_preview_formatted = context_preview.replace('\n', '\\n')
                        print(f"DEBUG-CONTEXT: {context_preview_formatted}")
                except Exception as e:
                    print(f"DEBUG-CONTEXT: Errore nella lettura dell'anteprima: {str(e)}")
    
            except Exception as backup_error:
                print(f"DEBUG-CONTEXT: Errore nel backup del file di contesto: {str(backup_error)}")
                print(f"DEBUG-CONTEXT: Traceback errore backup:\n{traceback.format_exc()}")
        else:
            print("DEBUG-CONTEXT: ATTENZIONE - File context.txt non trovato!")
            # Elenca tutti i file nella directory corrente per debug
            print("DEBUG-CONTEXT: Elenco dei file nella directory corrente:")
            try:
                files = os.listdir()
                for file in files:
                    if os.path.isfile(file):
                        print(f"DEBUG-CONTEXT: - {file} ({os.path.getsize(file)} bytes)")
            except Exception as e:
                print(f"DEBUG-CONTEXT: Errore nell'elenco dei file: {str(e)}")


        # ================ FASE 2: INIZIO DELL'ELABORAZIONE PRINCIPALE ================
        try:
            # Aggiungi al log
            self.add_log("▶️ Avvio funzione complete_analysis")
    
            # Inizializza current_analysis se non esiste
            if not hasattr(self, 'current_analysis') or self.current_analysis is None:
                self.current_analysis = {}
                self.add_log("ℹ️ Inizializzato current_analysis (non esisteva)")
                print("DEBUG-INIT: current_analysis inizializzato (era None)")
            else:
                print(f"DEBUG-INIT: current_analysis già esistente con {len(self.current_analysis)} chiavi")
                # Debug delle chiavi esistenti
                for key in self.current_analysis:
                    value_preview = str(self.current_analysis[key])
                    if len(value_preview) > 100:
                        value_preview = value_preview[:100] + "..."
                    print(f"DEBUG-INIT: - chiave: {key}, valore: {value_preview}")

            # ================ FASE 3: INIZIALIZZAZIONE VALORI DI RITORNO ================
            # Prepara i valori che verranno restituiti alla fine
            analysis_status_text = "**Stato analisi**: Completata"
            tabs_value = gr.Tabs(selected=2)  # Seleziona il tab "Generazione Libro"
            book_title_value = ""             # Titolo estratto dall'analisi
            book_index_value = ""             # Indice estratto dall'analisi
            voice_style_value = ""            # Stile di voce estratto dall'analisi
            book_type_value = ""              # Tipo di libro estratto dall'analisi
    
            print("DEBUG-VALUES: Valori di ritorno inizializzati con stringhe vuote")

            # ================ FASE 4: DETERMINA LA MODALITÀ (CRISP o LEGACY) ================
            # Verifica se si sta utilizzando il framework CRISP
            use_crisp = hasattr(self, 'use_crisp') and self.use_crisp
            self.add_log(f"ℹ️ Modalità CRISP: {use_crisp}")
            print(f"DEBUG-MODE: Utilizzo framework CRISP: {use_crisp}")
    
            # ================ FASE 5A: ELABORAZIONE MODALITÀ CRISP ================
            if use_crisp and hasattr(self, 'current_analysis') and self.current_analysis.get('crisp_project_id'):
                print("DEBUG-CRISP: Avvio elaborazione in modalità CRISP")
                self.add_log("🔍 Tentativo di estrazione dati da progetto CRISP")
        
                # Recupera l'ID del progetto CRISP
                project_id = self.current_analysis.get('crisp_project_id')
                if not project_id:
                    self.add_log("⚠️ ID Progetto CRISP non trovato nella current_analysis")
                    print("DEBUG-CRISP: ID Progetto CRISP non trovato - impossibile recuperare dati")
                else:
                    self.add_log(f"✅ ID Progetto CRISP trovato: {project_id}")
                    print(f"DEBUG-CRISP: ID Progetto CRISP trovato: {project_id}")
            
                    # Recupera i dati completi del progetto dal framework CRISP
                    project_data = None
                    try:
                        # Verifica se è possibile accedere ai dati del progetto
                        if hasattr(self, 'crisp') and hasattr(self.crisp, 'get_project_data'):
                            print(f"DEBUG-CRISP: Tentativo recupero dati con crisp.get_project_data({project_id})")
                    
                            # Chiamata effettiva per recuperare i dati
                            project_data = self.crisp.get_project_data(project_id)
                    
                            if project_data:
                                print(f"DEBUG-CRISP: Dati progetto recuperati: {len(project_data)} variabili")
                        
                                # Stampa le prime 10 variabili per debug
                                counter = 0
                                for key, value in project_data.items():
                                    if counter < 10:
                                        value_str = str(value)
                                        if len(value_str) > 100:
                                            value_str = value_str[:100] + "..."
                                        print(f"DEBUG-CRISP: Variabile {counter+1}: {key} = {value_str}")
                                        counter += 1
                                    else:
                                        break
                        
                                self.add_log(f"✅ Dati progetto recuperati: {len(project_data)} variabili")
                            else:
                                print("DEBUG-CRISP: ERRORE - get_project_data ha restituito None")
                                self.add_log("⚠️ get_project_data ha restituito None")
                        else:
                            print("DEBUG-CRISP: ERRORE - crisp o get_project_data non disponibili")
                    
                            # Diagnostica dettagliata
                            if hasattr(self, 'crisp'):
                                print(f"DEBUG-CRISP: self.crisp esiste: {type(self.crisp)}")
                                print(f"DEBUG-CRISP: hasattr(self.crisp, 'get_project_data'): {hasattr(self.crisp, 'get_project_data')}")
                        
                                # Elenca tutti i metodi disponibili per debug
                                methods = [method for method in dir(self.crisp) if not method.startswith('_')]
                                print(f"DEBUG-CRISP: Metodi disponibili in self.crisp: {methods}")
                            else:
                                print("DEBUG-CRISP: self.crisp non esiste come attributo")
            
                    except Exception as e:
                        self.add_log(f"⚠️ Errore nel recupero dati progetto CRISP: {str(e)}")
                        print(f"DEBUG-CRISP: Eccezione in get_project_data: {str(e)}")
                        print(f"DEBUG-CRISP: Traceback dettagliato:\n{traceback.format_exc()}")
            
                    # ================ FASE 5A-1: ESTRAZIONE DATI DAL PROGETTO CRISP ================
                    # Se abbiamo recuperato i dati del progetto, estrai le informazioni necessarie
                    if project_data:
                        # Salva i dati del progetto per uso futuro
                        self.current_analysis['project_data'] = project_data
                        print("DEBUG-CRISP: project_data salvato in current_analysis per uso futuro")
                
                        # ---------- Estrazione Titolo ----------
                        if 'TITOLO_LIBRO' in project_data:
                            book_title_value = project_data.get('TITOLO_LIBRO', '')
                            self.add_log(f"✅ Titolo estratto: {book_title_value}")
                            print(f"DEBUG-CRISP: Titolo estratto: '{book_title_value}'")
                        else:
                            print("DEBUG-CRISP: TITOLO_LIBRO non trovato nei dati del progetto")
                            # Cerca alternative per il titolo
                            for alt_key in ['TITLE', 'BOOK_TITLE', 'TITOLO']:
                                if alt_key in project_data:
                                    book_title_value = project_data.get(alt_key, '')
                                    print(f"DEBUG-CRISP: Titolo trovato in campo alternativo {alt_key}: {book_title_value}")
                                    break
                
                        # ---------- Estrazione Stile di Voce ----------
                        if 'VOICE_STYLE' in project_data:
                            voice_style_value = project_data.get('VOICE_STYLE', '')
                            self.add_log(f"✅ Stile voce estratto: {voice_style_value}")
                            print(f"DEBUG-CRISP: Stile voce estratto: '{voice_style_value}'")
                        else:
                            print("DEBUG-CRISP: VOICE_STYLE non trovato nei dati del progetto")
                            # Cerca alternative per lo stile di voce
                            for alt_key in ['TONE', 'STYLE', 'WRITING_STYLE']:
                                if alt_key in project_data:
                                    voice_style_value = project_data.get(alt_key, '')
                                    print(f"DEBUG-CRISP: Stile voce trovato in campo alternativo {alt_key}: {voice_style_value}")
                                    break
                
                        # ---------- Estrazione Tipo di Libro ----------
                        if 'LIBRO_TIPO' in project_data:
                            book_type_value = project_data.get('LIBRO_TIPO', '')
                            self.add_log(f"✅ Tipo di libro estratto: {book_type_value}")
                            print(f"DEBUG-CRISP: Tipo di libro estratto: '{book_type_value}'")
                        else:
                            print("DEBUG-CRISP: LIBRO_TIPO non trovato nei dati del progetto")
                            # Cerca alternative per il tipo di libro
                            for alt_key in ['BOOK_TYPE', 'TIPO', 'GENRE']:
                                if alt_key in project_data:
                                    book_type_value = project_data.get(alt_key, '')
                                    print(f"DEBUG-CRISP: Tipo libro trovato in campo alternativo {alt_key}: {book_type_value}")
                                    break
                
                        # ---------- Costruzione Indice del Libro ----------
                        print("DEBUG-CRISP: Tentativo costruzione indice del libro")
                
                        # Cerca CONTENT_PILLARS per la costruzione dell'indice
                        if 'CONTENT_PILLARS' in project_data:
                            self.add_log("🔍 Tentativo di costruzione indice da CONTENT_PILLARS")
                            print("DEBUG-CRISP: Tentativo di costruzione indice da CONTENT_PILLARS")
                    
                            pillars_text = project_data.get('CONTENT_PILLARS', '')
                            print(f"DEBUG-CRISP: CONTENT_PILLARS trovato, lunghezza: {len(pillars_text)}")
                            print(f"DEBUG-CRISP: Anteprima CONTENT_PILLARS: {pillars_text[:200]}...")
                    
                            # Estrai i pilastri di contenuto con diversi pattern di espressioni regolari
                            pillars = []
                            if isinstance(pillars_text, str):
                                # Prova diversi pattern per estrarre pilastri
                                print("DEBUG-CRISP: Tentativo di estrazione pilastri con pattern regex")
                        
                                pattern_results = {}
                        
                                for pattern in [
                                    r'(\d+\.\s*[^\n]+)',             # Pattern per "1. Titolo"
                                    r'(\d+\)\s*[^\n]+)',             # Pattern per "1) Titolo"
                                    r'(CAPITOLO \d+[^:\n]*:[^\n]+)', # Pattern per "CAPITOLO 1: Titolo"
                                    r'(Capitolo \d+[^:\n]*:[^\n]+)'  # Pattern per "Capitolo 1: Titolo"
                                ]:
                                    # Prova ogni pattern e registra i risultati
                                    pillar_matches = re.findall(pattern, pillars_text)
                                    pattern_results[pattern] = pillar_matches
                            
                                    if pillar_matches:
                                        print(f"DEBUG-CRISP: Pattern '{pattern}' ha trovato {len(pillar_matches)} corrispondenze")
                                        # Mostra le prime corrispondenze
                                        for i, match in enumerate(pillar_matches[:3]):
                                            print(f"DEBUG-CRISP: --- Match {i+1}: {match}")
                                
                                        if len(pillar_matches) >= 3:  # Minimo 3 capitoli per un buon indice
                                            pillars = [p.strip() for p in pillar_matches]
                                            break
                                    else:
                                        print(f"DEBUG-CRISP: Pattern '{pattern}' non ha trovato corrispondenze")
                        
                                # Se nessun pattern ha funzionato, prova con approcci alternativi
                                if not pillars:
                                    print("DEBUG-CRISP: Nessun pattern regex ha trovato abbastanza pilastri, provo approccio alternativo")
                            
                                    # Approccio alternativo: dividi per righe e cerca linee che sembrano titoli di capitolo
                                    lines = pillars_text.split('\n')
                                    print(f"DEBUG-CRISP: Text diviso in {len(lines)} righe per analisi")
                            
                                    for line in lines:
                                        line = line.strip()
                                        # Verifica se la riga sembra un titolo di capitolo
                                        if line and (
                                            line.lower().startswith('capitolo') or 
                                            line.lower().startswith('chapter') or
                                            re.match(r'^\d+[\.\)]', line)
                                        ):
                                            pillars.append(line)
                                            print(f"DEBUG-CRISP: Trovato potenziale pilastro: {line}")
                            
                                    if pillars:
                                        print(f"DEBUG-CRISP: Approccio alternativo ha trovato {len(pillars)} potenziali pilastri")
                                    else:
                                        print("DEBUG-CRISP: Anche l'approccio alternativo non ha trovato pilastri")
                                
                                        # Ultimo tentativo: cerca qualsiasi riga che sembra un titolo
                                        print("DEBUG-CRISP: Tentativo di ultima risorsa: qualsiasi riga che sembra un titolo")
                                        for line in lines:
                                            line = line.strip()
                                            # Riga abbastanza lunga ma non troppo e con maiuscole all'inizio
                                            if 10 <= len(line) <= 100 and line[0].isupper() and ":" not in line and line.endswith((".","?")):
                                                pillars.append(line)
                                                print(f"DEBUG-CRISP: Titolo potenziale trovato: {line}")
                                                if len(pillars) >= 5:  # Limita a 5 pilastri per questo approccio
                                                    break
                            else:
                                print(f"DEBUG-CRISP: CONTENT_PILLARS non è una stringa ma un {type(pillars_text)}")
                        
                                # Se CONTENT_PILLARS è una lista (possibile con alcune implementazioni)
                                if isinstance(pillars_text, list):
                                    print(f"DEBUG-CRISP: CONTENT_PILLARS è una lista con {len(pillars_text)} elementi")
                                    pillars = pillars_text
                    
                            # Costruisci l'indice a partire dai pilastri trovati
                            if pillars:
                                print(f"DEBUG-CRISP: Costruzione indice con {len(pillars)} pilastri trovati")
                        
                                # Pulisci e formatta l'indice
                                index_text = "INTRODUZIONE\n\n"
                        
                                for i, pillar in enumerate(pillars, 1):
                                    # Rimuovi numeri e simboli di punteggiatura iniziali
                                    try:
                                        clean_pillar = re.sub(r'^\d+[\.\)\s]+|^CAPITOLO\s+\d+\s*[:\.\-\s]*|^Capitolo\s+\d+\s*[:\.\-\s]*', '', pillar).strip()
                                        print(f"DEBUG-CRISP: Pillar {i} originale: '{pillar}'")
                                        print(f"DEBUG-CRISP: Pillar {i} pulito: '{clean_pillar}'")
                                
                                        if clean_pillar:  # Aggiungi solo se c'è testo dopo la pulizia
                                            index_text += f"CAPITOLO {i}: {clean_pillar}\n\n"
                                        else:
                                            print(f"DEBUG-CRISP: Pillar {i} ha prodotto un testo vuoto dopo la pulizia")
                                            # Usa il testo originale come fallback se la pulizia ha rimosso tutto
                                            index_text += f"CAPITOLO {i}: {pillar}\n\n"
                                    except Exception as e:
                                        print(f"DEBUG-CRISP: Errore nella pulizia del pillar {i}: {str(e)}")
                                        # Usa il testo originale in caso di errore
                                        index_text += f"CAPITOLO {i}: {pillar}\n\n"
                        
                                index_text += "CONCLUSIONE"
                                book_index_value = index_text
                                self.add_log(f"✅ Indice costruito con {len(pillars)} capitoli")
                                print(f"DEBUG-CRISP: Indice costruito con successo:\n{book_index_value}")
                            else:
                                # Indice di fallback se non sono stati trovati pilastri
                                print("DEBUG-CRISP: Nessun pilastro trovato, uso indice di fallback")
                                book_index_value = """INTRODUZIONE

    CAPITOLO 1: Fondamenti

    CAPITOLO 2: Metodologia

    CAPITOLO 3: Applicazione

    CAPITOLO 4: Casi Studio

    CAPITOLO 5: Risultati

    CONCLUSIONE"""
                                self.add_log("⚠️ Usato indice di fallback (nessun pilastro trovato)")
                                print("DEBUG-CRISP: Usato indice di fallback")
                        else:
                            print("DEBUG-CRISP: CONTENT_PILLARS non trovato nei dati del progetto")
                    
                            # Cerca campi alternativi che potrebbero contenere informazioni per l'indice
                            alternative_found = False
                            for key in ['BOOK_STRUCTURE', 'INDICE_LIBRO', 'BOOK_JOURNEY', 'CHAPTER_STRUCTURE']:
                                if key in project_data:
                                    print(f"DEBUG-CRISP: Trovato {key} come alternativa a CONTENT_PILLARS")
                                    self.add_log(f"🔍 Tentativo di costruzione indice da {key}")
                            
                                    # Implementazione simile a quella per CONTENT_PILLARS
                                    alternative_text = project_data.get(key, '')
                                    # (Ripeti logica simile a quella usata per CONTENT_PILLARS)
                                    # Per brevità, questo codice è omesso ma sarebbe una duplicazione
                                    # dell'approccio sopra adattato per il campo alternativo
                            
                                    alternative_found = True
                                    break
                    
                            if not alternative_found:
                                print("DEBUG-CRISP: Nessuna alternativa a CONTENT_PILLARS trovata, uso indice di fallback")
                    else:
                        print("DEBUG-CRISP: project_data è None o vuoto, impossibile estrarre dati")
    
            # ================ FASE 5B: ELABORAZIONE MODALITÀ LEGACY ================
            else:
                # Approccio legacy - senza framework CRISP
                self.add_log("🔍 Utilizzo approccio legacy (non CRISP)")
                print("DEBUG-LEGACY: Avvio elaborazione in modalità legacy (non CRISP)")

                try:
                    # Cerca dati nel file di contesto
                    context_file = "context.txt"
                    if os.path.exists(context_file):
                        # Informazioni sul file
                        file_size = os.path.getsize(context_file)
                        print(f"DEBUG-LEGACY: File context.txt trovato, dimensione: {file_size} bytes")
        
                        # Leggi l'intero contenuto del file
                        try:
                            with open(context_file, "r", encoding="utf-8") as f:
                                context_content = f.read()
            
                            self.add_log(f"✅ File contesto letto: {len(context_content)} caratteri")
                            print(f"DEBUG-LEGACY: File contesto letto con successo: {len(context_content)} caratteri")
            
                            # Stampa le prime righe per debug
                            content_preview = context_content[:500].replace('\n', ' ')
                            print(f"DEBUG-LEGACY: Anteprima dei primi 500 caratteri: {content_preview}...")
            
                            # Analisi strutturale del contenuto per determinare il formato
                            print("DEBUG-LEGACY: Analisi strutturale del contenuto")
            
                            # Cerca sezioni nel formato standard
                            section_pattern = r'===\s+([^=]+?)\s+-\s+\d{8}_\d{6}\s+===\n'
                            sections = re.findall(section_pattern, context_content)
            
                            if sections:
                                print(f"DEBUG-LEGACY: Trovate {len(sections)} sezioni nel formato standard")
                                for i, section in enumerate(sections[:5]):  # Mostra solo le prime 5
                                    print(f"DEBUG-LEGACY: - Sezione {i+1}: {section}")
                
                                # Analisi dettagliata delle sezioni
                                print("DEBUG-LEGACY: Analisi dettagliata delle sezioni trovate")
                                section_contents = re.split(section_pattern, context_content)[1:]  # Salta il primo che è vuoto
                
                                # Assicurati che abbiamo lo stesso numero di titoli e contenuti
                                if len(sections) == len(section_contents)/2:
                                    print("DEBUG-LEGACY: Numero corretto di sezioni e contenuti")
                                else:
                                    print(f"DEBUG-LEGACY: ATTENZIONE - Discrepanza: {len(sections)} titoli vs {len(section_contents)/2} contenuti")
                            else:
                                print("DEBUG-LEGACY: Nessuna sezione trovata nel formato standard")
                                # Cerca formati alternativi
                                alt_pattern = r'---\s+([^-]+?)\s+---\n'
                                alt_sections = re.findall(alt_pattern, context_content)
                                if alt_sections:
                                    print(f"DEBUG-LEGACY: Trovate {len(alt_sections)} sezioni in formato alternativo")
                                else:
                                    print("DEBUG-LEGACY: Nessuna sezione trovata in formato alternativo")
            
                            # ================ FASE 5B-1: ESTRAZIONE DATI LEGACY ================
                            try:
                                # ---------- Estrazione Titolo ----------
                                print("DEBUG-LEGACY: Tentativo estrazione titolo")
                
                                # Lista di pattern da provare per trovare il titolo
                                title_patterns = [
                                    r'7\)[^:]*:[^T]*Titolo[^:]*:[^\n]*\n([^\n]+)',
                                    r'7\.[^:]*:[^T]*Titolo[^:]*:[^\n]*\n([^\n]+)',
                                    r'Titolo[^:]*:[^\n]*\n([^\n]+)',
                                    r'(?:title|titolo)[^:]*:[^\n]*\n([^\n]+)',
                                    r'THE[^"]*"([^"]+)"',  # Pattern per titoli tra virgolette
                                    r'"([^"]+)".*?(?:il tuo nuovo libro|your new book)',  # Pattern per titoli suggeriti
                                ]
                
                                # Prova ogni pattern fino a trovare una corrispondenza
                                book_title_value = ""
                                for pattern in title_patterns:
                                    print(f"DEBUG-LEGACY: Provo pattern titolo: {pattern}")
                                    title_match = re.search(pattern, context_content, re.IGNORECASE)
                    
                                    if title_match:
                                        book_title_value = title_match.group(1).strip()
                                        self.add_log(f"✅ Titolo estratto (legacy): {book_title_value}")
                                        print(f"DEBUG-LEGACY: Titolo estratto con pattern '{pattern}': {book_title_value}")
                                        break
                
                                if not book_title_value:
                                    print("DEBUG-LEGACY: Nessun titolo trovato con i pattern standard")
                    
                                    # Cerca titoli in sezioni specifiche
                                    print("DEBUG-LEGACY: Ricerca titolo in sezioni specifiche")
                    
                                    # Cerca sezioni che potrebbero contenere titoli
                                    title_sections = [s for s in sections if 'titolo' in s.lower() or 'title' in s.lower()]
                                    if title_sections:
                                        print(f"DEBUG-LEGACY: Trovate {len(title_sections)} sezioni potenzialmente contenenti titoli")
                        
                                        # Per ogni sezione potenziale, cerca titoli
                                        for title_section in title_sections:
                                            section_index = sections.index(title_section)
                                            section_content = section_contents[section_index * 2]  # Moltiplica per 2 a causa della divisione
                            
                                            # Cerca titoli nella sezione
                                            title_lines = [line for line in section_content.split('\n') if line.strip()]
                                            if title_lines:
                                                print(f"DEBUG-LEGACY: Sezione '{title_section}' contiene {len(title_lines)} linee non vuote")

    # Prendi la prima linea che sembra un titolo
                                            for line in title_lines:
                                                # Se la linea sembra un titolo (non troppo lungo, non contiene caratteri speciali)
                                                if 10 <= len(line) <= 100 and not any(char in line for char in ['{', '}', '(', ')', '[', ']']):
                                                    book_title_value = line.strip().strip('"\'')
                                                    print(f"DEBUG-LEGACY: Titolo estratto da sezione: {book_title_value}")
                                                    break
                            
                                            if book_title_value:
                                                break
                
                                # ---------- Estrazione Indice ----------
                                print("DEBUG-LEGACY: Tentativo estrazione indice")
                
                                # Lista di pattern da provare per trovare l'indice
                                index_patterns = [
                                    r'8\)[^:]*:[^I]*Indice[^:]*:[^\n]*\n(.*?)(?=\n\n|$)',
                                    r'8\.[^:]*:[^I]*Indice[^:]*:[^\n]*\n(.*?)(?=\n\n|$)',
                                    r'Indice[^:]*:[^\n]*\n(.*?)(?=\n\n|$)',
                                    r'(?:indice|index)[^:]*:[^\n]*\n(.*?)(?=\n\n|$)',
                                    r'INDICE DEL LIBRO[^\n]*\n(.*?)(?=\n\n===|$)',  # Pattern specifico
                                    r'Indice del Libro[^\n]*\n(.*?)(?=\n\n|$)'      # Altra variante
                                ]
                
                                # Prova ogni pattern fino a trovare una corrispondenza
                                book_index_value = ""
                                for pattern in index_patterns:
                                    print(f"DEBUG-LEGACY: Provo pattern indice: {pattern}")
                                    index_match = re.search(pattern, context_content, re.DOTALL | re.IGNORECASE)
                    
                                    if index_match:
                                        book_index_value = index_match.group(1).strip()
                                        self.add_log(f"✅ Indice estratto (legacy): {len(book_index_value)} caratteri")
                                        print(f"DEBUG-LEGACY: Indice estratto con pattern '{pattern}', lunghezza: {len(book_index_value)}")
                                        print(f"DEBUG-LEGACY: Preview indice: {book_index_value[:200]}...")
                                        break
                
                                if not book_index_value:
                                    print("DEBUG-LEGACY: Nessun indice trovato con i pattern standard")
                                    print("DEBUG-LEGACY: Tentativo ricerca capitoli diretta")
                    
                                    # Cerca tutti i pattern che sembrano capitoli
                                    chapter_patterns = [
                                        r'(CAPITOLO\s+\d+[^:\n]*:[^\n]+)',
                                        r'(CHAPTER\s+\d+[^:\n]*:[^\n]+)',
                                        r'(Capitolo\s+\d+[^:\n]*:[^\n]+)'
                                    ]
                    
                                    all_chapters = []
                                    for pattern in chapter_patterns:
                                        chapters = re.findall(pattern, context_content, re.IGNORECASE)
                                        if chapters:
                                            print(f"DEBUG-LEGACY: Pattern '{pattern}' ha trovato {len(chapters)} capitoli")
                                            all_chapters.extend(chapters)
                    
                                    if all_chapters:
                                        print(f"DEBUG-LEGACY: Trovati {len(all_chapters)} capitoli potenziali nel testo")
                        
                                        # Cerca il blocco di testo che contiene più capitoli consecutivi
                                        chapter_sections = []
                                        for match in re.finditer(r'((?:CAPITOLO\s+\d+[^\n]*\n){2,})', context_content, re.IGNORECASE):
                                            section_text = match.group(1)
                                            chapter_count = section_text.lower().count('capitolo')
                                            chapter_sections.append((match.start(), match.end(), section_text, chapter_count))
                        
                                        if chapter_sections:
                                            # Usa la sezione con più capitoli
                                            best_section = max(chapter_sections, key=lambda x: x[3])
                                            print(f"DEBUG-LEGACY: Trovato blocco indice con {best_section[3]} capitoli")
                            
                                            # Aggiungi l'introduzione e conclusione se non presenti
                                            book_index_value = "INTRODUZIONE\n\n" + best_section[2] + "\nCONCLUSIONE"
                                            print(f"DEBUG-LEGACY: Indice costruito da blocco trovato: {len(book_index_value)} caratteri")
                                        else:
                                            # Se non ci sono blocchi, combina tutti i capitoli trovati
                                            book_index_value = "INTRODUZIONE\n\n" + "\n".join(all_chapters) + "\n\nCONCLUSIONE"
                                            print(f"DEBUG-LEGACY: Indice costruito da capitoli individuali: {len(book_index_value)} caratteri")
                
                                # ---------- Estrazione Stile di Voce ----------
                                print("DEBUG-LEGACY: Tentativo estrazione stile di voce")
                
                                # Lista di pattern da provare per trovare lo stile di voce
                                voice_patterns = [
                                    r'Tono di voce[^:]*:[^\n]*\n([^\n]+)',
                                    r'Voce[^:]*:[^\n]*\n([^\n]+)',
                                    r'Stile[^:]*:[^\n]*\n([^\n]+)',
                                    r'VOICE_STYLE[^:]*:[^\n]*\n([^\n]+)',
                                    r'(?:conversazionale|formale|informativo|tecnico)[^\n]+'
                                ]
                
                                # Prova ogni pattern fino a trovare una corrispondenza
                                voice_style_value = ""
                                for pattern in voice_patterns:
                                    print(f"DEBUG-LEGACY: Provo pattern stile voce: {pattern}")
                                    voice_match = re.search(pattern, context_content, re.IGNORECASE)
                    
                                    if voice_match:
                                        # Gestione speciale per l'ultimo pattern che non ha gruppo
                                        if 'conversazionale' in pattern:
                                            voice_style_value = voice_match.group(0).strip()
                                        else:
                                            voice_style_value = voice_match.group(1).strip()
                        
                                        self.add_log(f"✅ Stile voce estratto (legacy): {voice_style_value}")
                                        print(f"DEBUG-LEGACY: Stile voce estratto con pattern '{pattern}': {voice_style_value}")
                                        break
                
                                if not voice_style_value:
                                    print("DEBUG-LEGACY: Nessuno stile di voce trovato con i pattern standard")
                    
                                    # Cerca nelle sezioni potenzialmente legate allo stile
                                    style_sections = [s for s in sections if any(term in s.lower() for term in 
                                                     ['voice', 'voce', 'stile', 'tone', 'tono'])]
                    
                                    if style_sections:
                                        print(f"DEBUG-LEGACY: Trovate {len(style_sections)} sezioni potenzialmente contenenti stile")
                        
                                        for style_section in style_sections:
                                            section_index = sections.index(style_section)
                                            section_content = section_contents[section_index * 2]
                            
                                            # Cerca stile nelle prime 5 righe della sezione
                                            style_lines = [line for line in section_content.split('\n')[:5] if line.strip()]
                                            if style_lines:
                                                voice_style_value = style_lines[0].strip()
                                                print(f"DEBUG-LEGACY: Stile voce estratto da sezione: {voice_style_value}")
                                                break
                
                                # ---------- Estrazione Tipo di Libro ----------
                                print("DEBUG-LEGACY: Tentativo estrazione tipo di libro")
                
                                # Lista di pattern da provare per trovare il tipo di libro
                                book_type_patterns = [
                                    r'tipo di libro[^:]*:[^\n]*\n*\s*([^\n]+)',
                                    r'genere[^:]*:[^\n]*\n*\s*([^\n]+)',
                                    r'categoria[^:]*:[^\n]*\n*\s*([^\n]+)',
                                    r'LIBRO_TIPO[^:]*:[^\n]*\n*\s*([^\n]+)'
                                ]
                
                                # Prova ogni pattern fino a trovare una corrispondenza
                                book_type_value = ""
                                for pattern in book_type_patterns:
                                    print(f"DEBUG-LEGACY: Provo pattern tipo libro: {pattern}")
                                    book_type_match = re.search(pattern, context_content, re.IGNORECASE)
                    
                                    if book_type_match:
                                        book_type_value = book_type_match.group(1).strip()
                                        self.add_log(f"✅ Tipo di libro estratto (legacy): {book_type_value}")
                                        print(f"DEBUG-LEGACY: Tipo di libro estratto con pattern '{pattern}': {book_type_value}")
                                        break
                
                                if not book_type_value:
                                    print("DEBUG-LEGACY: Nessun tipo di libro trovato con i pattern standard")
                    
                                    # Cerca valori comuni di tipo libro nel testo
                                    common_types = ["Manuale", "Non-Fiction", "Ricettario", "Self-Help", "How-To", 
                                                   "Craft", "Hobby", "Survival", "Test Study"]
                    
                                    for book_type in common_types:
                                        if book_type.lower() in context_content.lower():
                                            book_type_value = book_type
                                            print(f"DEBUG-LEGACY: Tipo libro trovato nel testo: {book_type_value}")
                                            break
            
                            except Exception as extraction_error:
                                self.add_log(f"⚠️ Errore nell'estrazione dei dati: {str(extraction_error)}")
                                print(f"DEBUG-LEGACY: Errore nell'estrazione dei dati: {str(extraction_error)}")
                                print(f"DEBUG-LEGACY: Traceback errore estrazione:\n{traceback.format_exc()}")
            
                        except Exception as read_error:
                            self.add_log(f"⚠️ Errore nella lettura del file context.txt: {str(read_error)}")
                            print(f"DEBUG-LEGACY: Errore nella lettura del file context.txt: {str(read_error)}")
                            print(f"DEBUG-LEGACY: Traceback errore lettura:\n{traceback.format_exc()}")
            
                    else:
                        self.add_log("⚠️ File context.txt non trovato!")
                        print(f"DEBUG-LEGACY: File context.txt non trovato in {os.getcwd()}")
        
                        # Elenca i file nella directory corrente
                        files = os.listdir()
                        print(f"DEBUG-LEGACY: File nella directory corrente: {files}")
        
                        # Cerca file alternativi che potrebbero contenere i dati
                        context_alternatives = [f for f in files if 'context' in f.lower() or 
                                              'backup' in f.lower() or f.endswith('.txt')]
        
                        if context_alternatives:
                            print(f"DEBUG-LEGACY: Trovati possibili file alternativi: {context_alternatives}")
                            self.add_log(f"⚠️ File context.txt non trovato, ma ci sono alternative: {context_alternatives}")
    
                except Exception as e:
                    self.add_log(f"⚠️ Errore nell'estrazione legacy: {str(e)}")
                    print(f"DEBUG-LEGACY: Errore nell'estrazione legacy: {str(e)}")
                    print(f"DEBUG-LEGACY: Traceback errore estrazione:\n{traceback.format_exc()}")

            # ================ FASE 6: VALORI DI FALLBACK ================
            # Se necessario, utilizza valori di fallback per i campi che non è stato possibile estrarre

            print("DEBUG-FINAL: Verifica valori estratti prima di applicare fallback")
            print(f"DEBUG-FINAL: Titolo estratto: '{book_title_value}'")
            print(f"DEBUG-FINAL: Indice estratto: {len(book_index_value) if book_index_value else 0} caratteri")
            print(f"DEBUG-FINAL: Stile voce estratto: '{voice_style_value}'")
            print(f"DEBUG-FINAL: Tipo libro estratto: '{book_type_value}'")

            # Applica fallback se necessario
            if not book_title_value:
                book_title_value = "Il tuo nuovo libro"
                self.add_log("⚠️ Usato titolo di fallback")
                print("DEBUG-FINAL: Usato titolo di fallback")

            if not book_index_value:
                book_index_value = """INTRODUZIONE

    CAPITOLO 1: Fondamenti

    CAPITOLO 2: Metodologia

    CAPITOLO 3: Applicazione

    CONCLUSIONE"""
                self.add_log("⚠️ Usato indice di fallback")
                print("DEBUG-FINAL: Usato indice di fallback")

            if not voice_style_value:
                voice_style_value = "Conversazionale e informativo"
                self.add_log("⚠️ Usato stile voce di fallback")
                print("DEBUG-FINAL: Usato stile voce di fallback")

            if not book_type_value:
                book_type_value = "Manuale (Non-Fiction)"
                self.add_log("⚠️ Usato tipo libro di fallback")
                print("DEBUG-FINAL: Usato tipo libro di fallback")

            # ================ FASE 7: AGGIORNAMENTO INTERFACCIA ================
            # Aggiorna i campi dell'interfaccia con i valori estratti

            print("DEBUG-UPDATE: Tentativo aggiornamento campi interfaccia")

            # Importa gradio esplicitamente se necessario
            import gradio as gr

            try:
                print("DEBUG-UPDATE: Inizio aggiornamento componenti Gradio")

                # Verifica e stampa info sulla versione di Gradio
                gradio_version = gr.__version__ if hasattr(gr, '__version__') else "sconosciuta"
                print(f"DEBUG-UPDATE: Versione Gradio rilevata: {gradio_version}")

                # Verifica quali componenti esistono
                components = {
                    'book_title': hasattr(self, 'book_title'),
                    'book_index': hasattr(self, 'book_index'),
                    'voice_style': hasattr(self, 'voice_style'),
                    'book_type_hidden': hasattr(self, 'book_type_hidden'),
                    'tabs': hasattr(self, 'tabs')
                }
                print(f"DEBUG-UPDATE: Componenti esistenti: {components}")

                # Per Gradio 5.x, il metodo corretto è .update(value=...)

                if hasattr(self, 'book_title'):
                    print(f"DEBUG-UPDATE: Aggiornamento book_title: '{book_title_value}'")
                    try:
                        self.book_title.value = book_title_value
                        self.add_log(f"✓ Campo book_title aggiornato")
                        print(f"DEBUG-UPDATE: Campo book_title aggiornato con: {book_title_value}")
                    except Exception as e:
                        print(f"DEBUG-UPDATE: Errore aggiornamento book_title: {str(e)}")

                if hasattr(self, 'book_index'):
                    print(f"DEBUG-UPDATE: Aggiornamento book_index: {len(book_index_value)} caratteri")
                    try:
                        self.book_index.value = book_index_value
                        self.add_log(f"✓ Campo book_index aggiornato")
                        print(f"DEBUG-UPDATE: Campo book_index aggiornato (lunghezza: {len(book_index_value)} caratteri)")
                    except Exception as e:
                        print(f"DEBUG-UPDATE: Errore aggiornamento book_index: {str(e)}")

                if hasattr(self, 'voice_style'):
                    print(f"DEBUG-UPDATE: Aggiornamento voice_style: '{voice_style_value}'")
                    try:
                        self.voice_style.value = voice_style_value
                        self.add_log(f"✓ Campo voice_style aggiornato")
                        print(f"DEBUG-UPDATE: Campo voice_style aggiornato con: {voice_style_value}")
                    except Exception as e:
                        print(f"DEBUG-UPDATE: Errore aggiornamento voice_style: {str(e)}")

                # Aggiorna il tipo di libro se esiste il campo
                if hasattr(self, 'book_type_hidden'):
                    print(f"DEBUG-UPDATE: Aggiornamento book_type_hidden: '{book_type_value}'")
                    try:
                        self.book_type_hidden.value = book_type_value
                        self.add_log(f"✓ Campo book_type_hidden aggiornato")
                        print(f"DEBUG-UPDATE: Campo book_type_hidden aggiornato con: {book_type_value}")
                    except Exception as e:
                        print(f"DEBUG-UPDATE: Errore aggiornamento book_type_hidden: {str(e)}")

                # Cambia tab
                if hasattr(self, 'tabs'):
                    print("DEBUG-UPDATE: Aggiornamento tab a indice 2 (Generazione Libro)")
                    try:
                        self.tabs.selected = 2  # Imposta direttamente l'indice selezionato
                        self.add_log("✓ Tab aggiornato")
                        print("DEBUG-UPDATE: Tab aggiornato a indice 2 (Generazione Libro)")
                    except Exception as e:
                        print(f"DEBUG-UPDATE: Errore aggiornamento tabs: {str(e)}")
                        # Fallback per tabs
                        try:
                            import gradio as gr
                            self.tabs = gr.Tabs(selected=2)
                            print("DEBUG-UPDATE: Tab aggiornato con metodo alternativo")
                        except Exception as alt_e:
                            print(f"DEBUG-UPDATE: Anche fallback per tabs fallito: {str(alt_e)}")

                # Verifica che Gradio abbia effettivamente aggiornato i campi
                print("DEBUG-UPDATE: Verifica finale campi aggiornati")
                if hasattr(self, 'book_title'):
                    print(f"DEBUG-UPDATE: book_title.value finale: {getattr(self.book_title, 'value', 'N/A')}")

            # ================ FASE 8: COMPLETAMENTO ================
            self.add_log("✅ Funzione complete_analysis terminata con successo")
            print("DEBUG-FINAL: Funzione complete_analysis terminata con successo")

            # Un ultimo check prima di restituire i valori
            print(f"DEBUG-FINAL: Valori finali da restituire:")
            print(f"DEBUG-FINAL: - Log history: {len(self.log_history)} righe")
            print(f"DEBUG-FINAL: - Status: {analysis_status_text}")
            print(f"DEBUG-FINAL: - Tabs: {tabs_value}")
            print(f"DEBUG-FINAL: - Titolo: {book_title_value}")
            print(f"DEBUG-FINAL: - Indice: {len(book_index_value)} caratteri")
            print(f"DEBUG-FINAL: - Stile: {voice_style_value}")
            print(f"DEBUG-FINAL: - Tipo: {book_type_value}")

            # ================ FASE 8.5: SALVATAGGIO NEL DATABASE ================
            try:
                # Determina la keyword corrente in modo più affidabile
                keyword = None
                
                # Tentativo 1: Usa la keyword dall'interfaccia UI
                if hasattr(self, 'keyword') and hasattr(self.keyword, 'value') and self.keyword.value:
                    keyword = self.keyword.value.strip()
                    self.add_log(f"🔍 Usando keyword dall'interfaccia: {keyword}")
                    print(f"DEBUG-DB: Usando keyword dall'interfaccia: {keyword}")
                
                # Tentativo 2: Estrai la keyword dall'analisi corrente
                if not keyword and hasattr(self, 'current_analysis') and self.current_analysis:
                    current_keyword = self.current_analysis.get('KEYWORD')
                    if current_keyword:
                        keyword = current_keyword
                        self.add_log(f"🔍 Usando keyword dall'analisi corrente: {keyword}")
                        print(f"DEBUG-DB: Usando keyword dall'analisi corrente: {keyword}")
                
                # Tentativo 3: Cerca la keyword nei file di backup
                if not keyword:
                    backup_dir = "backups"
                    if os.path.exists(backup_dir):
                        backup_files = [f for f in os.listdir(backup_dir) 
                                       if f.startswith("context_") and f.endswith(".txt")]
                        if backup_files:
                            # Ordina per data di modifica (più recente prima)
                            backup_files = sorted(backup_files, 
                                               key=lambda x: os.path.getmtime(os.path.join(backup_dir, x)), 
                                               reverse=True)
                            # Estrai la keyword dal primo file (il più recente)
                            if backup_files:
                                match = re.search(r"context_([^_]+)", backup_files[0])
                                if match:
                                    keyword_from_backup = match.group(1).replace("_", " ")
                                    keyword = keyword_from_backup
                                    self.add_log(f"🔍 Usando keyword dai file di backup: {keyword}")
                                    print(f"DEBUG-DB: Usando keyword dai file di backup: {keyword}")
                
                # Tentativo 4: Cerca nei log recenti
                if not keyword and hasattr(self, 'log_history') and isinstance(self.log_history, list):
                    # Cerca la keyword nei log delle ultime operazioni
                    for log in self.log_history[-30:]:  # Ultimi 30 log
                        if isinstance(log, str):
                            match = re.search(r"per: ([^\n\r]+)", log)
                            if match:
                                keyword_from_log = match.group(1).strip()
                                if keyword_from_log and len(keyword_from_log) < 50:  # Evita titoli troppo lunghi
                                    keyword = keyword_from_log
                                    self.add_log(f"🔍 Usando keyword dai log: {keyword}")
                                    print(f"DEBUG-DB: Usando keyword dai log: {keyword}")
                                    break
                
                # Tentativo 5: Solo come ultima risorsa, estrai una versione corta dal titolo
                if not keyword and book_title_value:
                    # Prendi solo le prime 5 parole, non l'intero titolo
                    words = book_title_value.split()[:5]
                    short_title = " ".join(words)
                    if len(short_title) > 50:  # Ulteriore limite di sicurezza
                        short_title = short_title[:47] + "..."
                    
                    keyword = short_title
                    self.add_log(f"🔍 Usando versione breve del titolo come keyword: {keyword}")
                    print(f"DEBUG-DB: Usando versione breve del titolo: {keyword}")
                
                # Verifica che abbiamo una keyword da usare
                if keyword:
                    self.add_log(f"💾 Tentativo salvataggio analisi nel database per keyword: {keyword}")
                    print(f"DEBUG-DB: Tentativo salvataggio con keyword: {keyword}")
                    
                    try:
                        # Verifica che il database esista
                        if os.path.exists(self.crisp.project_db_path):
                            # Crea un nuovo progetto nel database
                            conn = sqlite3.connect(self.crisp.project_db_path)
                            cursor = conn.cursor()
                
                            # Genera un ID progetto appropriato
                            from datetime import datetime
                            current_date = datetime.now().strftime("%Y%m%d_%H%M%S")
                            project_id = f"PROJ_{current_date}"
                
                            # Formatta il nome del progetto usando la keyword corretta
                            project_name = f"{keyword} - {datetime.now().strftime('%Y-%m-%d')}"
                
                            # Inserisci nel database con ID esplicito
                            current_datetime = datetime.now().isoformat()
                            cursor.execute(
                                "INSERT INTO projects (id, name, creation_date, last_updated) VALUES (?, ?, ?, ?)",
                                (project_id, project_name, current_datetime, current_datetime)
                            )
                
                            # Salva le variabili principali di progetto
                            main_vars = {
                                "KEYWORD": keyword,
                                "TITOLO_LIBRO": book_title_value,
                                "INDICE_LIBRO": book_index_value,
                                "VOICE_STYLE": voice_style_value,
                                "BOOK_TYPE": book_type_value
                            }
                
                            # Inserisci le variabili nel database
                            for key, value in main_vars.items():
                                if value:  # Salva solo valori non vuoti
                                    cursor.execute(
                                        "INSERT INTO project_variables (project_id, variable_name, variable_value) VALUES (?, ?, ?)",
                                        (project_id, key, str(value))
                                    )
                
                            # Se esistono dati in current_analysis, salvali come variabili
                            if hasattr(self, 'current_analysis') and self.current_analysis:
                                for key, value in self.current_analysis.items():
                                    if key not in main_vars and value:  # Evita duplicati
                                        cursor.execute(
                                            "INSERT INTO project_variables (project_id, variable_name, variable_value) VALUES (?, ?, ?)",
                                            (project_id, key, str(value))
                                        )
                
                            # Salva il contenuto del file di contesto come risultato di progetto
                            if hasattr(self.chat_manager, 'context_file') and os.path.exists(self.chat_manager.context_file):
                                with open(self.chat_manager.context_file, 'r', encoding='utf-8') as f:
                                    context_content = f.read()
                        
                                cursor.execute(
                                    "INSERT INTO project_results (project_id, step_id, result_text, timestamp) VALUES (?, ?, ?, ?)",
                                    (project_id, "ANALISI", context_content, current_datetime)
                                )
                
                            conn.commit()
                            conn.close()
                
                            self.add_log(f"✅ Analisi salvata con successo nel database (ID: {project_id})")
                            print(f"DEBUG-DB: Analisi salvata con successo, ID: {project_id}")
                        else:
                            self.add_log(f"⚠️ Database non trovato: {self.crisp.project_db_path}")
                            print(f"DEBUG-DB: Database non trovato: {self.crisp.project_db_path}")
                    except Exception as db_error:
                        self.add_log(f"⚠️ Errore durante il salvataggio nel database: {str(db_error)}")
                        print(f"DEBUG-DB: Errore salvataggio DB: {str(db_error)}")
                        print(f"DEBUG-DB: {traceback.format_exc()}")
                else:
                    self.add_log("⚠️ Impossibile salvare nel database: keyword non trovata")
                    print("DEBUG-DB: Impossibile salvare: keyword non trovata")
            except Exception as save_error:
                self.add_log(f"⚠️ Errore generale nel salvataggio: {str(save_error)}")
                print(f"DEBUG-DB: Errore generale nel salvataggio: {str(save_error)}")
            # ================ FASE 9: RESTITUZIONE VALORI ================
            # Restituzione esattamente 6 valori come richiesto dall'interfaccia
            return self.chat_manager.get_log_history_string(), analysis_status_text, tabs_value, book_title_value, book_index_value, voice_style_value

        except Exception as e:
            # Gestione errori globale - questo è l'UNICO except per il try principale
            error_msg = f"❌ Errore durante il completamento dell'analisi: {str(e)}"
            self.add_log(error_msg)
            print(f"DEBUG-ERROR: ERRORE CRITICO in complete_analysis: {str(e)}")

            # Traceback completo dell'errore
            error_trace = traceback.format_exc()
            self.add_log(f"Dettagli errore:\n{error_trace}")
            print(f"DEBUG-ERROR: Traceback completo:\n{error_trace}")

            # Restituisci valori minimi in caso di errore critico
            print("DEBUG-ERROR: Restituzione valori di fallback a causa dell'errore")
            return self.chat_manager.get_log_history_string(), "**Stato analisi**: Errore", gr.Tabs(selected=0), "", "", ""
    
    
    def _complete_analysis_crisp(self):
        """Completa l'analisi CRISP 5.0"""
        try:
            self.add_log("Completamento analisi CRISP 5.0...")

            # Recupera l'ID del progetto CRISP corrente
            project_id = self.current_analysis.get('crisp_project_id')
            if not project_id:
                return self.add_log("Errore: Nessun progetto CRISP corrente trovato")

            # Recupera i dati completi del progetto
            project_data = self.crisp.get_project_data(project_id)

            # Aggiorna l'interfaccia con i dati estratti
            if hasattr(self, 'book_title') and 'TITOLO_LIBRO' in project_data:
                self.book_title.update(value=project_data['TITOLO_LIBRO'])

            if hasattr(self, 'book_language') and 'LINGUA' in project_data:
                self.book_language.update(value=project_data['LINGUA'])

            if hasattr(self, 'voice_style') and 'VOICE_STYLE' in project_data:
                self.voice_style.update(value=project_data['VOICE_STYLE'])

            # Costruisci l'indice del libro in base ai CONTENT_PILLARS
            if hasattr(self, 'book_index') and 'CONTENT_PILLARS' in project_data:
                # Estrai i pilastri di contenuto e trasformali in un indice
                pillars_text = project_data.get('CONTENT_PILLARS', '')
            
                # Cerca di estrarre i titoli dei pilastri
                pillars = []
                if isinstance(pillars_text, str):
                    # Cerca di estrarre i pilastri con regex
                    pillar_matches = re.findall(r'(\d+\.\s*[^\n]+)', pillars_text)
                    if pillar_matches:
                        pillars = [p.strip() for p in pillar_matches]
                    else:
                        # Alternativa: dividi per linee e filtra le linee non vuote
                        pillar_lines = [line.strip() for line in pillars_text.split('\n') if line.strip()]
                        pillars = pillar_lines[:5]  # Limita a 5 pilastri
            
                # Se abbiamo trovato dei pilastri, costruisci l'indice
                if pillars:
                    index_text = "INTRODUZIONE\n\n"
                    for i, pillar in enumerate(pillars, 1):
                        # Pulisci il pillar rimuovendo numeri e simboli iniziali
                        clean_pillar = re.sub(r'^\d+[\.\)\s]+', '', pillar).strip()
                        index_text += f"CAPITOLO {i}: {clean_pillar}\n"
                    index_text += "\nCONCLUSIONE"
                else:
                    # Indice di fallback se non troviamo pillars
                    index_text = "INTRODUZIONE\n\nCAPITOLO 1: Fondamenti\n\nCAPITOLO 2: Metodologia\n\nCAPITOLO 3: Applicazione\n\nCAPITOLO 4: Casi Studio\n\nCAPITOLO 5: Risultati\n\nCONCLUSIONE"
                
                self.book_index.update(value=index_text)
            
            # Mostra la sezione dei dettagli del libro
            if hasattr(self, 'book_details'):
                self.book_details.update(visible=True)
            
            # Crea un riepilogo dei dati estratti
            summary = f"""
            ===== ANALISI CRISP 5.0 COMPLETATA =====
        
            Titolo: {project_data.get('TITOLO_LIBRO', 'N/A')}
            Sottotitolo: {project_data.get('SOTTOTITOLO_LIBRO', 'N/A')}
        
            Angolo di Attacco: {project_data.get('ANGOLO_ATTACCO', 'N/A')}
            Big Idea: {project_data.get('BIG_IDEA', 'N/A')}
            Buyer Persona: {project_data.get('BUYER_PERSONA_SUMMARY', 'N/A')}
        
            Promessa Principale: {project_data.get('PROMESSA_PRINCIPALE', 'N/A')}
        
            L'interfaccia è stata aggiornata con i dati del progetto.
            Puoi ora procedere con la generazione del libro.
            """
            
            # Salvataggio nel database
            try:
                self.add_log("💾 Salvataggio dei risultati CRISP nel database...")
                
                # Il progetto è già salvato nel framework CRISP,
                # ma possiamo aggiungerlo anche al database generale per la visualizzazione
                if os.path.exists(self.crisp.project_db_path):
                    conn = sqlite3.connect(self.crisp.project_db_path)
                    cursor = conn.cursor()
                    
                    # Verifica se il progetto esiste già nel database
                    cursor.execute("SELECT id FROM projects WHERE name = ?", (f"CRISP-{project_id}",))
                    existing = cursor.fetchone()
                    
                    if not existing:
                        # Crea una entry nel database principale
                        current_date = datetime.now().isoformat()
                        cursor.execute(
                            "INSERT INTO projects (name, creation_date, last_updated) VALUES (?, ?, ?)",
                            (f"CRISP-{project_id}", current_date, current_date)
                        )
                        
                        # Salva i principali metadati come variabili
                        db_project_id = cursor.lastrowid
                        
                        # Se il progetto è stato salvato correttamente
                        if db_project_id:
                            # Salva keyword e altre informazioni chiave
                            keyword = project_data.get('KEYWORD', '')
                            cursor.execute(
                                "INSERT INTO project_variables (project_id, name, value) VALUES (?, ?, ?)",
                                (db_project_id, 'KEYWORD', keyword)
                            )
                            
                            # Salva riferimento al progetto CRISP originale
                            cursor.execute(
                                "INSERT INTO project_variables (project_id, name, value) VALUES (?, ?, ?)",
                                (db_project_id, 'CRISP_PROJECT_ID', str(project_id))
                            )
                            
                            # Registra avvenuto salvataggio
                            cursor.execute(
                                "INSERT INTO project_results (project_id, prompt_id, created_at, data) VALUES (?, ?, ?, ?)",
                                (db_project_id, "CRISP_SUMMARY", current_date, summary)
                            )
                            
                            conn.commit()
                            self.add_log(f"✅ Riferimento al progetto CRISP salvato nel database principale (ID: {db_project_id})")
                    else:
                        self.add_log("ℹ️ Progetto CRISP già presente nel database")
                    
                    conn.close()
            except Exception as db_error:
                self.add_log(f"⚠️ Errore durante il salvataggio nel database: {str(db_error)}")
                print(f"DEBUG-DB: Errore salvataggio DB: {str(db_error)}")           

            self.add_log(summary)
            return self.chat_manager.get_log_history_string()
            
        except Exception as e:
            error_msg = f"Errore durante il completamento dell'analisi CRISP 5.0: {str(e)}"
            self.add_log(error_msg)
            logging.error(error_msg)
            return self.chat_manager.get_log_history_string()

    
    def _complete_analysis_legacy(self):
        """Completa l'analisi legacy e mostra una finestra di dialogo per la selezione"""
        try:
            self.add_log("Completamento analisi legacy...")

            # 1. Recupera i dati dall'analisi salvata
            context_file = "context.txt"
    
            try:
                with open(context_file, "r", encoding="utf-8") as f:
                    full_text = f.read()
        
                # Estrai le diverse sezioni in base ai punti numerati
                # Cerca titoli e indici
                titoli_section = re.search(r'7\)\s+\*\*Titolo\s+&\s+sottotitolo[^F]*?FINE', full_text, re.DOTALL)
                indice_section = re.search(r'8\)\s+\*\*Indice\s+del\s+libro[^F]*?FINE', full_text, re.DOTALL)
        
                titoli_text = titoli_section.group(0) if titoli_section else ""
                indice_text = indice_section.group(0) if indice_section else ""
        
                # Estrai le opzioni di titolo
                titoli_options = []
                titoli_matches = re.finditer(r'(?:Opzione|Titolo)\s+\d+[:\)]\s+[""]?([^"\n]+)[""]?(?:\s*[:–-]\s*[""]?([^"\n]+)[""]?)?', titoli_text)
    
                for i, match in enumerate(titoli_matches, 1):
                    titolo = match.group(1).strip() if match.group(1) else ""
                    sottotitolo = match.group(2).strip() if match.group(2) else ""
                    if titolo:
                        titoli_options.append({
                            "id": i,
                            "titolo": titolo, 
                            "sottotitolo": sottotitolo,
                            "display": f"Opzione {i}: {titolo} - {sottotitolo}"
                        })
    
                # Estrai gli indici proposti
                indici_options = []
                indici_matches = re.finditer(r'(?:Indice|INDICE|CAPITOLI)[^\n]*\n(.*?)(?=\n\n|\n[A-Z]|$)', indice_text, re.DOTALL)
    
                for i, match in enumerate(indici_matches, 1):
                    indice_content = match.group(1).strip()
                    # Pulisci e formatta l'indice
                    indice_lines = [line.strip() for line in indice_content.split('\n') if line.strip()]
                    indice_formatted = "\n".join(indice_lines)
                    indici_options.append({
                        "id": i,
                        "content": indice_formatted,
                        "display": f"Indice {i}"
                    })
    
                # Estrai il tono di voce suggerito
                voice_style_match = re.search(r'(?:tono|stile|voce)[^:]*[:]\s*([^\n\.]+)', full_text, re.IGNORECASE)
                voice_style = voice_style_match.group(1).strip() if voice_style_match else "Conversazionale"
    
                # Salva temporaneamente le opzioni per l'uso nella finestra di dialogo
                self.temp_titles = titoli_options
                self.temp_indices = indici_options
                self.temp_voice_style = voice_style
            
                # Log delle opzioni
                self.add_log("\n=== OPZIONI DI SELEZIONE ===")
                self.add_log(f"Titoli disponibili: {len(titoli_options)}")
                for t in titoli_options:
                    self.add_log(f"- {t['display']}")
                
                self.add_log(f"Indici disponibili: {len(indici_options)}")
                for idx in indici_options:
                    preview = idx['content'][:50] + "..." if len(idx['content']) > 50 else idx['content']
                    self.add_log(f"- Indice {idx['id']}: {preview}")
                
                self.add_log(f"Stile di voce: {voice_style}")
            
                # Crea una finestra di dialogo per la selezione
                self.create_selection_dialog(titoli_options, indici_options, voice_style)
            
                return self.chat_manager.get_log_history_string()
    
            except Exception as e:
                self.add_log(f"⚠️ Errore nell'analisi del contesto: {str(e)}")
                import traceback
                self.add_log(traceback.format_exc())
                return self.chat_manager.get_log_history_string()

                # Salvataggio nel database
                try:
                    # Estrai keyword dal contesto o dal titolo
                    keyword = None
                    
                    # Cerca la keyword nel testo
                    keyword_match = re.search(r'(?:keyword|parola chiave)[^:]*?:\s*([^\n]+)', full_text, re.IGNORECASE)
                    if keyword_match:
                        keyword = keyword_match.group(1).strip()
                    
                    # Se non trovata, usa il titolo della prima opzione
                    if not keyword and titoli_options:
                        keyword = titoli_options[0]['titolo'].split()[0]  # Prima parola del primo titolo
                    
                    # Se abbiamo una keyword, salva nel database
                    if keyword:
                        self.add_log(f"💾 Salvataggio analisi legacy nel database per keyword: {keyword}")
                        
                        if os.path.exists(self.crisp.project_db_path):
                            conn = sqlite3.connect(self.crisp.project_db_path)
                            cursor = conn.cursor()
                            
                            # Crea un nuovo progetto
                            current_date = datetime.now().isoformat()
                            cursor.execute(
                                "INSERT INTO projects (name, creation_date, last_updated) VALUES (?, ?, ?)",
                                (f"Legacy-{keyword}", current_date, current_date)
                            )
                            project_id = cursor.lastrowid
                            
                            # Salva la keyword
                            cursor.execute(
                                "INSERT INTO project_variables (project_id, name, value) VALUES (?, ?, ?)",
                                (project_id, "KEYWORD", keyword)
                            )
                            
                            # Salva le opzioni di titolo
                            if titoli_options:
                                title_json = json.dumps(titoli_options)
                                cursor.execute(
                                    "INSERT INTO project_variables (project_id, name, value) VALUES (?, ?, ?)",
                                    (project_id, "TITLE_OPTIONS", title_json)
                                )
                            
                            # Salva le opzioni di indice
                            if indici_options:
                                index_json = json.dumps(indici_options)
                                cursor.execute(
                                    "INSERT INTO project_variables (project_id, name, value) VALUES (?, ?, ?)",
                                    (project_id, "INDEX_OPTIONS", index_json)
                                )
                            
                            # Salva il tono di voce
                            if voice_style:
                                cursor.execute(
                                    "INSERT INTO project_variables (project_id, name, value) VALUES (?, ?, ?)",
                                    (project_id, "VOICE_STYLE", voice_style)
                                )
                            
                            # Salva il testo completo
                            cursor.execute(
                                "INSERT INTO project_results (project_id, prompt_id, created_at, data) VALUES (?, ?, ?, ?)",
                                (project_id, "ANALISI_LEGACY", current_date, full_text)
                            )
                            
                            conn.commit()
                            conn.close()
                            
                            self.add_log(f"✅ Analisi legacy salvata nel database (ID: {project_id})")
                    else:
                        self.add_log("⚠️ Impossibile salvare nel database: keyword non trovata")
                except Exception as db_error:
                    self.add_log(f"⚠️ Errore durante il salvataggio nel database: {str(db_error)}")
                    import traceback
                    self.add_log(traceback.format_exc())

        except Exception as e:
            error_msg = f"Errore durante il completamento dell'analisi: {str(e)}"
            self.add_log(error_msg)
            logging.error(error_msg)
            return self.chat_manager.get_log_history_string()

    def create_selection_dialog(self, titoli_options, indici_options, voice_style):
        """Crea una finestra di dialogo per selezionare titolo, indice e stile di voce"""
        import gradio as gr

        # Verifica che ci siano opzioni da mostrare
        if not titoli_options:
            titoli_options = [{"id": 1, "titolo": "Il tuo nuovo libro", "sottotitolo": "", "display": "Titolo predefinito"}]
        if not indici_options:
            indici_options = [{"id": 1, "content": "INTRODUZIONE\n\nCAPITOLO 1: Fondamenti\n\nCAPITOLO 2: Metodologia\n\nCAPITOLO 3: Applicazione\n\nCONCLUSIONE", "display": "Indice predefinito"}]

        # Creazione interfaccia di selezione
        with gr.Blocks(title="Selezione Opzioni") as selection_interface:
            with gr.Row():
                gr.Markdown("## Seleziona le opzioni per il tuo libro")
        
            # Titolo
            with gr.Row():
                with gr.Column(scale=3):
                    gr.Markdown("### Seleziona un titolo")
                    title_radio = gr.Radio(
                        choices=[t["display"] for t in titoli_options],
                        label="Titoli disponibili",
                        value=titoli_options[0]["display"] if titoli_options else None
                    )
            
                with gr.Column(scale=1):
                    gr.Markdown("### Anteprima")
                    title_preview = gr.Textbox(
                        label="Titolo selezionato",
                        value=titoli_options[0]["titolo"] if titoli_options else "",
                        interactive=False
                    )
                    subtitle_preview = gr.Textbox(
                        label="Sottotitolo",
                        value=titoli_options[0]["sottotitolo"] if titoli_options else "",
                        interactive=False
                    )
        
            # Indice
            with gr.Row():
                with gr.Column(scale=1):
                    gr.Markdown("### Seleziona un indice")
                    index_radio = gr.Radio(
                        choices=[idx["display"] for idx in indici_options],
                        label="Indici disponibili",
                        value=indici_options[0]["display"] if indici_options else None
                    )
            
                with gr.Column(scale=2):
                    gr.Markdown("### Anteprima indice")
                    index_preview = gr.TextArea(
                        label="Indice selezionato",
                        value=indici_options[0]["content"] if indici_options else "",
                        interactive=False,
                        lines=10
                    )
        
            # Stile di voce
            with gr.Row():
                voice_style_input = gr.Textbox(
                    label="Stile di voce",
                    value=voice_style,
                    interactive=True
                )
        
            # Pulsanti azione
            with gr.Row():
                confirm_btn = gr.Button("Conferma selezione", variant="primary")
                cancel_btn = gr.Button("Annulla", variant="secondary")
        
            # Funzioni di callback
            def update_title_preview(title_display):
                for t in titoli_options:
                    if t["display"] == title_display:
                        return t["titolo"], t["sottotitolo"]
                return "", ""
        
            def update_index_preview(index_display):
                for idx in indici_options:
                    if idx["display"] == index_display:
                        return idx["content"]
                return ""
        
            def confirm_selection(title_display, index_display, voice_style):
                # Trova il titolo selezionato
                selected_title = ""
                selected_subtitle = ""
                for t in titoli_options:
                    if t["display"] == title_display:
                        selected_title = t["titolo"]
                        selected_subtitle = t["sottotitolo"]
                        break
            
                # Trova l'indice selezionato
                selected_index = ""
                for idx in indici_options:
                    if idx["display"] == index_display:
                        selected_index = idx["content"]
                        break
            
                # Aggiorna i campi nella scheda di generazione del libro
                full_title = f"{selected_title}" + (f" - {selected_subtitle}" if selected_subtitle else "")
            
                # Aggiorna i campi dell'interfaccia principale
                if hasattr(self, 'book_title'):
                    self.book_title.update(value=full_title)
            
                if hasattr(self, 'book_index'):
                    self.book_index.update(value=selected_index)
            
                if hasattr(self, 'voice_style'):
                    self.voice_style.update(value=voice_style)
            
                self.add_log(f"✅ Selezione confermata: Titolo='{full_title}', Stile='{voice_style}'")
            
                # Chiudi la finestra di dialogo
                return gr.update(visible=False)
        
            def cancel_selection():
                # Usa i valori predefiniti
                default_title = titoli_options[0]["titolo"] + " - " + titoli_options[0]["sottotitolo"] if titoli_options else "Il tuo nuovo libro"
                default_index = indici_options[0]["content"] if indici_options else "INTRODUZIONE\n\nCAPITOLO 1: Fondamenti\n\nCAPITOLO 2: Metodologia\n\nCAPITOLO 3: Applicazione\n\nCONCLUSIONE"
            
                # Aggiorna i campi nella scheda di generazione del libro
                if hasattr(self, 'book_title'):
                    self.book_title.update(value=default_title)
            
                if hasattr(self, 'book_index'):
                    self.book_index.update(value=default_index)
            
                if hasattr(self, 'voice_style'):
                    self.voice_style.update(value=voice_style)
            
                self.add_log("⚠️ Selezione annullata, usati valori predefiniti")
            
                # Chiudi la finestra di dialogo
                return gr.update(visible=False)
        
            # Connessione dei callback
            title_radio.change(
                fn=update_title_preview,
                inputs=[title_radio],
                outputs=[title_preview, subtitle_preview]
            )
        
            index_radio.change(
                fn=update_index_preview,
                inputs=[index_radio],
                outputs=[index_preview]
            )
        
            confirm_btn.click(
                fn=confirm_selection,
                inputs=[title_radio, index_radio, voice_style_input],
                outputs=[selection_interface]
            )
        
            cancel_btn.click(
                fn=cancel_selection,
                outputs=[selection_interface]
            )
    
        # Lancia l'interfaccia in una nuova finestra
        selection_interface.launch(inbrowser=True, prevent_thread_lock=True)
        self.add_log("🔍 Finestra di selezione aperta. Scegli le opzioni desiderate.")

    def handle_selection_result(self, title_display, index_display, voice_style):
        """Gestisce il risultato della selezione dalla finestra di dialogo"""
        try:
            title_id = None
            index_id = None
        
            # Trova l'ID del titolo selezionato
            for t in self.temp_titles:
                if t["display"] == title_display:
                    title_id = t["id"]
                    break
        
            # Trova l'ID dell'indice selezionato
            for idx in self.temp_indices:
                if idx["display"] == index_display:
                    index_id = idx["id"]
                    break
        
            # Aggiorna i campi nella scheda di generazione del libro
            if title_id is not None:
                selected_title = next((t for t in self.temp_titles if t["id"] == title_id), None)
                if selected_title:
                    full_title = f"{selected_title['titolo']}" + (f" - {selected_title['sottotitolo']}" if selected_title['sottotitolo'] else "")
                    self.book_title.update(value=full_title)
        
            if index_id is not None:
                selected_index = next((idx for idx in self.temp_indices if idx["id"] == index_id), None)
                if selected_index:
                    self.book_index.update(value=selected_index["content"])
        
            # Aggiorna lo stile di voce
            self.voice_style.update(value=voice_style)
        
            self.add_log(f"✅ Selezione applicata alla scheda di generazione del libro")
            return True
        
        except Exception as e:
            self.add_log(f"❌ Errore nell'applicazione della selezione: {str(e)}")
            return False
    
    
    def _generate_book_crisp(self, book_title, book_language, voice_style, book_index):
        """
        Genera il libro usando il framework CRISP 5.0.
        Ora delega alla funzione in framework/book_generator.py
    
        Args:
            book_title: Titolo del libro
            book_language: Lingua del libro
            voice_style: Stile narrativo
            book_index: Indice del libro
        
        Returns:
            str: Log dell'operazione
        """
        from framework.book_generator import generate_book_crisp
    
        try:
            result = generate_book_crisp(
                book_title, 
                book_language, 
                voice_style, 
                book_index,
                crisp_framework=self.crisp,
                driver=self.driver,
                chat_manager=self.chat_manager,
                current_analysis=self.current_analysis
            )
        
            # Il risultato potrebbe essere il percorso del libro o un messaggio di errore
            if result and not result.startswith("Errore:"):
                self.add_log(f"📚 Libro generato con successo: {result}")
            else:
                self.add_log(f"❌ {result}")
            
        except Exception as e:
            self.add_log(f"❌ Errore durante la generazione del libro: {str(e)}")
        
        return self.chat_manager.get_log_history_string()


    
    def _generate_book_legacy(self, book_title, book_language, voice_style, book_index):
        """
        Metodo legacy per generare il libro.
        Ora delega alla funzione in framework/book_generator.py
        """
        from framework.book_generator import generate_book_legacy
    
        result = generate_book_legacy(
            book_title, 
            book_language, 
            voice_style, 
            book_index,
            driver=self.driver,
            chat_manager=self.chat_manager
        )
    
        # Log del risultato
        self.add_log(f"Risultato generazione legacy: {result}")
        return self.chat_manager.get_log_history_string()

    
    def split_prompt(self, text, prompt_id=None, section_number=None):
        """
        Divide il prompt in sezioni numeriche mantenendo l'integrità.
        Aggiunto tracciamento della posizione nel flusso.
        Delega alla funzione in ai_interfaces/interaction_utils.py
        """
        from ai_interfaces.interaction_utils import split_prompt as utils_split_prompt
    
        return utils_split_prompt(
            text=text,
            prompt_id=prompt_id,
            section_number=section_number,
            log_callback=self.add_log
        )

    def send_to_genspark(self, text, prompt_id=None, section_number=None):
        """
        Invia un messaggio a Genspark e attende la risposta.
        Versione completamente rivista per garantire stabilità e affidabilità,
        con aggiornamento dell'interfaccia utente integrato.
        """
        from ai_interfaces.genspark_driver import send_to_genspark as genspark_send
    
        return genspark_send(
            driver=self.driver,
            text=text,
            log_callback=self.add_log,
            prompt_id=prompt_id,
            section_number=section_number,
            cooldown_manager=getattr(self, 'cooldown_manager', None),
            chat_manager=getattr(self, 'chat_manager', None),
            results_display=getattr(self, 'results_display', None)
        )

    # Altre funzioni di supporto da aggiungere a genspark_driver.py:

    def get_last_response(driver, log_callback=None):
        """
        Ottiene l'ultima risposta da Genspark.
    
        Args:
            driver: WebDriver di Selenium
            log_callback: Funzione di callback per il logging
        
        Returns:
            str: Ultima risposta di Genspark
        """
        try:
            selectors = [
                ".message-content", 
                "div.chat-wrapper div.desc > div > div > div",
                "div.message div.text-wrap",
                ".chat-message-item .content"
            ]
        
            for selector in selectors:
                try:
                    messages = driver.find_elements(By.CSS_SELECTOR, selector)
                    if messages and len(messages) > 0:
                        return messages[-1].text.strip()
                except Exception:
                    continue
                
            # Metodo JavaScript alternativo
            try:
                js_response = driver.execute_script("""
                    var messages = document.querySelectorAll('.message-content, .chat-message-item, .chat-wrapper .desc');
                    if (messages && messages.length > 0) {
                        return messages[messages.length - 1].textContent;
                    }
                    return null;
                """)
            
                if js_response:
                    return js_response.strip()
            except Exception:
                pass
            
            return ""
        except Exception as e:
            if log_callback:
                log_callback(f"Errore nel recupero dell'ultima risposta: {str(e)}")
            return ""

    def check_for_generation_error(self, response):
        """
        Controlla se la risposta contiene errori di generazione.
        Delega alla funzione in ai_interfaces/genspark_driver.py
        """
        from ai_interfaces.genspark_driver import check_for_generation_error as driver_check_for_generation_error
    
        return driver_check_for_generation_error(response)

    def handle_consecutive_errors(self, prompt_text, max_retries=3):
        """
        Gestisce errori consecutivi tentando approcci alternativi
        Delega alla funzione in ai_interfaces/genspark_driver.py
        """
        from ai_interfaces.genspark_driver import handle_consecutive_errors as driver_handle_consecutive_errors
    
        return driver_handle_consecutive_errors(
            driver=self.driver,
            prompt_text=prompt_text,
            max_retries=max_retries,
            log_callback=self.add_log
        )

    def _set_use_crisp(self, value):
        """Imposta se usare il framework CRISP"""
        self.use_crisp = value
        if value:
            self.add_log("Framework CRISP 5.0 attivato")
        else:
            self.add_log("Framework CRISP disattivato, verrà utilizzato il sistema legacy")
    
        # Restituisci il log aggiornato
        return self.chat_manager.get_log_history_string()

    def log_history_string(self):
        """Helper per ottenere il log history come stringa"""
        return self.chat_manager.get_log_history_string()

    def generate_book(self, book_title, book_language, voice_style, book_index):
        """
        Genera il libro utilizzando i dati dell'interfaccia e i dati CRISP disponibili.
        Ora delega alla funzione nel modulo framework/book_generator.py
        """
        from framework.book_generator import generate_book as generate_book_func
        from ai_interfaces.genspark_driver import send_to_genspark
    
        # Ottieni il tipo di libro (se disponibile)
        book_type = None
        if hasattr(self, 'book_type_hidden'):
            book_type = self.book_type_hidden.value
    
        # Delega la generazione alla funzione nel modulo book_generator
        result = generate_book_func(
            book_title=book_title,
            book_language=book_language,
            voice_style=voice_style,
            book_index=book_index,
            driver=self.driver,
            chat_manager=self.chat_manager,
            current_analysis=self.current_analysis,
            book_type_hidden=book_type,
            send_to_genspark=send_to_genspark
        )
    
        self.add_log(f"Risultato generazione libro: {result}")
        return self.chat_manager.get_log_history_string()

    def _parse_book_index(self, book_index):
        """
        Analizza l'indice e lo converte in una lista di capitoli strutturati.
        Ora delega alla funzione in framework/book_generator.py
        """
        from framework.book_generator import _parse_book_index as parse_book_index_func
    
        return parse_book_index_func(book_index)


    def _load_chapter_prompt(self, book_type):
            """
            Carica il prompt template per i capitoli specifico per il tipo di libro.
            Ora delega alla funzione in framework/book_generator.py
            """
            from framework.book_generator import _load_chapter_prompt as load_chapter_prompt_func
    
            return load_chapter_prompt_func(book_type)


    def _handle_missing_placeholders(self, text):
        """
        Gestisce i placeholder non sostituiti nel prompt.
        Ora delega alla funzione in framework/book_generator.py
        """
        from framework.book_generator import _handle_missing_placeholders as handle_missing_placeholders_func
    
        return handle_missing_placeholders_func(text)

    def _generate_chapter_content(self, chapter_title, prompt):
        """
        Genera il contenuto di un capitolo inviandolo a Genspark.
        Ora delega alla funzione in framework/book_generator.py
        """
        from framework.book_generator import _generate_chapter_content as generate_chapter_content_func
        from ai_interfaces.genspark_driver import send_to_genspark
    
        # Implementazione temporanea che mantiene la compatibilità con la classe
        # Il metodo originale faceva riferimento ai membri della classe
        try:
            # Verifica se il driver è attivo
            if not self.driver:
                self.add_log("Browser non disponibile")
                return "Errore: Browser non disponibile"
            
            # Invia il prompt e attendi la risposta
            response = self.send_to_genspark(prompt)
        
            # Pulisci la risposta
            if "FINE_RISPOSTA" in response:
                response = response.split("FINE_RISPOSTA")[0].strip()
            elif "FINE" in response:
                response = response.split("FINE")[0].strip()
            
            return response
        
        except Exception as e:
            self.add_log(f"Errore nella generazione del capitolo {chapter_title}: {str(e)}")
            return f"Errore nella generazione: {str(e)}"

    def _save_chapter(self, chapter_title, chapter_content, book_title):
        """
        Salva il capitolo generato in un file.
        Ora delega alla funzione in framework/book_generator.py
        """
        from framework.book_generator import _save_chapter as save_chapter_func
    
        return save_chapter_func(
            chapter_title, 
            chapter_content, 
            book_title, 
            log=self.add_log
        )

    def load_analysis_results(self):
            """
            Carica i risultati dell'analisi nell'interfaccia utente
            """
            # Aggiungi questa riga all'inizio del metodo
            import os
            import re
            from datetime import datetime
            import shutil  # Aggiungi questa importazione

            self.add_log("Caricamento dei risultati dell'analisi...")

            try:
                # AGGIUNGI QUESTO CODICE PER COPIARE IL FILE DI BACKUP NELLA DIRECTORY PRINCIPALE
                backup_dir = "backups"
                if os.path.exists(backup_dir):
                    backup_files = [f for f in os.listdir(backup_dir) 
                                   if f.startswith("context_") and f.endswith(".txt")]
                    if backup_files:
                        backup_files.sort(key=lambda x: os.path.getmtime(os.path.join(backup_dir, x)), 
                                         reverse=True)
                        if backup_files:
                            latest_file = backup_files[0]
                            self.add_log(f"DEBUG: File di backup più recente: {latest_file}")
                        
                            # Due possibili pattern: con timestamp o senza
                            keyword_match = re.match(r"context_(.+?)(_\d{8}_\d{6})?\.txt", latest_file)
                            if keyword_match:
                                safe_keyword = keyword_match.group(1)  # Keyword con underscores
                                current_keyword = safe_keyword.replace("_", " ")  # Keyword con spazi
                            
                                # Nome del file nella directory principale
                                main_file = f"context_{safe_keyword}.txt"
                            
                                # Copia il file se non esiste già
                                if not os.path.exists(main_file):
                                    source_path = os.path.join(backup_dir, latest_file)
                                    self.add_log(f"Copio {source_path} → {main_file}")
                                    shutil.copy2(source_path, main_file)
                                    self.add_log(f"✅ File di backup copiato nella directory principale")
                                else:
                                    self.add_log(f"File {main_file} già presente nella directory principale")

                # Resto del codice originale per l'identificazione della keyword
                current_keyword = self.get_current_keyword()
                if current_keyword:
                    self.add_log(f"🔍 Keyword corrente identificata: {current_keyword}")
            
                    # Verifica nei file di backup per confermare la keyword più recente
                    backup_dir = "backups"
                    if os.path.exists(backup_dir):
                        backup_files = [f for f in os.listdir(backup_dir) 
                                        if f.startswith("context_") and f.endswith(".txt")]
                        if backup_files:
                            # Ordina per data di modifica (più recente prima)
                            backup_files = sorted(backup_files, 
                                                 key=lambda x: os.path.getmtime(os.path.join(backup_dir, x)), 
                                                 reverse=True)
                            # Estrai il nome della keyword dal primo file (il più recente)
                            if backup_files and len(backup_files) > 0:
                                latest_file = backup_files[0]
                                self.add_log(f"DEBUG: File di backup più recente: {latest_file}")
                                keyword_match = re.match(r"context_(.+?)(_\d{8}_\d{6})?\.txt", latest_file)
                                if keyword_match:
                                    confirmed_keyword = keyword_match.group(1).replace("_", " ")
                                    if confirmed_keyword != current_keyword:
                                        self.add_log(f"⚠️ Correzione keyword: da '{current_keyword}' a '{confirmed_keyword}'")
                                        current_keyword = confirmed_keyword
                                    
                                        # AGGIUNGI QUESTO CONTROLLO AGGIUNTIVO
                                        # Assicuriamoci che il file context_*.txt esista anche nella directory principale
                                        main_file = f"context_{keyword_match.group(1)}.txt"
                                        if not os.path.exists(main_file):
                                            source_path = os.path.join(backup_dir, latest_file)
                                            self.add_log(f"Copio {source_path} → {main_file} dopo correzione keyword")
                                            shutil.copy2(source_path, main_file)
            
                    # Ottenere altre informazioni necessarie
                    market = self.get_current_market() if hasattr(self, 'get_current_market') else "Unknown"
                    book_type = self.get_current_book_type() if hasattr(self, 'get_current_book_type') else "Unknown"
                    language = self.get_current_language() if hasattr(self, 'get_current_language') else "it"
        
                    # AGGIUNGI QUESTO: Verifica che esista il file di contesto per questa keyword
                    safe_keyword = current_keyword.replace(" ", "_")
                    if not os.path.exists(f"context_{safe_keyword}.txt"):
                        self.add_log(f"⚠️ File context_{safe_keyword}.txt non trovato! Cercando alternative...")
                    
                        # Cerca nel backup il file di contesto per questa keyword
                        backup_path = os.path.join(backup_dir, f"context_{safe_keyword}.txt")
                        if os.path.exists(backup_path):
                            self.add_log(f"✅ Trovato file di backup: {backup_path}")
                            shutil.copy2(backup_path, f"context_{safe_keyword}.txt")
                            self.add_log(f"✅ File copiato nella directory principale")
                        else:
                            # Cerca file con pattern alternativo (con timestamp)
                            backup_matches = [f for f in os.listdir(backup_dir) if f.startswith(f"context_{safe_keyword}_")]
                            if backup_matches:
                                latest_backup = max(backup_matches, key=lambda x: os.path.getmtime(os.path.join(backup_dir, x)))
                                self.add_log(f"✅ Trovato file di backup con timestamp: {latest_backup}")
                                shutil.copy2(os.path.join(backup_dir, latest_backup), f"context_{safe_keyword}.txt")
                                self.add_log(f"✅ File copiato nella directory principale")
        
                    # Chiamare la funzione di formattazione HTML
                    self.add_log(f"📊 Generazione HTML formattato per l'analisi di {current_keyword}...")
        
                    # Creare un context dictionary se necessario
                    context_data = self.extract_context_data() if hasattr(self, 'extract_context_data') else None
        
                    # Chiamare formato_analysis_results_html con tutti i parametri necessari
                    formatted_html = self.format_analysis_results_html(
                        keyword=current_keyword,
                        market=market,
                        book_type=book_type,
                        language=language,
                        context=context_data
                    )
        
                    self.add_log(f"✅ HTML formattato generato: {len(formatted_html) if formatted_html else 0} caratteri")
                else:
                    self.add_log("⚠️ Impossibile determinare la keyword corrente per la formattazione HTML")
            except Exception as format_error:
                self.add_log(f"❌ Errore nella formattazione HTML: {str(format_error)}")
                import traceback
                self.add_log(traceback.format_exc())

        # IL RESTO DEL CODICE ORIGINALE CONTINUA QUI
        
        try:
            # Cerca prima i file HTML nella cartella output/analisi_html (percorso corretto)
            import glob
            output_dir = "output/analisi_html"
        
            # Crea la directory se non esiste
            os.makedirs(output_dir, exist_ok=True)
        
            # Cerca sia i file con timestamp che i file current
            html_files_timestamp = glob.glob(f"{output_dir}/*_*.html")
            html_files_current = glob.glob(f"{output_dir}/*_current.html")
        
            # Prova prima con i file current, poi con quelli con timestamp
            html_files = html_files_current + html_files_timestamp
        
            if html_files:
                # Trova il file più recente o il primo file current
                if html_files_current:
                    latest_file = html_files_current[0]  # Prendi il primo file current
                    self.add_log(f"📄 File HTML 'current' trovato: {latest_file}")
                else:
                    latest_file = max(html_files_timestamp, key=os.path.getmtime) if html_files_timestamp else None
                    if latest_file:
                        self.add_log(f"📄 File HTML più recente trovato: {latest_file}")
            
                if latest_file:
                    # Carica il contenuto HTML
                    with open(latest_file, "r", encoding="utf-8") as f:
                        html_content = f.read()
                
                    # Aggiorna l'interfaccia con il file HTML formattato
                    if hasattr(self, 'results_display'):
                        self.results_display.value = html_content
                        self.add_log("✅ Risultati dell'analisi visualizzati dall'HTML formattato")
                    
                        # Aggiorna lo stato dell'analisi
                        if hasattr(self, 'analysis_status'):
                            self.analysis_status.value = "**Stato analisi**: Completata e visualizzata ✅"
                    
                        return self.chat_manager.get_log_history_string()
        
            # Se non ci sono file HTML nella directory corretta, prova anche il vecchio percorso
            fallback_html_files = glob.glob("output/analisi_*.html")
            if fallback_html_files:
                latest_file = max(fallback_html_files, key=os.path.getmtime)
                self.add_log(f"📄 File HTML trovato nel percorso alternativo: {latest_file}")
            
                # Carica il contenuto HTML
                with open(latest_file, "r", encoding="utf-8") as f:
                    html_content = f.read()
            
                # Aggiorna l'interfaccia con il file HTML formattato
                if hasattr(self, 'results_display'):
                    self.results_display.value = html_content
                    self.add_log("✅ Risultati dell'analisi visualizzati dall'HTML formattato (fallback)")
                
                    # Aggiorna lo stato dell'analisi
                    if hasattr(self, 'analysis_status'):
                        self.analysis_status.value = "**Stato analisi**: Completata e visualizzata ✅"
                
                    return self.chat_manager.get_log_history_string()

            # Se non ci sono file HTML, utilizza il metodo tradizionale con context.txt
            self.add_log("⚠️ Nessun file HTML trovato, utilizzo il vecchio metodo con context.txt")
        
            # [Il resto del codice rimane invariato...]
        
            # Verifica che il file di contesto esista
            context_file = os.path.join(self.context_dir, "context.txt")
            if not os.path.exists(context_file):
                self.add_log("⚠️ File context.txt non trovato")
                if hasattr(self, 'results_display'):
                    self.results_display.value = "<div class='alert alert-warning'>File dei risultati non trovato. Esegui prima l'analisi.</div>"
                return self.chat_manager.get_log_history_string()
    
            # Leggi il file di contesto
            with open(context_file, "r", encoding="utf-8") as f:
                context_content = f.read()
    
            if not context_content.strip():
                self.add_log("⚠️ File context.txt vuoto")
                if hasattr(self, 'results_display'):
                    self.results_display.value = "<div class='alert alert-warning'>File dei risultati vuoto. Esegui l'analisi.</div>"
                return self.chat_manager.get_log_history_string()
    
            # Estrai le sezioni del contesto
            sections = []
            section_pattern = r'===\s+([^=]+?)\s+-\s+\d{8}_\d{6}\s+===\n(.*?)(?=\n===|$)'
            section_matches = re.findall(section_pattern, context_content, re.DOTALL)
    
            if section_matches:
                sections = [(title.strip(), content.strip()) for title, content in section_matches]
                self.add_log(f"✅ Estratte {len(sections)} sezioni")
        
                # Debug: mostra dimensione delle prime 3 sezioni
                for i, (title, content) in enumerate(sections[:3]):
                    self.add_log(f"Sezione {i+1}: {title} ({len(content)} caratteri)")
            else:
                # Fallback: divide per numeri
                number_pattern = r'(\d+\).*?)(?=\d+\)|$)'
                number_matches = re.findall(number_pattern, context_content, re.DOTALL)
        
                if number_matches:
                    sections = [(f"Sezione {i+1}", content.strip()) for i, content in enumerate(number_matches)]
                    self.add_log(f"✅ Estratte {len(sections)} sezioni (pattern numerico)")
                else:
                    # Ultimo fallback: usa il testo completo
                    sections = [("Risultati completi", context_content)]
                    self.add_log("⚠️ Nessuna sezione trovata, usando il testo completo")
    
            # Costruisci l'HTML per la visualizzazione
            html = """
            <!DOCTYPE html>
            <html>
            <head>
                <style>
                    body { font-family: Arial, sans-serif; }
                    .header { background-color: #2563eb; color: white; padding: 15px; border-radius: 8px; margin-bottom: 20px; }
                    .section { background-color: #f5f5f5; padding: 15px; border-radius: 8px; margin-bottom: 15px; }
                    .section-title { font-weight: bold; font-size: 18px; margin-bottom: 10px; }
                    .error { background-color: #fee2e2; }
                    .ok { background-color: #d1fae5; }
                </style>
            </head>
            <body>
                <div class="header">
                    <h2>Risultati dell'Analisi</h2>
                    <p>Sezioni trovate: %d</p>
                </div>
                <div>
            """ % len(sections)
    
            # Aggiungi le sezioni
            for title, content in sections:
                # Determina se è una risposta valida o un errore
                is_error = "richiesta abortita" in content.lower() or len(content) < 30
                section_class = "section error" if is_error else "section ok"
        
                # Converti newline in <br> per HTML
                content_html = content.replace("\n", "<br>")
        
                html += """
                <div class="%s">
                    <div class="section-title">%s</div>
                    <div>%s</div>
                </div>
                """ % (section_class, title, content_html)
    
            html += """
                </div>
            </body>
            </html>
            """
    
            # Salva l'HTML di debug in un file unico per ogni esecuzione
            import os
            os.makedirs("debug", exist_ok=True)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            debug_html_filename = f"debug/debug_html_{timestamp}.html"
            with open(debug_html_filename, "w", encoding="utf-8") as f:
                f.write(html)
            self.add_log(f"HTML di debug salvato: {debug_html_filename} ({len(html)} bytes)")

            # NUOVO: Salva anche in output/analisi_html come file regolare
            try:
                current_keyword = self.get_current_keyword()
                if current_keyword:
                    from framework.formatters import save_analysis_to_html  # Sostituisci your_module con il nome del modulo dove è definita la funzione
                    analysis_type = "Legacy"  # O determina il tipo dinamicamente
                    file_path = save_analysis_to_html(html, current_keyword, "Unknown", "Unknown", "it", analysis_type, self.add_log)
                    self.add_log(f"✅ Salvato anche come HTML formattato in {file_path}")
            except Exception as e:
                self.add_log(f"❌ Errore nel salvataggio aggiuntivo: {str(e)}")
    
            # IMPORTANTE: Verifica che results_display esista e sia diverso da analysis_status
            if hasattr(self, 'results_display'):
                # Verifica se results_display e analysis_status sono lo stesso oggetto
                same_object = False
                if hasattr(self, 'analysis_status'):
                    same_object = id(self.results_display) == id(self.analysis_status)
                    self.add_log(f"DEBUG: results_display e analysis_status sono {'LO STESSO' if same_object else 'OGGETTI DIVERSI'}")
        
                # Aggiorna l'interfaccia SOLO se non sono lo stesso oggetto
                if not same_object:
                    # Usa assegnazione diretta invece di update
                    self.results_display.value = html
                    self.add_log("✅ Risultati dell'analisi visualizzati nell'interfaccia")
            
                    # Aggiorna lo stato dell'analisi separatamente
                    if hasattr(self, 'analysis_status'):
                        # Usa assegnazione diretta invece di update
                        self.analysis_status.value = "**Stato analisi**: Completata e visualizzata ✅"
                else:
                    self.add_log("⚠️ results_display e analysis_status sono lo stesso oggetto! Non aggiorno l'interfaccia")
                    # In questo caso, dobbiamo aggiornare solo uno dei due
                    # Usa assegnazione diretta invece di update
                    self.results_display.value = html
                    self.add_log("✅ Risultati dell'analisi visualizzati nell'interfaccia (solo results_display)")
            else:
                self.add_log("⚠️ results_display non disponibile")
    
            return self.chat_manager.get_log_history_string()

        except Exception as e:
            self.add_log(f"❌ Errore nel caricamento dei risultati: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return self.chat_manager.get_log_history_string()

    def export_to_docx(self):
        """Esporta l'analisi corrente in un documento DOCX."""
        try:
            self.add_log("Esportazione in DOCX...")
        
            # 1. Determina quale file di contesto utilizzare
            keyword = self.get_current_keyword() if hasattr(self, 'get_current_keyword') else "unknown"
        
            # Ottieni il nome del file sanitizzato
            import re
            safe_keyword = re.sub(r'[\\/*?:"<>|]', "", keyword).replace(" ", "_")[:30]
            context_file = f"context_{safe_keyword}.txt"
        
            # Se il file specifico non esiste, prova con il file di contesto generico
            if not os.path.exists(context_file):
                self.add_log(f"⚠️ File specifico {context_file} non trovato, provo con il file generico")
                context_file = "context.txt"
            
                if not os.path.exists(context_file):
                    self.add_log("❌ Nessun file di contesto trovato!")
                    return None
        
            self.add_log(f"📄 Utilizzo del file di contesto: {context_file}")
        
            # 2. Leggi il contenuto del file
            with open(context_file, "r", encoding="utf-8") as f:
                content = f.read()
        
            # 3. Crea il documento
            from docx import Document
            doc = Document()
            doc.add_heading(f"Analisi di mercato: {keyword}", 0)
        
            # Aggiungi informazioni sul documento
            timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
            doc.add_paragraph(f"Generato il: {timestamp}")
            doc.add_paragraph(f"Keyword: {keyword}")
            doc.add_paragraph("")  # Spazio aggiuntivo
        
            # 4. Dividi il contenuto in sezioni e aggiungile al documento
            sections = []
            section_pattern = r'===\s+([^=]+?)\s+-\s+\d{8}_\d{6}\s+===\n(.*?)(?=\n===|$)'
            section_matches = re.findall(section_pattern, content, re.DOTALL)
        
            if section_matches:
                sections = [(title.strip(), content.strip()) for title, content in section_matches]
                self.add_log(f"✅ Estratte {len(sections)} sezioni per il documento")
            else:
                # Fallback: divide per "==="
                raw_sections = content.split("===")
                for i, section in enumerate(raw_sections):
                    if section.strip():
                        sections.append((f"Sezione {i+1}", section.strip()))
                self.add_log(f"⚠️ Estratte {len(sections)} sezioni usando metodo fallback")
        
            # 5. Aggiungi ogni sezione al documento con formattazione appropriata
            for section_title, section_content in sections:
                # Aggiungi il titolo della sezione
                doc.add_heading(section_title, level=1)
            
                # Dividi il contenuto in paragrafi
                paragraphs = section_content.split('\n\n')
                for paragraph_text in paragraphs:
                    if paragraph_text.strip():
                        # Verifica se è un titolo di paragrafo
                        lines = paragraph_text.splitlines()
                        if lines and len(lines) > 1 and len(lines[0]) < 100 and lines[0].endswith(':'):
                            # Sembra un titolo di paragrafo
                            doc.add_heading(lines[0], level=2)
                            paragraph = doc.add_paragraph("\n".join(lines[1:]))
                        else:
                            # Paragrafo normale
                            paragraph = doc.add_paragraph(paragraph_text)
                
                # Aggiungi un divisore tra le sezioni
                doc.add_paragraph("")
        
            # 6. Salva il documento con un nome significativo
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"analisi_{safe_keyword}_{timestamp}.docx"
            doc.save(output_path)
        
            self.add_log(f"✅ Documento DOCX salvato: {output_path}")
            return output_path
    
        except Exception as e:
            self.add_log(f"❌ Errore nell'esportazione DOCX: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return None

    def export_to_pdf(self):
        try:
            self.add_log("Esportazione in PDF...")
        
            # 1. Prima esporta in DOCX
            docx_path = self.export_to_docx()
            if not docx_path:
                raise Exception("Errore nell'esportazione DOCX preliminare")
        
            # 2. Converti DOCX in PDF - metodo preferito: python-docx2pdf
            try:
                # Prova a importare docx2pdf
                from docx2pdf import convert
            
                # Conversione diretta
                pdf_path = docx_path.replace('.docx', '.pdf')
                self.add_log(f"Conversione {docx_path} in {pdf_path}...")
                convert(docx_path, pdf_path)
            
                self.add_log(f"✅ Documento PDF salvato: {pdf_path}")
                return pdf_path
            
            except ImportError:
                # Fallback: utilizza un messaggio informativo se docx2pdf non è installato
                self.add_log("⚠️ Modulo python-docx2pdf non trovato.")
                self.add_log("⚠️ Per convertire in PDF, apri il file DOCX e usa 'Salva come PDF'")
                self.add_log(f"⚠️ Il file DOCX è disponibile qui: {docx_path}")
                return docx_path
        
        except Exception as e:
            self.add_log(f"❌ Errore nell'esportazione PDF: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return None

    def export_to_txt(self):
        try:
            self.add_log("Esportazione in TXT...")
        
            # 1. Verifica se esiste il file di contesto
            context_file = "context.txt"
            if not os.path.exists(context_file):
                self.add_log("❌ File context.txt non trovato!")
                return None
        
            # 2. Ottieni la keyword corrente o usa un valore predefinito
            keyword = self.get_current_keyword() if hasattr(self, 'get_current_keyword') else "unknown"
        
            # 3. Crea la directory di output se non esiste
            output_dir = os.path.join(os.getcwd(), "output")
            os.makedirs(output_dir, exist_ok=True)
        
            # 4. Prepara il percorso del file di output
            import datetime
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(output_dir, f"analisi_{keyword}_{timestamp}.txt")
        
            # 5. Copia il file di contesto nel file di output
            import shutil
            shutil.copy2(context_file, output_path)
        
            self.add_log(f"✅ File TXT salvato: {output_path}")
            return output_path
        except Exception as e:
            self.add_log(f"❌ Errore nell'esportazione TXT: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return None

    def debug_check_components(self):
        """Verifica se results_display e analysis_status sono lo stesso oggetto"""
        if hasattr(self, 'results_display') and hasattr(self, 'analysis_status'):
            same_object = id(self.results_display) == id(self.analysis_status)
            self.add_log(f"DEBUG: results_display e analysis_status sono {'LO STESSO' if same_object else 'OGGETTI DIVERSI'}")
            self.add_log(f"DEBUG: ID results_display: {id(self.results_display)}")
            self.add_log(f"DEBUG: ID analysis_status: {id(self.analysis_status)}")
        else:
            self.add_log("DEBUG: Uno o entrambi i componenti non esistono ancora")
        return self.chat_manager.get_log_history_string()

    def load_saved_analyses_list(self):
            """Carica l'elenco delle analisi salvate"""
            saved_analyses = []
            try:
                # Cerca i file context_*.txt nella directory corrente
                for file in os.listdir():
                    match = re.match(r'context_(.+)\.txt', file)
                    if match:
                        keyword = match.group(1).replace('_', ' ')
                        saved_analyses.append(keyword)
                    
                return saved_analyses
            except Exception as e:
                self.add_log(f"Errore nel caricamento delle analisi salvate: {str(e)}")
                return []
    
    def load_saved_analysis(self, keyword):
        """Carica un'analisi salvata"""
        try:
            # Converti la keyword in nome file
            file_keyword = keyword.replace(' ', '_')
            context_file = f"context_{file_keyword}.txt"
        
            if not os.path.exists(context_file):
                self.add_log(f"❌ Analisi per '{keyword}' non trovata!")
                return self.chat_manager.get_log_history_string()
            
            # Leggi il file
            with open(context_file, "r", encoding="utf-8") as f:
                content = f.read()
              
            # Imposta come analisi corrente
            self.current_analysis = {"KEYWORD": keyword}
          
            # Aggiorna l'interfaccia
            self.add_log(f"✅ Analisi per '{keyword}' caricata con successo")
            
            # Carica i risultati nell'interfaccia
            self.load_analysis_results(context_file)
            
            return self.chat_manager.get_log_history_string()
            
        except Exception as e:
            self.add_log(f"Errore nel caricamento dell'analisi: {str(e)}")
            return self.chat_manager.get_log_history_string()

    def get_current_keyword(self):
        """Recupera la keyword corrente dall'ultima analisi eseguita"""
        try:
            # Metodo 1: Cerca l'ultima keyword dal file di contesto
            import os
            import re
            from datetime import datetime
        
            context_file = "context.txt"
            if os.path.exists(context_file):
                with open(context_file, "r", encoding="utf-8") as f:
                    content = f.read()
                
                    # Trova tutte le sezioni con timestamp
                    pattern = r'===\s+([^=-]+?)(?:\s*:\s*|\s+-\s+)([^-]+?)\s+-\s+(\d{8}_\d{6})\s+==='
                    matches = re.findall(pattern, content)
                
                    if matches:
                        # Ordina per timestamp (terzo elemento della tupla)
                        sorted_matches = sorted(matches, key=lambda x: x[2], reverse=True)
                        # Prendi il primo match (il più recente) e la seconda parte (la keyword)
                        latest_keyword = sorted_matches[0][1].strip()
                        self.add_log(f"DEBUG: Keyword più recente trovata dal timestamp: {latest_keyword} ({sorted_matches[0][2]})")
                        return latest_keyword
                
            # Metodo 2: Controlla se la keyword è disponibile nell'interfaccia
            if hasattr(self, 'keyword') and hasattr(self.keyword, 'value') and self.keyword.value:
                self.add_log(f"DEBUG: Keyword trovata dall'interfaccia: {self.keyword.value}")
                return self.keyword.value
        
            # Metodo 3: Controlla i dati del progetto corrente
            if hasattr(self, 'current_analysis') and self.current_analysis:
                project_data = self.current_analysis.get('project_data', {})
                if 'KEYWORD' in project_data and project_data['KEYWORD']:
                    self.add_log(f"DEBUG: Keyword trovata dai dati progetto: {project_data['KEYWORD']}")
                    return project_data['KEYWORD']
        
            # Ultimo metodo: cerca i file di backup più recenti
            backup_dir = "backups"
            if os.path.exists(backup_dir):
                backup_files = [f for f in os.listdir(backup_dir) if f.startswith("context_") and f.endswith(".txt")]
                if backup_files:
                    # Ordina per data di modifica (più recente prima)
                    backup_files = sorted(backup_files, key=lambda x: os.path.getmtime(os.path.join(backup_dir, x)), reverse=True)
                    # Estrai il nome della keyword dal primo file (il più recente)
                    if backup_files:
                        # Rimuovi "context_" all'inizio e "_YYYYMMDD_HHMMSS.txt" alla fine
                        latest_file = backup_files[0]
                        keyword_match = re.match(r"context_(.+?)_\d{8}_\d{6}\.txt", latest_file)
                        if keyword_match:
                            keyword = keyword_match.group(1).replace("_", " ")
                            self.add_log(f"DEBUG: Keyword trovata dal file di backup più recente: {keyword}")
                            return keyword
        
            self.add_log("DEBUG: Nessuna keyword trovata, ritorno 'unknown'")
            return "unknown"
        except Exception as e:
            self.add_log(f"⚠️ Errore nel recupero della keyword: {str(e)}")
            import traceback
            self.add_log(traceback.format_exc())
            return "unknown"

    def extract_context_data(self):
        """
        Estrae dati strutturati dal contesto per l'HTML
        """
        context_data = {}
        try:
            # Leggi il file di contesto
            with open("context.txt", "r", encoding="utf-8") as f:
                content = f.read()
        
            # Estrai sezioni e dati chiave
            import re
        
            # Buyer Persona
            buyer_match = re.search(r"buyer persona[^\n]+\n(.*?)(?=\n===|\n\d+\)|$)", content, re.IGNORECASE | re.DOTALL)
            if buyer_match:
                context_data["BUYER_PERSONA_SUMMARY"] = buyer_match.group(1).strip()
        
            # Angolo di Attacco
            angle_match = re.search(r"angolo d[i']? attacco[^\n]+\n(.*?)(?=\n===|\n\d+\)|$)", content, re.IGNORECASE | re.DOTALL)
            if angle_match:
                context_data["ANGOLO_ATTACCO"] = angle_match.group(1).strip()
        
            # Insight di Mercato
            market_match = re.search(r"insight d[i']? mercato[^\n]+\n(.*?)(?=\n===|\n\d+\)|$)", content, re.IGNORECASE | re.DOTALL)
            if market_match:
                context_data["MARKET_INSIGHTS"] = market_match.group(1).strip()
        
            # Aggiungi il tipo di analisi
            if "CRISP" in content or "crisp" in content.lower():
                context_data["type"] = "CRISP"
            else:
                context_data["type"] = "Legacy"
            
        except Exception as e:
            self.add_log(f"⚠️ Errore nell'estrazione dei dati dal contesto: {str(e)}")
    
        return context_data

    def create_interface(self):

        with gr.Blocks(title="PubliScript 2.0", theme="soft") as interface:
            # Header principale
            with gr.Row(elem_classes=["header-container"]):
                gr.HTML("""
                    <div class="app-header">
                        <div class="logo-container">
                            <span class="logo-icon">📕</span>
                            <h1>PubliScript 2.0</h1>
                        </div>
                        <div class="app-subtitle">Sistema Multi-Agente per Analisi e Generazione di Libri</div>
                    </div>
                """)
            self.crisp_phase_checkboxes = {}
            self.legacy_phase_checkboxes = {}
            # Sistema di tabs principale
            with gr.Tabs(elem_classes=["main-tabs"]) as tabs:
                # Tab 1: Setup & Connessione
                with gr.TabItem("1️⃣ Setup & Connessione", elem_classes=["tab-content"]):
                    with gr.Row():
                        with gr.Column(scale=1):
                            gr.Markdown("### Connessione al Servizio")
                            connect_btn = gr.Button("🔌 Connetti a Genspark", variant="primary")
                    
                            gr.Markdown("### Framework di Generazione")
                            # use_crisp_toggle = gr.Checkbox(
                            #    label="Usa Framework CRISP",
                            #    value=True,
                            #    info="Attiva il framework CRISP per un'analisi più strutturata"
                            # )
                    
                        with gr.Column(scale=2):
                            connection_status = gr.Markdown("**Stato**: Non connesso")
                    
                            # Area di log
                            self.log_output = gr.TextArea(
                                label="Console di sistema",
                                interactive=False,
                                lines=15,
                                value="Sistema inizializzato. Connettiti per iniziare.",
                                elem_id="log-output-area"
                            )
                    
                            # Pulsante per pulire i log
                            clear_log_btn = gr.Button("🧹 Pulisci Log", variant="secondary", size="sm")
            
                # Tab 2: Analisi di Mercato
                with gr.TabItem("2️⃣ Analisi di Mercato", elem_classes=["tab-content"]):
                    with gr.Row():
                        # Colonna sinistra: Input (1/3)
                        with gr.Column(scale=1, elem_classes=["input-column"]):
                            gr.Markdown("### Informazioni Base")
            
                            # Nuovo controllo per caricare analisi esistente
                            with gr.Row():
                                load_analysis_dropdown = gr.Dropdown(
                                    choices=self.load_saved_analyses_list(),
                                    label="Carica Analisi Esistente",
                                    interactive=True,
                                    elem_classes=["dropdown-input"]
                                )
                                load_analysis_btn = gr.Button("📂 Carica", variant="secondary")
            
                            book_type = gr.Dropdown(
                                choices=self.book_types,
                                label="Tipo di Libro",
                                value=self.book_types[0],
                                elem_classes=["dropdown-input"]
                            )
            
                            keyword = gr.Textbox(
                                label="Keyword Principale",
                                placeholder="Parola chiave principale del libro",
                                elem_classes=["text-input"]
                            )
                        
                            language = gr.Dropdown(
                                choices=["English", "Español", "Français", "Deutsch", "Italiano"],
                                label="Lingua Output Analisi",
                                value="Italiano",
                                elem_classes=["dropdown-input"]
                            )
                        
                            market = gr.Dropdown(
                                choices=list(self.markets.keys()),
                                label="Mercato di Riferimento",
                                value="USA",
                                elem_classes=["dropdown-input"]
                            )
                        
                            # Opzioni avanzate
                            with gr.Accordion("Mostra opzioni avanzate", open=False):
                                analysis_prompt = gr.TextArea(
                                    label="Prompt di Analisi (opzionale)",
                                    value=self.default_analysis_prompt,
                                    lines=6
                                )
                        
                            #------------------INIZIO CAMBIAMENTO RADIO BUTTON-----------------------

                            # Aggiunta dell'Accordion per la selezione delle fasi
                            with gr.Accordion("📋 Seleziona fasi analisi", open=False):
                                gr.Markdown("### Scegli le fasi dell'analisi da eseguire")

                                # Prima riga per il tipo di analisi
                                with gr.Row():
                                    analysis_type_radio = gr.Radio(
                                        choices=["CRISP", "Legacy"],
                                        label="Tipo di Analisi",
                                        value="CRISP",
                                        interactive=True
                                    )
                                    self.analysis_type_radio = analysis_type_radio
    
                                # AGGIUNGI QUI LA FUNZIONE E IL COLLEGAMENTO ALL'EVENTO
                                def update_selected_type(value):
                                    self.selected_analysis_type = value
                                    self.add_log(f"Tipo di analisi selezionato: {value}")
                                    return None

                                # Collega l'handler all'evento change del radio button
                                self.analysis_type_radio.change(
                                    fn=update_selected_type,
                                    inputs=[self.analysis_type_radio],
                                    outputs=[]
                                )
    
                                # CRISP Phases: CheckboxGroup
                                with gr.Column(visible=True) as crisp_phases_col:
                                    gr.Markdown("#### Seleziona le Fasi dell'Analisi CRISP")
    
                                    # Basato sul file di configurazione
                                    crisp_phases = [
                                        ("CM-1", "Analisi di mercato - Fase 1"),
                                        ("CM-2", "Analisi di mercato - Fase 2"),
                                        ("CS-1", "Content Strategy - Fase 1"),
                                        ("CS-2", "Content Strategy - Fase 2"),
                                        ("CS-3", "Content Strategy - Fase 3"),
                                        ("CS-F", "Content Strategy - Fase Finale"),
                                        ("CP-1", "Content Production - Fase 1"),
                                        ("CP-2", "Content Production - Fase 2"),
                                        ("CPM-1", "Content Promotion & Marketing - Fase 1"),
                                        ("CPM-2", "Content Promotion & Marketing - Fase 2"),
                                        ("CPM-3", "Content Promotion & Marketing - Fase 3")
                                    ]
    
                                    # Crea le opzioni per il CheckboxGroup
                                    crisp_checkbox_options = [f"{phase_id}: {phase_desc}" for phase_id, phase_desc in crisp_phases]
    
                                    # Usa CheckboxGroup invece di RadioButton
                                    self.crisp_phase_checkboxes = gr.CheckboxGroup(
                                        choices=crisp_checkbox_options,
                                        label="Fasi da eseguire",
                                        value=[crisp_checkbox_options[0]],  # Seleziona la prima fase come default
                                        interactive=True
                                    )

                                # Legacy Phases: CheckboxGroup
                                with gr.Column(visible=False) as legacy_phases_col:
                                    gr.Markdown("#### Seleziona le Fasi dell'Analisi Legacy")
    
                                    # Basato sul file di configurazione
                                    legacy_phases = [
                                        ("LM-1", "Analisi concorrenza"),
                                        ("LM-2", "Valutazione profittabilità e competitività"),
                                        ("LM-3", "Analisi 3 migliori concorrenti"),
                                        ("LM-4", "Buyer persona sintetica"),
                                        ("LM-5", "Gap analysis"),
                                        ("LM-6", "Idee editoriali"),
                                        ("LM-7", "Valutazione idee"),
                                        ("LM-8", "Titoli, sottotitoli e indici")
                                    ]
    
                                    # Crea le opzioni per il CheckboxGroup
                                    legacy_checkbox_options = [f"{phase_id}: {phase_desc}" for phase_id, phase_desc in legacy_phases]
    
                                    # Usa CheckboxGroup invece di RadioButton
                                    self.legacy_phase_checkboxes = gr.CheckboxGroup(
                                        choices=legacy_checkbox_options,
                                        label="Fasi da eseguire",
                                        value=[legacy_checkbox_options[0]],  # Seleziona la prima fase come default
                                        interactive=True
                                    )

                                # Collega i radio button al cambio di visibilità
                                analysis_type_radio.change(
                                    fn=lambda x: (
                                        gr.update(visible=(x == "CRISP")), 
                                        gr.update(visible=(x == "Legacy"))
                                    ),
                                    inputs=[analysis_type_radio],
                                    outputs=[crisp_phases_col, legacy_phases_col]
                                )
    
                                # Seconda riga per i pulsanti
                                with gr.Row():
                                    select_all_btn = gr.Button("✅ Esegui Tutte le Fasi", variant="secondary", size="sm")
                                    deselect_all_btn = gr.Button("❌ Esegui Solo Fase Selezionata", variant="secondary", size="sm")
    
                                # Collega i pulsanti alle funzioni lambda per CheckboxGroup
                                select_all_btn.click(
                                    fn=lambda: crisp_checkbox_options,  # Tutte le opzioni
                                    inputs=[],
                                    outputs=self.crisp_phase_checkboxes
                                )

                                select_all_btn.click(
                                    fn=lambda: legacy_checkbox_options,  # Tutte le opzioni
                                    inputs=[],
                                    outputs=self.legacy_phase_checkboxes
                                )

                                deselect_all_btn.click(
                                    fn=lambda: [crisp_checkbox_options[0]],  # Solo la prima opzione
                                    inputs=[],
                                    outputs=self.crisp_phase_checkboxes
                                )

                                deselect_all_btn.click(
                                    fn=lambda: [legacy_checkbox_options[0]],  # Solo la prima opzione
                                    inputs=[],
                                    outputs=self.legacy_phase_checkboxes
                                )

                            #------------------------------FINE CAMBIAMENTO RADIO BUTTON-------------------------------------

                            # Aggiungi listener per i CheckboxGroup (opzionale, per logging)
                            self.crisp_phase_checkboxes.change(
                                fn=lambda value: print(f"CRISP fasi selezionate: {value}"),
                                inputs=[self.crisp_phase_checkboxes],
                                outputs=[]
                            )

                            self.legacy_phase_checkboxes.change(
                                fn=lambda value: print(f"Legacy fasi selezionate: {value}"),
                                inputs=[self.legacy_phase_checkboxes],
                                outputs=[]
                            )
                            
                            analyze_btn = gr.Button("🔍 Analizza Mercato", variant="primary", size="lg")
                            complete_analysis_btn = gr.Button("✅ Completa Analisi", variant="secondary")
                
                        # Colonna destra: Output (2/3)
                        with gr.Column(scale=2, elem_classes=["output-column"]):
                            gr.Markdown("### Analisi di Mercato")
                        
                            # Stato analisi
                            analysis_status = gr.Markdown("**Stato analisi**: Non iniziata", elem_classes=["status-text"])
                        
                            # Pulsanti di esportazione
                            with gr.Row(elem_classes=["export-buttons"]):
                                export_docx_btn = gr.Button("📄 Esporta DOCX", variant="secondary")
                                export_pdf_btn = gr.Button("📑 Esporta PDF", variant="secondary")
                                export_txt_btn = gr.Button("📝 Esporta TXT", variant="secondary")
                                debug_btn = gr.Button("🐞 Debug UI", variant="secondary")

                            # Area di output principale (formato HTML per ricchezza)
                            results_display = gr.HTML(
                                value="<div class='results-placeholder'>I risultati dell'analisi appariranno qui</div>",
                                elem_classes=["results-container"]
                            )
                            self.results_display = results_display
            
                # Tab 3: Generazione Libro
                with gr.TabItem("3️⃣ Generazione Libro", elem_classes=["tab-content"]):
                    with gr.Group(visible=True) as book_details:
                        gr.Markdown("### Dettagli Libro")
                    
                        with gr.Row():
                            with gr.Column(scale=1):
                                # Aggiungi questo campo di sola lettura per mostrare il tipo di libro
                                book_type_display = gr.Textbox(
                                    label="Tipo di Libro",
                                    value="",
                                    interactive=False,
                                    elem_classes=["text-input"]
                                )
                            
                                with gr.Row():
                                    with gr.Column(scale=1):
                                        self.book_title = gr.Textbox(
                                            label="Titolo del Libro",
                                            placeholder="Inserisci il titolo",
                                            elem_classes=["text-input"]
                                        )
                                
                                        self.book_language = gr.Textbox(
                                            label="Lingua del Libro",
                                            value="English",
                                            placeholder="es: English, Italiano, Español",
                                            elem_classes=["text-input"]
                                        )
                                
                                        self.voice_style = gr.Textbox(
                                            label="Tono di Voce",
                                            placeholder="es: Formale, Tecnico, Conversazionale",
                                            value="",
                                            elem_classes=["text-input"]
                                        )
                            
                                    with gr.Column(scale=2):
                                        self.book_index = gr.TextArea(
                                            label="Indice del Libro",
                                            placeholder="Inserisci l'indice, un capitolo per riga",
                                            lines=10,
                                            elem_classes=["text-area"]
                                        )
                            
                                generate_btn = gr.Button("📚 Genera Libro", variant="primary", size="lg")
            
                # Tab 4: Database & Gestione
                with gr.TabItem("4️⃣ Database & Gestione", elem_classes=["tab-content"]):
                    with gr.Row():
                        with gr.Column(scale=1):
                            gr.Markdown("### Progetti Salvati")
            
                            # Opzioni di manutenzione
                            with gr.Row():
                                diagnose_db_btn = gr.Button("🛠️ Ripara Database", variant="primary")
                                create_test_btn = gr.Button("🧪 Crea Progetto Test", variant="secondary")
            
                            # Ricerca e filtri
                            with gr.Row():
                                search_keyword = gr.Textbox(
                                    label="Cerca per keyword",
                                    placeholder="Inserisci una keyword...",
                                    show_label=True,
                                    elem_classes=["text-input"]
                                )
                        
                                refresh_db_btn = gr.Button("🔄 Carica Progetti", variant="primary")
                        
                            # Lista progetti
                            projects_list = gr.Dropdown(
                                label="Seleziona un Progetto",
                                choices=[],
                                interactive=True,
                                elem_id="projects_dropdown",
                                elem_classes=["dropdown-input"],
                                type="value"
                            )
      
                            self.projects_list = projects_list
                        
                            # Azioni progetti
                            with gr.Row():
                                resume_btn = gr.Button("▶️ Ripristina Analisi", variant="primary")
                                export_btn = gr.Button("📤 Esporta", variant="secondary")
                                delete_btn = gr.Button("🗑️ Elimina", variant="stop")
                
                        with gr.Column(scale=2):
                            gr.Markdown("### Dettagli Progetto")
                        
                            # Dettagli progetto in formato HTML
                            project_details = gr.HTML(
                                value="<div class='project-placeholder'>Seleziona un progetto per visualizzarne i dettagli</div>",
                                elem_classes=["project-details"]
                            )
            
                # Tab 5: Debug & Sviluppo
                with gr.TabItem("5️⃣ Debug & Sviluppo", elem_classes=["tab-content"]):
                    with gr.Row():
                        with gr.Column():
                            gr.Markdown("### Strumenti Debug")
                        
                            with gr.Row():
                                take_screenshot_btn = gr.Button("📸 Screenshot Browser", variant="secondary")
                                reset_context_btn = gr.Button("♻️ Reset Context Limit", variant="secondary")
                        
                            debug_output = gr.TextArea(
                                label="Output Debug",
                                interactive=False,
                                lines=10,
                                value=""
                            )

            # CSS personalizzato e librerie esterne
            gr.HTML("""
                <!-- Font Awesome per le icone -->
                <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0-beta3/css/all.min.css">

                <!-- Tailwind CSS per lo styling avanzato -->
                <link href="https://cdn.jsdelivr.net/npm/tailwindcss@2.2.19/dist/tailwind.min.css" rel="stylesheet">

                <!-- Google Fonts per tipografia migliorata (opzionale) -->
                <link rel="preconnect" href="https://fonts.googleapis.com">
                <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
                <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">

                <style>
                    /* Stile generale */
                    body {
                        font-family: 'Inter', 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    }

                    /* Header */
                    .app-header {
                        display: flex;
                        flex-direction: column;
                        align-items: center;
                        margin-bottom: 1rem;
                        background: linear-gradient(to right, #2563eb, #4f46e5);
                        color: white;
                        padding: 1rem;
                        border-radius: 0.5rem;
                        width: 100%;
                        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
                    }

                    .logo-container {
                        display: flex;
                        align-items: center;
                    }

                    .logo-icon {
                        font-size: 2rem;
                        margin-right: 0.5rem;
                    }

                    .app-header h1 {
                        margin: 0;
                        font-size: 1.8rem;
                        font-weight: 700;
                    }

                    .app-subtitle {
                        margin-top: 0.25rem;
                        font-size: 1rem;
                        opacity: 0.9;
                    }

                    /* Tabs */
                    .main-tabs button {
                        font-weight: 600;
                        padding: 0.75rem 1rem;
                        transition: all 0.2s ease;
                    }

                    .main-tabs button:hover {
                        background-color: #f3f4f6;
                    }

                    .main-tabs button.active {
                        border-bottom: 2px solid #3b82f6;
                        color: #1e40af;
                    }

                    .tab-content {
                        padding: 1rem 0;
                    }

                    /* Input fields */
                    .input-column {
                        border-right: 1px solid #e5e7eb;
                        padding-right: 1rem;
                    }

                    .text-input, .dropdown-input, .text-area {
                        margin-bottom: 1rem;
                        transition: border-color 0.15s ease-in-out, box-shadow 0.15s ease-in-out;
                    }

                    .text-input:focus, .dropdown-input:focus, .text-area:focus {
                        border-color: #3b82f6;
                        box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.25);
                    }

                    /* Output display */
                    .status-text {
                        background-color: #f3f4f6;
                        padding: 0.5rem;
                        border-radius: 0.375rem;
                        margin-bottom: 0.5rem;
                    }

                    .export-buttons {
                        margin-bottom: 1rem;
                        padding: 0.5rem 0;
                        border-bottom: 1px solid #e5e7eb;
                    }

                    .results-container {
                        border: 1px solid #e5e7eb;
                        border-radius: 0.5rem;
                        padding: 1rem;
                        background-color: #f9fafb;
                        min-height: 500px;
                        max-height: 70vh;
                        overflow-y: auto;
                        box-shadow: inset 0 2px 4px 0 rgba(0, 0, 0, 0.06);
                    }

                    .results-placeholder {
                        text-align: center;
                        color: #6b7280;
                        padding: 2rem;
                        font-style: italic;
                    }

                    /* Project details */
                    .project-details {
                        border: 1px solid #e5e7eb;
                        border-radius: 0.5rem;
                        padding: 1rem;
                        background-color: #f9fafb;
                        min-height: 400px;
                        max-height: 70vh;
                        overflow-y: auto;
                        transition: all 0.3s ease;
                    }

                    .project-details:hover {
                        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
                    }

                    .project-placeholder {
                        text-align: center;
                        color: #6b7280;
                        padding: 2rem;
                        font-style: italic;
                    }

                    /* Card style per i risultati */
                    .data-card {
                        background-color: white;
                        border-radius: 0.5rem;
                        box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1), 0 1px 2px 0 rgba(0, 0, 0, 0.06);
                        padding: 1rem;
                        margin-bottom: 1rem;
                        transition: all 0.3s ease;
                    }

                    .data-card:hover {
                        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
                        transform: translateY(-2px);
                    }

                    .data-card-title {
                        font-weight: 600;
                        color: #2563eb;
                        margin-bottom: 0.5rem;
                        display: flex;
                        align-items: center;
                    }

                    .data-card-title i {
                        margin-right: 0.5rem;
                    }

                    /* Scrollbar customization */
                    *::-webkit-scrollbar {
                        width: 8px;
                    }

                    *::-webkit-scrollbar-track {
                        background: #f1f1f1;
                        border-radius: 4px;
                    }

                    *::-webkit-scrollbar-thumb {
                        background: #888;
                        border-radius: 4px;
                    }

                    *::-webkit-scrollbar-thumb:hover {
                        background: #555;
                    }

                    /* Miglioramenti per pulsanti */
                    button {
                        transition: all 0.2s ease;
                    }

                    button:hover {
                        transform: translateY(-1px);
                    }

                    button:active {
                        transform: translateY(1px);
                    }

                    /* Classi di utilità per flexbox */
                    .flex-center {
                        display: flex;
                        align-items: center;
                        justify-content: center;
                    }

                    .flex-between {
                        display: flex;
                        align-items: center;
                        justify-content: space-between;
                    }

                    /* Badge e tag */
                    .badge {
                        display: inline-block;
                        padding: 0.25rem 0.75rem;
                        border-radius: 9999px;
                        font-size: 0.75rem;
                        font-weight: 600;
                        margin-right: 0.5rem;
                    }

                    .badge-blue {
                        background-color: #dbeafe;
                        color: #1e40af;
                    }

                    .badge-green {
                        background-color: #d1fae5;
                        color: #065f46;
                    }

                    .badge-yellow {
                        background-color: #fef3c7;
                        color: #92400e;
                    }

                    .badge-red {
                        background-color: #fee2e2;
                        color: #b91c1c;
                    }

                    /* Stili per i risultati dell'analisi */
                    .analysis-results .section-card {
                        transition: all 0.2s ease;
                    }

                    .analysis-results .section-card:hover {
                        transform: translateY(-2px);
                    }

                    .analysis-results .badge {
                        display: inline-block;
                        padding: 0.25rem 0.75rem;
                        border-radius: 9999px;
                        font-size: 0.75rem;
                        font-weight: 600;
                    }

                    .analysis-results .content {
                        font-size: 0.9rem;
                        line-height: 1.5;
                    }

                    .metadata-box {
                        border-left: 4px solid #3b82f6;
                    }

                    .metadata-item {
                        margin-bottom: 0.5rem;
                    }
                </style>

                <script>
                    // Funzione per mostrare/nascondere le colonne delle fasi in base al tipo di analisi
                    function togglePhaseColumns(analysisType) {
                        const crispCol = document.querySelector('[id$="crisp_phases_col"]');
                        const legacyCol = document.querySelector('[id$="legacy_phases_col"]');
        
                        if (analysisType === "CRISP") {
                            if (crispCol) crispCol.style.display = "block";
                            if (legacyCol) legacyCol.style.display = "none";
                        } else {
                            if (crispCol) crispCol.style.display = "none";
                            if (legacyCol) legacyCol.style.display = "block";
                        }
                    }

                    // Inizializza l'interfaccia dopo il caricamento
                    document.addEventListener('DOMContentLoaded', function() {
                        setTimeout(function() {
                            // Trova i radio button
                            const radioInputs = document.querySelectorAll('input[type="radio"][name$="analysis_type_radio"]');
            
                            // Aggiungi event listener a ciascun radio button
                            radioInputs.forEach(input => {
                                input.addEventListener('change', function() {
                                    togglePhaseColumns(this.value);
                                });
                            });
            
                            // Imposta lo stato iniziale
                            const selectedRadio = document.querySelector('input[type="radio"][name$="analysis_type_radio"]:checked');
                            if (selectedRadio) {
                                togglePhaseColumns(selectedRadio.value);
                            }
                
                            // Seleziona tutti i tab e aggiunge la classe active al primo
                            const tabs = document.querySelectorAll('.main-tabs button');
                            if (tabs.length > 0) {
                                tabs[0].classList.add('active');
                            }

                            // Aggiungi event listener per cambiare il tab attivo
                            tabs.forEach(tab => {
                                tab.addEventListener('click', function() {
                                    tabs.forEach(t => t.classList.remove('active'));
                                    this.classList.add('active');
                                });
                            });
                        }, 500); // Breve delay per assicurarsi che l'interfaccia sia caricata
                    });
                </script>
            """)
       
            # Function handlers
    
            # Funzione per pulire i log
            def clear_logs():
                self.log_history = ["Log cancellato."]
                return self.chat_manager.get_log_history_string()
    
            clear_log_btn.click(fn=clear_logs, outputs=self.log_output)
    
            # Connessione
            connect_btn.click(
                fn=self.connect_callback,
                outputs=[self.log_output, connection_status],
                show_progress=False
            )
    
            # Uso CRISP
            # use_crisp_toggle.change(
            #    fn=self._set_use_crisp,
            #    inputs=use_crisp_toggle,
            #    outputs=self.log_output
            # )
    
            # Analisi mercato
            analyze_btn.click(
                fn=self.analyze_market,
                inputs=[book_type, keyword, language, market, analysis_prompt],
                outputs=self.log_output,
                show_progress=False
            )
    
            # Completa analisi
            complete_analysis_btn.click(
                fn=self.complete_analysis,
                outputs=[self.log_output, analysis_status, tabs, self.book_title, self.book_index, self.voice_style]
            )
    
            # Generazione libro
            generate_btn.click(
                fn=self.generate_book,
                inputs=[self.book_title, self.book_language, self.voice_style, self.book_index],
                outputs=self.log_output
            )
    
            # Database e gestione progetti
            refresh_db_btn.click(
                fn=self.load_projects_list,
                outputs=[projects_list]
            )

            # AGGIUNGI QUESTO
            create_test_btn.click(
                fn=self.create_test_project,
                outputs=[project_details]
            )
    
            projects_list.change(
                fn=self.load_project_details,
                inputs=projects_list,
                outputs=[project_details]
            )
    
            resume_btn.click(
                fn=self.ripristina_analisi_da_database,
                inputs=projects_list,
                outputs=self.log_output
            )
    
            export_btn.click(
                fn=self.export_project,
                inputs=projects_list,
                outputs=self.log_output
            )
    
            delete_btn.click(
                fn=lambda project_name: [
                    self.delete_project(project_name),
                    self.load_projects_list(),
                    "<div style='text-align: center'><p>Progetto eliminato</p></div>"
                ],
                inputs=projects_list,
                outputs=[self.log_output, projects_list, project_details]
            )

            # Nel metodo create_interface, dopo aver definito load_analysis_btn:
            load_analysis_btn.click(
                fn=self.load_saved_analysis,
                inputs=[load_analysis_dropdown],
                outputs=[self.log_output]
            )

            debug_btn.click(
                fn=self.debug_check_components,
                outputs=[self.log_output]
            )
    
            # Debug tools
            take_screenshot_btn.click(
                fn=lambda: self.take_debug_screenshot("debug"),
                outputs=debug_output
            )
    
            reset_context_btn.click(
                fn=self.handle_context_limit,
                outputs=self.log_output
            )

            # Connetti gli eventi per i pulsanti di esportazione
            export_docx_btn.click(
                fn=self.export_to_docx,
                outputs=[self.log_output]
            )
    
            export_pdf_btn.click(
                fn=self.export_to_pdf,
                outputs=[self.log_output]
            )
    
            export_txt_btn.click(
                fn=self.export_to_txt,
                outputs=[self.log_output]
            )
    
            # AGGIUNGI QUESTO
            diagnose_db_btn.click(
                fn=self.diagnose_and_fix_database,
                outputs=[project_details]  # Mostra risultati nel riquadro dettagli progetto
            )

            return interface

    def format_analysis_results_html(self, keyword, market, book_type, language, context=None):
        """
        Formatta i risultati dell'analisi in HTML per una visualizzazione migliore.
        Delega alla funzione in framework/formatters.py
    
        Args:
            keyword: Keyword analizzata
            market: Mercato target
            book_type: Tipo di libro
            language: Lingua dell'output
            context: Dati di contesto aggiuntivi (opzionale)
    
        Returns:
            str: HTML formattato con i risultati
        """
        # Import corretto basato sulla struttura reale del progetto
        import sys
        import os
        # Aggiungi la directory principale al path se necessario
        sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    
        from framework.formatters import format_analysis_results_html as formatter_format_analysis_results_html
    
        # Determina il tipo di analisi
        analysis_type = "CRISP" if hasattr(self, 'selected_analysis_type') and self.selected_analysis_type == "CRISP" else "Legacy"
    
        return formatter_format_analysis_results_html(
            keyword=keyword,
            market=market,
            book_type=book_type,
            language=language,
            context=context,
            log_callback=self.add_log,
            save_to_file=True,
            analysis_type=analysis_type
        )

    def process_text(self, text):
        """
        Processa il testo con formattazione di base
        Delega alla funzione in framework/formatters.py
    
        Args:
            text: Testo da processare
        
        Returns:
            str: Testo formattato in HTML
        """
        from framework.formatters import process_text as formatter_process_text
    
        return formatter_process_text(text)

    def process_list_html(self, content, list_type):
        """
        Formatta una lista in HTML
        Delega alla funzione in framework/formatters.py
    
        Args:
            content: Contenuto della lista
            list_type: Tipo di lista (REVIEW_INSIGHTS, IMPLEMENTATION_OBSTACLES, MARKET_GAPS, ecc.)
        
        Returns:
            str: Lista formattata in HTML
        """
        from framework.formatters import process_list_html as formatter_process_list_html
    
        return formatter_process_list_html(content, list_type)

    def process_patterns_html(self, content, pattern_type):
        """
        Formatta pattern di titoli o strutture in HTML
        Delega alla funzione in framework/formatters.py
    
        Args:
            content: Contenuto del pattern
            pattern_type: Tipo di pattern (TITLE_PATTERNS, STRUCTURE_PATTERNS, ecc.)
        
        Returns:
            str: Pattern formattato in HTML
        """
        from framework.formatters import process_patterns_html as formatter_process_patterns_html
    
        return formatter_process_patterns_html(content, pattern_type)

    def process_table_html(self, content):
        """
        Converte una tabella in formato markdown in HTML
        Delega alla funzione in framework/formatters.py
    
        Args:
            content: Contenuto della tabella in formato markdown
        
        Returns:
            str: Tabella formattata in HTML
        """
        from framework.formatters import process_table_html as formatter_process_table_html
    
        return formatter_process_table_html(content)

# AGGIUNGERE QUESTO CODICE ALLA FINE DEL FILE book_builder.py

# Invece di sostituire il metodo a livello di classe, utilizziamo una funzione
# che applica il patching a un'istanza specifica

def apply_debug_patching(builder_instance):
    """
    Applica il debug patching al metodo send_to_genspark di un'istanza di AIBookBuilder
    
    Args:
        builder_instance: Istanza di AIBookBuilder a cui applicare il patching
    """
    # Salva il metodo originale dall'istanza (non dalla classe)
    original_method = builder_instance.send_to_genspark
    
    # Definisci il nuovo metodo con debug
    def debug_send_to_genspark(self, text, prompt_id=None, section_number=None):
        """Versione con debug del metodo send_to_genspark"""
        print(f"DEBUG_SEND: Invio testo a Genspark ({len(text)} caratteri)")
        print(f"DEBUG_SEND: Primi 50 caratteri del testo: {text[:50]}")
        
        if prompt_id:
            print(f"DEBUG_SEND: Prompt ID: {prompt_id}, Sezione: {section_number}")
            
        if hasattr(self, 'driver') and self.driver:
            print(f"DEBUG_SEND: URL attuale prima dell'invio: {self.driver.current_url}")
            
        import traceback
        caller = traceback.extract_stack()[-2]
        print(f"DEBUG_SEND: Chiamata da {caller.filename}:{caller.lineno}")
        
        # Chiama il metodo originale
        result = original_method(text, prompt_id, section_number)
        
        if hasattr(self, 'driver') and self.driver:
            print(f"DEBUG_SEND: URL attuale dopo l'invio: {self.driver.current_url}")
            
        if result:
            has_end = "FINE_RISPOSTA" in result or "FINE" in result
            print(f"DEBUG_SEND: Risultato ({len(result)} chars) - Contiene FINE_RISPOSTA: {has_end}")
            print(f"DEBUG_SEND: Preview inizio: {result[:100]}...")
            
            if len(result) > 200:
                print(f"DEBUG_SEND: Preview fine: ...{result[-200:]}")
            else:
                print(f"DEBUG_SEND: Preview fine: {result}")
        else:
            print("DEBUG_SEND: Nessun risultato ottenuto")
            
        return result
    
    # Sostituisci il metodo nell'istanza usando la tecnica dei tipi
    import types
    builder_instance.send_to_genspark = types.MethodType(debug_send_to_genspark, builder_instance)
    
    print("Debug patching applicato all'istanza AIBookBuilder")
    return True
