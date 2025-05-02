"""
Implementazione principale del framework CRISP.
"""

import re
import json
import logging
import sqlite3
import time
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from pathlib import Path
from datetime import datetime
from docx import Document

# Importazioni interne
from .crisp_utils import load_docx_content, extract_section, replace_variables, find_prompt_file, parse_prompt_data, save_result
from .crisp_extractors import extract_data_from_response

# crisp_framework.py
"""
Modulo che implementa il framework CRISP per la generazione di libri.
"""
# Configurazione logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger("crisp_framework")

class CRISPFramework:
    def __init__(self, prompt_dir="prompt_crisp", project_db_path=None, driver=None):
        """
        Inizializza il framework CRISP.
    
        Args:
            prompt_dir: Directory contenente i file prompt
            project_db_path: Percorso del database SQLite per i progetti
            driver: Reference al driver Selenium WebDriver
        """
        self.prompt_dir = Path(prompt_dir)
        self.project_db_path = project_db_path
        self.driver = driver  # Aggiungi questa linea
        self.prompt_cache = {}  # Cache per i prompt più usati
        
        # Carica la configurazione principale
        try:
            config_path = Path("Config.json")
            if config_path.exists():
                with open(config_path, 'r', encoding='utf-8') as f:
                    self.config = json.load(f)
                    logger.info(f"Configurazione caricata da {config_path}")
            else:
                self.config = {"prompt_dir": prompt_dir, "flow_config": "flow_config.json"}
                logger.warning(f"File Config.json non trovato, utilizzo configurazione predefinita")
        except Exception as e:
            logger.error(f"Errore nel caricamento della configurazione: {str(e)}")
            self.config = {"prompt_dir": prompt_dir, "flow_config": "flow_config.json"}
        
        # Inizializza una flow_map vuota
        self.flow_map = {}
        
        # Carica la flow_map direttamente dal file di configurazione
        self._load_flow_config()
        
        # Inizializza il database se specificato        
        if self.project_db_path:
            self._initialize_db()
            self._verify_database()
        
        logger.info(f"Framework CRISP inizializzato con directory prompt: {prompt_dir}")

    def _initialize_db(self):
        """Inizializza il database SQLite per i progetti CRISP."""
        try:
            conn = sqlite3.connect(self.project_db_path)
            cursor = conn.cursor()
            
            # Tabella dei progetti
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS projects (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                creation_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                status TEXT DEFAULT 'in_progress'
            )
            ''')
            
            # Tabella delle variabili di progetto
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS project_variables (
                project_id TEXT,
                variable_name TEXT,
                variable_value TEXT,
                FOREIGN KEY (project_id) REFERENCES projects (id),
                PRIMARY KEY (project_id, variable_name)
            )
            ''')
            
            # Tabella dei risultati
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS project_results (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_id TEXT,
                step_id TEXT,
                result_text TEXT,
                timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (project_id) REFERENCES projects (id)
            )
            ''')
   
            # Tabella delle risposte incrementali
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS incremental_responses (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_id TEXT,
                step_id TEXT,
                prompt_line TEXT,
 		response_text TEXT,
    		timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    		is_final BOOLEAN DEFAULT 0,
    		FOREIGN KEY (project_id) REFERENCES projects (id)	
	    )
	    ''')

            # Nuova tabella per le risposte ai prompt
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS prompt_responses (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_id TEXT,
                prompt_id TEXT,
                section_number INTEGER,
                response_text TEXT,
                timestamp TEXT,
                FOREIGN KEY (project_id) REFERENCES projects (id)
            )
            ''')

            conn.commit()
            conn.close()
            logger.info(f"Database inizializzato: {self.project_db_path}")
            
        except Exception as e:
            logger.error(f"Errore nell'inizializzazione del database: {str(e)}")
            raise

    def _verify_database(self):
        """Verifica che tutte le tabelle necessarie esistano nel database."""
        if not self.project_db_path:
            return
            
        try:
            conn = sqlite3.connect(self.project_db_path)
            cursor = conn.cursor()
            
            # Verifica tabella projects
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='projects'")
            if not cursor.fetchone():
                logger.warning("Tabella 'projects' non trovata nel database")
            
            # Verifica tabella project_variables
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='project_variables'")
            if not cursor.fetchone():
                logger.warning("Tabella 'project_variables' non trovata nel database")
            
            # Verifica tabella project_results
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='project_results'")
            if not cursor.fetchone():
                logger.warning("Tabella 'project_results' non trovata nel database")
            
            # Verifica tabella incremental_responses
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='incremental_responses'")
            if not cursor.fetchone():
                logger.warning("Tabella 'incremental_responses' non trovata nel database")
                
            conn.close()
            
        except Exception as e:
            logger.error(f"Errore nella verifica del database: {str(e)}")
        
   
    def _load_flow_config(self):
        """Carica la configurazione del flusso dal file flow_config.json."""
        try:
            flow_config_path = self.prompt_dir / self.config.get("flow_config", "flow_config.json")
        
            if flow_config_path.exists():
                with open(flow_config_path, 'r', encoding='utf-8') as f:
                    flow_config = json.load(f)
            
                # Debug del formato del file
                print(f"DEBUG - Flow config caricato: {json.dumps(flow_config, indent=2)[:200]}...")
            
                # Ottieni l'elenco dei file di prompt disponibili
                prompt_files = [f.stem for f in self.prompt_dir.glob("*.txt") if f.is_file()]
                logger.info(f"File di prompt trovati: {prompt_files}")
            
                # Verifica la struttura del flow_config
                steps_data = None
                if "steps" in flow_config:
                    # Formato semplice: { "steps": { ... } }
                    steps_data = flow_config["steps"]
                elif "flow" in flow_config and "steps" in flow_config["flow"]:
                    # Formato nidificato: { "flow": { "steps": { ... } } }
                    steps_data = flow_config["flow"]["steps"]
                else:
                    logger.warning("Struttura flow_config non riconosciuta")
                    return
                
                # Estrai la configurazione del flusso
                start_step = flow_config.get("start", None)
                if "flow" in flow_config and "start" in flow_config["flow"]:
                    start_step = flow_config["flow"]["start"]
            
                # Salva il punto di partenza del flusso
                if start_step:
                    self.start_step = start_step
                    logger.info(f"Start step impostato a: {self.start_step}")
            
                # Costruisci la flow_map dagli steps
                if steps_data and isinstance(steps_data, dict):
                    for step_id, step_info in steps_data.items():
                        # Verifica se esiste un file di prompt corrispondente
                        if step_id in prompt_files or step_id in ["CS5", "CS"]:  # CS e CS5 sono speciali
                            # Aggiungi lo step alla flow_map
                            next_step = step_info.get("next")
                            next_steps = [next_step] if next_step else []
                        
                            self.flow_map[step_id] = {
                                "name": step_info.get("description", f"Step {step_id}"),
                                "next": next_steps,
                                "required_vars": self._get_default_vars(step_id)
                            }
                            logger.info(f"Step {step_id} aggiunto alla flow_map")
                        else:
                            logger.warning(f"File di prompt {step_id}.txt non trovato per lo step {step_id}")
                else:
                    logger.warning("Dati degli steps non trovati o in formato non valido")
            
                # Carica anche gli estrattori dai flow_config
                if "extractors" in flow_config:
                    self.extractors = flow_config["extractors"]
                    logger.info(f"Estrattori caricati: {len(self.extractors)} configurazioni")
            
                logger.info(f"Flusso caricato da {flow_config_path}: {len(self.flow_map)} step")
                logger.info(f"Flow map contiene: {list(self.flow_map.keys())}")
            else:
                logger.warning(f"File di configurazione del flusso non trovato: {flow_config_path}")
    
        except Exception as e:
            logger.error(f"Errore nel caricamento della configurazione del flusso: {str(e)}")
            import traceback
            logger.debug(f"Traceback: {traceback.format_exc()}")
    
    def _get_default_vars(self, step_id):
        """Restituisce un insieme predefinito di variabili richieste per i prompt."""
        # Mapping di default per variabili richieste basato sui file effettivamente presenti
        default_vars = {
            # Nuovi ID dei prompt CRISP 5.0
            "CM-1": ["KEYWORD", "MERCATO"],
            "CM-2": ["KEYWORD", "MERCATO", "MARKET_INSIGHTS", "BESTSELLER_OVERVIEW"],
            "CS-1": ["KEYWORD", "MERCATO", "MARKET_INSIGHTS", "KEYWORD_DATA", "BESTSELLER_OVERVIEW", 
                    "STRUCTURE_PATTERNS", "TITLE_PATTERNS", "REVIEW_INSIGHTS", "MARKET_GAPS", 
                    "IMPLEMENTATION_OBSTACLES"],
            "CS-2": ["KEYWORD", "MERCATO", "MARKET_INSIGHTS", "REVIEW_INSIGHTS", 
                    "IMPLEMENTATION_OBSTACLES", "STRATEGIC_INSIGHTS"],
            "CS-3": ["KEYWORD", "MERCATO", "BUYER_PERSONA_SUMMARY", "STRATEGIC_INSIGHTS", 
                    "REVIEW_INSIGHTS", "MARKET_GAPS", "ONLINE_BEHAVIOR_INSIGHTS"],
            "CP-1": ["KEYWORD", "ANGOLO_ATTACCO", "BUYER_PERSONA_SUMMARY", "MARKET_GAPS", 
                    "IMPLEMENTATION_OBSTACLES", "STRUCTURE_PATTERNS"],
            "CP-2": ["KEYWORD", "ANGOLO_ATTACCO", "BUYER_PERSONA_SUMMARY", "ONLINE_BEHAVIOR_INSIGHTS", 
                    "TITLE_PATTERNS", "PROMESSA_PRINCIPALE", "BIG_IDEA", "BOOK_JOURNEY", "CONTENT_PILLARS"],
            "CPM-1": ["TITOLO_LIBRO", "SOTTOTITOLO_LIBRO", "KEYWORD", "MERCATO", "ANGOLO_ATTACCO", 
                     "BUYER_PERSONA_SUMMARY", "PROMESSA_PRINCIPALE", "ONLINE_BEHAVIOR_INSIGHTS", 
                     "REVIEW_INSIGHTS", "CONTENT_PILLARS"],
            "CPM-2": ["TITOLO_LIBRO", "SOTTOTITOLO_LIBRO", "KEYWORD", "MERCATO", "ANGOLO_ATTACCO", 
                     "BUYER_PERSONA_SUMMARY", "BIG_IDEA", "CONTENT_PILLARS", "BESTSELLER_OVERVIEW"],
            "CPM-3": ["TITOLO_LIBRO", "KEYWORD", "ANGOLO_ATTACCO", "BUYER_PERSONA_SUMMARY", 
                     "IMPLEMENTATION_OBSTACLES", "BOOK_JOURNEY", "CONTENT_PILLARS", "PROPRIETARY_METHOD"],
            "CS-F": ["KEYWORD", "MERCATO", "MARKET_INSIGHTS", "STRUCTURE_PATTERNS", "TITLE_PATTERNS", 
                    "REVIEW_INSIGHTS", "IMPLEMENTATION_OBSTACLES", "MARKET_GAPS", "STRATEGIC_INSIGHTS", 
                    "BUYER_PERSONA_SUMMARY", "ONLINE_BEHAVIOR_INSIGHTS", "ANGOLO_ATTACCO", 
                    "PROMESSA_PRINCIPALE", "USP_ELEMENTS", "BOOK_JOURNEY", "BIG_IDEA", "CONTENT_PILLARS", 
                    "PROPRIETARY_METHOD", "TITOLO_LIBRO", "SOTTOTITOLO_LIBRO", "VOICE_STYLE", 
                    "MARKETING_CLAIMS", "HEADLINE_OPTIONS", "AMAZON_DESCRIPTION", "COVER_BRIEF", 
                    "VISUAL_ELEMENTS", "BONUS_SYSTEM", "EMAIL_STRATEGY"],
        
            # Manteniamo anche i vecchi ID per retrocompatibilità
            "CM": ["KEYWORD", "MERCATO", "LINGUA"],
            "C1R": ["KEYWORD", "MERCATO", "LINGUA"],
            # ... altri vecchi ID se necessario
        }
    
        # Restituisci le variabili predefinite o un insieme vuoto
        return default_vars.get(step_id, [])

    def load_prompt(self, prompt_id):
        """
        Carica un prompt dal file DOCX e ne estrae le sezioni.
    
        Args:
            prompt_id: ID del prompt (es. "C1", "R2")
        
        Returns:
            dict: Dati del prompt (metadata, variables, content)
        """
        try:
            # Verifica se il prompt è in cache
            if prompt_id in self.prompt_cache:
                logger.info(f"Caricato prompt {prompt_id} dalla cache")
                return self.prompt_cache[prompt_id]
        
            # Trova il file del prompt
            prompt_file = find_prompt_file(prompt_id, self.prompt_dir)
            logger.info(f"Trovato file prompt {prompt_id}: {prompt_file}")
        
            # Carica il contenuto del file
            content = load_docx_content(prompt_file)
        
            # Analizza il contenuto del prompt
            prompt_data = parse_prompt_data(content)
        
            # Salva in cache
            self.prompt_cache[prompt_id] = prompt_data
        
            return prompt_data
        
        except Exception as e:
            logger.error(f"Errore caricamento prompt {prompt_id}: {str(e)}")
            raise ValueError(f"Impossibile caricare il prompt {prompt_id}: {str(e)}")

    def prepare_prompt(self, prompt_id, project_data):
        """
        Prepara un prompt sostituendo le variabili con i valori dal progetto.
    
        Args:
            prompt_id: ID del prompt (es. "C1", "R2")
            project_data: Dizionario con i dati del progetto
        
        Returns:
            str: Prompt pronto per essere inviato a Genspark
        """
        try:
            # Carica i dati del prompt
            prompt_data = self.load_prompt(prompt_id)
            prompt_content = prompt_data["content"]
            required_vars = prompt_data["variables"]
        
            # Verifica che tutte le variabili richieste siano disponibili
            missing_vars = []
            for var in required_vars:
                if var not in project_data:
                    missing_vars.append(var)
        
            if missing_vars:
                logger.warning(f"Variabili mancanti per {prompt_id}: {missing_vars}")
                # Opzionale: puoi sollevare un errore o fornire valori predefiniti
                for var in missing_vars:
                    project_data[var] = f"[MISSING: {var}]"
        
            # Sostituisci le variabili nel prompt
            final_prompt = replace_variables(prompt_content, project_data)
        
            logger.info(f"Preparato prompt {prompt_id} con {len(required_vars)} variabili")
            return final_prompt
        
        except Exception as e:
            logger.error(f"Errore preparazione prompt {prompt_id}: {str(e)}")
            raise ValueError(f"Impossibile preparare il prompt {prompt_id}: {str(e)}")

    def extract_data(self, result_text, prompt_id):
        """
        Estrae dati strutturati dal risultato di un prompt.
    
        Args:
            result_text: Testo risposta da Genspark
            prompt_id: ID del prompt che ha generato la risposta
        
        Returns:
            dict: Dati strutturati estratti dalla risposta
        """
        try:
            # Utilizza gli estrattori specializzati
            extracted_data = extract_data_from_response(result_text, prompt_id)
        
            logger.info(f"Estratti dati da risposta per {prompt_id}: {len(extracted_data)} elementi")
            return extracted_data
        
        except Exception as e:
            logger.error(f"Errore estrazione dati per {prompt_id}: {str(e)}")
            # In caso di errore, restituisci un dizionario con il testo originale
            return {"error": str(e), "raw_text": result_text}

    def execute_step(self, prompt_id, project_data, executor_func):
        """
        Esegue un singolo step del framework CRISP.
        """
        try:
            logger.info(f"Esecuzione step {prompt_id}")
        
            # Carica i dati del prompt
            prompt_data = self.load_prompt(prompt_id)
            content = prompt_data["content"]
        
            # Dividi in sezioni numerate
            sections = content.split("\n---\n")
            numbered_sections = []
            for section in sections:
                section = section.strip()
                if re.match(r'^\d+(\.\d+)?\.', section):
                    numbered_sections.append(section)
        
            if not numbered_sections:
                numbered_sections = [content]
            
            print(f"DEBUG - Trovate {len(numbered_sections)} sezioni numerate")
        
            all_responses = []
            extracted_data = {}
        
            # Processa ogni sezione
            for i, section in enumerate(numbered_sections):
                section_number = i + 1
                print(f"DEBUG - Processando sezione {section_number}/{len(numbered_sections)}")

                # Estrai il numero della sezione per un log più chiaro
                section_id = "unknown"
                section_id_match = re.match(r'^(\d+(\.\d+)?)\.', section)
                if section_id_match:
                    section_id = section_id_match.group(1)
                print(f"DEBUG - Processando sezione ID: {section_id}")

                # Identifica se questa è una delle sezioni OUTPUT
                is_market_insights = "MARKET_INSIGHTS" in section
                is_keyword_data = "KEYWORD_DATA" in section
                is_bestseller_overview = "BESTSELLER_OVERVIEW" in section
                is_output_section = is_market_insights or is_keyword_data or is_bestseller_overview or "DEFINIZIONE DEGLI OUTPUT" in section

                if is_output_section:
                    print(f"DEBUG - Rilevata sezione di output: {section_id}")
                    # Crea un contesto riassuntivo per la sezione di output
                    context_prompt = f"""
                    Riepilogo analisi per {project_data.get('KEYWORD', '')} su {project_data.get('MERCATO', '')}:
        
                    Dati principali raccolti:
                    {self._format_extracted_data(extracted_data)}
        
                    Ora procedi con la definizione dell'output richiesto, mantenendo la risposta MOLTO CONCISA.
                    """
        
                    # Invia il contesto
                    executor_func(context_prompt)
                    time.sleep(5)  # Breve pausa

                # Sostituisci le variabili nella sezione
                processed_section = replace_variables(section.strip(), project_data)
    
                # Aggiungi le sintesi appropriate in base al tipo di sezione
                if is_market_insights or is_keyword_data or is_bestseller_overview or "DEFINIZIONE DEGLI OUTPUT" in section:
                    # Ottieni la sintesi delle risposte precedenti
                    if project_data.get("PROJECT_ID"):
                        # Adatta l'estrazione in base al tipo di sezione
                        max_chars = 1500  # Riduci per rendere più concisa
                        if is_market_insights:
                            # Per MARKET_INSIGHTS, estrai principalmente dalle sezioni 1 e 5
                            previous_responses = self.get_targeted_responses_summary(
                                project_data["PROJECT_ID"],
                                prompt_id, 
                                target_sections=[1, 5],
                                max_chars=max_chars
                            )
                        elif is_keyword_data:
                            # Per KEYWORD_DATA, estrai principalmente dalla sezione 2
                            previous_responses = self.get_targeted_responses_summary(
                                project_data["PROJECT_ID"],
                                prompt_id, 
                                target_sections=[2],
                                max_chars=max_chars
                            )
                        elif is_bestseller_overview:
                            # Per BESTSELLER_OVERVIEW, estrai principalmente dalla sezione 4
                            previous_responses = self.get_targeted_responses_summary(
                                project_data["PROJECT_ID"],
                                prompt_id, 
                                target_sections=[4],
                                max_chars=max_chars
                            )
                        else:
                            # Per la sezione generale, usa tutte le risposte precedenti
                            previous_responses = self.get_previous_responses_summary(
                                project_data["PROJECT_ID"],
                                prompt_id,
                                max_chars=max_chars
                            )
            
                        # Aggiungi la sintesi al prompt con enfasi sulla brevità
                        processed_section = f"{processed_section}\n\nBasa la tua sintesi su queste analisi precedenti, ma mantieni la risposta dettagliata (massimo 400 parole):\n{previous_responses}"
    
                # Invia la sezione a Genspark
                response = executor_func(processed_section, prompt_id, section_number)
    
                # Salva la risposta nel database se il progetto ha un ID
                if response and project_data.get("PROJECT_ID"):
                    self.save_prompt_response(
                        project_data["PROJECT_ID"],
                        prompt_id,
                        section_number,
                        response
                    )

                if response:
                    # Pulisci la risposta
                    if "FINE_RISPOSTA" in response:
                        response = response.split("FINE_RISPOSTA")[0].strip()
                    elif "FINE" in response:
                        response = response.split("FINE")[0].strip()
                
                    # Estrai i dati
                    section_data = self.extract_data(response, prompt_id)
                    extracted_data.update(section_data)
                
                    all_responses.append(response)
                    print(f"DEBUG - Risposta sezione {section_number} completata: {len(response)} caratteri")
                else:
                    print(f"DEBUG - Nessuna risposta per sezione {section_number}")
                
                    # Se è la sezione OUTPUT e non abbiamo risposta, prova un approccio semplificato
                    if is_output_section:
                        print("DEBUG - Tentativo semplificato per sezione OUTPUT")
                        simplified_prompt = """
                        Fornisci una sintesi strutturata dei dati analizzati, includendo:
                        1. MARKET_INSIGHTS
                        2. KEYWORD_DATA
                        3. BESTSELLER_OVERVIEW
                    
                        Usa i dati già analizzati nelle sezioni precedenti.
                        """
                        retry_response = executor_func(simplified_prompt, prompt_id, section_number)
                        if retry_response:
                            all_responses.append(retry_response)
                            section_data = self.extract_data(retry_response, prompt_id)
                            extracted_data.update(section_data)
            
                # Pausa tra le sezioni
                time.sleep(5)
        
            # Unisci tutte le risposte
            full_response = "\n\n".join(all_responses)

            # AGGIUNGI QUESTI DEBUG QUI
            print(f"DEBUG - prompt_id: {prompt_id}")
            print(f"DEBUG - numbered_sections contiene: {[section[:30] + '...' for section in numbered_sections]}")
            test_match = any(re.search(r'6\.\d+\.', section.strip()) for section in numbered_sections)
            print(f"DEBUG - Test match per sottosezioni 6.x: {test_match}")
            for i, section in enumerate(numbered_sections):
                if re.search(r'6\.\d+\.', section.strip()):
                    print(f"DEBUG - Trovata sezione 6.x: {i} - {section[:50]}...")
            # FINE DEBUG AGGIUNTIVI
           
            # Verifica se ci sono sottosezioni 6.x che potrebbero contenere dati chiave
            is_cm1_with_subsections = (prompt_id == "CM-1" and 
                                      any(re.search(r'6\.\d+\.', section.strip()) for section in numbered_sections))

            if is_cm1_with_subsections:
                print(f"DEBUG - Rilevate sottosezioni 6.x in CM-1, applicando estrazione speciale")
                # Estrai esplicitamente i dati dalle key sections
                extracted_data = {}
    
                # Cerca nelle risposte per trovare dati chiave
                for response in all_responses:
                    if "MARKET_INSIGHTS" in response:
                        extracted_data["MARKET_INSIGHTS"] = response.strip()
                    if "KEYWORD_DATA" in response:
                        extracted_data["KEYWORD_DATA"] = response.strip()
                    if "BESTSELLER_OVERVIEW" in response:
                        extracted_data["BESTSELLER_OVERVIEW"] = response.strip()
    
                # Aggiungi fallback se necessario
                if "MARKET_INSIGHTS" not in extracted_data:
                    extracted_data["MARKET_INSIGHTS"] = "Analisi di mercato completata"
                if "BESTSELLER_OVERVIEW" not in extracted_data:
                    extracted_data["BESTSELLER_OVERVIEW"] = "Panoramica bestseller completata"
                if "KEYWORD_DATA" not in extracted_data:
                    extracted_data["KEYWORD_DATA"] = "Dati keyword analizzati"
    
                print(f"DEBUG_EXTRACT: Dati estratti manualmente: {list(extracted_data.keys())}")
            else:
                # Usa l'estrattore standard
                extracted_data = self.extract_data(full_response, prompt_id)

            # Aggiungi sempre valori di fallback per variabili critiche del CM-1
            if prompt_id == "CM-1":
                if "MARKET_INSIGHTS" not in extracted_data or not extracted_data["MARKET_INSIGHTS"]:
                    extracted_data["MARKET_INSIGHTS"] = "Analisi di mercato completata"
                if "BESTSELLER_OVERVIEW" not in extracted_data or not extracted_data["BESTSELLER_OVERVIEW"]:
                    extracted_data["BESTSELLER_OVERVIEW"] = "Panoramica bestseller completata"
                if "KEYWORD_DATA" not in extracted_data or not extracted_data["KEYWORD_DATA"]:
                    extracted_data["KEYWORD_DATA"] = "Dati keyword analizzati"
    
                # Aggiungi queste variabili anche a project_data (questa potrebbe essere la parte mancante!)
                project_data["MARKET_INSIGHTS"] = extracted_data["MARKET_INSIGHTS"] 
                project_data["BESTSELLER_OVERVIEW"] = extracted_data["BESTSELLER_OVERVIEW"]
                project_data["KEYWORD_DATA"] = extracted_data["KEYWORD_DATA"]

                print(f"DEBUG_EXTRACT: Valori finali per variabili critiche CM-1:")
                print(f"MARKET_INSIGHTS: {extracted_data.get('MARKET_INSIGHTS', 'MANCANTE')[:50]}...")
                print(f"BESTSELLER_OVERVIEW: {extracted_data.get('BESTSELLER_OVERVIEW', 'MANCANTE')[:50]}...")
                print(f"KEYWORD_DATA: {extracted_data.get('KEYWORD_DATA', 'MANCANTE')[:50]}...")
        
            # Salva nel database se disponibile
            if self.project_db_path and "PROJECT_ID" in project_data:
                self._save_result_to_db(
                    project_data["PROJECT_ID"],
                    prompt_id,
                    full_response,
                    extracted_data
                )
        
            # Aggiungi questo appena prima del return
            print(f"DEBUG_FINAL: Variabili cruciali in project_data alla fine di execute_step:")
            print(f"MARKET_INSIGHTS in project_data: {'MARKET_INSIGHTS' in project_data}")
            print(f"BESTSELLER_OVERVIEW in project_data: {'BESTSELLER_OVERVIEW' in project_data}")
            print(f"KEYWORD_DATA in project_data: {'KEYWORD_DATA' in project_data}")

            return project_data, full_response, extracted_data
        
        except Exception as e:
            logger.error(f"Errore nell'esecuzione dello step {prompt_id}: {str(e)}")
            raise

    def _format_extracted_data(self, data):
        """Formatta i dati estratti in un formato leggibile per il prompt di pulizia."""
        summary = []
        for key, value in data.items():
            if isinstance(value, str):
                # Prendi solo i primi 100 caratteri per ogni valore
                summary.append(f"- {key}: {value[:100]}...")
            elif isinstance(value, (list, dict)):
                summary.append(f"- {key}: {str(value)[:100]}...")
        return "\n".join(summary)
   
    def execute_flow(self, initial_data, executor_func, start_step="CM", max_steps=50):
        """
        Esegue l'intero flusso CRISP usando la funzione executor fornita.
        """
        project_data = initial_data.copy()
    
        if "PROJECT_ID" not in project_data:
            project_data["PROJECT_ID"] = f"PROJ_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    
        execution_history = []
        current_step = start_step
        steps_executed = 0
    
        while current_step and steps_executed < max_steps:
            try:
                logger.info(f"Esecuzione flusso - Step corrente: {current_step}")
            
                # Carica il prompt per vedere quante sezioni ha
                prompt_data = self.load_prompt(current_step)
                content = prompt_data["content"]
                sections = content.split("\n---\n")
                numbered_sections = [s.strip() for s in sections if re.match(r'^\d+\.', s.strip())]
            
                print(f"DEBUG_FLOW: Step {current_step} ha {len(numbered_sections)} sezioni")
            
                # Esegui lo step e verifica che tutte le sezioni siano state processate
                project_data, result, extracted = self.execute_step(
                    current_step, 
                    project_data, 
                    lambda prompt_text: executor_func(prompt_text, current_step)
                )
            
                # Verifica che abbiamo ricevuto risposte per tutte le sezioni
                responses = result.split("\n\n")
                if len(responses) < len(numbered_sections):
                    print(f"DEBUG_FLOW: ERRORE - Solo {len(responses)} risposte ricevute per {len(numbered_sections)} sezioni")
                    raise Exception(f"Step {current_step} non completato: mancano {len(numbered_sections) - len(responses)} sezioni")
            
                # Registra l'esecuzione solo se tutte le sezioni sono state processate
                execution_history.append({
                    "step_id": current_step,
                    "step_name": self.flow_map[current_step]["name"],
                    "timestamp": datetime.now().isoformat(),
                    "result_length": len(result),
                    "sections_completed": len(responses),
                    "total_sections": len(numbered_sections),
                    "extracted_data_keys": list(extracted.keys())
                })
            
                # Determina il prossimo step SOLO se questo è stato completato
                next_steps = self.flow_map[current_step].get("next", [])
                next_step = None
            
                for step in next_steps:
                    required_vars = self.flow_map[step].get("required_vars", [])
                    if all(var in project_data for var in required_vars):
                        next_step = step
                        logger.info(f"Prossimo step selezionato: {next_step}")
                        break
            
                current_step = next_step
                steps_executed += 1
            
            except Exception as e:
                logger.error(f"Errore nell'esecuzione del flusso al step {current_step}: {str(e)}")
                execution_history.append({
                    "step_id": current_step,
                    "step_name": self.flow_map.get(current_step, {}).get("name", "Unknown"),
                    "timestamp": datetime.now().isoformat(),
                    "error": str(e)
                })
                break
    
        return project_data, execution_history
    
    def _save_result_to_db(self, project_id, step_id, result_text, extracted_data):
            # Prima del salvataggio
            print(f"DEBUG: Salvataggio risultato nel DB - Project ID: {project_id}, Step ID: {step_id}")
            print(f"DEBUG: Lunghezza risultato: {len(result_text)}")
            print(f"DEBUG: Preview risultato: {result_text[:200]}...")
    
            # Codice esistente...
            try:
                conn = sqlite3.connect(self.project_db_path)
                cursor = conn.cursor()
        
                # Salva il risultato
                cursor.execute(
                    'INSERT INTO project_results (project_id, step_id, result_text) VALUES (?, ?, ?)',
                    (project_id, step_id, result_text)
                )
        
                # Salva i dati estratti come variabili del progetto
                for var_name, var_value in extracted_data.items():
                    # Converti valori non stringa in JSON
                    if not isinstance(var_value, str):
                        var_value = json.dumps(var_value)
            
                    # Verifica se la variabile esiste già
                    cursor.execute(
                        'SELECT COUNT(*) FROM project_variables WHERE project_id = ? AND name = ?',
                        (project_id, var_name)
                    )
                    exists = cursor.fetchone()[0] > 0
            
                    if exists:
                        # Aggiorna la variabile esistente
                        cursor.execute(
                            'UPDATE project_variables SET value = ? WHERE project_id = ? AND name = ?',
                            (var_value, project_id, var_name)
                        )
                    else:
                        # Inserisce una nuova variabile
                        cursor.execute(
                            'INSERT INTO project_variables (project_id, name, value) VALUES (?, ?, ?)',
                            (project_id, var_name, var_value)
                        )
        
                conn.commit()
                conn.close()
                print(f"DEBUG: Salvataggio nel DB completato con successo")
        
            except Exception as e:
                print(f"DEBUG: Errore nel salvataggio dei risultati nel database: {str(e)}")
                logger.error(f"Errore nel salvataggio dei risultati nel database: {str(e)}")

    def save_incremental_response(self, project_id, step_id, prompt_line, response_text, is_final=False):
        """
        Salva una risposta incrementale nel database.

        Args:
            project_id: ID del progetto
            step_id: ID dello step (es. CM-1, CS-2, etc.)
            prompt_line: La riga del prompt che ha generato questa risposta
            response_text: Testo della risposta
            is_final: Indica se questa è l'ultima risposta per lo step
        """
        # Aggiungi debug per vedere lo stack di chiamate
        import traceback
        stack_trace = traceback.format_stack()
        print(f"DEBUG - save_incremental_response chiamato con step_id={step_id}")
        print(f"DEBUG - Provenienza chiamata (ultime 3 chiamate):")
        for line in stack_trace[-4:-1]:  # Mostra solo le ultime 3 chiamate per brevità
            print(f"DEBUG - STACK: {line.strip()}")
    
        if not self.project_db_path:
            logger.warning("Impossibile salvare la risposta incrementale: nessun database specificato")
            print(f"DEBUG - Database path non specificato, impossibile salvare risposta per {step_id}")
            return
    
        try:
            # Verifica che step_id sia in un formato valido (per debugging)
            if step_id is None:
                print(f"DEBUG - ATTENZIONE: step_id è None!")
                step_id = "unknown"  # Fallback per evitare errori SQL
            
            if step_id and not re.match(r'^[A-Z]+-\d+$|^[A-Z]+-[A-Z]$', step_id):
                warning_msg = f"ID del prompt potenzialmente in formato non standard: {step_id}"
                logger.warning(warning_msg)
                print(f"DEBUG - {warning_msg}")
        
            # Stampa informazioni dettagliate su tutti i parametri
            print(f"DEBUG - project_id: {project_id}")
            print(f"DEBUG - step_id: {step_id}")
            print(f"DEBUG - prompt_line (primi 50 char): {prompt_line[:50]}...")
            print(f"DEBUG - response_text (lunghezza): {len(response_text)} caratteri")
            print(f"DEBUG - response_text (primi 50 char): {response_text[:50]}...")
            print(f"DEBUG - is_final: {is_final}")
        
            # Verifica tracce di CM-1 o CM-2 nella prompt_line o response_text
            if "CM-1" in prompt_line or "CM-1" in response_text[:200]:
                print(f"DEBUG - Trovato riferimento a CM-1 ma step_id è {step_id}")
            if "CM-2" in prompt_line or "CM-2" in response_text[:200]:
                print(f"DEBUG - Trovato riferimento a CM-2 ma step_id è {step_id}")
        
            conn = sqlite3.connect(self.project_db_path)
            cursor = conn.cursor()
        
            # Verifica stato attuale del database
            cursor.execute(
                "SELECT COUNT(*) FROM incremental_responses WHERE project_id = ? AND step_id = ?", 
                (project_id, step_id)
            )
            existing_count = cursor.fetchone()[0]
            print(f"DEBUG - Risposte esistenti per step_id={step_id}: {existing_count}")
        
            # Stampa log dettagliato per debug
            short_response = response_text[:100] + "..." if len(response_text) > 100 else response_text
            logger.debug(f"Salvando risposta per step {step_id}: {short_response}")
        
            # Salva la risposta incrementale
            insert_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            print(f"DEBUG - Timestamp inserimento: {insert_time}")
        
            cursor.execute(
                '''INSERT INTO incremental_responses 
                   (project_id, step_id, prompt_line, response_text, timestamp, is_final) 
                   VALUES (?, ?, ?, ?, datetime('now'), ?)''',
                (project_id, step_id, prompt_line, response_text, is_final)
            )
        
            # Verifica che l'inserimento sia avvenuto con successo
            cursor.execute("SELECT last_insert_rowid()")
            last_id = cursor.fetchone()[0]
            print(f"DEBUG - Risposta inserita con ID: {last_id}")
        
            conn.commit()
            conn.close()
        
            success_msg = f"Risposta incrementale salvata per {step_id} (progetto {project_id})"
            logger.info(success_msg)
            print(f"DEBUG - {success_msg}")
        
            # Controlla se altre risposte con step_id diverso sono state salvate recentemente
            try:
                conn = sqlite3.connect(self.project_db_path)
                cursor = conn.cursor()
                cursor.execute(
                    """SELECT step_id, COUNT(*) 
                       FROM incremental_responses 
                       WHERE project_id = ? 
                       AND timestamp > datetime('now', '-5 minutes')
                       GROUP BY step_id""",
                    (project_id,)
                )
                recent_steps = cursor.fetchall()
                conn.close()
            
                if recent_steps and len(recent_steps) > 1:
                    print(f"DEBUG - Ultimi 5 minuti, risposte per step diversi: {recent_steps}")
            except Exception:
                pass  # Ignora eventuali errori nella query di debug
        
        except Exception as e:
            error_msg = f"Errore nel salvataggio della risposta incrementale: {str(e)}"
            logger.error(error_msg)
            print(f"DEBUG - ERRORE: {error_msg}")
            print(f"DEBUG - Traceback completo:")
            traceback.print_exc()

    def save_incremental_response_wrapper(self, project_id, step_id, prompt_line, response_text, is_final=False):
        """Wrapper temporaneo per debug"""
        # Determina quale sia l'ID del prompt attualmente in elaborazione
        current_prompt_in_process = None
    
        # Usa traceback per determinare da dove viene chiamata questa funzione
        import traceback
        stack = traceback.extract_stack()
        for frame in reversed(stack):
            if 'execute_step' in frame.name:
                # Se la chiamata proviene da execute_step, cerca nei frame superiori
                break
    
        # Usa l'ID rilevato dalla traccia se disponibile e diverso
        if current_prompt_in_process and current_prompt_in_process != step_id:
            logger.warning(f"Mismatch negli ID dei prompt: passato {step_id}, ma in elaborazione {current_prompt_in_process}")
            # Usa l'ID corretto
            step_id = current_prompt_in_process
    
        # Chiama la funzione originale
        return self.save_incremental_response(project_id, step_id, prompt_line, response_text, is_final)
    
    def create_project(self, project_name, initial_data=None):
        """
        Crea un nuovo progetto nel database.
        
        Args:
            project_name: Nome del progetto
            initial_data: Dizionario con i dati iniziali del progetto
            
        Returns:
            str: ID del progetto creato
        """
        if not self.project_db_path:
            raise ValueError("Database path non specificato")
        
        try:
            # Genera un ID progetto
            project_id = f"PROJ_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            conn = sqlite3.connect(self.project_db_path)
            cursor = conn.cursor()
            
            # Crea il progetto
            cursor.execute(
                'INSERT INTO projects (id, name) VALUES (?, ?)',
                (project_id, project_name)
            )
            
            # Salva i dati iniziali
            if initial_data:
                for var_name, var_value in initial_data.items():
                    # Converti valori non stringa in JSON
                    if not isinstance(var_value, str):
                        var_value = json.dumps(var_value)
                    
                    cursor.execute(
                        'INSERT INTO project_variables (project_id, variable_name, variable_value) VALUES (?, ?, ?)',
                        (project_id, var_name, var_value)
                    )
            
            conn.commit()
            conn.close()
            
            logger.info(f"Progetto creato: {project_id} - {project_name}")
            return project_id
            
        except Exception as e:
            logger.error(f"Errore nella creazione del progetto: {str(e)}")
            raise
    
    def get_project_data(self, project_id):
        """
        Recupera i dati di un progetto dal database.
    
        Args:
            project_id: ID del progetto
        
        Returns:
            dict: Dati del progetto
        """
        # All'inizio del metodo
        print(f"DEBUG: Tentativo di recupero dati progetto {project_id}")
    
        if not self.project_db_path:
            raise ValueError("Database path non specificato")
    
        try:
            conn = sqlite3.connect(self.project_db_path)
            cursor = conn.cursor()
        
            # Recupera le informazioni di base del progetto
            cursor.execute(
                'SELECT name, creation_date, status FROM projects WHERE id = ?',
                (project_id,)
            )
            project_info = cursor.fetchone()
        
            if not project_info:
                conn.close()
                raise ValueError(f"Progetto {project_id} non trovato")
        
            # Crea il dizionario base del progetto
            project_data = {
                "PROJECT_ID": project_id,
                "PROJECT_NAME": project_info[0],
                "CREATION_DATE": project_info[1],
                "STATUS": project_info[2]
            }
        
            # Recupera le variabili del progetto
            cursor.execute(
                'SELECT name, value FROM project_variables WHERE project_id = ?',
                (project_id,)
            )
        
            for var_name, var_value in cursor.fetchall():
                # Tenta di convertire JSON in oggetti Python
                try:
                    project_data[var_name] = json.loads(var_value)
                except (json.JSONDecodeError, TypeError):
                    project_data[var_name] = var_value
        
            conn.close()
        
            # Dopo aver recuperato i dati
            if project_data:
                print(f"DEBUG: Dati progetto recuperati - {len(project_data)} variabili")
                # Mostra alcune variabili chiave
                for key in ['TITOLO_LIBRO', 'CONTENT_PILLARS', 'VOICE_STYLE']:
                    if key in project_data:
                        value = project_data[key]
                        preview = str(value)[:100] + ('...' if len(str(value)) > 100 else '')
                        print(f"DEBUG: {key}: {preview}")
        
            logger.info(f"Dati progetto {project_id} recuperati: {len(project_data)} variabili")
            return project_data
        
        except Exception as e:
            logger.error(f"Errore nel recupero dei dati del progetto: {str(e)}")
            raise
    
    def view_incremental_responses(self, project_id, step_id):
        """
    	Recupera tutte le risposte incrementali per uno specifico step.
    
        Args:
            project_id: ID del progetto
            step_id: ID dello step
        
        Returns:
            list: Lista di tuple (prompt_line, response_text, timestamp)
        """
        try:
            conn = sqlite3.connect(self.project_db_path)
            cursor = conn.cursor()
        
            cursor.execute(
                '''SELECT prompt_line, response_text, timestamp 
                   FROM incremental_responses 
                   WHERE project_id = ? AND step_id = ?
                   ORDER BY id''',
                (project_id, step_id)
            )
        
            responses = cursor.fetchall()
            conn.close()
        
            return responses
        
        except Exception as e:
            logger.error(f"Errore nel recupero delle risposte incrementali: {str(e)}")
            return []

    def export_complete_document(self, project_id, output_format="docx"):
        """
        Esporta tutte le risposte di un progetto in un documento completo.
    
        Args:
            project_id: ID del progetto
            output_format: Formato di output (docx, pdf, txt)
        
        Returns:
            str: Percorso del file esportato
        """
        try:
            # Recupera tutte le risposte del progetto
            conn = sqlite3.connect(self.project_db_path)
            cursor = conn.cursor()
        
            # Ottieni nome del progetto
            cursor.execute('SELECT name FROM projects WHERE id = ?', (project_id,))
            project_name = cursor.fetchone()[0]
        
            # Ottieni tutte le risposte in ordine
            cursor.execute(
                '''SELECT step_id, result_text FROM project_results 
                   WHERE project_id = ? ORDER BY id''',
                (project_id,)
            )
            results = cursor.fetchall()
            conn.close()
        
            # Crea directory di output se non esiste
            output_dir = Path("Output") / project_id
            output_dir.mkdir(parents=True, exist_ok=True)
        
            # Percorso del file di output
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"{project_name}_{timestamp}.{output_format}"
        
            # Crea il documento
            if output_format == "docx":
                doc = Document()
                doc.add_heading(project_name, 0)
            
                for step_id, result_text in results:
                    doc.add_heading(f"Risultati {step_id}", 1)
                    doc.add_paragraph(result_text)
                    doc.add_page_break()
                
                doc.save(output_path)
            
            elif output_format == "txt":
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(f"PROGETTO: {project_name}\n\n")
                
                    for step_id, result_text in results:
                        f.write(f"=== RISULTATI {step_id} ===\n\n")
                        f.write(result_text)
                        f.write("\n\n" + "="*50 + "\n\n")
        
            # Altri formati possono essere aggiunti qui
        
            logger.info(f"Documento completo esportato: {output_path}")
            return str(output_path)
        
        except Exception as e:
            logger.error(f"Errore nell'esportazione del documento completo: {str(e)}")
            return None
    
    def get_completed_phases(self, project_id):
        """
        Recupera le fasi completate per un progetto specifico.
    
        Args:
            project_id: ID del progetto
        
        Returns:
            list: Lista delle fasi completate (prompt_id)
        """
        try:
            conn = sqlite3.connect(self.project_db_path)
            cursor = conn.cursor()
        
            # Recupera le fasi completate
            cursor.execute("""
                SELECT DISTINCT prompt_id 
                FROM project_results 
                WHERE project_id = ?
                ORDER BY id
            """, (project_id,))
        
            phases = [row[0] for row in cursor.fetchall()]
            conn.close()
        
            return phases
        
        except Exception as e:
            logger.error(f"Errore nel recupero delle fasi completate: {str(e)}")
            return []

    def get_execution_history(self, project_id):
        """
        Recupera la storia dell'esecuzione di un progetto.
        
        Args:
            project_id: ID del progetto
            
        Returns:
            list: Storia dell'esecuzione
        """
        if not self.project_db_path:
            raise ValueError("Database path non specificato")
        
        try:
            conn = sqlite3.connect(self.project_db_path)
            cursor = conn.cursor()
            
            # Recupera i risultati del progetto
            cursor.execute(
                'SELECT step_id, result_text, timestamp FROM project_results WHERE project_id = ? ORDER BY id',
                (project_id,)
            )
            
            execution_history = []
            for step_id, result_text, timestamp in cursor.fetchall():
                step_name = self.flow_map.get(step_id, {}).get("name", "Unknown")
                execution_history.append({
                    "step_id": step_id,
                    "step_name": step_name,
                    "timestamp": timestamp,
                    "result_length": len(result_text)
                })
            
            conn.close()
            
            logger.info(f"Storia esecuzione progetto {project_id} recuperata: {len(execution_history)} step")
            return execution_history
            
        except Exception as e:
            logger.error(f"Errore nel recupero della storia dell'esecuzione: {str(e)}")
            raise
	
    def view_incremental_responses(self, project_id, step_id):
        """
        Recupera tutte le risposte incrementali per uno specifico step.
        
        Args:
            project_id: ID del progetto
            step_id: ID dello step
            
        Returns:
            list: Lista di tuple (prompt_line, response_text, timestamp)
        """
        try:
            conn = sqlite3.connect(self.project_db_path)
            cursor = conn.cursor()
            
            cursor.execute(
                '''SELECT prompt_line, response_text, timestamp 
                   FROM incremental_responses 
                   WHERE project_id = ? AND step_id = ?
                   ORDER BY id''',
                (project_id, step_id)
            )
            
            responses = cursor.fetchall()
            conn.close()
            
            return responses
            
        except Exception as e:
            logger.error(f"Errore nel recupero delle risposte incrementali: {str(e)}")
            return []

    def export_complete_document(self, project_id, output_format="docx"):
        """
        Esporta tutte le risposte di un progetto in un documento completo.
        
        Args:
            project_id: ID del progetto
            output_format: Formato di output (docx, pdf, txt)
            
        Returns:
            str: Percorso del file esportato
        """
        try:
            # Recupera tutte le risposte del progetto
            conn = sqlite3.connect(self.project_db_path)
            cursor = conn.cursor()
            
            # Ottieni nome del progetto
            cursor.execute('SELECT name FROM projects WHERE id = ?', (project_id,))
            project_info = cursor.fetchone()
            
            if not project_info:
                conn.close()
                raise ValueError(f"Progetto {project_id} non trovato")
                
            project_name = project_info[0]
            
            # Ottieni tutte le risposte in ordine
            cursor.execute(
                '''SELECT step_id, result_text FROM project_results 
                   WHERE project_id = ? ORDER BY id''',
                (project_id,)
            )
            results = cursor.fetchall()
            
            # Crea directory di output se non esiste
            output_dir = Path("Output") / project_id
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Percorso del file di output
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"{project_name}_{timestamp}.{output_format}"
            
            # Crea il documento
            if output_format == "docx":
                doc = Document()
                doc.add_heading(project_name, 0)
                
                for step_id, result_text in results:
                    doc.add_heading(f"Risultati {step_id}", 1)
                    doc.add_paragraph(result_text)
                    doc.add_page_break()
                    
                doc.save(output_path)
                
            elif output_format == "txt":
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(f"PROGETTO: {project_name}\n\n")
                    
                    for step_id, result_text in results:
                        f.write(f"=== RISULTATI {step_id} ===\n\n")
                        f.write(result_text)
                        f.write("\n\n" + "="*50 + "\n\n")
            
            conn.close()
            logger.info(f"Documento completo esportato: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Errore nell'esportazione del documento completo: {str(e)}")
            return None

    def save_prompt_response(self, project_id, prompt_id, section_number, response_text):
        """Salva la risposta di un prompt nel database."""
        conn = sqlite3.connect(self.project_db_path)
        cursor = conn.cursor()
    
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
        cursor.execute('''
        INSERT INTO prompt_responses (project_id, prompt_id, section_number, response_text, timestamp)
        VALUES (?, ?, ?, ?, ?)
        ''', (project_id, prompt_id, section_number, response_text, timestamp))
    
        conn.commit()
        conn.close()
        logger.info(f"Risposta salvata per {prompt_id}, sezione {section_number}")

    def get_previous_responses_summary(self, project_id, prompt_id, max_chars=2000):
        """Recupera e sintetizza le risposte precedenti per un determinato prompt."""
        conn = sqlite3.connect(self.project_db_path)
        cursor = conn.cursor()
    
        # Estrai le risposte precedenti per questo progetto
        cursor.execute('''
        SELECT prompt_id, section_number, response_text
        FROM prompt_responses
        WHERE project_id = ? AND NOT (prompt_id = ? AND section_number = 6)
        ORDER BY prompt_id, section_number
        ''', (project_id, prompt_id))
    
        responses = cursor.fetchall()
        conn.close()
    
        summary = "SINTESI DELLE ANALISI PRECEDENTI:\n\n"
        current_length = len(summary)
    
        for resp_prompt_id, section_number, response_text in responses:
            # Estrai solo le prime righe o frasi importanti dalla risposta
            response_summary = self._extract_key_points(response_text)
        
            section_summary = f"Da {resp_prompt_id}-{section_number}: {response_summary}\n\n"
        
            # Controlla se aggiungendo questa sezione supereremmo il limite di caratteri
            if current_length + len(section_summary) > max_chars:
                # Tronca e aggiungi indicatore
                remaining_chars = max_chars - current_length - 20
                if remaining_chars > 0:
                    section_summary = section_summary[:remaining_chars] + "...\n\n"
                else:
                    break
        
            summary += section_summary
            current_length += len(section_summary)
    
        return summary

    def get_targeted_responses_summary(self, project_id, prompt_id, target_sections=None, max_chars=2000):
        """Recupera e sintetizza le risposte precedenti per un determinato prompt, focalizzandosi su specifiche sezioni."""
        conn = sqlite3.connect(self.project_db_path)
        cursor = conn.cursor()
    
        # Costruisci la query in base alle sezioni target
        query = '''
        SELECT prompt_id, section_number, response_text
        FROM prompt_responses
        WHERE project_id = ? AND NOT (prompt_id = ? AND section_number = 6)
        '''
    
        params = [project_id, prompt_id]
    
        # Se ci sono sezioni target specifiche, aggiungi la condizione alla query
        if target_sections and len(target_sections) > 0:
            section_placeholders = ','.join(['?' for _ in target_sections])
            query += f' AND section_number IN ({section_placeholders})'
            params.extend(target_sections)
    
        query += ' ORDER BY prompt_id, section_number'
    
        # Esegui la query
        cursor.execute(query, params)
        responses = cursor.fetchall()
        conn.close()
    
        summary = "SINTESI MIRATA DELLE ANALISI PRECEDENTI:\n\n"
        current_length = len(summary)
    
        for resp_prompt_id, section_number, response_text in responses:
            # Estrai solo le prime righe o frasi importanti dalla risposta
            response_summary = self._extract_key_points(response_text)
        
            section_summary = f"Da {resp_prompt_id}-{section_number}: {response_summary}\n\n"
        
            # Controlla se aggiungendo questa sezione supereremmo il limite di caratteri
            if current_length + len(section_summary) > max_chars:
                # Tronca e aggiungi indicatore
                remaining_chars = max_chars - current_length - 20
                if remaining_chars > 0:
                    section_summary = section_summary[:remaining_chars] + "...\n\n"
                else:
                    break
        
            summary += section_summary
            current_length += len(section_summary)
    
        return summary

    def _extract_key_points(self, text, max_points=5, max_chars_per_point=500):
        """Estrae i punti chiave da un testo di risposta."""
        # Implementazione semplificata - in realtà vorremmo estrarre i punti più importanti
        lines = text.split("\n")
        key_points = []
    
        for line in lines:
            line = line.strip()
            # Salta linee vuote o troppo corte
            if len(line) < 10:
                continue
        
            # Tronca linee troppo lunghe
            if len(line) > max_chars_per_point:
                line = line[:max_chars_per_point] + "..."
        
            key_points.append(line)
            if len(key_points) >= max_points:
                break
    
        return "\n".join(key_points)

# AGGIUNGERE QUESTO CODICE ALLA FINE DEL FILE crisp_framework.py - VERSIONE ALTERNATIVA

# Salva il riferimento al metodo originale
original_execute_step = CRISPFramework.execute_step

# Definisci la funzione di debug
def debug_execute_step(self, prompt_id, project_data, executor_func):
    """Versione con debug del metodo execute_step"""
    print(f"DEBUG_STEP: Inizio execute_step per prompt {prompt_id}")
    print(f"DEBUG_STEP: Dati disponibili per prompt {prompt_id}:")
    for key, value in project_data.items():
        value_preview = str(value)[:50] + "..." if value and len(str(value)) > 50 else value
        print(f"DEBUG_STEP:   - {key}: {type(value)} {value_preview}")

    # Verifichiamo se ci sono dati mancanti per questo prompt
    if hasattr(self, 'flow_map') and prompt_id in self.flow_map:
        required_vars = self.flow_map[prompt_id].get('required_vars', [])
        missing_vars = [var for var in required_vars if var not in project_data]
        if missing_vars:
            print(f"DEBUG_STEP: ATTENZIONE! Variabili mancanti per {prompt_id}: {missing_vars}")

    try:
        result = original_execute_step(self, prompt_id, project_data, executor_func)
        print(f"DEBUG_STEP: Fine execute_step per prompt {prompt_id}")
        return result
    except Exception as e:
        print(f"DEBUG_STEP: ERRORE in execute_step per prompt {prompt_id}: {str(e)}")
        raise

# Sostituisci il metodo originale con quello di debug
CRISPFramework.execute_step = debug_execute_step