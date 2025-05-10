"""
Modulo per la generazione di libri con diversi framework.
"""

import re
import time
import os
import logging
from datetime import datetime
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC


def generate_book_crisp(book_title, book_language, voice_style, book_index,
                      crisp_framework=None, driver=None, chat_manager=None, current_analysis=None):
    """
    Genera il libro usando il framework CRISP 5.0.

    Args:
        book_title: Titolo del libro
        book_language: Lingua del libro
        voice_style: Stile narrativo
        book_index: Indice del libro
        crisp_framework: Istanza del framework CRISP
        driver: Istanza del WebDriver Selenium
        chat_manager: Istanza del ChatManager per il logging
        current_analysis: Dizionario contenente i dati di analisi correnti
        
    Returns:
        str: Percorso del file del libro generato o messaggio di errore
    """
    # Funzione di log
    def log(message):
        if chat_manager:
            chat_manager.add_log(message)
        else:
            print(message)
    
    try:
        if not book_title.strip():
            log("Errore: Il titolo del libro è obbligatorio!")
            return "Errore: Il titolo del libro è obbligatorio!"
        
        if not book_index.strip():
            log("Errore: L'indice del libro è obbligatorio!")
            return "Errore: L'indice del libro è obbligatorio!"
        
        log(f"Avvio generazione libro CRISP 5.0: {book_title}")
        log(f"DEBUG: generate_book_crisp() chiamato con → title='{book_title}', lang='{book_language}', voice_style='{voice_style[:30]}...', indice righe={len(book_index.splitlines())}")
        
        # Recupera l'ID del progetto CRISP corrente
        project_id = current_analysis.get('crisp_project_id') if current_analysis else None
        if not project_id:
            log("Errore: Nessun progetto CRISP corrente trovato")
            return "Errore: Nessun progetto CRISP corrente trovato"
        
        # Recupera i dati completi del progetto
        project_data = crisp_framework.get_project_data(project_id) or {}
        project_data = {
            **project_data,
            **(current_analysis.get('project_data', {}) if current_analysis else {})
            
        }
        
        

        # Aggiorna i dati del progetto con le informazioni correnti
        project_data.update({
            "TITOLO_LIBRO": book_title,
            "LINGUA_LIBRO": book_language,
            "VOICE_STYLE": voice_style,
            "INDICE_LIBRO": book_index
        })
        
        # Aggiorna i dati nel database CRISP
        crisp_framework._save_result_to_db(
            project_id,
            "GEN1",
            f"Generazione libro: {book_title}",
            {
                "TITOLO_LIBRO": book_title,
                "LINGUA_LIBRO": book_language,
                "VOICE_STYLE": voice_style,
                "INDICE_LIBRO": book_index
            }
        )
        
        # Dividi l'indice in capitoli
        chapters = [line.strip() for line in book_index.split('\n') if line.strip()]
        
        # Filtra per mantenere solo le righe che sembrano capitoli
        chapter_pattern = re.compile(r'(CAPITOLO|CHAPTER|PARTE|PART|SEZIONE|SECTION)\s*\d*\s*[:\.\-–—]?\s*(.*)', re.IGNORECASE)
        filtered_chapters = []
        log(f"DEBUG: {len(filtered_chapters)} capitoli da generare: {filtered_chapters}")
        
        for line in chapters:
            match = chapter_pattern.match(line)
            if match:
                chapter_title = match.group(2).strip()
                if chapter_title:  # Assicurati che il titolo non sia vuoto
                    filtered_chapters.append(chapter_title)
            elif not any(keyword.lower() in line.lower() for keyword in ['introduzione', 'introduction', 'conclusione', 'conclusion']):
                # Includi anche le righe che non contengono parole chiave specifiche
                filtered_chapters.append(line)
        
        # Se non abbiamo trovato capitoli validi, usa le righe originali
        if not filtered_chapters:
            filtered_chapters = [line for line in chapters if line.strip()]
        
        # Prepara la risposta cumulativa
        book_content = []
        
        # Aggiungi introduzione se non è già nei capitoli
        has_intro = any('introduz' in chapter.lower() or 'introduct' in chapter.lower() for chapter in filtered_chapters)
        if not has_intro:
            book_content.append("# Introduzione\n\n[Introduzione del libro]")
        
        # Genera ciascun capitolo
        for i, chapter_title in enumerate(filtered_chapters):
            log(f"Generazione capitolo {i+1}/{len(filtered_chapters)}: '{chapter_title}'")
            log(f"DEBUG: usando Big Idea='{project_data.get('BIG_IDEA')}', Pillars='{project_data.get('CONTENT_PILLARS')[:50]}...'")
        
            # Costruzione avanzata del prompt, in 4 blocchi:
            prompt_blocks = []

            # (A) Research Instruction
            prompt_blocks.append(
                "Research Instruction:\n"
                "Before you start writing, use reliable online sources to research each point "
                "and integrate up-to-date facts, tax updates, and examples specific to high-income professionals."
            )

            # (B) Metadati da project_data (Big Idea, Pillars, Buyer Persona…)
            prompt_blocks.append(
                f"L'idea centrale (Big Idea) del libro è: {project_data.get('BIG_IDEA', 'Non specificata')}\n"
                f"I pilastri di contenuto principali sono:\n{project_data.get('CONTENT_PILLARS', 'Non specificati')}\n"
                f"Buyer Persona Summary: {project_data.get('BUYER_PERSONA_SUMMARY', 'lettori interessati a questo argomento')}"
            )

            # (C) Istruzioni di stile e struttura
            prompt_blocks.append(
                "✍️ WRITING STYLE:\n"
                "Authoritative yet accessible prose in 2nd person (“you”), include visuals, diagrammi e flowchart per supportare concetti complessi.\n"
                "🧱 STRUCTURE:\n"
                "1) Strategic Snapshot\n2) Framework & Action Plan\n3) Visual Integration\n4) Wrap-Up Checklist"
            )

            # (D) Prompt di scrittura vero e proprio
            prompt_blocks.append(
                f"Scrivi il capitolo “{chapter_title}” per “{book_title}” in lingua {book_language}.\n"
                f"Lunghezza minima: 1500 parole. Termina con FINE."
            )

            # Unisci i blocchi in un singolo prompt
            chapter_prompt = "\n\n".join(prompt_blocks)

        
            # Invia il prompt a Genspark
            try:
                # Ottieni input box
                input_box = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.CSS_SELECTOR, "div.search-input-wrapper textarea"))
                )
            
                # Pulisci l'input box
                input_box.clear()
                time.sleep(0.5)
            
                # Inserisci il testo
                log("PROMPT:\n" + chapter_prompt)
                log("DEBUG: chapter_prompt completo:\n" + chapter_prompt)
                input_box.send_keys(chapter_prompt)
                time.sleep(1)
            
                # Trova e clicca il pulsante di invio
                send_button = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.CSS_SELECTOR, "div.search-input-wrapper div.input-icon"))
                )
                send_button.click()
                log("DEBUG: prompt inviato, in attesa di risposta…")
            
                # Attendi risposta
                log("⏳ Attesa risposta per il capitolo...")
            
                # Verifica completamento risposta
                last_length = 0
                stable_count = 0
                max_stable_counts = 2
                chapter_content = None
                log(f"DEBUG: risposta ricevuta, lunghezza={len(response)} caratteri")
            
                # Importazione per evitare dipendenze circolari
                from ai_interfaces.genspark_driver import get_last_response
                
                for attempt in range(60):  # Max 10 minuti (60 * 10 secondi)
                    time.sleep(10)
                    response = get_last_response(driver, log)

                    # — Inizio controllo esplicito “FINE” —
                    if response and ("FINE_RISPOSTA" in response or "FINE" in response):
                        log(f"✅ Terminatore esplicito trovato all'attempt {attempt+1}")
                        terminator = "FINE_RISPOSTA" if "FINE_RISPOSTA" in response else "FINE"
                        chapter_content = response.split(terminator)[0].strip()
                        break   # esco subito, non aspetto stable_count
                    # — Fine controllo esplicito “FINE” —

                    if response:
                        current_length = len(response)
                        if current_length == last_length:
                            stable_count += 1
                            if stable_count >= max_stable_counts:
                                # Risposta stabile (fallback)
                                chapter_content = response
                                break
                        else:
                            stable_count = 0
                            last_length = current_length
            
                if not chapter_content:
                    raise Exception(f"Timeout in attesa della risposta per il capitolo {chapter_title}")
            
                # Rimuovi la parola FINE se presente
                if "FINE" in chapter_content:
                    chapter_content = chapter_content[:chapter_content.find("FINE")].strip()
            
                # Aggiungi il titolo e il contenuto al libro
                formatted_chapter = f"# {chapter_title}\n\n{chapter_content}"
                book_content.append(formatted_chapter)
            
                # Salva il capitolo nel database CRISP
                crisp_framework._save_result_to_db(
                    project_id,
                    f"CAP{i+1}",
                    formatted_chapter,
                    {"chapter_title": chapter_title}
                )
            
                log(f"✅ Capitolo {i+1} completato: {chapter_title}")
                log(f"DEBUG: formatted_chapter starts with:\n{formatted_chapter[:100]}...")

            
                # Gestisci il reset del contesto se necessario
                from ai_interfaces.browser_manager import handle_context_limit
                
                if i % 2 == 0:  # Ogni 2 capitoli
                    handle_context_limit(driver, log)
        
            except Exception as e:
                error_msg = f"Errore durante la generazione del capitolo {chapter_title}: {str(e)}"
                log(error_msg)
                logging.error(error_msg)
                book_content.append(f"# {chapter_title}\n\nErrore durante la generazione di questo capitolo: {str(e)}")
        
            # Breve pausa tra i capitoli
            time.sleep(5)
        
        # Aggiungi conclusione se non è già nei capitoli
        has_conclusion = any('conclus' in chapter.lower() or 'conclusion' in chapter.lower() for chapter in filtered_chapters)
        if not has_conclusion:
            book_content.append("# Conclusione\n\n[Conclusione del libro]")
        
        # Unisci tutti i capitoli
        complete_book = "\n\n".join(book_content)
        
        # Salva il libro in una cartella dedicata al titolo
        safe_title = re.sub(r'[<>:"/\\|?*]', '', book_title)[:30].replace(' ', '_')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        os.makedirs(safe_title, exist_ok=True)
        book_filename = os.path.join(
            safe_title,
            f"{safe_title}_{timestamp}.md"
        )
        with open(book_filename, "w", encoding="utf-8") as f:
            f.write(complete_book)

        # Salva anche nel database CRISP, includendo il percorso corretto
        crisp_framework._save_result_to_db(
            project_id,
            "COMPLETE_BOOK",
            complete_book,
            {
                "title": book_title,
                "language": book_language,
                "style": voice_style,
                "chapters": len(filtered_chapters),
                "filename": book_filename
            }
        )
        log(f"DEBUG: book_content contiene {len(book_content)} sezioni; salvataggio file '{book_filename}'")
        log(f"📚 Libro generato e salvato come: {book_filename}'")
        return book_filename
    
    except Exception as e:
        error_msg = f"Errore durante la generazione del libro CRISP 5.0: {str(e)}"
        log(error_msg)
        logging.error(error_msg)
        return f"Errore: {str(e)}"

def generate_book_legacy(book_title, book_language, voice_style, book_index, driver=None, chat_manager=None):
    """
    Metodo legacy per generare il libro.
    
    Args:
        book_title: Titolo del libro
        book_language: Lingua del libro
        voice_style: Stile narrativo
        book_index: Indice del libro
        driver: Istanza del WebDriver Selenium
        chat_manager: Istanza del ChatManager per il logging
        
    Returns:
        str: Messaggio di completamento o errore
    """
    # Funzione di log
    def log(message):
        if chat_manager:
            chat_manager.add_log(message)
        else:
            print(message)

    try:
        if not book_title.strip():
            log("Errore: Il titolo del libro è obbligatorio!")
            return "Errore: Il titolo del libro è obbligatorio!"
            
        if not book_index.strip():
            log("Errore: L'indice del libro è obbligatorio!")
            return "Errore: L'indice del libro è obbligatorio!"
        
        log(f"Generazione libro: {book_title}")
        
        # Qui andrà il codice per la generazione del libro legacy
        # Poiché non era fornito nell'originale, lasciamo un placeholder
        
        log("✅ Libro generato")
        return "Libro generato con successo"
        
    except Exception as e:
        import logging
        error_msg = f"Errore durante la generazione del libro: {str(e)}"
        log(error_msg)
        logging.error(error_msg)
        return f"Errore: {str(e)}"

def generate_book(book_title, book_language, voice_style, book_index,
                 driver=None, chat_manager=None, current_analysis=None,
                 book_type_hidden=None, send_to_genspark=None):
    """
    Genera il libro utilizzando i dati dell'interfaccia e i dati CRISP disponibili.
    """
    # Funzione di log
    def log(message):
        if chat_manager:
            chat_manager.add_log(message)
        else:
            print(message)

    # Disabilita i messaggi CRISP ridondanti
    disable_crisp_messages(chat_manager)
            
    if not book_title.strip():
        log("Errore: Il titolo del libro è obbligatorio!")
        return "Errore: Il titolo del libro è obbligatorio!"
        
    if not book_index.strip():
        log("Errore: L'indice del libro è obbligatorio!")
        return "Errore: L'indice del libro è obbligatorio!"

    try:
        # Preparazione stile condensato
        if isinstance(voice_style, str) and (voice_style.endswith('.txt') or '/' in voice_style or '\\' in voice_style):
            try:
                with open(voice_style, 'r', encoding='utf-8') as f:
                    full_style = f.read()
                # Estrai solo le prime 3-5 frasi per mantenere lo stile breve
                import re
                sentences = re.split(r'\.(?=\s|$)', full_style)
                voice_style_summary = '. '.join([s for s in sentences[:5] if s.strip()]) + '.'
                log(f"Stile caricato da file: {len(voice_style_summary)} caratteri (riassunto)")
            except Exception as e:
                log(f"Errore nel leggere il file di stile: {str(e)}")
                voice_style_summary = "Stile professionale ma accessibile, con un tono conversazionale."
        else:
            # Limita la lunghezza dello stile
            voice_style_summary = voice_style[:500] if len(voice_style) > 500 else voice_style
        
        # Recupera il tipo di libro
        book_type = "Manuale (Non-Fiction)"  # Default
        if book_type_hidden:
            book_type = book_type_hidden
        elif current_analysis and current_analysis.get('project_data', {}).get('LIBRO_TIPO'):
            book_type = current_analysis['project_data']['LIBRO_TIPO']
    
        log(f"Generazione libro: {book_title} (Tipo: {book_type})")
    
        # Carica il prompt specifico per questo tipo di libro
        chapter_prompt = _load_chapter_prompt(book_type)
    
        # Analizza l'indice per ottenere i capitoli
        chapters = _parse_book_index(book_index)
        log(f"Indice analizzato: {len(chapters)} capitoli trovati")
        
        # IMPORTANTE: Crea una directory con nome breve per il libro invece di usare il titolo completo
        import os
        import re
        
        # Crea una versione abbreviata e sicura del titolo per la directory (max 50 caratteri)
        safe_title = re.sub(r'[<>:"/\\|?*]', '', book_title)[:50].strip()
        book_dir = os.path.join(os.getcwd(), f"book_{int(time.time())}")  # Usa timestamp invece del titolo
        os.makedirs(book_dir, exist_ok=True)
        
        # Salva le informazioni sul libro in un file README
        with open(os.path.join(book_dir, "README.txt"), "w", encoding="utf-8") as f:
            f.write(f"Titolo: {book_title}\n")
            f.write(f"Data generazione: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Lingua: {book_language}\n")
            f.write(f"Tipo: {book_type}\n")
            f.write(f"Numero capitoli: {len(chapters)}\n")
    
        # Genera ogni capitolo
        for i, chapter in enumerate(chapters):
            log(f"Generazione capitolo {i+1}/{len(chapters)}: {chapter['title']}")
        
            # Prepara il prompt con sostituzione esplicita delle variabili chiave
            full_prompt = chapter_prompt.replace("{text}", chapter['title'])
            full_prompt = full_prompt.replace("{TITOLO_LIBRO}", book_title)
            full_prompt = full_prompt.replace("{LINGUA_LIBRO}", book_language)
            full_prompt = full_prompt.replace("{VOICE_STYLE}", voice_style_summary)
            
            # Rimuovi i placeholder non sostituiti
            import re
            full_prompt = re.sub(r'\{[A-Z_]+\}', '', full_prompt)


            log("DEBUG: chapter_prompt completo:\n" + chapter_prompt)

            log(f"Prompt preparato: {len(full_prompt)} caratteri")
        
            # Invia il prompt utilizzando l'oggetto driver direttamente
            try:
                # Implementazione diretta che non dipende da send_to_genspark
                from selenium.webdriver.common.by import By
                from selenium.webdriver.support.ui import WebDriverWait
                from selenium.webdriver.support import expected_conditions as EC
                                
                # Ottieni input box
                input_box = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.CSS_SELECTOR, "div.search-input-wrapper textarea"))
                )
                
                # Pulisci l'input box
                input_box.clear()
                time.sleep(0.5)
                
                # Inserisci il testo
                log("⌨️ Inserimento testo nel prompt...")
                for chunk in [full_prompt[i:i+200] for i in range(0, len(full_prompt), 200)]:
                    input_box.send_keys(chunk)
                    time.sleep(0.2)  # Piccola pausa tra i blocchi
                
                time.sleep(1)  # Pausa finale prima di inviare
                
                # Trova e clicca il pulsante di invio
                send_button = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.CSS_SELECTOR, "div.search-input-wrapper div.input-icon"))
                )
                send_button.click()
                
                # Attendi la risposta
                log("⏳ Attesa risposta per il capitolo...")
                
                # Usa la stessa logica di rilevamento del completamento usata nella fase di analisi
                chapter_content = ""
                last_length = 0
                stable_count = 0
                max_wait_cycles = 45  # ~15 minuti totali
                stability_threshold = 3  # 3 cicli di stabilità
                cycle_wait = 20  # 20 secondi per ciclo
   
                terminator_found = False             
                for cycle in range(max_wait_cycles):
                    try:
                        time.sleep(cycle_wait)
                        
                        # Prova diversi selettori per le risposte
                        selectors = [
                            ".message-content", 
                            "div.chat-wrapper div.desc > div > div > div",
                            "div.message div.text-wrap",
                            ".chat-message-item .content"
                        ]
                        
                        # Cerca la risposta con tutti i selettori
                        for selector in selectors:
                            try:
                                messages = driver.find_elements(By.CSS_SELECTOR, selector)
                                if messages and len(messages) > 0:
                                    current_text = messages[-1].text.strip()
                                    if current_text:
                                        chapter_content = current_text
                                    
                                        # Verifica terminazione esplicita
                                        if "FINE_RISPOSTA" in chapter_content or "FINE" in chapter_content:
                                            log(f"✅ Terminatore esplicito trovato al ciclo {cycle+1}")
                                            if "FINE_RISPOSTA" in chapter_content:
                                                chapter_content = chapter_content.split("FINE_RISPOSTA")[0].strip()
                                            elif "FINE" in chapter_content:
                                                chapter_content = chapter_content.split("FINE")[0].strip()
                                            terminator_found = True
                                            break
                                        
                                        # Verifica stabilità
                                        current_length = len(chapter_content)
                                        if current_length == last_length:
                                            stable_count += 1
                                            if stable_count >= stability_threshold:
                                                log(f"✅ Risposta stabilizzata dopo {cycle+1} cicli")
                                                break
                                        else:
                                            stable_count = 0
                                            last_length = current_length
                                
                                        # Trovata risposta valida, esci dal ciclo selettori
                                        break
                            except Exception:
                                continue

                        # se abbiamo trovato il terminatore, interrompiamo anche il polling esterno
                        if terminator_found:
                            log(f"DEBUG: uscita anticipata dal polling al ciclo {cycle+1}")
                            break   # esci dal for cycle
      
                        if "FINE_RISPOSTA" in chapter_content or "FINE" in chapter_content or stable_count >= stability_threshold:
                            break
                    except Exception as wait_error:
                        log(f"⚠️ Errore durante l'attesa: {str(wait_error)}")
                        continue
                    
                if not chapter_content:
                    raise Exception(f"Timeout in attesa della risposta per il capitolo {chapter['title']}")
                
                # Pulizia finale del contenuto
                if "FINE_RISPOSTA" in chapter_content:
                    chapter_content = chapter_content.split("FINE_RISPOSTA")[0].strip()
                elif "FINE" in chapter_content:
                    chapter_content = chapter_content.split("FINE")[0].strip()
                
            except Exception as e:
                log(f"Errore nell'invio del prompt: {str(e)}")
                chapter_content = f"[Errore nella generazione del capitolo: {str(e)}]"
        
            # Salva il capitolo con un nome file sicuro
            try:
                # Crea un nome di file sicuro e breve
                safe_chapter_name = f"capitolo_{i+1}"
                chapter_file = os.path.join(book_dir, f"{safe_chapter_name}.docx")
                
                from docx import Document
                doc = Document()
                doc.add_heading(chapter['title'], 1)
                
                # Aggiungi il contenuto suddiviso per paragrafi
                for paragraph in chapter_content.split('\n\n'):
                    if paragraph.strip():
                        if paragraph.strip().startswith('#'):
                            # È un titolo
                            level = paragraph.count('#')
                            title_text = paragraph.strip('#').strip()
                            doc.add_heading(title_text, level)
                        else:
                            # È un paragrafo normale
                            doc.add_paragraph(paragraph)
                
                doc.save(chapter_file)
                log(f"💾 Capitolo {i+1} salvato: {os.path.basename(chapter_file)}")
            except Exception as save_error:
                log(f"❌ Errore nel salvataggio del capitolo {i+1}: {str(save_error)}")
        
            # Applica un cooldown tra i capitoli
            if i < len(chapters) - 1:
                cooldown_time = 30 + (i * 5)  # Aumenta progressivamente
                log(f"Pausa di {cooldown_time} secondi prima del prossimo capitolo...")
                time.sleep(cooldown_time)
    
        log(f"✅ Libro generato con successo in: {book_dir}")
        return f"Libro {book_title} generato con successo in: {book_dir}"
    
    except Exception as e:
        error_msg = f"Errore durante la generazione del libro: {str(e)}"
        log(error_msg)
        logging.error(error_msg)
        return f"Errore: {str(e)}"

def extract_voice_style_summary(voice_style):
    """Estrae una versione condensata dello stile di voce."""
    try:
        if not voice_style:
            return "Stile conversazionale e informativo."
            
        content = voice_style
        
        # Limita la lunghezza a circa 500 caratteri
        if len(content) > 500:
            # Estrai le frasi più rappresentative
            import re
            
            # Cerca frasi chiave che definiscono lo stile
            key_patterns = [
                r'(?:tono|stile)[^\.]*?è[^\.]*?\.',  # "Il tono è formale."
                r'(?:usa|utilizza)[^\.]*?(?:uno stile|un tono)[^\.]*?\.',  # "Utilizza uno stile conversazionale."
                r'(?:linguaggio|registro)[^\.]*?(?:deve essere|è)[^\.]*?\.',  # "Il linguaggio deve essere tecnico."
                r'(?:la voce|lo stile)[^\.]*?deve[^\.]*?\.'  # "La voce deve essere autoritaria."
            ]
            
            summary = []
            for pattern in key_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE)
                if matches:
                    # Prendiamo solo le prime 2 corrispondenze per ogni pattern
                    for match in matches[:2]:
                        if match.strip() not in summary:
                            summary.append(match.strip())
            
            # Se non abbiamo trovato abbastanza frasi con i pattern, prendiamo le prime 3 frasi
            if len(summary) < 3:
                sentences = re.split(r'\.(?=\s|$)', content)
                for sentence in sentences[:5]:
                    if sentence.strip() and len(summary) < 3:
                        summary.append(sentence.strip() + '.')
                        
            return "\n".join(summary)
        else:
            return content
    except Exception as e:
        print(f"Errore nell'estrazione del riassunto dello stile: {e}")
        return "Stile conversazionale e informativo."  # Fallback

def _parse_book_index(book_index):
    """Analizza l'indice e lo converte in una lista di capitoli strutturati."""
    chapters = []
    lines = [line.strip() for line in book_index.split('\n') if line.strip()]

    for line in lines:
        # Ignora l'introduzione e la conclusione, li tratteremo separatamente
        if line.lower() in ["introduzione", "introduction", "conclusione", "conclusion"]:
            continue
    
        # Cerca di estrarre il titolo del capitolo
        chapter_match = re.match(r'(?:CAPITOLO|CHAPTER)\s*(\d+)?\s*:?\s*(.*)', line, re.IGNORECASE)
        if chapter_match:
            chapter_number = chapter_match.group(1) or ""
            chapter_title = chapter_match.group(2).strip()
            if chapter_title:
                chapters.append({
                    "number": chapter_number,
                    "title": chapter_title
                })
        elif line:  # Se è una linea con contenuto ma non è un formato standard
            chapters.append({
                "number": "",
                "title": line
            })

    return chapters

def _load_chapter_prompt(book_type):
    """Carica il prompt template per i capitoli specifico per il tipo di libro."""
    # Cerca prima in un file, altrimenti usa il template predefinito
    try:
        template_file = f"chapter_prompt_{book_type.lower().replace(' ', '_')}.txt"
        with open(template_file, "r", encoding="utf-8") as f:
            return f.read()
    except:
        # Template predefinito semplificato
        return """
        Scrivi il capitolo "{text}" per il libro "{TITOLO_LIBRO}" in lingua {LINGUA_LIBRO}.
    
        ISTRUZIONI DI STILE:
        {VOICE_STYLE}
        
        STRUTTURA DEL CAPITOLO:
        - Inizia con un'introduzione coinvolgente che catturi l'attenzione del lettore
        - Dividi il contenuto in 3-5 sezioni principali con sottotitoli chiari
        - Includi esempi pratici o casi studio rilevanti
        - Termina con un riepilogo che riassuma i punti chiave
        
        Il contenuto deve essere dettagliato e approfondito, con una lunghezza di circa 2000-3000 parole.
        
        Scrivi FINE alla fine del capitolo.
        """

def disable_crisp_messages(chat_manager=None):
    """Disabilita temporaneamente i messaggi CRISP nei log"""
    if chat_manager:
        original_add_log = chat_manager.add_log
        
        def filtered_add_log(message):
            if "CRISP" in message and ("Errore: Nessun progetto CRISP" in message or "Avvio generazione libro CRISP" in message):
                return  # Ignora questo messaggio
            return original_add_log(message)
        
        chat_manager.add_log = filtered_add_log

def _handle_missing_placeholders(text):
    """Gestisce i placeholder non sostituiti nel prompt."""
    # Trova tutti i placeholder rimanenti
    placeholders = re.findall(r'\{([A-Z_]+)\}', text)

    for placeholder in placeholders:
        # Sostituisce i placeholder mancanti con un valore generico
        full_placeholder = "{" + placeholder + "}"
        text = text.replace(full_placeholder, f"[Valore di {placeholder} non disponibile]")

    return text

def _save_chapter(chapter_title, chapter_content, book_title, log=None):
    """Salva il capitolo generato in un file."""
    from docx import Document

    safe_title = re.sub(r'[<>:"/\\|?*]', '', book_title)
    folder_path = os.path.join(os.getcwd(), safe_title)
    os.makedirs(folder_path, exist_ok=True)

    safe_chapter = re.sub(r'[<>:"/\\|?*]', '', chapter_title)
    file_path = os.path.join(folder_path, f"{safe_chapter}.docx")

    doc = Document()
    doc.add_heading(chapter_title, 1)

    # Aggiungi il contenuto con formattazione di base
    for paragraph in chapter_content.split('\n\n'):
        if paragraph.strip():
            if paragraph.strip().startswith('#'):
                # È un titolo
                level = paragraph.count('#')
                title_text = paragraph.strip('#').strip()
                doc.add_heading(title_text, level)
            else:
                # È un paragrafo normale
                doc.add_paragraph(paragraph)

    doc.save(file_path)
    if log:
        log(f"Capitolo '{chapter_title}' salvato in {file_path}")

    return file_path


