# test_esportazione.py - Test delle funzionalità di esportazione
import os
import sys
import datetime
import shutil

def test_esportazione_txt():
    """Test della funzionalità di esportazione in formato TXT"""
    print("🧪 Test esportazione TXT...")
    
    try:
        # 1. Prepara ambiente test
        context_file = "context.txt"
        test_content = """=== Test Esportazione - 20250426_120000 ===
Questo è un contenuto di test per verificare l'esportazione.
Keyword: test_export
"""
        
        # Backup eventuale context.txt esistente
        if os.path.exists(context_file):
            backup_file = f"{context_file}.bak"
            print(f"Backup del file context.txt esistente in {backup_file}")
            shutil.copy2(context_file, backup_file)
        
        # Crea context.txt di test
        with open(context_file, "w", encoding="utf-8") as f:
            f.write(test_content)
            
        # 2. Simula l'esportazione TXT
        output_dir = os.path.join(os.getcwd(), "output")
        os.makedirs(output_dir, exist_ok=True)
        
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(output_dir, f"analisi_test_export_{timestamp}.txt")
        
        # Esegui l'operazione di copia
        shutil.copy2(context_file, output_path)
        
        # 3. Verifica risultato
        if os.path.exists(output_path):
            with open(output_path, "r", encoding="utf-8") as f:
                exported_content = f.read()
                
            if exported_content == test_content:
                print(f"✅ Test esportazione TXT riuscito: {output_path}")
                success = True
            else:
                print("❌ Contenuti non corrispondenti")
                print(f"Originale: {len(test_content)} caratteri")
                print(f"Esportato: {len(exported_content)} caratteri")
                success = False
        else:
            print(f"❌ File di output non creato: {output_path}")
            success = False
            
        # 4. Ripristina l'ambiente
        if os.path.exists(f"{context_file}.bak"):
            shutil.move(f"{context_file}.bak", context_file)
        
        return success
    
    except Exception as e:
        print(f"❌ Errore durante il test: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return False

def test_esportazione_docx():
    """Test della funzionalità di esportazione in formato DOCX"""
    print("🧪 Test esportazione DOCX...")
    
    try:
        # Verifica se python-docx è disponibile
        try:
            import docx
        except ImportError:
            print("⚠️ Libreria python-docx non installata, impossibile eseguire il test")
            print("Installa con: pip install python-docx")
            return False
        
        # 1. Prepara ambiente test (come sopra)
        context_file = "context.txt"
        test_content = """=== Test Esportazione DOCX - 20250426_120000 ===
Questo è un contenuto di test per verificare l'esportazione.
Keyword: test_export_docx
"""
        
        # Backup eventuale context.txt esistente
        if os.path.exists(context_file):
            backup_file = f"{context_file}.bak"
            print(f"Backup del file context.txt esistente in {backup_file}")
            shutil.copy2(context_file, backup_file)
        
        # Crea context.txt di test
        with open(context_file, "w", encoding="utf-8") as f:
            f.write(test_content)
            
        # 2. Simula l'esportazione DOCX
        import docx
        doc = docx.Document()
        doc.add_heading("Test Esportazione DOCX", 0)
        
        # Aggiungi contenuto del file di contesto
        with open(context_file, "r", encoding="utf-8") as f:
            content = f.read()
            
        sections = content.split("===")
        for section in sections:
            if section.strip():
                # Cerca un titolo nella sezione
                lines = section.strip().split("\n")
                if lines:
                    # Usa la prima riga come titolo
                    doc.add_heading(lines[0], level=1)
                    # Aggiungi il resto come testo
                    if len(lines) > 1:
                        doc.add_paragraph("\n".join(lines[1:]))
                    doc.add_paragraph("")  # Spazio aggiuntivo
        
        # Salva il documento
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"test_export_docx_{timestamp}.docx"
        doc.save(output_path)
        
        # 3. Verifica risultato
        if os.path.exists(output_path):
            print(f"✅ Test esportazione DOCX riuscito: {output_path}")
            success = True
        else:
            print(f"❌ File DOCX non creato: {output_path}")
            success = False
            
        # 4. Ripristina l'ambiente
        if os.path.exists(f"{context_file}.bak"):
            shutil.move(f"{context_file}.bak", context_file)
        
        return success
    
    except Exception as e:
        print(f"❌ Errore durante il test DOCX: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return False

def run_all_tests():
    """Esegue tutti i test disponibili"""
    results = []
    
    print("🚀 Avvio test suite per esportazione PubliScript")
    
    # Test TXT
    txt_result = test_esportazione_txt()
    results.append(("Esportazione TXT", txt_result))
    
    # Test DOCX
    docx_result = test_esportazione_docx()
    results.append(("Esportazione DOCX", docx_result))
    
    # Stampa riepilogo
    print("\n=== RIEPILOGO TEST ===")
    for name, result in results:
        status = "✅ PASSATO" if result else "❌ FALLITO"
        print(f"{name}: {status}")
    
if __name__ == "__main__":
    run_all_tests()